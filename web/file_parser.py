#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文件解析器 - 支持 PDF/DOCX/TXT/MD 文件读取
用于 PPT 生成的多源文件融合
"""

import os
import json
from typing import Optional, Dict, List
from pathlib import Path


class FileParser:
    """多格式文件解析器"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.md', '.markdown']
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB 上限
    MAX_CONTENT_LENGTH = 100000  # 提取最多 10 万字符
    
    @staticmethod
    def parse_file(file_path: str) -> Optional[Dict[str, any]]:
        """
        解析单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            {
                "success": bool,
                "filename": str,
                "format": str,
                "content": str,  # 提取的文本内容
                "char_count": int,
                "error": str (if failed)
            }
        """
        if not os.path.exists(file_path):
            return {"success": False, "error": f"文件不存在: {file_path}"}
        
        file_size = os.path.getsize(file_path)
        if file_size > FileParser.MAX_FILE_SIZE:
            return {"success": False, "error": f"文件过大 ({file_size/1024/1024:.1f}MB > 50MB)"}
        
        file_ext = Path(file_path).suffix.lower()
        filename = os.path.basename(file_path)
        
        if file_ext not in FileParser.SUPPORTED_FORMATS:
            return {"success": False, "error": f"不支持的格式: {file_ext}"}
        
        try:
            if file_ext == '.pdf':
                content = FileParser._parse_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = FileParser._parse_docx(file_path)
            elif file_ext in ['.txt', '.md', '.markdown']:
                content = FileParser._parse_text(file_path)
            else:
                return {"success": False, "error": "未知格式"}
            
            # 截断超长内容
            if len(content) > FileParser.MAX_CONTENT_LENGTH:
                content = content[:FileParser.MAX_CONTENT_LENGTH] + "\n\n[内容已截断]"
            
            return {
                "success": True,
                "filename": filename,
                "format": file_ext.lstrip('.'),
                "content": content,
                "char_count": len(content)
            }
        
        except Exception as e:
            return {
                "success": False,
                "filename": filename,
                "error": f"解析失败: {str(e)}"
            }
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """PDF 文本提取（尝试 pypdf → PyPDF2 → pdfplumber）"""
        content = []

        # 优先: pypdf（PyPDF2 的继任者，纯Python，已作为 pypdf 包发布）
        for pkg_name, mod_name in [("pypdf", "pypdf"), ("PyPDF2", "PyPDF2")]:
            try:
                mod = __import__(mod_name)
                PdfReader = getattr(mod, "PdfReader")
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    for page_num, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            content.append(f"[第 {page_num + 1} 页]\n{text}")
                if content:
                    return "\n\n".join(content)
            except ImportError:
                continue
            except Exception:
                break  # 库可用但解析失败，尝试 pdfplumber

        # 回退: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        content.append(f"[第 {page_num + 1} 页]\n{text}")
            return "\n\n".join(content)
        except ImportError:
            raise ImportError("需要安装 pypdf 或 pdfplumber: pip install pypdf")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """DOCX 文本提取"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 也提取表格
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_data))
            if table_content:
                content.append("\n".join(table_content))
        
        return "\n".join(content)
    
    @staticmethod
    def _parse_text(file_path: str) -> str:
        """纯文本/Markdown 读取"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def batch_parse(file_paths: List[str]) -> List[Dict]:
        """
        批量解析多个文件，融合为统一格式
        
        Returns:
            [
                {
                    "filename": str,
                    "format": str,
                    "content": str,
                    "success": bool,
                    "error": str (if failed)
                },
                ...
            ]
        """
        results = []
        for path in file_paths:
            result = FileParser.parse_file(path)
            results.append(result)
        return results
    
    @staticmethod
    def merge_contents(parse_results: List[Dict]) -> str:
        """
        将多个文件的内容合并为统一的参考材料格式
        
        Args:
            parse_results: batch_parse 的返回结果
            
        Returns:
            合并后的文本（带来源标记）
        """
        merged = []
        for i, result in enumerate(parse_results, 1):
            if result.get("success"):
                filename = result.get("filename", f"文件{i}")
                format_type = result.get("format", "unknown")
                content = result.get("content", "")
                
                merged.append(
                    f"【来源文件 {i}】{filename} ({format_type})\n"
                    f"{'=' * 60}\n"
                    f"{content}\n"
                    f"{'=' * 60}\n"
                )
        
        return "\n\n".join(merged)
    
    @staticmethod
    def sanitize_file_path(file_path: str) -> Optional[str]:
        """
        检查文件路径是否安全（防止路径遍历攻击）
        
        返回规范化的绝对路径，或 None if 不安全
        """
        try:
            abs_path = os.path.abspath(file_path)
            
            # 确保文件在允许的目录内
            allowed_dirs = [
                os.path.abspath(os.path.dirname(__file__)),  # web/ 目录
                os.path.abspath(os.path.join(os.path.dirname(__file__), '..')),  # 项目根目录
            ]
            
            if not any(abs_path.startswith(d) for d in allowed_dirs):
                return None
            
            return abs_path
        except:
            return None
