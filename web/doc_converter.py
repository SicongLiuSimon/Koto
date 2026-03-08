#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
doc_converter.py — 多格式文档转换器
将 .doc / .pdf / .txt / .md / .rtf / .odt 转换为 .docx，
以便 DocumentFeedbackSystem（标注系统）进行处理。

支持格式:
  .doc   → 先尝试 python-docx 直读（部分老版 .doc 实为 OOXML），
            再试 LibreOffice CLI，最后用 docx2txt 提取文本重建
  .pdf   → pypdf 提取文本 → 重建 .docx
  .txt   → 直接读取 → 重建 .docx
  .md    → Markdown 解析标题/段落 → 重建 .docx（带样式）
  .rtf   → striprtf 提取文本 → 重建 .docx
  .odt   → LibreOffice CLI（优先）→ 文本提取回退
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

# ──────────────────────────────────────────────────────────────────────────────
# 支持的输入格式 & 对应的 MIME 类型（供前端 accept 属性参考）
# ──────────────────────────────────────────────────────────────────────────────
SUPPORTED_INPUT_EXTS = {
    ".docx", ".doc",
    ".pdf",
    ".txt", ".md", ".markdown",
    ".rtf",
    ".odt",
}

ACCEPT_ATTR = (
    ".docx,.doc,.pdf,.txt,.md,.markdown,.rtf,.odt,"
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document,"
    "application/msword,application/pdf,text/plain,text/markdown,"
    "application/rtf,application/vnd.oasis.opendocument.text"
)


# ──────────────────────────────────────────────────────────────────────────────
# 公开接口
# ──────────────────────────────────────────────────────────────────────────────

def needs_conversion(file_ext: str) -> bool:
    """返回该扩展名是否需要在标注前转换为 .docx"""
    return file_ext.lower() in SUPPORTED_INPUT_EXTS - {".docx"}


def convert_to_docx(
    source_path: str,
    output_dir: Optional[str] = None,
) -> Tuple[str, str]:
    """
    将 source_path 转换为 .docx，返回 (docx_path, warning_message)。
    warning_message 非空时表示转换有降级（如格式丢失）。

    output_dir 默认为 source_path 同级目录的临时文件夹。

    Raises:
        ValueError: 不支持的格式
        RuntimeError: 所有转换方法均失败
    """
    source_path = os.path.abspath(source_path)
    ext = Path(source_path).suffix.lower()
    stem = Path(source_path).stem

    if ext == ".docx":
        return source_path, ""          # 已经是 docx，无需转换

    if ext not in SUPPORTED_INPUT_EXTS:
        raise ValueError(f"不支持的格式：{ext}（支持：{', '.join(sorted(SUPPORTED_INPUT_EXTS))}）")

    out_dir = output_dir or os.path.join(os.path.dirname(source_path), "_converted")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{stem}_converted.docx")

    converters = {
        ".doc":      _convert_doc,
        ".pdf":      _convert_pdf,
        ".txt":      _convert_txt,
        ".md":       _convert_md,
        ".markdown": _convert_md,
        ".rtf":      _convert_rtf,
        ".odt":      _convert_odt,
    }

    convert_fn = converters[ext]
    return convert_fn(source_path, out_path)


# ──────────────────────────────────────────────────────────────────────────────
# 内部：各格式转换器
# ──────────────────────────────────────────────────────────────────────────────

def _build_docx_from_text(text: str, out_path: str, title: str = "") -> str:
    """
    从纯文本重建一个带基本样式的 .docx 文件。
    空行分段；以 # 开头的行映射到标题样式。
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    for line in text.splitlines():
        stripped = line.rstrip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # Markdown-style headings inside plain text
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(11)

    doc.save(out_path)
    return out_path


def _try_libreoffice(source_path: str, out_dir: str) -> Optional[str]:
    """尝试用 LibreOffice --headless 转换，成功返回生成的 .docx 路径否则返回 None"""
    for soffice in ("soffice", "libreoffice", r"C:\Program Files\LibreOffice\program\soffice.exe"):
        if not (shutil.which(soffice) or os.path.exists(soffice)):
            continue
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", out_dir, source_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                expected = os.path.join(out_dir, Path(source_path).stem + ".docx")
                if os.path.exists(expected):
                    return expected
        except Exception:
            pass
    return None


# ── .doc ──────────────────────────────────────────────────────────────────────

def _convert_doc(source_path: str, out_path: str) -> Tuple[str, str]:
    warning = ""

    # 方法 1: 直接用 python-docx（部分 .doc 实为 OOXML 包，只是扩展名错误）
    try:
        from docx import Document
        doc = Document(source_path)
        doc.save(out_path)
        print(f"[DocConverter] .doc → .docx via python-docx ✓")
        return out_path, warning
    except Exception:
        pass

    # 方法 2: LibreOffice CLI
    lo_result = _try_libreoffice(source_path, os.path.dirname(out_path))
    if lo_result:
        if lo_result != out_path:
            shutil.move(lo_result, out_path)
        print(f"[DocConverter] .doc → .docx via LibreOffice ✓")
        return out_path, warning

    # 方法 3: docx2txt 文本提取 → 重建
    try:
        import docx2txt
        text = docx2txt.process(source_path)
        if text and text.strip():
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            warning = "⚠️ `.doc` 文件格式较旧，已提取文本并重建为 `.docx`（原有图片/表格样式可能丢失）。"
            print(f"[DocConverter] .doc → .docx via docx2txt text extraction ✓ (degraded)")
            return out_path, warning
    except Exception:
        pass

    # 方法 4: 以 UTF-8 / Latin-1 强行读取二进制文本（最后手段）
    try:
        with open(source_path, 'rb') as f:
            raw = f.read()
        # 尝试提取可读 ASCII / UTF-8 片段 (最少8字符，避免太多短碎片)
        text_fragments = re.findall(rb'[\x20-\x7e\n\r\t]{8,}', raw)
        # 过滤：只保留字母比例 > 35% 的有意义片段（排除乱码）
        _meaningful = []
        _seen = set()
        for _frag in text_fragments:
            _s = _frag.decode('latin-1').strip()
            if not _s:
                continue
            _alpha = sum(1 for c in _s if c.isalpha())
            if _alpha / max(len(_s), 1) < 0.35:
                continue  # 乱码（字母比例过低）
            _key = _s[:60].lower()
            if _key in _seen:
                continue  # 去重
            _seen.add(_key)
            _meaningful.append(_s)
            if len(_meaningful) >= 800:  # 限制最多800个片段，防止爆炸
                break
        text = "\n".join(_meaningful)
        if len(text.strip()) > 50:
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            warning = "⚠️ `.doc` 格式无法完整解析，仅提取了可读文本片段，部分内容可能缺失。"
            print(f"[DocConverter] .doc → .docx via raw binary extraction ✓ (highly degraded, {len(_meaningful)} fragments)")
            return out_path, warning
    except Exception:
        pass

    raise RuntimeError(
        "无法转换 `.doc` 文件。请用 Microsoft Word 另存为 `.docx` 后重新上传，"
        "或安装 LibreOffice 以支持自动转换。"
    )


# ── .pdf ──────────────────────────────────────────────────────────────────────

def _convert_pdf(source_path: str, out_path: str) -> Tuple[str, str]:
    warning = "⚠️ PDF 已转换为可编辑 `.docx`（纯文本重建，原有图片/复杂排版已忽略）。"

    # 方法 1: LibreOffice（最佳质量）
    lo_result = _try_libreoffice(source_path, os.path.dirname(out_path))
    if lo_result:
        if lo_result != out_path:
            shutil.move(lo_result, out_path)
        print(f"[DocConverter] .pdf → .docx via LibreOffice ✓")
        return out_path, "⚠️ PDF 已由 LibreOffice 转换为 .docx（排版可能略有偏差）。"

    # 方法 2: pypdf
    try:
        import pypdf
        pages = []
        with open(source_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[第 {i+1} 页]\n{text}")
        text = "\n\n".join(pages)
        if text.strip():
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            print(f"[DocConverter] .pdf → .docx via pypdf ✓ ({len(pages)} pages)")
            return out_path, warning
    except Exception as e:
        print(f"[DocConverter] pypdf failed: {e}")

    raise RuntimeError("PDF 文本提取失败。请确保文件不是扫描件（图片 PDF）。")


# ── .txt ──────────────────────────────────────────────────────────────────────

def _convert_txt(source_path: str, out_path: str) -> Tuple[str, str]:
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            with open(source_path, encoding=enc) as f:
                text = f.read()
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            print(f"[DocConverter] .txt → .docx ✓ (encoding={enc})")
            return out_path, ""
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise RuntimeError(f"TXT 转换失败: {e}") from e
    raise RuntimeError("无法以任何已知编码读取该 TXT 文件。")


# ── .md ───────────────────────────────────────────────────────────────────────

def _convert_md(source_path: str, out_path: str) -> Tuple[str, str]:
    for enc in ("utf-8", "utf-8-sig", "gbk"):
        try:
            with open(source_path, encoding=enc) as f:
                text = f.read()
            # _build_docx_from_text 已支持 Markdown 标题语法
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            print(f"[DocConverter] .md → .docx ✓")
            return out_path, ""
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise RuntimeError(f"Markdown 转换失败: {e}") from e
    raise RuntimeError("无法读取该 Markdown 文件（编码问题）。")


# ── .rtf ──────────────────────────────────────────────────────────────────────

def _convert_rtf(source_path: str, out_path: str) -> Tuple[str, str]:
    warning = "⚠️ RTF 已转换为 `.docx`（文本重建，原有复杂格式可能丢失）。"

    # 方法 1: LibreOffice
    lo_result = _try_libreoffice(source_path, os.path.dirname(out_path))
    if lo_result:
        if lo_result != out_path:
            shutil.move(lo_result, out_path)
        return out_path, "⚠️ RTF 已由 LibreOffice 转换为 .docx。"

    # 方法 2: striprtf
    try:
        from striprtf.striprtf import rtf_to_text
        with open(source_path, encoding='latin-1') as f:
            rtf_content = f.read()
        text = rtf_to_text(rtf_content)
        if text and text.strip():
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            print(f"[DocConverter] .rtf → .docx via striprtf ✓")
            return out_path, warning
    except Exception as e:
        print(f"[DocConverter] striprtf failed: {e}")

    raise RuntimeError("RTF 转换失败，请安装 LibreOffice 或手动另存为 .docx。")


# ── .odt ──────────────────────────────────────────────────────────────────────

def _convert_odt(source_path: str, out_path: str) -> Tuple[str, str]:
    # ODT 方法 1: LibreOffice（首选）
    lo_result = _try_libreoffice(source_path, os.path.dirname(out_path))
    if lo_result:
        if lo_result != out_path:
            shutil.move(lo_result, out_path)
        print(f"[DocConverter] .odt → .docx via LibreOffice ✓")
        return out_path, ""

    # 方法 2: ODT 是 ZIP，直接提取 content.xml 文本
    try:
        import zipfile, xml.etree.ElementTree as ET
        with zipfile.ZipFile(source_path) as z:
            with z.open("content.xml") as xf:
                tree = ET.parse(xf)
        ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
        paragraphs = [
            "".join(n.text or "" for n in p.iter())
            for p in tree.findall(".//text:p", ns)
        ]
        text = "\n".join(paragraphs)
        if text.strip():
            _build_docx_from_text(text, out_path, title=Path(source_path).stem)
            print(f"[DocConverter] .odt → .docx via ZIP/XML ✓")
            return out_path, "⚠️ ODT 已提取文本并重建为 .docx（原有格式可能丢失）。"
    except Exception as e:
        print(f"[DocConverter] ODT ZIP extraction failed: {e}")

    raise RuntimeError("ODT 转换失败，请安装 LibreOffice 以支持自动转换。")
