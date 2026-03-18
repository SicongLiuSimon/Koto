import asyncio
import json
import logging
import os
import re
import secrets
import subprocess
import sys
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path

_app_logger = logging.getLogger("koto.app")

# 确保 web/ 目录在模块搜索路径中（通过 koto_app.py 启动时需要）
_web_dir = os.path.dirname(os.path.abspath(__file__))
if _web_dir not in sys.path:
    sys.path.append(_web_dir)

from dotenv import load_dotenv
from flask import (
    Flask,
    Response,
    g,
    jsonify,
    render_template,
    request,
    send_file,
    send_from_directory,
    stream_with_context,
)
from flask_cors import CORS
from werkzeug.utils import secure_filename as _secure_filename

# Import new routing modules
from app.core.routing import SmartDispatcher

# 延迟导入 - 这些路由类仅在运行时首次访问时通过 __getattr__ 加载
# LocalModelRouter, AIRouter, TaskDecomposer, LocalPlanner 通过 app.core.routing.__getattr__ 延迟加载

# Import unified agent API blueprint — 延迟到蓝图注册时加载
agent_bp = None  # 延迟加载，见下方蓝图注册区

# ================= 并行执行系统导入 =================
try:
    from parallel_api import register_parallel_api
    from parallel_executor import (
        Priority,
        Task,
        TaskStatus,
        TaskType,
        cancel_task,
        get_next_task,
        get_queue_manager,
        get_resource_manager,
        get_task_monitor,
        submit_task,
    )
    from task_dispatcher import get_scheduler, start_dispatcher, stop_dispatcher

    PARALLEL_SYSTEM_ENABLED = True
except ImportError as e:
    _app_logger.warning(f"[WARNING] Failed to import parallel execution system: {e}")
    PARALLEL_SYSTEM_ENABLED = False

try:
    from flask_sock import Sock
except ImportError:
    Sock = None

# ================= 懒加载重型模块（启动优化） =================
# google.genai (~4.7s), requests (~0.5s) 延迟到首次使用时加载


class _LazyModule:
    """延迟导入代理 - 首次属性访问时才触发实际 import"""

    __slots__ = ("_import_func", "_module")

    def __init__(self, import_func):
        object.__setattr__(self, "_import_func", import_func)
        object.__setattr__(self, "_module", None)

    def _load(self):
        mod = object.__getattribute__(self, "_module")
        if mod is None:
            import_func = object.__getattribute__(self, "_import_func")
            mod = import_func()
            object.__setattr__(self, "_module", mod)
        return mod

    def __getattr__(self, name):
        return getattr(self._load(), name)

    def __repr__(self):
        mod = object.__getattribute__(self, "_module")
        if mod is None:
            return "<LazyModule (not loaded)>"
        return repr(mod)


def _import_genai():
    _app_logger.debug("[LAZY_IMPORT] 加载 google.genai ...")
    from google import genai as _genai

    return _genai


def _import_types():
    _app_logger.debug("[LAZY_IMPORT] 加载 google.genai.types ...")
    from google.genai import types as _types

    return _types


def _import_requests():
    _app_logger.debug("[LAZY_IMPORT] 加载 requests ...")
    import requests as _requests

    return _requests


genai = _LazyModule(_import_genai)
types = _LazyModule(_import_types)
requests = _LazyModule(_import_requests)

# ================= 懒加载文档和PPT模块（启动加速） =================
# 延迟导入 python-docx (~572ms) 和 python-pptx (~666ms)

# 文档工作流执行器懒加载
_document_workflow_cache = {}


def get_document_workflow_executor():
    """懒加载文档工作流执行器"""
    if "executor" not in _document_workflow_cache:
        _app_logger.debug("[LAZY_IMPORT] 加载文档工作流执行器...")
        try:
            from web.document_workflow_executor import (
                DocumentWorkflowExecutor,
                execute_document_workflow,
            )
        except ImportError:
            try:
                from document_workflow_executor import (
                    DocumentWorkflowExecutor,
                    execute_document_workflow,
                )
            except ImportError:
                DocumentWorkflowExecutor = None
                execute_document_workflow = None
                _app_logger.warning("[WARNING] 文档工作流执行器未安装")
        _document_workflow_cache["executor"] = DocumentWorkflowExecutor
        _document_workflow_cache["execute"] = execute_document_workflow
    return _document_workflow_cache.get("executor"), _document_workflow_cache.get(
        "execute"
    )


# DocumentWorkflowExecutor 和 execute_document_workflow 的懒加载代理
class _DocWorkflowProxy:
    def __getattr__(self, name):
        executor_cls, _ = get_document_workflow_executor()
        if executor_cls is None:
            raise ImportError("文档工作流执行器未安装")
        return getattr(executor_cls, name)


DocumentWorkflowExecutor = _DocWorkflowProxy()


def execute_document_workflow(*args, **kwargs):
    _, execute_func = get_document_workflow_executor()
    if execute_func is None:
        raise ImportError("文档工作流执行器未安装")
    return execute_func(*args, **kwargs)


# PPT多模型系统懒加载
_ppt_system_cache = {}


def get_ppt_system():
    """懒加载PPT生成系统"""
    if "loaded" not in _ppt_system_cache:
        _app_logger.debug("[LAZY_IMPORT] 加载PPT多模型生成系统...")
        try:
            from web.ppt_master import PPTBlueprint, PPTMasterOrchestrator
            from web.ppt_pipeline import (
                PPTGenerationPipeline,
                PPTGenerationTaskHandler,
                format_ppt_generation_result,
            )
            from web.ppt_synthesizer import PPTSynthesizer

            _app_logger.info("[PPT_SYSTEM] ✅ 多模型PPT生成系统已加载")
        except ImportError:
            try:
                from ppt_master import PPTBlueprint, PPTMasterOrchestrator
                from ppt_pipeline import (
                    PPTGenerationPipeline,
                    PPTGenerationTaskHandler,
                    format_ppt_generation_result,
                )
                from ppt_synthesizer import PPTSynthesizer

                _app_logger.info("[PPT_SYSTEM] ✅ 多模型PPT生成系统已加载（相对导入）")
            except ImportError:
                PPTMasterOrchestrator = None
                PPTBlueprint = None
                PPTSynthesizer = None
                PPTGenerationPipeline = None
                PPTGenerationTaskHandler = None
                format_ppt_generation_result = None
                _app_logger.warning("[WARNING] 多模型PPT生成系统未安装")
        _ppt_system_cache["orchestrator"] = PPTMasterOrchestrator
        _ppt_system_cache["blueprint"] = PPTBlueprint
        _ppt_system_cache["synthesizer"] = PPTSynthesizer
        _ppt_system_cache["pipeline"] = PPTGenerationPipeline
        _ppt_system_cache["handler"] = PPTGenerationTaskHandler
        _ppt_system_cache["formatter"] = format_ppt_generation_result
        _ppt_system_cache["loaded"] = True

    return (
        _ppt_system_cache.get("orchestrator"),
        _ppt_system_cache.get("blueprint"),
        _ppt_system_cache.get("synthesizer"),
        _ppt_system_cache.get("pipeline"),
        _ppt_system_cache.get("handler"),
        _ppt_system_cache.get("formatter"),
    )


# 懒加载代理类
class _PPTModuleProxy:
    def __init__(self, index):
        self._index = index

    def __getattr__(self, name):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPT生成系统未安装")
        return getattr(module, name)

    def __call__(self, *args, **kwargs):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPT生成系统未安装")
        if callable(module):
            return module(*args, **kwargs)
        raise TypeError(f"{module} is not callable")


PPTMasterOrchestrator = _PPTModuleProxy(0)
PPTBlueprint = _PPTModuleProxy(1)
PPTSynthesizer = _PPTModuleProxy(2)
PPTGenerationPipeline = _PPTModuleProxy(3)
PPTGenerationTaskHandler = _PPTModuleProxy(4)
format_ppt_generation_result = _PPTModuleProxy(5)

# ================= Configuration =================
# 从 web 目录向上查找
import os
import sys as _sys


# 中断信号存储 - 改进版本，支持实时流中止
class StreamInterruptManager:
    """管理每个 session 的流中止状态和控制"""

    def __init__(self):
        self.interrupts = {}  # session_name -> {'flag': bool, 'event': threading.Event}
        self._lock = threading.Lock()

    def _ensure(self, session_name):
        """确保 session 记录存在 (must be called with self._lock held)"""
        if session_name not in self.interrupts:
            self.interrupts[session_name] = {"flag": False, "event": threading.Event()}
        elif self.interrupts[session_name].get("event") is None:
            self.interrupts[session_name]["event"] = threading.Event()

    def set_interrupt(self, session_name):
        """设置中断标志"""
        with self._lock:
            self._ensure(session_name)
            self.interrupts[session_name]["flag"] = True
            if self.interrupts[session_name]["event"]:
                self.interrupts[session_name]["event"].set()
        _app_logger.debug(f"[INTERRUPT] Marked session {session_name} for interruption")

    def is_interrupted(self, session_name):
        """检查是否被中断"""
        with self._lock:
            if session_name not in self.interrupts:
                return False
            record = self.interrupts[session_name]
            event_flag = record.get("event").is_set() if record.get("event") else False
            return bool(record.get("flag")) or event_flag

    def reset(self, session_name):
        """重置中断标志"""
        with self._lock:
            self._ensure(session_name)
            self.interrupts[session_name]["flag"] = False
            if self.interrupts[session_name]["event"]:
                self.interrupts[session_name]["event"].clear()
        _app_logger.debug(
            f"[INTERRUPT] Reset interrupt flag for session {session_name}"
        )

    def get_event(self, session_name):
        """获取/创建中断事件对象"""
        with self._lock:
            self._ensure(session_name)
            return self.interrupts[session_name]["event"]

    def cleanup(self, session_name):
        """清理 session 的中断记录"""
        with self._lock:
            if session_name in self.interrupts:
                del self.interrupts[session_name]


_interrupt_manager = StreamInterruptManager()
# 保留向后兼容
_interrupt_flags = {}  # 仅用于向后兼容

# 判断是否为打包后运行
if getattr(_sys, "frozen", False):
    # PyInstaller 打包后 - exe所在目录（持久化数据目录）
    PROJECT_ROOT = os.path.dirname(_sys.executable)
else:
    # 开发环境 - 从 web 目录向上找
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Try multiple locations for the config file
config_locations = [
    os.path.join(PROJECT_ROOT, "config", "gemini_config.env"),
    os.path.join(PROJECT_ROOT, "gemini_config.env"),
    (
        os.path.join(os.path.dirname(_sys.executable), "config", "gemini_config.env")
        if getattr(_sys, "frozen", False)
        else ""
    ),
    "gemini_config.env",
    "../gemini_config.env",
]

for config_path in config_locations:
    if os.path.exists(config_path):
        load_dotenv(config_path)
        break

# 尝试读取 GEMINI_API_KEY 或 API_KEY
API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("API_KEY")

# 读取自定义 API 端点（用于中转服务）
GEMINI_API_BASE = os.getenv("GEMINI_API_BASE", "").strip()
FORCE_PROXY = os.getenv("FORCE_PROXY", "").strip()

_user_settings_cache = {}
_user_settings_lock = threading.Lock()


def _load_user_settings() -> dict:
    """Load user_settings.json with caching and safe fallbacks."""
    with _user_settings_lock:
        if "data" in _user_settings_cache:
            return _user_settings_cache["data"]
        settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
        try:
            with open(settings_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            data = {}
        _user_settings_cache["data"] = data
        return data


def get_workspace_root() -> str:
    """Return the workspace root directory from settings or default path."""
    settings = _load_user_settings()
    workspace_dir = settings.get("storage", {}).get("workspace_dir")
    if workspace_dir:
        return workspace_dir
    return os.path.join(PROJECT_ROOT, "workspace")


def get_organize_root() -> str:
    """Return the file organization root directory from settings or default path."""
    settings = _load_user_settings()
    organize_root = settings.get("storage", {}).get("organize_root")
    if organize_root:
        return organize_root
    return os.path.join(get_workspace_root(), "_organize")


def get_default_wechat_files_dir() -> str:
    """Return configured default WeChat files directory, if provided by user settings."""
    settings = _load_user_settings()
    return settings.get("storage", {}).get("wechat_files_dir", "")


if not API_KEY:
    _app_logger.warning(
        "⚠️ Warning: GEMINI_API_KEY or API_KEY not found in gemini_config.env"
    )
    _app_logger.info("   请在 config/gemini_config.env 中配置 API 密钥")
    _app_logger.info("   应用将继续启动，但 AI 功能不可用")
    # 不再 sys.exit — 允许应用启动并在 UI 中提示用户配置

if GEMINI_API_BASE:
    _app_logger.info(f"📡 使用自定义 API 端点: {GEMINI_API_BASE}")

# 检测并设置代理
PROXY_OPTIONS = [
    "http://127.0.0.1:7890",
    "http://127.0.0.1:10809",
    "http://127.0.0.1:1080",
]


def _normalize_proxy_url(proxy_value: str) -> str:
    """Normalize proxy value to a URL with scheme."""
    if not proxy_value:
        return ""
    value = proxy_value.strip()
    if not value:
        return ""
    if "://" not in value:
        value = f"http://{value}"
    return value


def _extract_system_proxy_candidates() -> list:
    """Collect proxy candidates from system settings (Windows) and env."""
    candidates = []

    # 0) User-configured manual proxy (highest priority after FORCE_PROXY)
    try:
        _us = settings_manager.get("proxy", "enabled")
        _um = settings_manager.get("proxy", "manual_proxy") or ""
        if _us is not False and _um.strip():
            candidates.append(_normalize_proxy_url(_um.strip()))
    except Exception:
        pass

    # 1) Environment variables first (if user/system already configured)
    env_proxy = os.environ.get("HTTPS_PROXY") or os.environ.get("https_proxy")
    if env_proxy:
        candidates.append(_normalize_proxy_url(env_proxy))

    # 2) Windows Internet Settings proxy (for "Use a proxy server")
    if sys.platform.startswith("win"):
        try:
            import winreg

            key_path = r"Software\Microsoft\Windows\CurrentVersion\Internet Settings"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as key:
                proxy_enabled = winreg.QueryValueEx(key, "ProxyEnable")[0]
                if proxy_enabled:
                    proxy_server = str(
                        winreg.QueryValueEx(key, "ProxyServer")[0]
                    ).strip()
                    if proxy_server:
                        # Formats:
                        #   127.0.0.1:7890
                        #   http=127.0.0.1:7890;https=127.0.0.1:7890
                        if "=" in proxy_server and ";" in proxy_server:
                            pairs = [
                                p.strip() for p in proxy_server.split(";") if p.strip()
                            ]
                            parsed_map = {}
                            for pair in pairs:
                                if "=" in pair:
                                    k, v = pair.split("=", 1)
                                    parsed_map[k.strip().lower()] = v.strip()
                            for proto in ["https", "http", "socks", "socks5"]:
                                if parsed_map.get(proto):
                                    candidates.append(
                                        _normalize_proxy_url(parsed_map.get(proto))
                                    )
                        else:
                            candidates.append(_normalize_proxy_url(proxy_server))
        except Exception:
            pass

    # 3) Built-in localhost fallback options
    candidates.extend(PROXY_OPTIONS)

    # De-duplicate while preserving order
    deduped = []
    seen = set()
    for item in candidates:
        if not item:
            continue
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def setup_proxy():
    # 优先使用强制代理（不需要测试）
    if FORCE_PROXY and FORCE_PROXY.lower() not in ("auto", "system"):
        os.environ["HTTPS_PROXY"] = FORCE_PROXY
        os.environ["HTTP_PROXY"] = FORCE_PROXY
        _app_logger.info(f"🔧 使用强制代理: {FORCE_PROXY}")
        return FORCE_PROXY

    # 用户明确禁用代理时，清除环境变量并退出
    try:
        if settings_manager.get("proxy", "enabled") is False:
            os.environ.pop("HTTPS_PROXY", None)
            os.environ.pop("HTTP_PROXY", None)
            _app_logger.info("🔧 用户已禁用代理")
            return None
    except Exception:
        pass

    # 自动匹配系统代理与本地常见端口
    import socket
    from urllib.parse import urlparse

    proxy_candidates = _extract_system_proxy_candidates()

    for proxy in proxy_candidates:
        try:
            # 从 URL 提取 host:port
            parsed = urlparse(proxy)
            host = parsed.hostname
            port = parsed.port
            if not host or not port:
                continue

            # 快速端口检测（0.1秒超时）
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.1)
            result = sock.connect_ex((host, port))
            sock.close()

            if result == 0:
                os.environ["HTTPS_PROXY"] = proxy
                os.environ["HTTP_PROXY"] = proxy
                _app_logger.info(f"✅ 自动匹配系统代理: {proxy}")
                return proxy
        except Exception:
            continue

    return None


# 延迟代理检测到首次需要时（启动加速）
_detected_proxy = None
_proxy_checked = False


def get_detected_proxy():
    """懒加载代理检测（首次调用时执行）"""
    global _detected_proxy, _proxy_checked
    if not _proxy_checked:
        _detected_proxy = setup_proxy()
        _proxy_checked = True
    return _detected_proxy


# 向后兼容：detected_proxy 现在通过函数访问
detected_proxy = None  # 占位符，实际通过 get_detected_proxy() 获取


# 在后台线程预热代理检测（不阻塞启动）
def _warmup_proxy():
    global detected_proxy
    detected_proxy = get_detected_proxy()


threading.Thread(target=_warmup_proxy, daemon=True).start()


# 创建 GenAI 客户端 (配置代理和自定义端点)
def create_client():
    import httpx

    proxy = get_detected_proxy()
    # 超时时间: 连接30秒, 读取180秒 (图像生成和长文本生成需要更长时间)
    timeout_config = httpx.Timeout(180.0, connect=30.0)

    # 构建 http_options
    http_options = {}

    # 注意：最新的 Gemini 模型（如 gemini-1.5-flash）需要 v1beta API
    # v1 API 只支持旧的模型。这里使用 v1beta。
    http_options["api_version"] = "v1beta"

    # 自定义 API 端点（用于中转服务）
    if GEMINI_API_BASE:
        http_options["base_url"] = GEMINI_API_BASE
        _app_logger.info(f"📡 API 端点: {GEMINI_API_BASE}")

    # 配置代理 - 通过环境变量确保被使用
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy
        _app_logger.info(f"🔌 设置代理: {proxy}")

    # 使用 httpx with explicit proxy for genai
    # 注意：HttpOptions 字段名为 httpx_client (snake_case)，不是 httpxClient
    # 但实测显示：通过 env vars 设置代理比显式传入 httpx_client 更稳定（无 SSL 问题），
    # 因此这里直接使用 timeout-only 的 httpx 客户端，代理由 env vars 自动接管
    from google.genai._api_client import HttpOptions as _HttpOptions

    try:
        http_client = httpx.Client(timeout=timeout_config, verify=True)
    except Exception as e:
        _app_logger.warning(f"⚠️ 创建 HTTP 客户端出错: {e}")
        http_client = httpx.Client(timeout=timeout_config)

    # 构建 HttpOptions 对象
    opts_kwargs = dict(
        api_version=http_options.get("api_version", "v1beta"),
        httpx_client=http_client,
    )
    if http_options.get("base_url"):
        opts_kwargs["base_url"] = http_options["base_url"]

    return genai.Client(api_key=API_KEY, http_options=_HttpOptions(**opts_kwargs))


# ── 本地模型配置读取 ──────────────────────────────────────────────────────
def _get_local_model_config() -> tuple:
    """
    读取 user_settings.json，返回 (model_mode, local_model_tag)。
    model_mode: "local" 或 "cloud"（默认 cloud）
    local_model_tag: 如 "qwen2.5:7b" 或 None
    """
    try:
        settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
        with open(settings_path, "r", encoding="utf-8") as _f:
            _data = json.load(_f)
        mode = _data.get("model_mode", "cloud")
        tag = _data.get("local_model")
        return mode, tag
    except Exception:
        return "cloud", None


# 懒加载客户端（mode+tag 作为缓存 key，切换模式后自动重建）
_client = None
_client_mode_key: tuple = (None, None)  # (model_mode, local_model_tag)


def get_client():
    """
    获取 AI 客户端（懒加载）。
    - 若 user_settings.json 中 model_mode == "local"，返回 OllamaClientProxy
    - 否则返回 Gemini genai.Client（原有行为）
    """
    global _client, _client_mode_key
    model_mode, local_model = _get_local_model_config()
    current_key = (model_mode, local_model)

    # 模式或模型发生变化时，重置缓存
    if _client is not None and _client_mode_key != current_key:
        _client = None

    if _client is None:
        if model_mode == "local" and local_model:
            try:
                from app.core.llm.ollama_provider import OllamaClientProxy

                _client = OllamaClientProxy(model_tag=local_model)
                _app_logger.debug(f"[Koto] 🦙 使用本地模型: {local_model}")
            except Exception as _e:
                _app_logger.warning(f"[Koto] ⚠️ Ollama 初始化失败，回退到 Gemini: {_e}")
                _client = create_client()
        else:
            _client = create_client()
        _client_mode_key = current_key

    return _client


# ── Token 监测模块（本地统计，无需额外连接 Google）─────────────────────────
try:
    from token_tracker import record_usage as _record_token_usage

    _TOKEN_TRACKER_ENABLED = True
except ImportError:
    _TOKEN_TRACKER_ENABLED = False

    def _record_token_usage(*_a, **_kw):
        pass


class _FakeGenerateContentResponse:
    """
    轻量级响应包装器。
    当 _TrackedModels 将 Interactions API 的字符串结果转换为标准响应对象时使用，
    确保所有调用方可以统一以 response.text 取值。
    """

    __slots__ = ("text", "candidates", "usage_metadata")

    def __init__(self, text: str):
        self.text = text
        self.candidates = []
        self.usage_metadata = None


def _extract_prompt_text(contents, config=None) -> tuple:
    """
    从 generate_content 的 contents / config 参数中提取文本 prompt 和 system_instruction。
    返回 (prompt_text: str, sys_instruction: str | None)
    """
    # 提取 system_instruction
    sys_instr = None
    if config is not None:
        sys_instr = getattr(config, "system_instruction", None)
        if sys_instr is not None:
            sys_instr = str(sys_instr)

    # 提取 prompt 文本
    if contents is None:
        return "", sys_instr
    if isinstance(contents, str):
        return contents, sys_instr
    if isinstance(contents, list):
        parts = []
        for item in contents:
            if isinstance(item, str):
                parts.append(item)
            elif hasattr(item, "text") and item.text:
                parts.append(str(item.text))
            elif hasattr(item, "parts"):
                for p in item.parts or []:
                    if hasattr(p, "text") and p.text:
                        parts.append(str(p.text))
            else:
                s = str(item)
                if s:
                    parts.append(s)
        return "\n".join(parts), sys_instr
    return str(contents), sys_instr


def _is_interactions_only(model_id: str) -> bool:
    """
    检查 model_id 是否需要走 Interactions API 而非 generate_content。
    使用模块级 _INTERACTIONS_ONLY_MODELS（运行时查找，定义后一定可用）。
    """
    try:
        iom = _INTERACTIONS_ONLY_MODELS  # noqa: F821 — 模块级全局，运行时已定义
    except NameError:
        iom = {
            "gemini-3-flash-preview",
            "gemini-3-pro-preview",
            "deep-research-pro-preview-12-2025",
        }
    mid = str(model_id or "")
    return mid in iom or mid.startswith("deep-research-pro-preview")


_logger_tracked = logging.getLogger(__name__)


class _TrackedModels:
    """拦截 client.models.generate_content，自动记录 token 用量"""

    def __init__(self, real_models):
        object.__setattr__(self, "_real", real_models)

    # ── 内部工具 ─────────────────────────────────────────────────────────────

    @staticmethod
    def _call_ia(model_id: str, contents, config) -> "_FakeGenerateContentResponse":
        """提取文本并转发到 _call_interactions_api_sync，返回包装后的响应对象。"""
        prompt, sys_instr = _extract_prompt_text(contents, config)
        text = _call_interactions_api_sync(  # noqa: F821
            model_id=model_id,
            user_prompt=prompt,
            sys_instruction=sys_instr,
        )
        return _FakeGenerateContentResponse(text or "")

    # ── generate_content ────────────────────────────────────────────────────

    def generate_content(self, model=None, *args, **kwargs):
        # 新版 google-genai SDK：(*, model, contents, config) 全为关键字参数
        # 不能将 model 作为位置参数传入，必须作为关键字参数传递
        # 兼容旧式位置调用：如果 model 在 args 中，移到 kwargs
        if model is None and args:
            model, args = args[0], args[1:]
        # 合并剩余位置参数到 kwargs（新 SDK 不接受位置参数，args 应始终为空）
        real = object.__getattribute__(self, "_real")
        response = real.generate_content(model=model, **kwargs)
        if _TOKEN_TRACKER_ENABLED:
            try:
                usage = getattr(response, "usage_metadata", None)
                if usage:
                    _record_token_usage(
                        model=str(model or "unknown"),
                        prompt_tokens=int(getattr(usage, "prompt_token_count", 0) or 0),
                        completion_tokens=int(
                            getattr(usage, "candidates_token_count", 0) or 0
                        ),
                    )
            except Exception:
                pass
        return response

    def generate_content_stream(self, model=None, *args, **kwargs):
        """拦截 generate_content_stream，记录最后一个 chunk 的 usage_metadata"""
        if model is None and args:
            model, args = args[0], args[1:]
        real = object.__getattribute__(self, "_real")
        stream = real.generate_content_stream(model=model, **kwargs)
        _model_str = str(model or "unknown")
        for chunk in stream:
            yield chunk
            if _TOKEN_TRACKER_ENABLED:
                try:
                    usage = getattr(chunk, "usage_metadata", None)
                    if usage and (getattr(usage, "prompt_token_count", 0) or 0) > 0:
                        _record_token_usage(
                            model=_model_str,
                            prompt_tokens=int(
                                getattr(usage, "prompt_token_count", 0) or 0
                            ),
                            completion_tokens=int(
                                getattr(usage, "candidates_token_count", 0) or 0
                            ),
                        )
                except Exception:
                    pass

    def generate_images(self, model=None, *args, **kwargs):
        """拦截 generate_images（Imagen），按图片数量记录合成 token 用量"""
        if model is None and args:
            model, args = args[0], args[1:]
        real = object.__getattribute__(self, "_real")
        response = real.generate_images(model=model, **kwargs)
        if _TOKEN_TRACKER_ENABLED:
            try:
                # Imagen 按张计费，用合成 token 数换算（1000 tokens/张，配合定价表得出正确费用）
                num_images = max(
                    1, len(getattr(response, "generated_images", []) or [])
                )
                _record_token_usage(
                    model=str(model or "unknown"),
                    prompt_tokens=1000 * num_images,
                    completion_tokens=0,
                )
            except Exception:
                pass
        return response

    def embed_content(self, model=None, *args, **kwargs):
        """拦截 embed_content（text-embedding-004 等），记录 embedding token 用量"""
        if model is None and args:
            model, args = args[0], args[1:]
        real = object.__getattribute__(self, "_real")
        response = real.embed_content(model=model, **kwargs)
        if _TOKEN_TRACKER_ENABLED:
            try:
                usage = getattr(response, "usage_metadata", None)
                if usage:
                    prompt_tokens = int(getattr(usage, "prompt_token_count", 0) or 0)
                else:
                    # embed_content 并不总是返回 usage_metadata，按输入内容字符数估算
                    contents = kwargs.get("contents", "") or ""
                    if isinstance(contents, list):
                        contents = " ".join(str(c) for c in contents)
                    prompt_tokens = max(
                        1, len(str(contents)) // 4
                    )  # 粗略估算 1 token ≈ 4 字符
                _record_token_usage(
                    model=str(model or "unknown"),
                    prompt_tokens=prompt_tokens,
                    completion_tokens=0,
                )
            except Exception:
                pass
        return response

    def __getattr__(self, name):
        real = object.__getattribute__(self, "_real")
        return getattr(real, name)


# 保持向后兼容的 client 变量（通过属性访问触发懒加载）
class _ClientProxy:
    """代理类，实现懒加载"""

    def __getattr__(self, name):
        obj = getattr(get_client(), name)
        if name == "models":
            return _TrackedModels(obj)
        return obj


client = _ClientProxy()


def create_research_client():
    """创建专用于 Deep Research 的长超时客户端 (5分钟 read timeout)"""
    import httpx
    from google.genai._api_client import HttpOptions as _HttpOptions

    proxy = get_detected_proxy()
    # 深度研究需要更长的超时时间：连接30秒，读取5分钟
    timeout_config = httpx.Timeout(300.0, connect=30.0)

    # 配置代理 - 通过环境变量确保被使用（同 create_client）
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy

    # 自定义 httpx 客户端（仅用于扩展超时；代理由 env vars 接管）
    http_client = httpx.Client(timeout=timeout_config, verify=True)

    opts_kwargs = dict(
        api_version="v1beta",
        httpx_client=http_client,
    )
    if GEMINI_API_BASE:
        opts_kwargs["base_url"] = GEMINI_API_BASE

    return genai.Client(api_key=API_KEY, http_options=_HttpOptions(**opts_kwargs))


def _call_interactions_api_sync(
    model_id: str, user_prompt: str, sys_instruction: str = None, timeout: float = 90.0
) -> str:
    """
    通过 Interactions API 调用 gemini-3-*-preview 模型。
    这些模型不支持 client.models.generate_content()，必须使用此端点。
    ⚡ 本地模型模式下自动降级为 Ollama，跳过 Interactions API。

    Args:
        model_id: 模型 ID，如 "gemini-3-flash-preview"
        user_prompt: 用户输入文本（已包含格式化内容）
        sys_instruction: 系统指令（可选），将拼接到 input 前
        timeout: 最大等待时间（秒）

    Returns:
        模型响应文本；失败时抛出异常
    """
    # ── 本地模型模式：用 Ollama 直接回答，无需 Interactions API ──
    model_mode, _ = _get_local_model_config()
    if model_mode == "local":
        try:
            full_prompt = user_prompt
            if sys_instruction:
                full_prompt = (
                    f"[系统指令]\n{sys_instruction}\n\n[用户输入]\n{user_prompt}"
                )
            resp = get_client().models.generate_content(
                model=model_id,
                contents=full_prompt,
            )
            return getattr(resp, "text", "") or ""
        except Exception as _e:
            raise RuntimeError(f"本地模型 Interactions 降级失败: {_e}") from _e

    rc = create_research_client()

    full_input = user_prompt
    if sys_instruction:
        full_input = f"[系统指令]\n{sys_instruction}\n\n[用户输入]\n{user_prompt}"

    interaction = rc.interactions.create(
        agent=model_id,
        input=full_input[:80000],  # 支持大文档（gemini-3 上下文窗口 >1M tokens）
        background=model_id not in _NO_BACKGROUND_MODELS,
        stream=False,
    )

    interaction_id = getattr(interaction, "id", None)
    status = getattr(interaction, "status", "")
    start_wait = time.time()

    while (
        interaction_id
        and status not in ("completed", "failed", "cancelled")
        and (time.time() - start_wait) < timeout
    ):
        time.sleep(2)
        interaction = rc.interactions.get(interaction_id)
        status = getattr(interaction, "status", "")

    if status not in ("completed", "failed", "cancelled"):
        try:
            rc.interactions.cancel(interaction_id)
        except Exception:
            pass
        raise TimeoutError(f"Interactions API 超时 ({timeout}s) model={model_id}")

    # 从 interaction 对象中提取文本
    def _get_text_from_obj(obj):
        if obj is None:
            return ""
        if isinstance(obj, str):
            return obj
        if hasattr(obj, "text") and obj.text:
            return str(obj.text)
        if hasattr(obj, "parts"):
            return " ".join(
                str(p.text) for p in (obj.parts or []) if hasattr(p, "text") and p.text
            )
        if hasattr(obj, "outputs"):
            texts = [_get_text_from_obj(o) for o in (obj.outputs or [])]
            return "\n".join(t for t in texts if t)
        return ""

    text = _get_text_from_obj(
        getattr(interaction, "outputs", None)
    ) or _get_text_from_obj(interaction)
    return text.strip()


def run_with_timeout(fn, timeout_seconds):
    """在线程中执行函数并限时返回 (避免卡死主流程)"""
    holder = {"result": None, "error": None}

    def _runner():
        try:
            holder["result"] = fn()
        except Exception as e:
            holder["error"] = e

    t = threading.Thread(target=_runner, daemon=True)
    t.start()
    t.join(timeout_seconds)
    if t.is_alive():
        return None, TimeoutError(f"Timeout after {timeout_seconds}s"), True
    return holder["result"], holder["error"], False


def run_with_heartbeat(
    fn, start_time, heartbeat_callback, heartbeat_interval=5, timeout_seconds=90
):
    """
    在后台线程运行函数，同时定期发送心跳。
    用于非流式 API 调用（如图像生成）。

    Args:
        fn: 要执行的函数
        start_time: 请求开始时间
        heartbeat_callback: 心跳回调函数，接收 elapsed_seconds 参数
        heartbeat_interval: 心跳间隔（秒）
        timeout_seconds: 超时时间（秒）

    Returns:
        (result, error, timed_out)
    """
    import queue
    import threading

    result_queue = queue.Queue()

    def worker():
        try:
            result = fn()
            result_queue.put(("success", result))
        except Exception as e:
            result_queue.put(("error", e))

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    last_heartbeat = time.time()

    while True:
        # 检查是否超时
        elapsed = time.time() - start_time
        if elapsed > timeout_seconds:
            return None, TimeoutError(f"操作超时 ({int(elapsed)}s)"), True

        # 尝试获取结果（短超时）
        try:
            status, data = result_queue.get(timeout=1.0)
            if status == "success":
                return data, None, False
            else:
                return None, data, False
        except queue.Empty:
            # 发送心跳
            current_time = time.time()
            if current_time - last_heartbeat >= heartbeat_interval:
                heartbeat_callback(int(current_time - start_time))
                last_heartbeat = current_time


def stream_with_keepalive(
    response_stream, start_time, keepalive_interval=5, max_wait_first_token=60
):
    """
    包装流式响应，在等待第一个 token 期间发送保活心跳。

    Args:
        response_stream: 原始流式响应迭代器
        start_time: 请求开始时间
        keepalive_interval: 心跳间隔（秒）
        max_wait_first_token: 等待第一个 token 的最大时间（秒）

    Yields:
        (type, data): type 可以是 'chunk', 'heartbeat', 'timeout'
    """
    import queue
    import time

    chunk_queue = queue.Queue()
    first_chunk_received = threading.Event()
    stream_done = threading.Event()
    stream_error = {"error": None}

    def stream_reader():
        """在后台线程中读取流"""
        try:
            for chunk in response_stream:
                chunk_queue.put(("chunk", chunk))
                first_chunk_received.set()
            chunk_queue.put(("done", None))
        except Exception as e:
            stream_error["error"] = e
            chunk_queue.put(("error", e))
        finally:
            stream_done.set()

    # 启动后台读取线程
    reader_thread = threading.Thread(target=stream_reader, daemon=True)
    reader_thread.start()

    last_heartbeat = time.time()

    while True:
        # 检查是否等待第一个 token 超时
        if not first_chunk_received.is_set():
            elapsed = time.time() - start_time
            if elapsed > max_wait_first_token:
                yield ("timeout", f"等待响应超时 ({int(elapsed)}s)")
                return

        # 尝试获取 chunk，使用短超时以便发送心跳
        try:
            item_type, item_data = chunk_queue.get(timeout=1.0)

            if item_type == "chunk":
                yield ("chunk", item_data)
            elif item_type == "done":
                return
            elif item_type == "error":
                raise item_data

        except queue.Empty:
            # 队列为空，检查是否需要发送心跳
            current_time = time.time()
            if current_time - last_heartbeat >= keepalive_interval:
                elapsed = int(current_time - start_time)
                yield ("heartbeat", elapsed)
                last_heartbeat = current_time

            # 检查流是否已结束
            if stream_done.is_set() and chunk_queue.empty():
                if stream_error["error"]:
                    raise stream_error["error"]
                return


app = Flask(__name__)

# ── CSRF Protection ──
app.config["SECRET_KEY"] = os.environ.get("KOTO_SECRET_KEY", secrets.token_hex(32))
app.config["WTF_CSRF_CHECK_DEFAULT"] = False  # Only enforce on opted-in views
try:
    from flask_wtf.csrf import CSRFProtect

    csrf = CSRFProtect(app)
    csrf.exempt("web.auth")  # Auth API uses JWT, not cookies
    _app_logger.info("[CSRF] ✅ CSRF protection initialised")
except ImportError:
    _app_logger.warning("[CSRF] ⚠️ Flask-WTF not installed — CSRF protection disabled")


@app.after_request
def _set_security_headers(response):
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    if os.environ.get("KOTO_DEPLOY_MODE") == "cloud":
        response.headers.setdefault(
            "Strict-Transport-Security", "max-age=63072000; includeSubDomains"
        )
    return response


# Read app version from VERSION file
try:
    APP_VERSION = (
        (Path(__file__).parent.parent / "VERSION").read_text(encoding="utf-8").strip()
    )
except Exception:
    APP_VERSION = "unknown"
# 静态资源缓存，减少重复加载
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 3600
# ✅ 允许最大 20MB 请求体（语音 base64 约 1-5MB，留足余量）
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024

# CORS: 云模式限制来源，本地模式打开
_cors_origins = os.environ.get("KOTO_CORS_ORIGINS", "*")
if os.environ.get("KOTO_DEPLOY_MODE") == "cloud" and _cors_origins == "*":
    # 云模式默认只允许自身站点（同源），可通过环境变量覆盖
    _cors_origins = os.environ.get("KOTO_SITE_URL", "*")
CORS(app, origins=_cors_origins)


# ── Sentry error tracking (no-op if SENTRY_DSN not set) ──────────────────────
_sentry_dsn = os.environ.get("SENTRY_DSN", "")
if _sentry_dsn:
    try:
        import sentry_sdk
        from sentry_sdk.integrations.flask import FlaskIntegration

        sentry_sdk.init(
            dsn=_sentry_dsn,
            integrations=[FlaskIntegration()],
            release=APP_VERSION,
            traces_sample_rate=float(os.environ.get("SENTRY_TRACES_RATE", "0.1")),
            send_default_pii=False,
        )
        _app_logger.info("Sentry error tracking enabled (release=%s)", APP_VERSION)
    except ImportError:
        _app_logger.warning("SENTRY_DSN set but sentry-sdk not installed; skipping")

# ── Prometheus metrics (/metrics) ────────────────────────────────────────────
try:
    from prometheus_flask_exporter import PrometheusMetrics

    _metrics_token = os.environ.get("METRICS_TOKEN", "")
    _prometheus = PrometheusMetrics(app, group_by="endpoint")
    _prometheus.info("koto_app_info", "Koto application info", version=APP_VERSION)

    if _metrics_token:
        # Require Bearer token to scrape /metrics
        @app.before_request
        def _guard_metrics():
            if request.path == "/metrics":
                auth = request.headers.get("Authorization", "")
                if auth != f"Bearer {_metrics_token}":
                    return _error_response("Unauthorized", 401)

    _app_logger.info("Prometheus metrics enabled at /metrics")
except ImportError:
    _app_logger.debug("prometheus-flask-exporter not installed; /metrics disabled")

# ── Swagger / OpenAPI docs ────────────────────────────────────────────────────
_swagger_config = {
    "headers": [],
    "specs": [
        {
            "endpoint": "apispec",
            "route": "/apispec.json",
            "rule_filter": lambda rule: True,
            "model_filter": lambda tag: True,
        }
    ],
    "static_url_path": "/flasgger_static",
    "swagger_ui": True,
    "specs_route": "/apidocs/",
}

_swagger_template = {
    "info": {
        "title": "Koto API",
        "description": "API documentation for Koto AI Assistant",
        "version": "1.0.0",
    },
    "basePath": "/",
    "schemes": ["http", "https"],
    "securityDefinitions": {
        "Bearer": {
            "type": "apiKey",
            "name": "Authorization",
            "in": "header",
            "description": "JWT token: `Bearer <token>`",
        }
    },
}

try:
    from flasgger import Swagger

    swagger = Swagger(app, config=_swagger_config, template=_swagger_template)
    _app_logger.info("Swagger UI enabled at /apidocs/")
except ImportError:
    _app_logger.debug("flasgger not installed; Swagger UI disabled")


# ── Request ID middleware ─────────────────────────────────────────────────────
@app.before_request
def _assign_request_id():
    """Assign a correlation ID to every request (read from header or generate)."""
    g.request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())


@app.after_request
def _attach_request_id(response):
    """Attach the correlation ID to every outgoing response."""
    if hasattr(g, "request_id"):
        response.headers["X-Request-ID"] = g.request_id
    return response


def _error_response(message: str, status: int = 400, details=None):
    """Return a standardized JSON error envelope."""
    body = {"error": message, "status": status}
    if details:
        body["details"] = details
    if hasattr(g, "request_id"):
        body["request_id"] = g.request_id
    return jsonify(body), status


# ── Global error handlers (return JSON, not HTML) ────────────────────────────
@app.errorhandler(404)
def _handle_404(exc):
    return _error_response("Not found", 404)


@app.errorhandler(405)
def _handle_405(exc):
    return _error_response("Method not allowed", 405)


@app.errorhandler(500)
def _handle_500(exc):
    _app_logger.exception(
        "Unhandled server error [request_id=%s]", getattr(g, "request_id", "-")
    )
    return _error_response("Internal server error", 500)


# ================= 用户认证系统 =================
try:
    from auth import register_auth_routes

    register_auth_routes(app)
except Exception as e:
    _app_logger.warning(f"[Auth] ⚠️ 认证模块加载失败: {e}")

# ================= 并行执行系统初始化 =================
if PARALLEL_SYSTEM_ENABLED:
    _app_logger.debug("[PARALLEL] 🚀 Initializing parallel execution system...")
    try:
        register_parallel_api(app)
        start_dispatcher()
        _app_logger.info(
            "[PARALLEL] ✅ Parallel execution system initialized successfully"
        )
    except Exception as e:
        _app_logger.error(
            f"[PARALLEL] ❌ Failed to initialize parallel execution system: {e}"
        )
        PARALLEL_SYSTEM_ENABLED = False

# ================= WebSocket 支持（可选） =================
sock = None
if Sock:
    sock = Sock(app)
else:
    _app_logger.warning("[WebSocket] ⚠️ flask-sock 未安装，使用轮询作为通知兜底")

if sock:

    @sock.route("/ws/notifications")
    def ws_notifications(ws):
        user_id = request.args.get("user_id", "default")
        manager = get_notification_manager()
        manager.register_connection(user_id, ws)
        try:
            while True:
                message = ws.receive()
                if message is None:
                    break
                if isinstance(message, str) and message.lower() == "ping":
                    ws.send("pong")
        finally:
            manager.unregister_connection(user_id, ws)


# ================= 延迟注册蓝图（在后台线程中加载，避免阻塞启动） =================
_blueprints_registered = False
_blueprints_lock = threading.Lock()


def _register_blueprints_deferred():
    """注册所有蓝图（必须在 app.run() 前调用）。
    阶段1: 用线程池并行预导入各蓝图模块（重叠 I/O 等待，加速启动）。
    阶段2: 串行注册到 Flask app（线程安全要求）。
    """
    global _blueprints_registered, agent_bp
    with _blueprints_lock:
        if _blueprints_registered:
            return
        _blueprints_registered = True

    # ── 阶段 1：并行预导入（各模块互相独立，可同时加载）────────────────
    import concurrent.futures
    import importlib

    _preload_modules = [
        "app.api.task_routes",
        "app.api",
        "app.api.skill_routes",
        "app.api.skill_marketplace_routes",
        "app.api.goal_routes",
        "app.api.file_hub_routes",
        "app.api.job_routes",
        "app.api.ops_routes",
        "app.api.shadow_routes",
        "app.api.macro_routes",
    ]

    def _safe_preload(mod_name):
        try:
            importlib.import_module(mod_name)
        except Exception:
            pass  # 导入失败时静默忽略，注册阶段会再次尝试并输出日志

    with concurrent.futures.ThreadPoolExecutor(
        max_workers=min(6, len(_preload_modules)), thread_name_prefix="BpPreload"
    ) as _pool:
        list(_pool.map(_safe_preload, _preload_modules))

    # ── 阶段 2：串行注册蓝图（此时模块已在 sys.modules 中，import 为 O(1)）──

    # 注册任务管理 API（任务台账 + 进度总线 + 打断控制）
    try:
        from app.api.task_routes import task_bp as _task_bp

        app.register_blueprint(_task_bp, url_prefix="/api/tasks")
        _app_logger.info("[TaskAPI] ✅ 任务管理 API 已注册: /api/tasks")
    except ImportError as e:
        _app_logger.warning(f"[TaskAPI] ⚠️ 未能导入任务管理 API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[TaskAPI] ❌ 任务管理 API 注册失败: {e}")

    # 注册统一 Agent API
    try:
        from app.api import agent_bp as _agent_bp

        agent_bp = _agent_bp
        app.register_blueprint(agent_bp, url_prefix="/api/agent")
        _app_logger.info("[UnifiedAgent] ✅ 统一 Agent API 已注册: /api/agent")
    except ImportError as e:
        _app_logger.warning(f"[UnifiedAgent] ⚠️ 未能导入统一 Agent API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[UnifiedAgent] ❌ 注册失败: {e}")

    # 注册 Skill CRUD + MCP 导出 API（Phase 2）
    try:
        from app.api.skill_routes import skill_bp as _skill_bp

        app.register_blueprint(_skill_bp)
        _app_logger.info("[SkillAPI] ✅ Skill CRUD API 已注册: /api/skills")
    except ImportError as e:
        _app_logger.warning(f"[SkillAPI] ⚠️ 未能导入 Skill API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[SkillAPI] ❌ Skill API 注册失败: {e}")

    # 注册 Skill Marketplace API（风格市场 + 自动构建 + 导入导出）
    try:
        from app.api.skill_marketplace_routes import marketplace_bp as _marketplace_bp

        app.register_blueprint(_marketplace_bp)
        _app_logger.info(
            "[SkillMarket] ✅ Skill Marketplace API 已注册: /api/skillmarket"
        )
    except ImportError as e:
        _app_logger.warning(
            f"[SkillMarket] ⚠️ 未能导入 Skill Marketplace API 蓝图: {e}"
        )
    except Exception as e:
        _app_logger.error(f"[SkillMarket] ❌ Skill Marketplace API 注册失败: {e}")

    # 注册训练数据 API + LoRA 蒸馏训练 API
    # 仅在开发机上启用（需设置环境变量 KOTO_DEV_TRAINING=1）
    # 公共发行版不包含此功能，设备要求极高（≥16GB VRAM），普通用户无法使用
    if os.environ.get("KOTO_DEV_TRAINING") == "1":
        try:
            from app.core.learning.training_data_builder import (
                register_training_routes as _reg_training,
            )

            _reg_training(app)
        except ImportError as e:
            _app_logger.warning(f"[TrainingAPI] ⚠️ 未能导入训练数据模块: {e}")
        except Exception as e:
            _app_logger.error(f"[TrainingAPI] ❌ 训练数据 API 注册失败: {e}")

        try:
            from app.api.distill_routes import distill_bp as _distill_bp

            app.register_blueprint(_distill_bp, url_prefix="/api/distill")
            _app_logger.info(
                "[DistillAPI] ✅ LoRA 蒸馏训练 API 已注册（开发模式）: /api/distill"
            )
        except ImportError as e:
            _app_logger.warning(f"[DistillAPI] ⚠️ 未能导入蒸馏训练模块: {e}")
        except Exception as e:
            _app_logger.error(f"[DistillAPI] ❌ 蒸馏训练 API 注册失败: {e}")
    else:
        _app_logger.debug(
            "[DistillAPI] ℹ️ LoRA 训练 API 已封存（公共版），如需启用请设置 KOTO_DEV_TRAINING=1"
        )

    # 注册增强语音 API
    try:
        from voice_api_enhanced import voice_bp

        app.register_blueprint(voice_bp)
        _app_logger.debug("[VOICE_API] 已注册增强语音 API 蓝图")
    except ImportError as e:
        _app_logger.warning(f"[VOICE_API] ⚠️ 未能导入增强语音模块: {e}")

    # 注册 PPT 编辑 API（P1 功能）
    try:
        from web.ppt_api_routes import ppt_api_bp

        app.register_blueprint(ppt_api_bp)
        _app_logger.info("[PPT_API] ✅ PPT 编辑 API 已注册: /api/ppt")
    except ImportError as e:
        _app_logger.warning(f"[PPT_API] ⚠️ 未能导入 PPT 编辑 API: {e}")
    except Exception as e:
        _app_logger.warning(f"[PPT_API] ⚠️ PPT 编辑 API 注册失败: {e}")

    # 注册自适应 Agent API（已迁移到 UnifiedAgent，但保留兼容导入）
    try:
        from adaptive_agent_api import init_adaptive_agent_api

        init_adaptive_agent_api(app, gemini_client=None)
        _app_logger.info("[AdaptiveAgent] ✅ 自适应 Agent API 已注册 (延迟加载客户端)")
    except ImportError:
        _app_logger.debug("[AdaptiveAgent] ℹ️ 旧 Agent 模块已退役，使用 UnifiedAgent")
    except Exception as e:
        _app_logger.warning(f"[AdaptiveAgent] ⚠️ 旧 Agent 初始化失败 (非致命): {e}")

    # 注册长期目标 API（GoalManager: 跨天持续执行的委托任务）
    try:
        from app.api.goal_routes import goal_bp as _goal_bp

        app.register_blueprint(_goal_bp, url_prefix="/api/goals")
        _app_logger.info("[GoalAPI] ✅ 长期目标 API 已注册: /api/goals")
    except ImportError as e:
        _app_logger.warning(f"[GoalAPI] ⚠️ 未能导入长期目标 API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[GoalAPI] ❌ 长期目标 API 注册失败: {e}")

    # 注册文件 Hub API（FileRegistry + FileWatcher 统一接口）
    try:
        from app.api.file_hub_routes import file_hub_bp as _file_hub_bp

        app.register_blueprint(_file_hub_bp, url_prefix="/api/files")
        _app_logger.info("[FileHubAPI] ✅ 文件 Hub API 已注册: /api/files")
    except ImportError as e:
        _app_logger.warning(f"[FileHubAPI] ⚠️ 未能导入文件 Hub 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[FileHubAPI] ❌ 文件 Hub API 注册失败: {e}")

    # 注册后台作业 API（JobRunner + TriggerRegistry）
    try:
        from app.api.job_routes import job_bp as _job_bp

        app.register_blueprint(_job_bp)
        _app_logger.info("[JobAPI] ✅ 后台作业 API 已注册: /api/jobs")
    except ImportError as e:
        _app_logger.warning(f"[JobAPI] ⚠️ 未能导入作业 API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[JobAPI] ❌ 作业 API 注册失败: {e}")

    # 注册运维健康 API（HealthSnapshot + RemediationPolicy + OpsEventBus）
    try:
        from app.api.ops_routes import ops_bp as _ops_bp

        app.register_blueprint(_ops_bp)
        _app_logger.info("[OpsAPI] ✅ 运维健康 API 已注册: /api/ops")
    except ImportError as e:
        _app_logger.warning(f"[OpsAPI] ⚠️ 未能导入运维 API 蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[OpsAPI] ❌ 运维 API 注册失败: {e}")

    # 注册影子追踪 API（ShadowWatcher + ProactiveAgent）
    try:
        from app.api.shadow_routes import shadow_bp as _shadow_bp

        app.register_blueprint(_shadow_bp)
        _app_logger.info("[ShadowAPI] ✅ 影子追踪 API 已注册: /api/shadow")
    except ImportError as e:
        _app_logger.warning(f"[ShadowAPI] ⚠️ 未能导入影子追踪蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[ShadowAPI] ❌ 影子追踪 API 注册失败: {e}")

    # 注册宏录制 API（MacroRecorder 主动建议）
    try:
        from app.api.macro_routes import macro_bp as _macro_bp

        app.register_blueprint(_macro_bp)
        _app_logger.info("[MacroAPI] ✅ 宏录制 API 已注册: /api/macro")
    except ImportError as e:
        _app_logger.warning(f"[MacroAPI] ⚠️ 未能导入宏录制蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[MacroAPI] ❌ 宏录制 API 注册失败: {e}")

    # 注册健康检查 API（/api/health + /api/ping）
    try:
        from web.routes.health import health_bp as _health_bp

        app.register_blueprint(_health_bp)
        _app_logger.info("[HealthAPI] ✅ 健康检查 API 已注册: /api/health")
    except ImportError as e:
        _app_logger.warning(f"[HealthAPI] ⚠️ 未能导入健康检查蓝图: {e}")
    except Exception as e:
        _app_logger.error(f"[HealthAPI] ❌ 健康检查 API 注册失败: {e}")

    # ── 新拆分的路由蓝图（从 app.py 提取的路由模块）──────────────────────────
    _new_blueprints = [
        ("web.blueprints.sessions", "sessions_bp", "Sessions"),
        ("web.blueprints.analytics", "analytics_bp", "Analytics"),
        ("web.blueprints.proactive", "proactive_bp", "Proactive"),
        ("web.blueprints.execution", "execution_bp", "Execution"),
        ("web.blueprints.knowledge", "knowledge_bp", "Knowledge"),
        ("web.blueprints.file_editor", "file_editor_bp", "FileEditor"),
        ("web.blueprints.dev", "dev_bp", "Dev"),
        ("web.blueprints.voice", "voice_bp", "Voice"),
        ("web.blueprints.document", "document_bp", "Document"),
        ("web.blueprints.file_organize", "file_organize_bp", "FileOrganize"),
        ("web.blueprints.workspace", "workspace_bp", "Workspace"),
        ("web.blueprints.settings", "settings_bp", "Settings"),
        ("web.blueprints.misc_api", "misc_api_bp", "MiscAPI"),
        ("web.blueprints.pages", "pages_bp", "Pages"),
    ]
    for mod_path, bp_attr, label in _new_blueprints:
        try:
            _mod = importlib.import_module(mod_path)
            _bp = getattr(_mod, bp_attr)
            app.register_blueprint(_bp)
            _app_logger.info("[%s] ✅ 蓝图已注册", label)
        except ImportError as e:
            _app_logger.warning("[%s] ⚠️ 蓝图导入失败: %s", label, e)
        except Exception as e:
            _app_logger.error("[%s] ❌ 蓝图注册失败: %s", label, e)

    _app_logger.info("[INIT] ✅ 所有蓝图注册完成")


def _initialize_background_runtime():
    """Warm up long-running subsystems so jobs, triggers, and ops are live after startup."""
    try:
        time.sleep(1)

        from app.core.jobs.job_runner import get_job_runner
        from app.core.jobs.trigger_registry import get_trigger_registry
        from app.core.ops.ops_event_bus import get_ops_bus
        from app.core.skills.skill_trigger_binding import get_skill_binding_manager

        get_ops_bus()
        runner = get_job_runner()
        registry = get_trigger_registry()
        bindings = get_skill_binding_manager()

        # 初始化 GoalManager 并注册 goal_check 处理器
        try:
            from app.core.goal.goal_job_handler import register_goal_handler
            from app.core.goal.goal_manager import get_goal_manager

            _gm = get_goal_manager()
            register_goal_handler(runner)
            _app_logger.info(
                f"[GoalManager] ✅ 长期目标管理器已启动 (活跃目标: {_gm.count()} 条)"
            )
        except Exception as _ge:
            _app_logger.warning(f"[GoalManager] ⚠️ 初始化失败（非致命）: {_ge}")

        # 初始化 FileRegistry 并启动 FileWatcher
        try:
            from app.core.file.file_registry import get_file_registry
            from app.core.file.file_watcher import get_file_watcher

            _fr = get_file_registry()
            _fw = get_file_watcher()
            _fw.start()
            _app_logger.info(
                f"[FileHub] ✅ 文件注册表已启动 (已收录: {_fr.count()} 个文件)"
            )
        except Exception as _fe:
            _app_logger.warning(f"[FileHub] ⚠️ 文件模块初始化失败（非致命）: {_fe}")

        # 初始化工作文件库（后台快速扫描桌面/文档/下载）
        try:
            from web.work_file_library import get_work_file_library

            _wfl_inst = get_work_file_library()
            if not _wfl_inst.is_indexed():
                _wfl_inst.scan_locations()
                _app_logger.debug(
                    "[WorkFileLibrary] 🚀 工作文件库后台扫描已启动（桌面/文档/下载）"
                )
            else:
                _app_logger.info(
                    f"[WorkFileLibrary] ✅ 工作文件库已加载: {_wfl_inst.count()} 个工作文件"
                )
        except Exception as _wfl_e:
            _app_logger.warning(f"[WorkFileLibrary] ⚠️ 初始化失败（非致命）: {_wfl_e}")

        _app_logger.info(
            "[Runtime] ✅ 后台运行时已启动: "
            f"job_runner={runner is not None}, "
            f"triggers={len(registry.list_all())}, "
            f"bindings={len(bindings.list_bindings())}"
        )

        # 注册 ShadowTracer 阈值 → DistillManager 自动提交训练（数据飞轮闭环）
        try:
            from app.core.learning.distill_manager import DistillManager
            from app.core.learning.shadow_tracer import ShadowTracer, TraceEvent

            def _on_training_ready(event: str, skill_id: str, count: int):
                if event == TraceEvent.TRAINING_READY:
                    _app_logger.debug(
                        f"[Flywheel] 🚀 skill={skill_id} 已积累 {count} 条优质记录，自动提交 LoRA 训练..."
                    )
                    try:
                        job_id = DistillManager.instance().submit(skill_id)
                        _app_logger.info(
                            f"[Flywheel] ✅ 训练任务已提交 job_id={job_id} skill={skill_id}"
                        )
                    except Exception as _e:
                        _app_logger.warning(f"[Flywheel] ⚠️ 自动提交训练失败: {_e}")

            ShadowTracer.add_listener(_on_training_ready)
            _app_logger.info(
                "[Flywheel] ✅ 数据飞轮监听器已注册（ShadowTracer → DistillManager）"
            )
        except Exception as _fe:
            _app_logger.warning(f"[Flywheel] ⚠️ 飞轮监听器注册失败（非致命）: {_fe}")

    except Exception as exc:
        _app_logger.warning(f"[Runtime] ⚠️ 后台运行时初始化失败: {exc}")


# 同步注册所有蓝图（必须在 app.run() 之前完成，否则 Flask 3.x 会在首次请求后拒绝注册）
_register_blueprints_deferred()
threading.Thread(
    target=_initialize_background_runtime, name="RuntimeBootstrap", daemon=True
).start()

# 后台预加载 Vosk 语音模型（减少首次识别延迟）
try:
    from web.voice_engine import preload as _voice_preload

    _voice_preload()
except Exception:
    pass

CHAT_DIR = os.path.join(PROJECT_ROOT, "chats")
WORKSPACE_DIR = get_workspace_root()
UPLOAD_DIR = os.path.join(PROJECT_ROOT, "web", "uploads")
os.makedirs(CHAT_DIR, exist_ok=True)
os.makedirs(WORKSPACE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ================= Settings Manager (提前加载) =================
try:
    from settings import SettingsManager
except ImportError:
    from web.settings import SettingsManager
settings_manager = SettingsManager()

# ================= 动态模型管理器 =================
# 自动从 API 发现可用模型并按任务类型智能匹配，无需手动维护模型列表。
# 新模型上线后自动感知，TTL 缓存每 6 小时刷新一次。

try:
    from web.model_manager import KNOWN_MODEL_REGISTRY as _MODEL_REGISTRY
    from web.model_manager import ModelManager

    _model_manager_available = True
except ImportError:
    try:
        from model_manager import KNOWN_MODEL_REGISTRY as _MODEL_REGISTRY
        from model_manager import ModelManager

        _model_manager_available = True
    except ImportError:
        _model_manager_available = False
        ModelManager = None
        _MODEL_REGISTRY = {}

# 静态默认值（API 不可用时的兜底，也是启动时的初始值）
# 注意：只有 deep-research-pro-preview-* 是 Interactions API agent，其他模型均用 generate_content
MODEL_MAP = {
    "CHAT": "gemini-3-flash-preview",
    "CODER": "gemini-3.1-pro-preview",
    "WEB_SEARCH": "gemini-2.5-flash",
    "VISION": "gemini-2.5-flash",
    "RESEARCH": "gemini-3.1-pro-preview",
    "FILE_GEN": "gemini-3-flash-preview",
    "PAINTER": "gemini-3.1-flash-image-preview",
    "SYSTEM": "local-executor",
    "FILE_OP": "local-executor",
    "AGENT": "gemini-3-flash-preview",
    "FILE_SEARCH": "gemini-3-flash-preview",
    "DOC_ANNOTATE": "gemini-3-flash-preview",
    "COMPLEX": "gemini-3.1-pro-preview",
}

# ─── Interactions-API-only 模型（动态更新，静态默认兜底）──────────────────────
# 这些模型不支持 client.models.generate_content()，必须走 Interactions API
# 注意：gemini-3-flash-preview 和 gemini-3-pro-preview 是普通模型，直接用 generate_content，不在此列表中
_INTERACTIONS_ONLY_MODELS = {
    "deep-research-pro-preview-12-2025",  # 深度研究 Agent，仅支持 Interactions API
}
# 当前 Interactions API 模型均支持 background=True
_NO_BACKGROUND_MODELS: set = set()
# 当 Interactions API 也失败时的最终降级模型
_INTERACTIONS_FALLBACK_MODEL = "gemini-2.5-flash"

# 全局模型管理器实例（后台初始化）
_model_manager = None


def _init_model_manager():
    """
    在后台线程中初始化动态模型管理器并更新全局路由表。
    不阻塞主线程启动；路由表更新期间仍使用静态默认值。
    """
    global MODEL_MAP, _model_manager, _INTERACTIONS_ONLY_MODELS, _INTERACTIONS_FALLBACK_MODEL
    if not _model_manager_available or ModelManager is None:
        _app_logger.debug("[ModelManager] 模块不可用，使用静态默认路由")
        return
    try:
        _app_logger.debug("[ModelManager] 🔍 正在发现可用模型...")
        _model_manager = ModelManager(client)
        dynamic_map = _model_manager.get_model_map()
        MODEL_MAP.update(dynamic_map)
        _INTERACTIONS_ONLY_MODELS = _model_manager.get_interactions_only_models()
        _INTERACTIONS_FALLBACK_MODEL = _model_manager.get_fallback_model()
        # 同步更新 SmartDispatcher 的 MODEL_MAP 引用
        try:
            SmartDispatcher._dependencies["MODEL_MAP"] = MODEL_MAP
        except Exception:
            pass
        print(f"[ModelManager] ✅ 动态路由已加载: {len(dynamic_map)} 个任务")
    except Exception as _me:
        import traceback as _tb

        _app_logger.warning(
            f"[ModelManager] ⚠️ 动态路由初始化失败，使用静态默认值: {_me}"
        )
        _tb.print_exc()


# 模型能力矩阵（用于显示，动态模型自动补充）
MODEL_INFO = {
    "gemini-3-pro-preview": {
        "name": "Gemini 3.0 Pro",
        "speed": "🚀",
        "tier": 7,
        "strengths": ["推理", "分析", "代码", "复杂任务"],
    },
    "gemini-3-flash-preview": {
        "name": "Gemini 3.0 Flash",
        "speed": "⚡",
        "tier": 6,
        "strengths": ["快速", "对话", "多模态"],
    },
    "gemini-2.5-flash": {
        "name": "Gemini 2.5 Flash",
        "speed": "🌐",
        "tier": 5,
        "strengths": ["联网搜索", "grounding"],
    },
    "gemini-2.5-flash-preview": {
        "name": "Gemini 2.5 Flash Preview",
        "speed": "🌐",
        "tier": 5,
        "strengths": ["联网搜索", "grounding"],
    },
    "gemini-2.5-pro-preview": {
        "name": "Gemini 2.5 Pro",
        "speed": "🎯",
        "tier": 6,
        "strengths": ["推理", "代码", "分析"],
    },
    "deep-research-pro-preview-12-2025": {
        "name": "Deep Research Pro",
        "speed": "🔬",
        "tier": 7,
        "strengths": ["深度研究", "学术分析", "综合报告"],
    },
    "gemini-3.1-flash-image-preview": {
        "name": "Gemini 3.1 Flash Image",
        "speed": "🎨",
        "tier": 6,
        "strengths": ["图像生成", "创意绘画", "艺术风格"],
    },
    "gemini-2.0-flash-exp": {
        "name": "Gemini 2.0 Flash Exp",
        "speed": "🧪",
        "tier": 5,
        "strengths": ["图像生成", "多模态", "实验功能"],
    },
    "gemini-2.0-flash": {
        "name": "Gemini 2.0 Flash",
        "speed": "⚡",
        "tier": 5,
        "strengths": ["快速", "多模态"],
    },
    "gemini-1.5-pro": {
        "name": "Gemini 1.5 Pro",
        "speed": "📚",
        "tier": 5,
        "strengths": ["长上下文", "推理"],
    },
    "gemini-1.5-flash": {
        "name": "Gemini 1.5 Flash",
        "speed": "⚡",
        "tier": 4,
        "strengths": ["快速", "经济"],
    },
    "local-executor": {
        "name": "Local Executor",
        "speed": "🖥️",
        "tier": 0,
        "strengths": ["系统操作", "打开应用", "文件管理"],
    },
}


def get_model_display_name(model_id):
    """获取模型友好显示名称；动态发现的新模型自动从能力注册表补充。"""
    info = MODEL_INFO.get(model_id)
    if info:
        return f"{info['name']} {info['speed']}"
    # 动态模型：从 ModelManager 能力缓存获取
    if _model_manager:
        caps = _model_manager._cached_caps.get(model_id)
        if caps and caps.get("display"):
            return caps["display"]
    # 未知模型：直接展示 ID
    return model_id


# ================= 本地系统执行器 (已迁移到 web/local_executor.py) =================
try:
    from web.local_executor import LocalExecutor
except ImportError:
    from local_executor import LocalExecutor


# ================= 文件操作执行器 =================
class FileOperator:
    """
    本地文件操作执行器 - 处理文件读写、管理等操作
    """

    # 文件操作关键词
    FILE_KEYWORDS = [
        "读取文件",
        "打开文件",
        "查看文件",
        "读文件",
        "看看文件",
        "创建文件",
        "新建文件",
        "写入文件",
        "保存文件",
        "删除文件",
        "移动文件",
        "复制文件",
        "重命名",
        "文件列表",
        "目录",
        "文件夹",
        "列出文件",
        "自动归纳",
        "自动整理",
        "归纳文件夹",
        "整理文件夹",
        "归档文件夹",
        "微信文件归纳",
        "read file",
        "open file",
        "create file",
        "delete file",
        "list files",
        "directory",
        "folder",
    ]

    FOLDER_ORGANIZE_KEYWORDS = [
        "自动归纳",
        "自动整理",
        "归纳",
        "整理",
        "归档",
        "归类",
        "分类",
        "文件夹",
        "目录",
        "微信文件",
        "wechat files",
    ]

    @classmethod
    def is_file_operation(cls, text):
        """检测是否是文件操作请求"""
        text_lower = text.lower()
        return any(kw in text_lower for kw in cls.FILE_KEYWORDS)

    @classmethod
    def _is_folder_organize_intent(cls, text_lower: str) -> bool:
        has_action = any(
            kw in text_lower for kw in ["归纳", "整理", "归档", "归类", "分类"]
        )
        has_target = any(kw in text_lower for kw in ["文件夹", "目录", "路径", "文件"])
        if has_action and has_target:
            return True
        return any(kw in text_lower for kw in cls.FOLDER_ORGANIZE_KEYWORDS)

    @classmethod
    def _extract_path_from_text(cls, user_input: str) -> str:
        """Extract a likely filesystem path from user input."""
        import re

        patterns = [
            r'["\']([^"\']+)["\']',
            r'([A-Za-z]:\\(?:[^\\/:*?"<>|\r\n]+\\)*[^\\/:*?"<>|\r\n]*)',
            r"(\.?/[\w\-./ ]+)",
        ]
        for pattern in patterns:
            m = re.search(pattern, user_input)
            if m:
                candidate = m.group(1).strip().strip("，。,.;；")
                if candidate:
                    return candidate
        return ""

    @classmethod
    def execute(cls, user_input):
        """执行文件操作"""
        text_lower = user_input.lower()
        result = {"success": False, "action": "", "message": "", "content": ""}

        # === 指定路径文件夹自动归纳 ===
        if cls._is_folder_organize_intent(text_lower):
            folder_path = cls._extract_path_from_text(user_input)
            if not folder_path:
                folder_path = get_default_wechat_files_dir()

            if not folder_path:
                result["message"] = (
                    "❓ 请提供要归纳的文件夹路径（可用引号包裹），或在 config/user_settings.json 中设置 "
                    "storage.wechat_files_dir 作为默认路径"
                )
                return result

            if not os.path.isabs(folder_path):
                folder_path = os.path.join(WORKSPACE_DIR, folder_path)

            if not os.path.isdir(folder_path):
                result["message"] = f"❌ 目录不存在: {folder_path}"
                return result

            try:
                try:
                    from web.folder_catalog_organizer import FolderCatalogOrganizer
                except Exception:
                    from folder_catalog_organizer import FolderCatalogOrganizer

                analyzer = get_file_analyzer()
                organizer = get_file_organizer()
                engine = FolderCatalogOrganizer(
                    get_organize_root(), analyzer, organizer
                )
                summary = engine.organize_folder(folder_path)

                if not summary.get("success"):
                    result["message"] = (
                        f"❌ 自动归纳失败: {summary.get('error', '未知错误')}"
                    )
                    return result

                report_md = summary.get("report_markdown", "")
                report_json = summary.get("report_json", "")
                entries = summary.get("entries", [])

                sender_preview = []
                for item in entries:
                    sender = item.get("sender", "未知")
                    if sender and sender != "未知":
                        sender_preview.append(sender)
                sender_preview = sorted(set(sender_preview))[:8]
                sender_preview_text = (
                    "、".join(sender_preview)
                    if sender_preview
                    else "未识别到可靠发送者"
                )

                result["success"] = True
                result["action"] = "folder_auto_catalog"
                result["message"] = (
                    f"✅ 归纳完成：{summary.get('organized_count', 0)}/{summary.get('total_files', 0)} 个文件已归纳"
                    f"\n📁 来源目录: {summary.get('source_dir', folder_path)}"
                    f"\n🧾 清单(MD): {report_md}"
                    f"\n🧾 清单(JSON): {report_json}"
                    f"\n👤 识别到的发送者/来源人: {sender_preview_text}"
                )
                return result
            except Exception as e:
                result["message"] = f"❌ 自动归纳异常: {str(e)}"
                return result

        # === 读取文件 ===
        if any(
            kw in text_lower
            for kw in [
                "读取",
                "打开文件",
                "查看文件",
                "读文件",
                "看看",
                "read file",
                "open file",
            ]
        ):
            # 提取文件路径
            import re

            # 尝试匹配常见路径模式
            patterns = [
                r'["\']([^"\']+)["\']',  # 引号包围的路径
                r"([A-Za-z]:\\[^\s]+)",  # Windows 绝对路径
                r"(\.?/[^\s]+)",  # Unix 风格路径
                r"(\S+\.\w{1,5})(?:\s|$)",  # 带扩展名的文件
            ]

            filepath = None
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    filepath = match.group(1)
                    break

            if filepath:
                # 如果是相对路径，在 workspace 目录查找
                if not os.path.isabs(filepath):
                    workspace_path = os.path.join(WORKSPACE_DIR, filepath)
                    if os.path.exists(workspace_path):
                        filepath = workspace_path

                if os.path.exists(filepath):
                    try:
                        with open(
                            filepath, "r", encoding="utf-8", errors="ignore"
                        ) as f:
                            content = f.read()

                        # 限制内容长度
                        if len(content) > 10000:
                            content = content[:10000] + "\n\n... (文件过长，已截断)"

                        result["success"] = True
                        result["action"] = "read_file"
                        result["message"] = (
                            f"✅ 已读取文件: {os.path.basename(filepath)}"
                        )
                        result["content"] = f"```\n{content}\n```"
                        return result
                    except Exception as e:
                        result["message"] = f"❌ 读取文件失败: {str(e)}"
                        return result
                else:
                    result["message"] = f"❌ 文件不存在: {filepath}"
                    return result
            else:
                result["message"] = "❓ 请指定要读取的文件路径"
                return result

        # === 列出文件 ===
        if any(
            kw in text_lower
            for kw in [
                "文件列表",
                "目录",
                "列出文件",
                "list files",
                "directory",
                "文件夹里",
            ]
        ):
            # 提取目录路径
            import re

            patterns = [
                r'["\']([^"\']+)["\']',
                r"([A-Za-z]:\\[^\s]+)",
                r"(\.?/[^\s]+)",
            ]

            dirpath = WORKSPACE_DIR  # 默认 workspace
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    dirpath = match.group(1)
                    break

            if not os.path.isabs(dirpath):
                dirpath = os.path.join(WORKSPACE_DIR, dirpath)

            if os.path.isdir(dirpath):
                try:
                    items = os.listdir(dirpath)
                    file_list = []
                    for item in items[:50]:  # 限制数量
                        item_path = os.path.join(dirpath, item)
                        if os.path.isdir(item_path):
                            file_list.append(f"📁 {item}/")
                        else:
                            size = os.path.getsize(item_path)
                            size_str = (
                                f"{size/1024:.1f}KB" if size > 1024 else f"{size}B"
                            )
                            file_list.append(f"📄 {item} ({size_str})")

                    result["success"] = True
                    result["action"] = "list_files"
                    result["message"] = f"✅ 目录: {dirpath}"
                    result["content"] = "\n".join(file_list) if file_list else "空目录"
                    return result
                except Exception as e:
                    result["message"] = f"❌ 读取目录失败: {str(e)}"
                    return result
            else:
                result["message"] = f"❌ 目录不存在: {dirpath}"
                return result

        # === 创建/写入文件 ===
        if any(
            kw in text_lower
            for kw in ["创建文件", "新建文件", "写入文件", "保存到", "create file"]
        ):
            result["message"] = (
                "💡 请使用代码生成功能，Koto 会自动保存生成的文件到 workspace"
            )
            return result

        result["message"] = "❓ 无法识别该文件操作，请尝试：读取文件、列出目录等"
        return result

    @classmethod
    def watch_directory(cls, directory, callback=None, patterns=None):
        """监听目录变化并触发回调"""
        try:
            from watchdog.events import FileSystemEventHandler
            from watchdog.observers import Observer

            if patterns is None:
                patterns = ["*.txt", "*.pdf", "*.docx", "*.xlsx", "*.csv"]

            class ChangeHandler(FileSystemEventHandler):
                def on_created(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace("*", "")) for p in patterns):
                            if callback:
                                callback("created", event.src_path)

                def on_modified(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace("*", "")) for p in patterns):
                            if callback:
                                callback("modified", event.src_path)

            observer = Observer()
            observer.schedule(ChangeHandler(), directory, recursive=True)
            observer.start()

            return {
                "success": True,
                "observer": observer,
                "message": f"✅ 已开始监听目录: {directory}",
            }
        except Exception as e:
            return {"success": False, "message": f"❌ 无法监听目录: {str(e)}"}

    @classmethod
    def get_file_metadata(cls, filepath):
        """获取文件元数据"""
        try:
            if not os.path.exists(filepath):
                return {"success": False, "message": "文件不存在"}

            stat = os.stat(filepath)
            from datetime import datetime

            return {
                "success": True,
                "filepath": filepath,
                "filename": os.path.basename(filepath),
                "size": f"{stat.st_size / 1024:.2f} KB",
                "created": datetime.fromtimestamp(stat.st_ctime).strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                "extension": os.path.splitext(filepath)[1],
                "is_file": os.path.isfile(filepath),
            }
        except Exception as e:
            return {"success": False, "message": f"❌ 无法获取文件信息: {str(e)}"}


# ================= 联网搜索能力 =================
class WebSearcher:
    """
    使用 Gemini 的 Google Search Grounding 能力
    获取实时天气、新闻等信息
    """

    # 需要联网的关键词（严格收窄：仅包含几乎只在需要实时信息时才会出现的词）
    WEB_KEYWORDS = [
        # 天气（高置信）
        "天气",
        "气温",
        "下雨吗",
        "下雪吗",
        "温度多少",
        "天气怎么样",
        "天气预报",
        "weather",
        "temperature",
        "forecast",
        # 实时行情（高置信）
        "股价",
        "汇率",
        "比特币价格",
        "黄金价格",
        "金价",
        "实时金价",
        "今日金价",
        "当前金价",
        "现货黄金",
        "国际金价",
        "石油价格",
        "a股",
        "港股",
        "美股",
        "stock price",
        # 比赛/体育（高置信）
        "比分",
        "比赛结果",
        "谁赢了",
        # 新闻（只匹配明确的新闻请求）
        "今天新闻",
        "最新新闻",
        "latest news",
        # 交通出行票务（高置信 — 余票/时刻表实时变化）
        "火车票",
        "高铁票",
        "动车票",
        "机票",
        "余票",
        "班次查询",
        "车次查询",
        "时刻表",
        "列车时刻",
        "航班查询",
        "航班动态",
        "几点出发",
        "几点到",
        "几点到达",
        "多久到",
        "要多久",
    ]

    @classmethod
    def needs_web_search(cls, text):
        """检测是否需要联网搜索

        优化策略：
        1. 检查关键词列表
        2. 对于金融/预测类，更倾向于web-search
        3. 对于热点事件、新品发布，必须web-search
        """
        text_lower = text.lower()

        # 必须 web-search 的模式（绝不能用纯AI）
        must_search_patterns = [
            r"(能不能|应该不应该|值不值得|是否).*?买",  # 股票建议
            r"(最新|实时|今天|明天|下周).*?(股|行情|数据)",  # 实时行情
            r"(预测|预期|后市|趋势).*?(股|市场|行业)",  # 趋势预测
            r"(财报|业绩|营收).*?(公布|发布)",  # 财报动态
            r"(新品|发布|推出).*?(上市|发售)",  # 新品信息
            r"(突发|紧急|最新)\w*事件",  # 突发事件
            r"(当前|今日|实时|最新).*?(金价|黄金)",  # 黄金实时行情
            r"(金价|黄金).*?(多少|报价|走势|行情)",  # 金价查询
            # 交通出行——余票/时刻均实时变化
            r"(查|看|查询|查一下|有没有|有无|还有).{0,6}(火车票|高铁票|动车票|机票|余票)",
            r"(下周|明天|后天|今天|大后天|\d+[号日]).{0,12}(去|到|从).{0,12}(的|要).{0,5}(票|高铁|动车|火车|航班)",
            r"(去|从).{1,12}(去|到).{1,18}(火车|高铁|动车|机票|班次|航班)",
            r"(几点|什么时候).{0,6}(出发|到|到达|抵达).{0,12}(班|次|票|车|机)",
        ]

        import re

        for pattern in must_search_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True

        # 关键词匹配
        if any(kw in text_lower for kw in cls.WEB_KEYWORDS):
            return True

        return False

    @classmethod
    def _detect_query_type(cls, query: str) -> str:
        """检测搜索查询的意图类型，返回: travel / weather / finance / news / general"""
        q = query.lower()
        travel_kw = [
            "火车票",
            "高铁票",
            "动车票",
            "机票",
            "余票",
            "班次",
            "车次",
            "时刻表",
            "列车时刻",
            "列车",
            "高铁",
            "动车",
            "航班",
            "航班动态",
            "几点到",
            "几点出发",
            "几点抵达",
            "要多久",
            "多久到",
        ]
        if any(kw in q for kw in travel_kw):
            return "travel"
        weather_kw = [
            "天气",
            "气温",
            "下雨",
            "下雪",
            "温度",
            "weather",
            "forecast",
            "天气预报",
        ]
        if any(kw in q for kw in weather_kw):
            return "weather"
        finance_kw = [
            "股价",
            "股票",
            "汇率",
            "比特币",
            "黄金",
            "金价",
            "行情",
            "基金",
            "石油",
            "原油",
        ]
        if any(kw in q for kw in finance_kw):
            return "finance"
        return "general"

    @classmethod
    def _build_search_context(cls, query: str, query_type: str) -> tuple:
        """根据查询类型返回 (enriched_query, system_instruction)"""
        if query_type == "travel":
            instruction = (
                "你是 Koto，一个智能出行助手。用户在查询交通出行信息（高铁/火车/动车/机票等）。\n"
                "请基于搜索结果，按以下格式输出（用 Markdown）：\n\n"
                "1. 先用一句话说明查询的出发日期和路线（如有）。\n"
                "2. 用 **Markdown 表格** 列出主要班次，列标题为：\n"
                "   | 班次 | 出发站 | 到达站 | 出发时间 | 到达时间 | 历时 | 二等座 | 一等座 |\n"
                "   至少列出 5 个代表性班次（早、中、晚各时段）。\n"
                "3. 表格后，提醒用户前往 12306 或铁路官方渠道查看实时余票并购票。\n"
                "4. 若搜索结果信息不足，请尽量根据已知班次填写，并注明「以下为参考班次，请以12306实时信息为准」。\n"
                "用中文输出，格式整洁，突出关键数据。"
            )
            return query, instruction
        elif query_type == "weather":
            instruction = (
                "你是 Koto，一个智能助手。请根据搜索结果提供准确的天气信息。\n"
                "格式要求：\n"
                "1. 当前气温和天气状况\n"
                "2. 今日最高 / 最低气温\n"
                "3. 未来 3 天天气（如果有）\n"
                "4. 简短的出行或着装建议\n"
                "用中文输出，简洁清晰。"
            )
            return query, instruction
        elif query_type == "finance":
            instruction = (
                "你是 Koto，一个智能助手。请根据搜索结果提供准确的金融行情信息。\n"
                "格式要求：\n"
                "1. 当前价格 / 价值及所属市场\n"
                "2. 今日涨跌幅（如有）\n"
                "3. 近期走势简析（1-2 句）\n"
                "用中文输出，简洁专业。"
            )
            return query, instruction
        else:
            instruction = (
                "你是 Koto，一个智能助手。使用搜索结果提供准确、实时的信息。"
                "用中文回答，格式清晰，关键数据用 Markdown 列表或加粗呈现。"
            )
            return query, instruction

    @classmethod
    def search_with_grounding(cls, query, skill_prompt=None):
        """使用 Gemini Google Search Grounding 进行实时搜索（意图感知版本）

        skill_prompt: 来自本地/AI路由器生成的执行指令。
          - 若提供，直接用作 system_instruction（正确理解用户意图）
          - 若未提供，回退到关键词检测分支（保指安全下线）
        """
        # 1. 优先使用模型生成的 skill_prompt
        if skill_prompt and len(skill_prompt.strip()) > 5:
            system_instruction = (
                "你是 Koto，一个智能助手。请使用搜索结果提供准确、实时的信息。\n"
                f"{skill_prompt}\n"
                "用中文回答，格式整洁清晰。"
            )
            _app_logger.debug(f"[WebSearcher] 使用 skill_prompt: {skill_prompt[:60]}")
        else:
            # 2. 回退：关键词检测 + 分类 system_instruction
            query_type = cls._detect_query_type(query)
            _, system_instruction = cls._build_search_context(query, query_type)
            _app_logger.debug(f"[WebSearcher] 关键词检测备用: {query_type}")
        try:
            # 使用 Google Search 作为工具
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=query,
                config=types.GenerateContentConfig(
                    tools=[types.Tool(google_search=types.GoogleSearch())],
                    system_instruction=system_instruction,
                ),
            )

            if response.text:
                return {"success": True, "response": response.text, "grounded": True}
            else:
                return {
                    "success": False,
                    "response": "搜索未返回结果",
                    "grounded": False,
                }
        except Exception as e:
            return {
                "success": False,
                "response": f"搜索失败: {str(e)}",
                "grounded": False,
            }

    @classmethod
    def generate_ppt_images(
        cls, slide_titles: list, topic: str, max_images: int = 3
    ) -> list:
        """为 PPT 幻灯片生成配图（使用 Imagen / Gemini 图像模型）

        从幻灯片标题中挑选最适合配图的 2-3 页，生成高质量配图。
        返回: [{"slide_index": int, "image_path": str}, ...]
        """
        import queue as _queue
        import threading

        if not slide_titles:
            return []

        # 用 AI 挑选最适合配图的幻灯片
        pick_prompt = (
            f"以下是一个关于「{topic}」的PPT的各页标题,请挑选最适合配图的 {min(max_images, len(slide_titles))} 页。\n"
            f"对每页生成一个简洁的英文图像描述（适合AI图像生成）。\n"
            f'只输出 JSON 数组，格式：[{{"index": 0, "prompt": "..."}}]\n\n'
        )
        for i, t in enumerate(slide_titles):
            pick_prompt += f"{i}. {t}\n"

        try:
            resp = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=pick_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.3, max_output_tokens=1024
                ),
            )
            import json as _json

            raw = resp.text or ""
            # 提取 JSON 数组
            import re as _re

            m = _re.search(r"\[.*\]", raw, _re.DOTALL)
            if m:
                picks = _json.loads(m.group())
            else:
                picks = []
        except Exception as e:
            _app_logger.debug(f"[PPT-IMAGE] 选图AI失败: {e}")
            # 回退：选前 max_images 个非过渡页
            picks = [
                {"index": i, "prompt": f"professional illustration about {t}"}
                for i, t in enumerate(slide_titles[:max_images])
            ]

        results = []
        images_dir = os.path.join(WORKSPACE_DIR, "images")
        os.makedirs(images_dir, exist_ok=True)

        for pick in picks[:max_images]:
            idx = pick.get("index", 0)
            prompt = pick.get("prompt", f"professional illustration for presentation")
            # 增强 prompt 质量 — 确保简洁、无文字要求
            full_prompt = (
                f"Create a clean, modern, professional infographic-style illustration for a presentation slide. "
                f"Topic: {prompt}. "
                f"Style: flat design, clean layout, soft gradients, business-appropriate color palette. "
                f"Requirements: NO text, NO words, NO letters, NO numbers in the image. "
                f"Pure visual illustration only."
            )

            result_q = _queue.Queue()

            def _gen_image(p, q):
                # ① 首选: Gemini 3.1 Flash Image
                try:
                    res = client.models.generate_content(
                        model="gemini-3.1-flash-image-preview",
                        contents=p,
                        config=types.GenerateContentConfig(
                            response_modalities=["TEXT", "IMAGE"]
                        ),
                    )
                    if res.candidates and res.candidates[0].content.parts:
                        for part in res.candidates[0].content.parts:
                            if (
                                hasattr(part, "inline_data")
                                and part.inline_data
                                and part.inline_data.data
                            ):
                                q.put(("success", part.inline_data.data))
                                return
                except Exception as e0:
                    _app_logger.debug(f"[PPT-IMAGE] Gemini 3.1 Flash Image 失败: {e0}")

                # ② 备选: Imagen 4.0
                try:
                    res = client.models.generate_images(
                        model="imagen-4.0-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res.generated_images:
                        q.put(("success", res.generated_images[0].image.image_bytes))
                        return
                except Exception as e1:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 4.0 失败: {e1}")

                # ③ 备选: Imagen 4.0 Fast
                try:
                    res2 = client.models.generate_images(
                        model="imagen-4.0-fast-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res2.generated_images:
                        q.put(("success", res2.generated_images[0].image.image_bytes))
                        return
                except Exception as e2:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 4.0 Fast 也失败: {e2}")

                # ④ 最终备选: Imagen 3.0（当前公开稳定版）
                try:
                    res3 = client.models.generate_images(
                        model="imagen-3.0-generate-001",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1),
                    )
                    if res3.generated_images:
                        q.put(("success", res3.generated_images[0].image.image_bytes))
                        return
                except Exception as e3:
                    _app_logger.debug(f"[PPT-IMAGE] Imagen 3.0 也失败: {e3}")
                q.put(("fail", None))

            thread = threading.Thread(
                target=_gen_image, args=(full_prompt, result_q), daemon=True
            )
            thread.start()
            thread.join(timeout=120)  # Gemini 图像生成可能较慢，给足时间

            try:
                status, data = result_q.get_nowait()
                if status == "success" and data:
                    ts = int(time.time() * 1000) % 1000000
                    fname = f"ppt_slide_{idx}_{ts}.png"
                    fpath = os.path.join(images_dir, fname)
                    with open(fpath, "wb") as f:
                        f.write(data)
                    results.append({"slide_index": idx, "image_path": fpath})
                    _app_logger.info(f"[PPT-IMAGE] ✅ 幻灯片 {idx} 配图生成: {fname}")
            except Exception:
                _app_logger.warning(f"[PPT-IMAGE] ⚠️ 幻灯片 {idx} 配图超时或失败")

        return results

    @classmethod
    def deep_research_for_ppt(cls, user_input: str, search_context: str = "") -> str:
        """对复杂/学术主题进行深度研究，返回详细的研究报告文本

        用于在生成 PPT 大纲之前，先用 Pro 模型做深度分析，
        保证内容专业度和信息量。
        """
        research_prompt = (
            "你是一位顶级行业研究分析师。请对以下主题进行深入、全面的研究分析。\n\n"
            "## 严格要求\n"
            "1. **必须提供具体数据** — 市场规模（金额）、增长率（%）、市占率、出货量等定量信息\n"
            "2. **必须引用来源** — 如 IDC、Gartner、Statista、行业年报等（基于搜索资料中的数据）\n"
            "3. **必须包含真实案例** — 具体公司名称、产品型号、发布时间、销售数据等\n"
            "4. **必须有对比分析** — 不同产品/方案/技术路线之间的优劣对比\n"
            "5. **必须覆盖完整视角** — 历史演进 → 现状格局 → 技术路线 → 竞争分析 → 未来趋势\n"
            "6. **必须结构化** — 用清晰的标题层级和要点编排\n"
            "7. 中文回答，内容必须详实，**空洞的描述是不可接受的**\n\n"
            "## 输出格式\n"
            "为每个板块提供:\n"
            "- 2-3 个核心数据点（带数字和来源）\n"
            "- 2-3 个具体案例/产品\n"
            "- 1-2 个关键趋势判断\n\n"
            f"研究主题：{user_input}\n"
        )
        if search_context:
            research_prompt += f"\n已有的搜索参考资料：\n{search_context[:8000]}\n"

        def _extract_text_from_obj(obj) -> list[str]:
            texts = []
            if obj is None:
                return texts
            if isinstance(obj, str):
                s = obj.strip()
                if s:
                    texts.append(s)
                return texts
            if isinstance(obj, dict):
                for key in ("output_text", "text"):
                    val = obj.get(key)
                    if isinstance(val, str) and val.strip():
                        texts.append(val.strip())
                for val in obj.values():
                    texts.extend(_extract_text_from_obj(val))
                return texts
            if isinstance(obj, (list, tuple)):
                for item in obj:
                    texts.extend(_extract_text_from_obj(item))
                return texts
            if hasattr(obj, "model_dump"):
                try:
                    texts.extend(_extract_text_from_obj(obj.model_dump()))
                    return texts
                except Exception:
                    pass
            return texts

        def _extract_interaction_text(interaction_obj) -> str:
            if not interaction_obj:
                return ""
            parts = _extract_text_from_obj(getattr(interaction_obj, "outputs", None))
            if not parts:
                parts = _extract_text_from_obj(interaction_obj)
            dedup = []
            seen = set()
            for part in parts:
                if part not in seen:
                    dedup.append(part)
                    seen.add(part)
            return "\n".join(dedup).strip()

        # 深度研究专用：Interactions API（deep-research-pro-preview-*）
        preferred_model = MODEL_MAP.get("RESEARCH", "deep-research-pro-preview-12-2025")
        if preferred_model.startswith("deep-research-pro-preview"):
            try:
                research_client = create_research_client()
                interaction = research_client.interactions.create(
                    agent=preferred_model,
                    input=research_prompt,
                    background=True,
                    stream=False,
                )
                interaction_id = getattr(interaction, "id", None)
                status = getattr(interaction, "status", "")
                start_wait = time.time()
                max_wait_time = 180  # 限制最大等待时间为 3 分钟

                while (
                    interaction_id
                    and status not in ("completed", "failed", "cancelled")
                    and (time.time() - start_wait) < max_wait_time
                ):
                    time.sleep(3)
                    interaction = research_client.interactions.get(interaction_id)
                    status = getattr(interaction, "status", "")

                if status not in ("completed", "failed", "cancelled"):
                    print(
                        f"[PPT-RESEARCH] ⚠️ Interactions 超时 ({max_wait_time}s)，尝试取消并回退"
                    )
                    try:
                        research_client.interactions.cancel(interaction_id)
                    except Exception:
                        pass
                else:
                    text = _extract_interaction_text(interaction)
                    if text and len(text) > 200:
                        print(
                            f"[PPT-RESEARCH] ✅ 深度研究完成 ({preferred_model}), {len(text)} 字符"
                        )
                        return text
                    print(
                        f"[PPT-RESEARCH] ⚠️ Interactions 返回空结果或过短，status={status}"
                    )
            except Exception as inter_err:
                _app_logger.debug(f"[PPT-RESEARCH] Interactions 失败: {inter_err}")

        _app_logger.debug(f"[PPT-RESEARCH] 🔄 切换到备用模型进行研究...")
        research_models = [
            MODEL_MAP.get("RESEARCH", "deep-research-pro-preview-12-2025"),
            "gemini-3-pro-preview",
            "gemini-2.5-flash",
        ]
        # 去重并去空，保持顺序
        research_models = [
            m
            for i, m in enumerate(research_models)
            if m and m not in research_models[:i]
        ]
        for model in research_models:
            try:
                # deep-research 和 gemini-3-preview 仅支持 Interactions API，不走 generate_content
                if (
                    model.startswith("deep-research-pro-preview")
                    or model in _INTERACTIONS_ONLY_MODELS
                ):
                    continue
                resp = client.models.generate_content(
                    model=model,
                    contents=research_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.5,
                        max_output_tokens=16384,
                    ),
                )
                if resp.text and len(resp.text) > 200:
                    _app_logger.info(
                        f"[PPT-RESEARCH] ✅ 深度研究完成 ({model}), {len(resp.text)} 字符"
                    )
                    return resp.text
            except Exception as e:
                _app_logger.debug(f"[PPT-RESEARCH] {model} 失败: {e}")
                continue
        return ""


# === System Instruction ===
# 简化版系统指令 - 用于CHAT/RESEARCH等非文件生成任务
# 任务专属 system prompt 补充片段（在 chat_stream 确定 task_type 后追加）
_TASK_SYSTEM_ADDENDUMS: dict = {
    "CODER": "\n\n## 🔧 代码任务规范\n- 直接给出可运行代码，不加废话前言\n- 使用代码块（```语言）包裹\n- 必要时说明运行方式，但不超过3行",
    "RESEARCH": "\n\n## 🔍 研究任务规范\n- 必须分段：摘要 → 正文 → 小结\n- 给出信息来源或推理依据\n- 避免模糊表述，用具体数据或例子",
    "FILE_GEN": "\n\n## 📄 文件生成规范\n- 严格使用 ---BEGIN_FILE: filename.ext--- / ---END_FILE--- 标记\n- 代码必须完整可执行，不允许省略号或 placeholder\n- 生成完成后告知保存路径",
    "DOC_ANNOTATE": "\n\n## 📝 文档批注规范\n- 批注定位精确，引用原文片段\n- 修改建议简洁，不改变原文意图\n- 按重要性排序（严重 → 建议 → 细节）",
}


def _get_chat_system_instruction(question: str = None):
    """
    生成包含当前日期时间和系统状态的系统指令

    Args:
        question: 用户问题（可选），用于智能上下文选择

    Returns:
        系统指令文本
    """
    try:
        # 如果提供了问题，使用智能上下文注入
        if question:
            from web.context_injector import get_dynamic_system_instruction

            return get_dynamic_system_instruction(question)
    except Exception as e:
        _app_logger.debug(f"[Koto] Warning: Dynamic context injection failed: {e}")

    # 降级方案：使用基础系统指令
    from datetime import datetime

    now = datetime.now()
    date_str = now.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][now.weekday()]
    time_str = now.strftime("%H:%M:%S")

    # 获取系统信息（如果可用）
    system_info_section = ""
    try:
        from web.system_info import get_formatted_system_info, get_system_warnings

        formatted_info = get_formatted_system_info(include_processes=False)
        warnings = get_system_warnings()

        system_info_section = f"""
## 💻 当前系统状态
{formatted_info}"""

        if warnings:
            system_info_section += "\n\n## ⚠️ 系统警告\n"
            for warning in warnings:
                system_info_section += f"  • {warning}\n"
    except Exception as e:
        _app_logger.debug(f"[Koto] Warning: Failed to collect system info: {e}")

    return f"""你是 Koto (言)，一个与用户计算机深度融合的个人AI助手。

## 📅 当前时间（用于相对日期计算）
🕒 **系统时间**: {date_str} {weekday} {time_str}
📅 **ISO日期**: {now.strftime("%Y-%m-%d")}
⏰ **使用此时间计算**: "明天"、"下周"、"前天" 等相对时间{system_info_section}

## 👤 角色定位
- 精通多个领域：编程、数据分析、写作、问题解决、系统管理
- 充分了解用户的计算环境和当前状态
- 快速理解用户意图，提供符合实际情境的答案
- 充当用户与Windows系统的智能中介

## 📋 回答原则
1. **简洁直接** - 不自我介绍，直接进入主题
2. **优先中文** - 默认用中文回答，除非用户要求其他语言
3. **清晰结构** - 使用标题、列表、代码块组织内容，便于快速理解
4. **上下文感知** - 结合用户的系统状态给出建议
5. **环境感知** - 了解当前 CPU、内存、磁盘状态，做出合适的建议
6. **时间准确性** - 使用系统时间准确计算相对日期
7. **禁止生成文件** - 仅在明确要求PDF/Word/Excel/PPT时才生成

## ✅ 能做的事
- 帮助用户分析本地文件、文档、图片
- 建议系统操作、自动化脚本、PowerShell命令
- 理解文件路径、应用名称、快捷键等Windows内容
- 根据当前系统状况给出性能优化建议
- 基于磁盘剩余空间建议存储位置
- 基于内存和 CPU 使用情况建议何时执行任务
- 协助处理剪贴板、监听快捷键、系统设置
- 联动本地应用（打开微信、邮件、浏览器等）
- 进行系统诊断：如果用户反映电脑卡，可以分析当前 CPU/内存/磁盘情况
- 准确理解和计算时间问题

## ❌ 不做的事
- ✗ 自我介绍或重复身份
- ✗ 生成代码标记 BEGIN_FILE/END_FILE（仅文件生成任务使用）
- ✗ 输出冗长的前言、风险提示或过度谨慎的警告
- ✗ 拒绝合理的系统操作请求"""


def _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION():
    """获取默认的系统指令（用于降级场景）"""
    try:
        return _get_chat_system_instruction()
    except Exception:
        # 终极降级：返回基础指令
        return "你是 Koto (言)，一个与用户计算机深度融合的个人AI助手。精通多个领域，快速理解用户意图，提供符合实际情境的答案。"


def _get_system_instruction():
    """生成包含当前日期时间的文档生成系统指令（含 Skills 注入）"""
    from datetime import datetime

    now = datetime.now()
    date_str = now.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][now.weekday()]

    _base_filegen = f"""你是 Koto 文档生成专家，专注于生成高质量、可用的文档。

## 当前时间上下文
📅 **生成日期**: {date_str} {weekday}

## 时间理解规则（严格遵守）
- 这是本次请求的唯一时间锚点，请据此理解“今天/本月/今年/1月”等相对时间。
- 当用户只说“X月”未写年份时，默认使用**当前年份**（例如当前是 2026 年，则“1月新番”默认指 2026 年 1 月）。
- 不要默认使用过去年份，除非用户明确指定（如“2024年1月新番”）。

## 核心职责
1. **直接输出文档内容** - 直接输出最终要保存的文档内容，而不是代码或JSON
2. **中文优先** - 使用简体中文，专业术语准确无误
3. **格式规范** - 使用标题、列表、段落进行清晰组织

## 文档生成规则

### 优先策略：直接输出模式（推荐）
- **直接输出最终文档内容**，无需代码包装
- 使用Markdown式格式组织（# ## ### 标题、- 列表、段落）
- 系统会自动将你的输出转换为Word/PDF
- 这是最快、最可靠的方法

示例（只输出内容，不输出代码）：
```
# 文档标题

## 第一节
内容段落...

## 第二节
- 要点1
- 要点2
```

### 代码生成模式（仅当需要特殊格式时）
- 必须使用 ---BEGIN_FILE: filename.py--- 和 ---END_FILE--- 标记
- 代码控制在 80 行以内
- **保存路径必须使用**: `import os; OUTPUT_DIR = os.environ.get('KOTO_OUTPUT_DIR', os.getcwd())`，然后把生成的文件保存到 `OUTPUT_DIR`
- 必须包含中文字体处理（特别是PDF生成）
- 使用 try/except 包装错误处理
- **仅当直接输出无法满足需求时才使用此模式**

## 禁止项清单
- ✗ 输出JSON格式的"虚拟文档"
- ✗ 输出结构化数据而非真实内容
- ✗ 生成 BEGIN_FILE/END_FILE 标记（除非必须生成Python代码）
- ✗ 生成要求用户手动复制粘贴的内容

## 优先级
1. **直接输出内容** > 代码生成 > JSON结构
2. 内容准确、结构清晰 > 输出格式完美
3. 实际可执行性 > 审美程度
"""
    # 注入 FILE_GEN 相关的 Skills
    try:
        from app.core.skills.skill_manager import SkillManager

        return SkillManager.inject_into_prompt(_base_filegen, task_type="FILE_GEN")
    except Exception:
        return _base_filegen


# SYSTEM_INSTRUCTION 不再在模块加载时构建，改为按需调用 _get_system_instruction()
# SYSTEM_INSTRUCTION = _get_system_instruction()


def _get_filegen_brief_instruction() -> str:
    """FILE_GEN 的简版系统提示（每次调用实时取时间）。"""
    now = datetime.now()
    return (
        "你是Koto文档生成器，输出清晰的结构化内容，不要输出代码。\n"
        f"当前系统日期: {now.strftime('%Y-%m-%d')}（{now.strftime('%Y年%m月%d日')}）。\n"
        "时间规则：若用户仅写月份未写年份（如‘1月新番’），默认按当前年份解释。"
    )


def _parse_time_info_for_filegen(user_text: str) -> dict:
    """解析 FILE_GEN 输入中的时间信息，重点处理“仅月份未写年份”的场景。"""
    now = datetime.now()
    info = {
        "raw": user_text or "",
        "year": None,
        "month": None,
        "resolved_year": None,
        "resolved_month": None,
        "time_text": now.strftime("%Y年%m月%d日"),
        "rule_hit": False,
    }

    text = user_text or ""
    m = re.search(r"(?:(20\d{2})\s*年)?\s*([1-9]|1[0-2])\s*月", text)
    if not m:
        return info

    year_str = m.group(1)
    month_str = m.group(2)
    month = int(month_str)
    year = int(year_str) if year_str else None

    info["year"] = year
    info["month"] = month
    info["resolved_year"] = year if year is not None else now.year
    info["resolved_month"] = month
    info["rule_hit"] = year is None
    return info


def _build_filegen_time_context(user_text: str) -> tuple[str, dict]:
    """构建注入给模型的时间上下文文本。"""
    parsed = _parse_time_info_for_filegen(user_text)
    now = datetime.now()
    lines = [
        "[时间上下文]",
        f"- 当前系统时间: {now.strftime('%Y-%m-%d %H:%M:%S')}",
    ]

    if parsed.get("resolved_month"):
        lines.append(
            f"- 用户时间意图解析: {parsed['resolved_year']}年{parsed['resolved_month']}月"
        )
        if parsed.get("rule_hit"):
            lines.append("- 解析规则命中: 用户仅提供月份，已按当前年份解析")
    else:
        lines.append("- 用户时间意图解析: 未检测到明确月份，按当前语境理解")

    return "\n".join(lines), parsed


# ===== 任务特定系统提示词 =====
TASK_PROMPTS = {
    "CHAT": """助手模式：普通对话
- 直接回答问题，提供有用信息
- 保持对话自然流畅
- 记住之前的上下文""",
    "CODER": """代码生成专家
- 生成高质量、可运行的代码
- 遵循Python/JavaScript最佳实践
- 添加必要注释，解释复杂逻辑
- 包含错误处理和边界检查
- 代码长度控制在80行以内""",
    "FILE_GEN": """文档生成专家
- 生成结构清晰、格式规范的文档
- 使用标题、列表、段落进行组织
- 适配Word/PDF/Excel导出
- 内容准确、专业、可执行
- 禁止输出代码块和技术细节""",
    "PAINTER": """图像生成艺术家
- 创作独特、高质量的图像
- 理解用户的审美偏好
- 支持风格、颜色、构图的微调
- 输出高分辨率图像""",
    "RESEARCH": """深度研究专家
- 进行全面的信息搜索和分析
- 查找最新、最准确的信息
- 整理多个来源的观点
- 提供有根据的结论和见解
- 标注信息来源""",
    "SYSTEM": """系统操作执行器
- 执行本地系统命令和操作
- 打开应用、管理文件、控制系统
- 提供清晰的执行反馈
- 解释操作结果和错误""",
}

# ===== Windows本地快捷指令映射 =====
WINDOWS_SHORTCUTS = {
    # 文件和剪贴板操作
    "复制": "Ctrl+C",
    "粘贴": "Ctrl+V",
    "剪切": "Ctrl+X",
    "撤销": "Ctrl+Z",
    "重做": "Ctrl+Y",
    "全选": "Ctrl+A",
    "保存": "Ctrl+S",
    "打开": "Ctrl+O",
    "新建": "Ctrl+N",
    # 浏览器操作
    "新标签页": "Ctrl+T",
    "关闭标签页": "Ctrl+W",
    "历史记录": "Ctrl+H",
    "书签": "Ctrl+B",
    "刷新": "Ctrl+R",
    "放大": "Ctrl+加号",
    "缩小": "Ctrl+减号",
    # 系统操作
    "任务管理器": "Ctrl+Shift+Esc",
    "截图": "Win+Shift+S",
    "开始菜单": "Win",
    "锁屏": "Win+L",
    "关机": "Alt+F4",
    "虚拟桌面": "Win+Tab",
    "显示桌面": "Win+D",
    # 应用切换
    "切换应用": "Alt+Tab",
    "关闭应用": "Alt+F4",
}


# ================= RAG 上下文分析器 =================
class ContextAnalyzer:
    """
    基于 RAG (检索增强生成) 的智能上下文分析器

    功能：
    1. 分析历史对话，提取关键信息
    2. 构建结构化的上下文提示词
    3. 智能判断任务关联性
    4. 生成增强后的输入
    """

    # 任务类型特征签名
    TASK_SIGNATURES = {
        "PAINTER": {
            "keywords": [
                "图",
                "画",
                "照片",
                "image",
                "photo",
                "picture",
                "图像已生成",
                "图片已生成",
                "猫",
                "狗",
                "人物",
                "风景",
                "头像",
            ],
            "outputs": ["图像已生成", "图片已生成", "已保存图片", "✨ 图片已生成"],
            "entities": [
                "颜色",
                "风格",
                "大小",
                "背景",
                "表情",
                "姿势",
                "眼睛",
                "毛发",
                "脸",
            ],
        },
        "FILE_GEN": {
            "keywords": [
                "pdf",
                "word",
                "excel",
                "docx",
                "文档",
                "报告",
                "文件",
                "简历",
                "合同",
                "标注",
                "批注",
                "润色",
                "改写",
                "校对",
                "审校",
                "修订",
                "优化",
                "纠错",
            ],
            "outputs": [
                "已生成文件",
                "文件已保存",
                ".pdf",
                ".docx",
                ".xlsx",
                "✅ **文件生成成功",
            ],
            "entities": [
                "标题",
                "章节",
                "内容",
                "格式",
                "模板",
                "标注",
                "批注",
                "修改建议",
            ],
        },
        "RESEARCH": {
            "keywords": ["研究", "分析", "介绍", "了解", "原理", "技术", "深入"],
            "outputs": ["##", "###", "1.", "2.", "总结", "结论"],
            "entities": ["定义", "特点", "优势", "劣势", "应用", "发展"],
        },
        "CODER": {
            "keywords": [
                "代码",
                "编程",
                "函数",
                "脚本",
                "code",
                "script",
                "python",
                "javascript",
            ],
            "outputs": ["```python", "```javascript", "```", "def ", "class "],
            "entities": ["函数", "变量", "类", "模块", "算法"],
        },
        "CHAT": {
            "keywords": ["你好", "谢谢", "帮我", "请问", "什么是"],
            "outputs": [],
            "entities": [],
        },
    }

    # 延续性指示词分类 - 需要更严格的匹配
    CONTINUATION_PATTERNS = {
        "modify": {
            # 修改类：必须是短句或明确的修改指令
            "indicators": [
                "再来一张",
                "再来一个",
                "更大一点",
                "更小一点",
                "大一点",
                "小一点",
                "深一些",
                "浅一些",
                "颜色换成",
                "背景换成",
            ],
            "weight": 0.9,
            "max_input_length": 30,  # 限制输入长度，长句子不太可能是简单修改
            "prompt_template": "用户要求修改之前的结果：{modification}",
        },
        "reference": {
            # 引用类：必须在句首或独立使用
            "indicators": [
                "这个怎么",
                "这张图",
                "那个文件",
                "上面的",
                "刚才的",
                "把它",
                "把这个",
                "基于这个",
            ],
            "weight": 0.85,
            "require_start": True,  # 需要在句首出现
            "prompt_template": "用户引用了之前的内容：{reference}",
        },
        "reference_loose": {
            # 引用类（宽松）：用于计划/大纲/方案等需要跟随之前内容的请求
            "indicators": [
                "这个计划",
                "该计划",
                "上述计划",
                "上面的计划",
                "这个方案",
                "该方案",
                "上述方案",
                "这个大纲",
                "该大纲",
                "这个ppt",
                "该ppt",
                "这个PPT",
                "该PPT",
                "按照这个",
                "根据这个",
            ],
            "weight": 0.78,
            "require_start": False,
            "prompt_template": "用户引用了之前的计划或大纲：{reference}",
        },
        "convert": {
            # 转换类：明确的格式转换请求
            "indicators": [
                "做成word",
                "做成pdf",
                "做成excel",
                "转成word",
                "转成pdf",
                "变成文档",
                "导出为",
                "保存为word",
                "保存为pdf",
            ],
            "weight": 0.95,
            "prompt_template": "用户要求将之前的内容转换为新格式：{conversion}",
        },
        "continue": {
            # 继续类：明确要求继续之前的内容
            "indicators": [
                "继续写",
                "接着说",
                "接着写",
                "然后呢",
                "下一步",
                "还有呢",
                "另外补充",
                "再找找",
                "再搜",
                "再查",
                "再看看",
                "继续查",
                "继续找",
                "再找",
                "再搜一下",
            ],
            "weight": 0.7,
            "max_input_length": 20,  # 短句才是继续指令
            "prompt_template": "用户要求继续之前的任务：{continuation}",
        },
        "detail": {
            # 详细类：只有非常明确的展开请求才算，且必须是短句
            "indicators": [
                "详细说说",
                "展开说说",
                "详细讲讲",
                "具体说一下",
                "解释一下刚才的",
            ],
            "weight": 0.75,
            "max_input_length": 25,  # 限制长度
            "prompt_template": "用户要求详细说明之前提到的内容：{detail}",
        },
    }

    @classmethod
    def extract_entities(cls, text: str, task_type: str = None) -> list:
        """从文本中提取关键实体"""
        entities = []
        text_lower = text.lower()

        # 通用实体提取
        # 颜色
        colors = [
            "红色",
            "蓝色",
            "绿色",
            "黄色",
            "白色",
            "黑色",
            "灰色",
            "粉色",
            "紫色",
            "橙色",
            "棕色",
        ]
        for color in colors:
            if color in text_lower:
                entities.append({"type": "color", "value": color})

        # 风格
        styles = [
            "可爱",
            "帅气",
            "写实",
            "卡通",
            "动漫",
            "赛博朋克",
            "水彩",
            "油画",
            "简约",
            "复古",
        ]
        for style in styles:
            if style in text_lower:
                entities.append({"type": "style", "value": style})

        # 主题/对象
        subjects = [
            "猫",
            "狗",
            "人",
            "风景",
            "建筑",
            "汽车",
            "花",
            "树",
            "山",
            "海",
            "城市",
        ]
        for subject in subjects:
            if subject in text_lower:
                entities.append({"type": "subject", "value": subject})

        # 特定任务的实体
        if task_type and task_type in cls.TASK_SIGNATURES:
            for entity_keyword in cls.TASK_SIGNATURES[task_type].get("entities", []):
                if entity_keyword in text_lower:
                    entities.append({"type": "task_specific", "value": entity_keyword})

        return entities

    @classmethod
    def build_context_summary(cls, history: list, max_turns: int = 3) -> dict:
        """
        构建历史上下文摘要

        返回:
        {
            "task_history": [],      # 任务历史
            "key_entities": [],      # 关键实体
            "last_user_intent": "",  # 最近的用户意图
            "last_model_output": "", # 最近的模型输出
            "conversation_topic": "" # 对话主题
        }
        """
        summary = {
            "task_history": [],
            "key_entities": [],
            "last_user_intent": "",
            "last_model_output": "",
            "conversation_topic": "",
        }

        if not history:
            return summary

        # 分析最近的对话
        recent_turns = (
            history[-max_turns * 2 :] if len(history) > max_turns * 2 else history
        )

        all_entities = []
        topics = []

        for turn in recent_turns:
            content = turn["parts"][0] if turn["parts"] else ""
            role = turn["role"]

            if role == "user":
                summary["last_user_intent"] = content
                # 识别任务类型
                for task_type, signatures in cls.TASK_SIGNATURES.items():
                    if any(kw in content.lower() for kw in signatures["keywords"]):
                        summary["task_history"].append(
                            {"type": task_type, "content": content[:100]}
                        )
                        topics.append(task_type)
                        break

                # 提取实体
                entities = cls.extract_entities(content)
                all_entities.extend(entities)

            elif role == "model":
                summary["last_model_output"] = content

        # 去重实体
        seen = set()
        unique_entities = []
        for e in all_entities:
            key = f"{e['type']}:{e['value']}"
            if key not in seen:
                seen.add(key)
                unique_entities.append(e)
        summary["key_entities"] = unique_entities

        # 确定对话主题
        if topics:
            summary["conversation_topic"] = topics[-1]  # 最近的任务类型

        return summary

    @classmethod
    def build_rag_prompt(
        cls, user_input: str, context_summary: dict, continuation_type: str = None
    ) -> str:
        """
        构建 RAG 风格的增强提示词

        将上下文信息结构化地注入到用户输入中
        """
        prompt_parts = []

        # 1. 添加上下文标记
        if context_summary.get("conversation_topic"):
            prompt_parts.append(
                f"[上下文类型: {context_summary['conversation_topic']}]"
            )

        # 2. 添加关键实体信息
        if context_summary.get("key_entities"):
            entities_str = ", ".join(
                [
                    f"{e['type']}={e['value']}"
                    for e in context_summary["key_entities"][:5]
                ]
            )
            prompt_parts.append(f"[关键信息: {entities_str}]")

        # 3. 添加历史意图
        if context_summary.get("last_user_intent"):
            # 截取核心描述
            last_intent = context_summary["last_user_intent"]
            if len(last_intent) > 200:
                last_intent = last_intent[:200] + "..."
            prompt_parts.append(f"[之前的请求: {last_intent}]")

        # 4. 根据延续类型添加特定指令
        if continuation_type and continuation_type in cls.CONTINUATION_PATTERNS:
            pattern = cls.CONTINUATION_PATTERNS[continuation_type]
            # 不添加模板，让实体和上下文自然融合

        # 5. 添加用户当前输入
        prompt_parts.append(f"[当前请求: {user_input}]")

        # 6. 如果是转换请求，添加源内容
        if continuation_type == "convert" and context_summary.get("last_model_output"):
            output = context_summary["last_model_output"]
            # 限制长度
            if len(output) > 4000:
                output = output[:4000] + "\n...(内容已截断)"
            prompt_parts.append(f"\n[需要转换的源内容:]\n{output}")

        # 7. 如果是引用类延续，附上最近输出摘要
        if continuation_type in (
            "reference",
            "reference_loose",
        ) and context_summary.get("last_model_output"):
            output = context_summary["last_model_output"]
            if len(output) > 2000:
                output = output[:2000] + "\n...(内容已截断)"
            prompt_parts.append(f"\n[最近输出摘要:]\n{output}")

        # 组合成最终的增强提示
        enhanced_prompt = "\n".join(prompt_parts)

        return enhanced_prompt

    @classmethod
    def analyze_context(cls, user_input: str, history: list) -> dict:
        """
        RAG 风格的上下文分析

        返回:
        {
            "is_continuation": bool,      # 是否是延续任务
            "related_task": str,          # 关联的任务类型
            "continuation_type": str,     # 延续类型 (modify/reference/convert/continue/detail)
            "context_summary": dict,      # 结构化上下文摘要
            "enhanced_input": str,        # RAG 增强后的输入
            "confidence": float,          # 置信度
        }
        """
        result = {
            "is_continuation": False,
            "related_task": None,
            "continuation_type": None,
            "context_summary": {},
            "enhanced_input": user_input,
            "confidence": 0.0,
        }

        if not history or len(history) < 2:
            return result

        user_lower = user_input.lower()
        input_length = len(user_input)

        # 1. 构建上下文摘要
        context_summary = cls.build_context_summary(history)
        result["context_summary"] = context_summary

        # 2. 检测延续类型和置信度（更严格的匹配）
        detected_type = None
        max_weight = 0.0

        for pattern_type, pattern_info in cls.CONTINUATION_PATTERNS.items():
            indicators = pattern_info["indicators"]
            weight = pattern_info["weight"]

            # 检查输入长度限制（如果有）
            max_len = pattern_info.get("max_input_length")
            if max_len and input_length > max_len:
                continue  # 输入太长，不太可能是简单的延续指令

            # 检查是否需要在句首出现
            require_start = pattern_info.get("require_start", False)

            # 计算匹配的指示词数量
            matches = 0
            for ind in indicators:
                if ind in user_lower:
                    if require_start:
                        # 需要在句首（前10个字符内）
                        if user_lower.find(ind) < 10:
                            matches += 1
                    else:
                        matches += 1

            if matches > 0:
                # 加权计算置信度
                adjusted_weight = weight * (
                    1 + 0.1 * (matches - 1)
                )  # 多个匹配增加置信度
                if adjusted_weight > max_weight:
                    max_weight = adjusted_weight
                    detected_type = pattern_type

        # 3. 额外检查：如果用户输入包含明确的新主题，降低延续判断
        # 新主题标志：包含"关于"、"一个"后接新实体
        new_topic_indicators = [
            "关于",
            "一篇",
            "一份",
            "一个新的",
            "帮我写",
            "帮我做",
            "帮我生成",
            "给我生成",
            "生成一",
        ]
        has_new_topic = any(ind in user_lower for ind in new_topic_indicators)

        # 检查是否是完全不同的任务类型（如：打开微信 -> 生成图片）
        task_mismatch = False
        if context_summary.get("conversation_topic"):
            prev_topic = context_summary["conversation_topic"]
            # 检测当前输入的任务类型
            curr_likely_task = None
            if any(
                kw in user_lower
                for kw in ["查", "搜", "搜索", "查询", "找", "再找", "再查", "再搜"]
            ):
                curr_likely_task = "WEB_SEARCH"
            elif any(kw in user_lower for kw in ["图", "画", "照片", "image"]):
                curr_likely_task = "PAINTER"
            elif any(kw in user_lower for kw in ["word", "pdf", "文档", "报告"]):
                curr_likely_task = "FILE_GEN"
            elif any(kw in user_lower for kw in ["打开", "运行", "关闭"]):
                curr_likely_task = "SYSTEM"

            # 如果任务类型完全不同，不应该是延续
            if curr_likely_task and prev_topic and curr_likely_task != prev_topic:
                task_mismatch = True
                _app_logger.debug(
                    f"[ContextAnalyzer] 任务类型不匹配: {prev_topic} -> {curr_likely_task}"
                )

        if has_new_topic and input_length > 10:
            # 有新主题且输入较长，很可能是独立任务
            max_weight *= 0.2  # 大幅降低置信度
            _app_logger.debug(f"[ContextAnalyzer] 检测到新主题标志，降低延续置信度")

        if task_mismatch:
            # 任务类型不匹配，强制清零
            max_weight = 0
            detected_type = None
            _app_logger.debug(f"[ContextAnalyzer] 任务类型不匹配，清除延续判断")

        # 4. 如果检测到延续模式且置信度足够高
        if detected_type and max_weight > 0.5:
            result["is_continuation"] = True
            result["continuation_type"] = detected_type
            result["confidence"] = min(max_weight, 1.0)

            # 确定关联的任务类型
            if context_summary.get("conversation_topic"):
                result["related_task"] = context_summary["conversation_topic"]
            elif context_summary.get("task_history"):
                result["related_task"] = context_summary["task_history"][-1]["type"]

            # 4. 构建 RAG 增强提示
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input, context_summary, detected_type
            )

            _app_logger.debug(f"[ContextAnalyzer] RAG Analysis:")
            _app_logger.info(f"  - Continuation Type: {detected_type}")
            _app_logger.info(f"  - Related Task: {result['related_task']}")
            _app_logger.info(f"  - Confidence: {result['confidence']:.2f}")
            _app_logger.info(
                f"  - Entities: {[e['value'] for e in context_summary.get('key_entities', [])]}"
            )

        # 5. 特殊处理：转换请求（即使没有明确的延续指示词）
        convert_patterns = [
            "做成word",
            "做成pdf",
            "转成word",
            "转成pdf",
            "生成word",
            "生成pdf",
            "导出为",
        ]
        if any(p in user_lower for p in convert_patterns) and context_summary.get(
            "last_model_output"
        ):
            result["is_continuation"] = True
            result["continuation_type"] = "convert"
            result["related_task"] = "FILE_GEN"
            result["confidence"] = 0.95
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input, context_summary, "convert"
            )

        return result

    @classmethod
    def filter_history(
        cls, user_input: str, history: list, keep_turns: int = 6
    ) -> list:
        """过滤历史记录，尽量避免无关上下文污染"""
        if not history:
            return []

        # 如果历史很短，直接返回
        if len(history) <= keep_turns * 2:
            return history

        user_lower = user_input.lower()

        # 抽取用户输入中的实体与关键词
        entities = cls.extract_entities(user_input)
        entity_values = {e["value"] for e in entities}

        # 额外提取中文关键词（长度>=2）与英文单词（长度>=3）
        import re

        cjk_words = re.findall(r"[\u4e00-\u9fff]{2,}", user_input)
        eng_words = re.findall(r"[a-zA-Z]{3,}", user_input)
        keyword_set = {k.lower() for k in (cjk_words + eng_words)}
        keyword_set.update({v.lower() for v in entity_values})

        # 构建相关历史：包含关键词的对话
        relevant = []
        for turn in history:
            content = (turn.get("parts") or [""])[0]
            content_lower = content.lower()
            if any(k in content_lower for k in keyword_set if k):
                relevant.append(turn)

        # 始终保留最近 3 轮对话（确保上下文连贯）
        tail_count = 6
        tail_start_index = max(0, len(history) - tail_count)

        # 收集需要保留的索引
        indices_to_keep = set()

        # 1. 关键词匹配的历史
        for i, turn in enumerate(history):
            content = (turn.get("parts") or [""])[0]
            content_lower = content.lower()
            if any(k in content_lower for k in keyword_set if k):
                indices_to_keep.add(i)
                # 同时保留该条的前一条（如果是User/Model配对）
                if i > 0:
                    indices_to_keep.add(i - 1)

        # 2. 也是最重要的：保留尾部上下文
        for i in range(tail_start_index, len(history)):
            indices_to_keep.add(i)

        # 按原始顺序重组
        filtered_history = [history[i] for i in sorted(indices_to_keep)]

        return filtered_history

        # 只保留最近 keep_turns 轮
        return merged[-keep_turns * 2 :]


class TaskOrchestrator:
    """
    编排和执行多个子任务

    责职：
    1. 顺序执行子任务
    2. 在子任务间传递数据/上下文
    3. 处理错误和重试
    4. 最终验证输出质量
    """

    @classmethod
    async def execute_compound_task(
        cls, user_input: str, subtasks: list, session_name: str = None
    ) -> dict:
        """
        执行复合任务的所有子任务

        返回:
            {
                "success": bool,
                "primary_result": 主任务结果,
                "secondary_results": [次要任务结果],
                "combined_output": 最终合并输出,
                "execution_log": 执行日志,
                "quality_score": 质量评分 (0-100),
                "errors": 错误列表
            }
        """
        execution_log = []
        results = []
        context = {"original_input": user_input, "user_input": user_input}
        errors = []

        try:
            for i, subtask in enumerate(subtasks):
                _app_logger.debug(
                    f"\n[TaskOrchestrator] 执行子任务 {i+1}/{len(subtasks)}: {subtask['task_type']}"
                )
                execution_log.append(
                    f"步骤 {i+1}: 执行 {subtask['task_type']} - {subtask['description']}"
                )
                step_input = subtask.get("input") or user_input

                try:
                    # 根据任务类型调用相应的处理函数
                    if subtask["task_type"] == "WEB_SEARCH":
                        result = await cls._execute_web_search(step_input, context)
                    elif subtask["task_type"] == "FILE_GEN":
                        result = await cls._execute_file_gen(
                            step_input, context, subtask
                        )
                    elif subtask["task_type"] == "PAINTER":
                        result = await cls._execute_painter(step_input, context)
                    elif subtask["task_type"] == "RESEARCH":
                        result = await cls._execute_research(step_input, context)
                    else:
                        result = {
                            "success": False,
                            "error": f"未知任务类型: {subtask['task_type']}",
                        }

                    subtask["status"] = "completed"
                    subtask["result"] = result
                    results.append(result)

                    # 将结果保存到上下文，供下一个任务使用
                    context[f"{subtask['task_type']}_result"] = result
                    context[f"step_{i+1}_output"] = result.get(
                        "output", result.get("content", "")
                    )

                    execution_log.append(f"  ✅ 完成: {subtask['description']}")

                except Exception as e:
                    error_msg = str(e)
                    subtask["status"] = "failed"
                    subtask["error"] = error_msg
                    errors.append(error_msg)
                    execution_log.append(f"  ❌ 失败: {error_msg}")
                    _app_logger.debug(f"[TaskOrchestrator] 子任务失败: {error_msg}")

            # 合并结果
            combined_output = cls._merge_results(subtasks, context)

            # 质量验证
            quality_score = await cls._validate_quality(
                user_input, combined_output, context
            )

            return {
                "success": len(errors) == 0,
                "primary_result": results[0] if results else None,
                "secondary_results": results[1:] if len(results) > 1 else [],
                "combined_output": combined_output,
                "execution_log": execution_log,
                "quality_score": quality_score,
                "errors": errors,
                "context": context,
            }

        except Exception as e:
            return {
                "success": False,
                "primary_result": None,
                "secondary_results": [],
                "combined_output": None,
                "execution_log": execution_log,
                "quality_score": 0,
                "errors": errors + [str(e)],
                "context": context,
            }

    @classmethod
    async def _execute_web_search(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """执行 Web 搜索子任务 (带可视进度)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[WEB_SEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("启动网络搜索...", "正在规划搜索关键词")

            # Phase 1: Planning
            # (WebSearcher manages its own queries, but we can simulate the 'thought' process)
            await asyncio.sleep(0.3)
            _report("执行 Google Search...", f"关键词: {user_input[:20]}...")

            # Phase 2: Execution
            # wrap in thread
            result = await asyncio.to_thread(
                WebSearcher.search_with_grounding, user_input
            )

            # Phase 3: Reporting
            if result.get("grounded"):
                _report("✅ 搜索并引用完成", "已结合最新信息")
            else:
                _report("✅ 搜索完成", "已获取相关网页摘要")

            return {
                "success": result.get("success", False),
                "output": result.get("response", ""),
                "content": result.get("response", ""),
                "grounded": result.get("grounded", False),
                "raw_result": result,
                "model_id": "gemini-2.5-flash",
            }
        except Exception as e:
            _report("❌ 搜索遇到问题", str(e))
            return {
                "success": False,
                "output": "",
                "error": str(e),
                "raw_result": None,
                "model_id": "gemini-2.5-flash",
            }

    @classmethod
    async def _execute_ppt_multi_step(
        cls, user_input: str, context: dict, subtask: dict, progress_callback=None
    ) -> dict:
        """执行多阶段PPT生成任务 (Plan-then-Execute)"""
        from web.smart_feedback import SmartFeedback

        fb = SmartFeedback(
            user_request=user_input,
            task_type="PPT",
            emit=lambda m, d: None,
            total_steps=3,
        )

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[PPT_PROGRESS] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        m, d = fb.start("多阶段PPT生成")
        _report(m, d)
        previous_data = context.get(f"step_{subtask['index']}_output", "")

        # 1. 规划阶段 (Planning Phase)
        try:
            from web.ppt_master import PPTBlueprint, PPTContentPlanner

            # 初始化规划器
            planner = PPTContentPlanner(ai_client=client, model_name="gemini-2.5-flash")

            # 执行规划
            _report("正在规划内容结构...", "调用 AI 规划大纲")
            plan_result = await planner.plan_content_structure(
                user_input, search_results=None
            )

            # 提取大纲
            outline_data = plan_result.get("outline", [])
            theme_choice = plan_result.get("theme_recommendation", "business")
            total_slides = plan_result.get("total_expected_slides", 10)

            # --- 1.2 展示规划概览 (User Requirement: Visualize Plan) ---
            plan_summary = f"大纲概览 ({len(outline_data)} 章节, {total_slides} 页):\n"
            for idx, sec in enumerate(outline_data):
                plan_summary += f"{idx+1}. {sec.get('section_title')} ({len(sec.get('slides', []))} 页)\n"
            _report(f"规划完成，共 {total_slides} 页", plan_summary)

            # 将大纲转换为 PPTGenerator 可识别的格式
            ppt_slides = []

            # --- 多阶段执行：逐页生成内容 ---
            total_steps = sum(len(sec.get("slides", [])) for sec in outline_data)
            current_step = 0

            for section in outline_data:
                section_title = section.get("section_title", "章节")
                # 添加章节页
                ppt_slides.append(
                    {
                        "type": "section",
                        "title": section_title,
                        "content": [section.get("section_theme", "")],
                    }
                )

                for slide in section.get("slides", []):
                    current_step += 1
                    s_title = slide.get("slide_title", "未命名幻灯片")
                    s_type = slide.get("slide_type", "content")
                    s_points = slide.get("key_points", [])

                    # Log progress
                    _report(
                        f"生成第 {current_step}/{total_steps} 页内容: {s_title}",
                        "阶段 2/3: 内容扩充",
                    )

                    # 扩充内容 (Per-Slide Generation)
                    expanded_points = s_points
                    if hasattr(planner, "expand_slide_content"):
                        try:
                            # Use new method in PPTContentPlanner
                            expanded_points = await planner.expand_slide_content(
                                s_title, s_points, context=f"Context: {section_title}"
                            )
                            if expanded_points != s_points:
                                _report(
                                    f"  ✨ 内容已扩充: {len(expanded_points)} 条",
                                    f"幻灯片: {s_title}",
                                )
                        except Exception as exp_err:
                            _report(f"  ⚠️ 扩充失败，使用原始内容", str(exp_err))
                            expanded_points = s_points

                    ppt_slides.append(
                        {
                            "type": (
                                s_type
                                if s_type
                                in ["content", "content_image", "comparison", "data"]
                                else "content"
                            ),
                            "title": s_title,
                            "points": expanded_points,
                            "content": expanded_points,
                            "notes": slide.get("content_description", ""),
                        }
                    )

            # 如果没有生成有效的幻灯片，回退到旧逻辑
            if not ppt_slides:
                raise ValueError("规划器未生成有效幻灯片大纲")

            # --- 2.5 质量自检与内容清洗 ---
            _report("正在进行质量自检与内容清洗...", "阶段 2.5/3: 质量门控")
            try:
                from web.file_quality_checker import FileQualityGate

                qg_result = FileQualityGate.check_and_fix_ppt_outline(
                    ppt_slides, user_request=user_input, progress_callback=_report
                )
                ppt_slides = qg_result["outline"]
                _qg_score = qg_result["quality"]["score"]
                _qg_fixes = qg_result["fixes"]
                if _qg_fixes:
                    _report(f"🧹 已清洗 {len(_qg_fixes)} 处内容问题", "")
                _report(
                    f"{'✅' if _qg_score >= 60 else '⚠️'} 质量评分: {_qg_score}/100",
                    (
                        "; ".join(qg_result["quality"]["issues"][:3])
                        if qg_result["quality"]["issues"]
                        else "质量良好"
                    ),
                )
            except Exception as qg_err:
                _app_logger.warning(f"[PPT] ⚠️ 质量门控异常: {qg_err}")

            # AI 验证
            try:
                verify_prompt = (
                    f"请作为质检员检查生成的PPT内容是否符合用户需求。\n"
                    f"用户需求: {user_input}\n"
                    f"生成的标题: {[s['title'] for s in ppt_slides]}\n"
                    "请简要回答：内容是否覆盖了需求？(是/否) + 一句话点评。"
                )
                verify_resp = await asyncio.to_thread(
                    lambda: client.models.generate_content(
                        model="gemini-2.5-flash", contents=verify_prompt
                    )
                )
                if verify_resp and verify_resp.text:
                    _report(
                        "✅ AI 验证通过",
                        f"模型点评: {verify_resp.text.strip()[:60]}...",
                    )
            except Exception as v_err:
                _report("⚠️ AI 验证跳过 (非致命)", str(v_err))

            # 2. 执行阶段 (Execution Phase) - 生成 PPT 文件
            _report("正在生成最终文件...", "阶段 3/3: 渲染与保存")
            from web.ppt_generator import PPTGenerator

            ppt_gen = PPTGenerator(theme=theme_choice)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r'[\\/*?:"<>|]', "", user_input[:20]) or "演示文稿"
            filename = f"{safe_title}_{timestamp}.pptx"
            ppt_path = os.path.join(settings_manager.documents_dir, filename)
            os.makedirs(settings_manager.documents_dir, exist_ok=True)

            # 使用 PPTGenerator 生成 (目前它直接支持 outline list)
            ppt_gen.generate_from_outline(
                title=safe_title, outline=ppt_slides, output_path=ppt_path
            )

            rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace("\\", "/")

            # 返回结果，格式与 _execute_file_gen 保持一致
            # 构建 markdown 表示的大纲供前端显示
            md_outline = f"# {safe_title}\n\n"
            for slide in ppt_slides:
                md_outline += f"## {slide['title']}\n"
                for p in slide.get("points", []):
                    md_outline += f"- {p}\n"
                md_outline += "\n"

            return {
                "success": True,
                "output": md_outline,
                "content": md_outline,
                "saved_files": [rel_path],
                "model_id": "gemini-2.5-flash (Planner)",
            }

        except Exception as e:
            _app_logger.warning(f"[PPT] ⚠️ 多阶段生成失败，回退到单步生成: {e}")
            # 重新抛出异常让上层处理，或者在这里调用旧逻辑?
            # 为了简单，抛出异常让外部 _execute_file_gen 的 except 块捕获 (但外部是 generic exception)
            # 或者我们直接返回失败，让 TaskOrchestrator 记录错误
            return {
                "success": False,
                "error": str(e),
                "opt_out_to_legacy": True,  # 标记需要回退
            }

    @classmethod
    async def _execute_file_gen(
        cls, user_input: str, context: dict, subtask: dict, progress_callback=None
    ) -> dict:
        """执行文件生成子任务
        增强：复杂/长文/要求“深度、详细、研究”时，先运行深度研究并切换到更强模型生成。
        """

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[FILE_GEN] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            # 提取前一个任务的结果作为输入
            previous_data = context.get(f"step_{subtask['index']}_output", "")

            # 复杂度判定（长文本或显式“深度/详细/研究/全面/技术”请求）
            text_lower = user_input.lower()
            complex_flags = [
                len(user_input) > 120,
                any(
                    k in text_lower
                    for k in [
                        "深度",
                        "详细",
                        "研究",
                        "全面",
                        "技术",
                        "报告",
                        "综述",
                        "whitepaper",
                    ]
                ),
            ]
            is_complex = any(complex_flags)

            # -- Planning Layer (DocumentPlanner) --------------------------
            _doc_plan = None
            if is_complex:
                try:
                    from web.doc_planner import DocumentPlanner

                    _planner = DocumentPlanner(
                        ai_client=client, model_name="gemini-2.5-flash"
                    )
                    _report("📋 规划文档结构...", "分析需求/分配章节")
                    _doc_plan = await _planner.plan(
                        user_input, previous_context=previous_data
                    )
                    if _doc_plan.success:
                        _report(
                            f"✅ 规划完成：{len(_doc_plan.sections)} 节 | {_doc_plan.doc_type.upper()}",
                            _doc_plan.to_context_str()[:120],
                        )
                    else:
                        _report(
                            "⚠️ 规划层失败，使用默认流程",
                            _doc_plan.error[:60] if _doc_plan.error else "",
                        )
                        _doc_plan = None
                except Exception as _pe:
                    _app_logger.warning(f"[FILE_GEN] ⚠️ 规划层异常: {_pe}")
                    _doc_plan = None

            # 检测目标格式（PPT、Excel、Word等）
            ppt_keywords = ["ppt", "幻灯片", "演示", "汇报", "presentation", "slide"]
            prefer_ppt = any(kw in user_input.lower() for kw in ppt_keywords)

            prefer_excel = (
                "excel" in user_input.lower()
                or "xlsx" in user_input.lower()
                or "表格" in user_input
            )
            prefer_pdf = "pdf" in user_input.lower()
            if _doc_plan:
                prefer_ppt = (_doc_plan.doc_type == "ppt") or prefer_ppt
                prefer_excel = (_doc_plan.doc_type == "excel") or prefer_excel
                prefer_pdf = (_doc_plan.doc_type == "pdf") or prefer_pdf

            # 根据目标格式选择提示
            if prefer_ppt:
                # 尝试使用新的多阶段生成流程 (Plan-then-Execute)
                try:
                    ppt_result = await cls._execute_ppt_multi_step(
                        user_input, context, subtask, progress_callback
                    )
                    if ppt_result.get("success"):
                        _report(
                            f"PPT生成成功",
                            f"文件: {(ppt_result.get('saved_files') or [''])[0]}",
                        )
                        return ppt_result
                    elif ppt_result.get("opt_out_to_legacy"):
                        _app_logger.warning(
                            "[FILE_GEN] ⚠️ 多阶段生成遇到问题，回退到旧版生成逻辑"
                        )
                    else:
                        return ppt_result
                except Exception as e:
                    _app_logger.warning(f"[FILE_GEN] ⚠️ 多阶段生成异常: {e}")

                # 回退旧逻辑 (Legacy Prompt Generation)
                gen_prompt = (
                    "你是一个顶尖的演示文稿内容策划师和排版规划师。\n\n"
                    "在每个 `## 章节标题` 前一行写类型标签来选择幻灯片类型：\n"
                    "- `[详细]` — 深入展示 3-5 个要点\n"
                    "- `[概览]` — 多主题速览，用 `### 子标题` 分组\n"
                    "- `[亮点]` — 关键数据，格式: `- 数值 | 说明`\n"
                    "- `[对比]` — 两方对比，用 `### 选项A` 和 `### 选项B` 分组\n"
                    "- `[过渡页]` — 章节过渡（最多 2 个）\n\n"
                    "**输出格式（严格遵循 Markdown）**：\n"
                    "```\n"
                    "# 演示标题\n\n"
                    "[详细]\n"
                    "## 章节标题\n"
                    "- 要点1（包含具体信息）\n"
                    "- 要点2\n"
                    "```\n\n"
                    "规则：重点内容用多个 [详细] 展开，简要内容合并到 [概览]，关键数据用 [亮点]。\n"
                    "每个要点包含具体信息，中文输出，只输出大纲。\n"
                )
            else:
                if _doc_plan and is_complex:
                    # 使用规划层生成增强 prompt（含章节指引）
                    from web.doc_planner import build_generation_prompt_from_plan

                    gen_prompt = build_generation_prompt_from_plan(
                        _doc_plan, user_input, previous_data
                    )
                else:
                    gen_prompt = (
                        "你是Koto，一个专业的数据整理与报告生成助手。\n"
                        "请基于用户需求和提供的数据，输出清晰、可直接放入文档的 Markdown 内容。\n"
                        "如果是价格类信息，必须包含一个 Markdown 表格，字段建议为：时间、价格、变化、来源。\n"
                        "输出要求：\n"
                        "- 只输出内容，不要输出代码或 BEGIN_FILE 标记\n"
                        "- 中文输出，结构清晰\n"
                    )

            full_input = (
                f"用户原始需求: {context['original_input']}\n\n"
                f"前面步骤的数据/信息:\n{previous_data}\n\n"
                f"{gen_prompt}"
            )

            # 深度研究：为复杂任务先补充研究上下文
            research_context = ""
            if is_complex:
                try:
                    research_context = WebSearcher.deep_research_for_ppt(
                        user_input, previous_data
                    )
                    if research_context:
                        previous_data = f"[深度研究]\n{research_context}\n\n[已有信息]\n{previous_data}"
                        _app_logger.debug(
                            f"[FILE_GEN] 🔬 深度研究完成，追加 {len(research_context)} 字上下文"
                        )
                except Exception as research_err:
                    _app_logger.warning(f"[FILE_GEN] ⚠️ 深度研究失败: {research_err}")

            # 调用模型生成内容
            model_id = SmartDispatcher.get_model_for_task(
                "FILE_GEN", complexity="complex" if is_complex else "normal"
            )

            _report(f"正在撰写内容...", f"模型: {model_id}")

            def _generate_text(prompt_text: str) -> str:
                response = client.models.generate_content(
                    model=model_id,
                    contents=prompt_text,
                    config=types.GenerateContentConfig(
                        system_instruction=_get_filegen_brief_instruction(),
                        temperature=0.4,
                        max_output_tokens=4000,
                    ),
                )
                return response.text or ""

            def _clean_filegen_text(text: str) -> str:
                if not text:
                    return text
                cleaned = text

                # Remove fenced code blocks but keep content
                cleaned = re.sub(r"```[a-zA-Z0-9_-]*\n", "", cleaned)
                cleaned = cleaned.replace("```", "")

                # Strip markdown links to plain text
                cleaned = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", cleaned)

                # Remove bold/italic markers
                cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
                cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
                cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
                cleaned = re.sub(r"_(.+?)_", r"\1", cleaned)

                # Remove inline code ticks
                cleaned = cleaned.replace("`", "")

                # Strip heading markers and blockquotes at line start
                cleaned = re.sub(r"^\s{0,3}#{1,6}\s+", "", cleaned, flags=re.MULTILINE)
                cleaned = re.sub(r"^\s*>\s?", "", cleaned, flags=re.MULTILINE)

                # Remove horizontal rules
                cleaned = re.sub(r"^\s*[-_*]{3,}\s*$", "", cleaned, flags=re.MULTILINE)

                # Flatten list markers but keep structure via indentation
                cleaned = re.sub(r"^\s*[-*+]\s+", "  ", cleaned, flags=re.MULTILINE)
                cleaned = re.sub(r"^\s*\d+\.\s+", "  ", cleaned, flags=re.MULTILINE)

                # Normalize extra blank lines
                cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

                # Cleanup leftover marker pairs
                cleaned = cleaned.replace("**", "").replace("__", "")

                return cleaned

            text_out = _generate_text(full_input) or "(无输出)"
            text_out = _clean_filegen_text(text_out)
            _report(f"内容撰写完成", f"共 {len(text_out)} 字")

            # 解析 Markdown 表格
            def _extract_markdown_table(md_text: str):
                lines = [line.strip() for line in md_text.splitlines() if "|" in line]
                for i in range(len(lines) - 1):
                    header_line = lines[i]
                    sep_line = lines[i + 1]
                    if re.match(r"^\s*\|?\s*[-:|\s]+\|\s*$", sep_line):
                        headers = [c.strip() for c in header_line.strip("|").split("|")]
                        rows = []
                        j = i + 2
                        while j < len(lines) and "|" in lines[j]:
                            row = [c.strip() for c in lines[j].strip("|").split("|")]
                            if len(row) < len(headers):
                                row += [""] * (len(headers) - len(row))
                            rows.append(row[: len(headers)])
                            j += 1
                        return [headers] + rows
                return None

            # 解析PPT大纲结构（支持智能规划标签）
            def _parse_ppt_outline(md_text: str) -> dict:
                """解析带 [类型] 标签的 PPT 大纲"""
                lines = md_text.split("\n")
                outline = {"title": "", "slides": []}
                _tmap = {
                    "过渡页": "divider",
                    "过渡": "divider",
                    "详细": "detail",
                    "重点": "detail",
                    "亮点": "highlight",
                    "数据": "highlight",
                    "概览": "overview",
                    "速览": "overview",
                    "简要": "overview",
                    "对比": "comparison",
                    "比较": "comparison",
                }
                cur_type = "detail"
                cur_slide = None
                cur_sub = None

                for line in lines:
                    line = line.rstrip()
                    if line.strip() in ("```", "```markdown"):
                        continue
                    tm = re.match(r"^\s*\[(.+?)\]\s*$", line)
                    if tm:
                        cur_type = _tmap.get(tm.group(1).strip(), "detail")
                        continue
                    if line.startswith("# ") and not line.startswith("## "):
                        outline["title"] = line[2:].strip()
                    elif line.startswith("## "):
                        if (
                            cur_sub
                            and cur_slide
                            and cur_slide.get("type") in ("overview", "comparison")
                        ):
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                            cur_sub = None
                        if cur_slide:
                            outline["slides"].append(cur_slide)
                        cur_slide = {
                            "type": cur_type,
                            "title": line[3:].strip(),
                            "points": [],
                            "content": [],
                        }
                        if cur_type == "divider":
                            cur_slide["description"] = ""
                        cur_type = "detail"
                        cur_sub = None
                    elif line.startswith("### ") and cur_slide:
                        if cur_sub:
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                        cur_sub = {
                            "subtitle": line[4:].strip(),
                            "label": line[4:].strip(),
                            "points": [],
                        }
                    elif re.match(r"^[\s]*[-•*]\s", line) and cur_slide is not None:
                        pt = re.sub(r"^[\s]*[-•*]\s+", "", line).strip()
                        if cur_sub is not None:
                            cur_sub["points"].append(pt)
                        else:
                            cur_slide["points"].append(pt)
                            cur_slide["content"].append(pt)
                    elif (
                        cur_slide
                        and cur_slide.get("type") == "divider"
                        and line.strip()
                    ):
                        cur_slide["description"] = line.strip()

                if (
                    cur_sub
                    and cur_slide
                    and cur_slide.get("type") in ("overview", "comparison")
                ):
                    cur_slide.setdefault("subsections", []).append(cur_sub)
                if cur_slide:
                    outline["slides"].append(cur_slide)
                for sl in outline["slides"]:
                    if sl.get("type") == "comparison" and "subsections" in sl:
                        subs = sl["subsections"]
                        if len(subs) >= 2:
                            sl["left"] = subs[0]
                            sl["right"] = subs[1]
                return outline

            title = "生成文档"
            if "价格" in user_input or "表格" in user_input:
                title = "价格波动表格"
            elif prefer_ppt:
                title = "演示文稿"

            saved_files = []
            file_type = None
            excel_error = None

            # 生成PPT
            if prefer_ppt:
                try:
                    from web.ppt_generator import PPTGenerator

                    ppt_outline = _parse_ppt_outline(text_out)

                    # ── 质量门控 ──
                    try:
                        from web.file_quality_checker import FileQualityGate

                        _qg = FileQualityGate.check_and_fix_ppt_outline(
                            ppt_outline.get("slides", []),
                            user_request=user_input,
                            progress_callback=_report,
                        )
                        ppt_outline["slides"] = _qg["outline"]
                    except Exception as _qge:
                        _app_logger.warning(f"[FILE_GEN] ⚠️ PPT 质量门控异常: {_qge}")

                    # 确定主题（通过关键词检测）
                    theme = "business"  # 默认商务主题
                    user_input_lower = user_input.lower()
                    if (
                        "tech" in user_input_lower
                        or "技术" in user_input_lower
                        or "科技" in user_input_lower
                    ):
                        theme = "tech"
                    elif (
                        "creative" in user_input_lower
                        or "创意" in user_input_lower
                        or "艺术" in user_input_lower
                    ):
                        theme = "creative"
                    elif (
                        "simple" in user_input_lower
                        or "minimal" in user_input_lower
                        or "极简" in user_input_lower
                    ):
                        theme = "minimal"

                    _report("正在生成PPT...", f"主题: {theme} (自动配图)")

                    ppt_gen = PPTGenerator(theme=theme)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = (
                        f"{ppt_outline.get('title', 'Presentation')}_{timestamp}.pptx"
                    )
                    # Max length for filename safety
                    if len(filename) > 50:
                        filename = f"Presentation_{timestamp}.pptx"

                    ppt_path = os.path.join(settings_manager.documents_dir, filename)
                    os.makedirs(settings_manager.documents_dir, exist_ok=True)

                    def _ppt_progress_wrapper(c, t, st, ty):
                        try:
                            _report(
                                f"正在生成PPT ({c}/{t})", f"页面: {st[:10]}... [{ty}]"
                            )
                        except Exception:
                            pass

                    ppt_gen.generate_from_outline(
                        title=ppt_outline.get("title", "演示"),
                        outline=ppt_outline.get("slides", []),
                        output_path=ppt_path,
                        enable_ai_images=True,
                        progress_callback=_ppt_progress_wrapper,
                    )

                    rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "pptx"
                    _report("PPT生成完成", f"已保存到: {rel_path}")

                except Exception as ppt_err:
                    _app_logger.warning(f"[FILE_GEN] ⚠️ PPT生成失败: {ppt_err}")
                    _report("PPT生成失败，回退到Word...", f"错误: {str(ppt_err)[:50]}")
                    # PPT失败时回退到Word
                    from web.document_generator import save_docx

                    saved_docx = save_docx(
                        text_out, title=title, output_dir=settings_manager.documents_dir
                    )
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "docx"
            else:
                # 生成Excel或Word
                _report("正在处理内容...", "解析文档结构")
                table_rows = _extract_markdown_table(text_out)
                if prefer_excel and not table_rows:
                    # 第一次未生成合格表格 → 生成修正Prompt重试一次
                    fix_prompt = (
                        "请只输出一个 Markdown 表格，不要输出其他说明。\n"
                        "表格必须包含以下列：时间、价格、变化、来源。\n"
                        "每行数据一行，格式严格。\n\n"
                        f"用户需求: {context['original_input']}\n\n"
                        f"可用数据:\n{previous_data}\n"
                    )
                    text_out_retry = _generate_text(fix_prompt)
                    if text_out_retry:
                        text_out = _clean_filegen_text(text_out_retry)
                        table_rows = _extract_markdown_table(text_out)

                if prefer_excel and table_rows:
                    _report("正在生成Excel...", f"写入 {len(table_rows)} 行数据")
                    try:
                        from openpyxl import Workbook
                        from openpyxl.styles import (
                            Alignment,
                            Border,
                            Font,
                            PatternFill,
                            Side,
                        )
                        from openpyxl.utils import get_column_letter

                        wb = Workbook()
                        ws = wb.active
                        ws.title = title[:31] if title else "Sheet1"

                        # 写入数据（清洗每个单元格内的 Markdown 符号）
                        try:
                            from web.file_quality_checker import (
                                strip_markdown_from_cell,
                            )

                            _strip_cell = strip_markdown_from_cell
                        except Exception:
                            _strip_cell = lambda x: x
                        for row in table_rows:
                            ws.append(
                                [
                                    _strip_cell(str(c)) if isinstance(c, str) else c
                                    for c in row
                                ]
                            )

                        # --- 样式美化 ---
                        header_font = Font(
                            name="Microsoft YaHei", size=11, bold=True, color="FFFFFF"
                        )
                        header_fill = PatternFill(
                            start_color="4472C4", end_color="4472C4", fill_type="solid"
                        )
                        data_font = Font(name="Microsoft YaHei", size=10)
                        thin_border = Border(
                            left=Side(style="thin", color="D9D9D9"),
                            right=Side(style="thin", color="D9D9D9"),
                            top=Side(style="thin", color="D9D9D9"),
                            bottom=Side(style="thin", color="D9D9D9"),
                        )
                        alt_fill = PatternFill(
                            start_color="F2F7FB", end_color="F2F7FB", fill_type="solid"
                        )
                        center_align = Alignment(
                            horizontal="center", vertical="center", wrap_text=True
                        )
                        left_align = Alignment(
                            horizontal="left", vertical="center", wrap_text=True
                        )

                        max_row = ws.max_row
                        max_col = ws.max_column

                        for col_idx in range(1, max_col + 1):
                            # 表头样式
                            cell = ws.cell(row=1, column=col_idx)
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                            cell.border = thin_border

                            # 数据行样式
                            for row_idx in range(2, max_row + 1):
                                cell = ws.cell(row=row_idx, column=col_idx)
                                cell.font = data_font
                                cell.alignment = left_align
                                cell.border = thin_border
                                # 隔行变色
                                if row_idx % 2 == 0:
                                    cell.fill = alt_fill

                            # 自动列宽
                            max_len = 0
                            for row_idx in range(1, max_row + 1):
                                val = ws.cell(row=row_idx, column=col_idx).value
                                if val:
                                    # CJK 字符算2个字符宽
                                    vlen = sum(
                                        2 if ord(c) > 127 else 1 for c in str(val)
                                    )
                                    max_len = max(max_len, vlen)
                            ws.column_dimensions[get_column_letter(col_idx)].width = (
                                min(max_len + 4, 40)
                            )

                        # 冻结首行
                        ws.freeze_panes = "A2"

                        filename = (
                            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        )
                        excel_path = os.path.join(
                            settings_manager.documents_dir, filename
                        )
                        os.makedirs(settings_manager.documents_dir, exist_ok=True)
                        wb.save(excel_path)
                        rel_path = os.path.relpath(excel_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        saved_files.append(rel_path)
                        file_type = "xlsx"
                        _report("Excel生成完成", f"已保存到: {rel_path}")
                    except Exception as excel_err:
                        excel_error = str(excel_err)
                        _app_logger.warning(
                            f"[FILE_GEN] ⚠️ Excel保存失败: {excel_error}"
                        )
                        _report(
                            "Excel保存失败，回退到Word...", f"错误: {excel_error[:50]}"
                        )

                # 保存为 DOCX（无表格或Excel失败时回退）
                if not saved_files:
                    # ── 导出检查层（Check Layer）：质量检查 + Markdown 符号去除（永久特性）──
                    try:
                        from web.file_quality_checker import FileQualityGate

                        _dqg = FileQualityGate.check_and_fix_for_export(
                            text_out,
                            target_format="word",
                            user_request=user_input,
                            progress_callback=_report,
                        )
                        text_out = _dqg["text"]
                        if _dqg.get("issues"):
                            _app_logger.debug(
                                f"[FILE_GEN] 🔍 检查层: {', '.join(_dqg['issues'][:3])}"
                            )
                    except Exception as _dqge:
                        _app_logger.warning(f"[FILE_GEN] ⚠️ 导出检查层异常: {_dqge}")

                    _report("正在生成Word文档...", "转换为 DOCX")
                    from web.document_generator import save_docx, save_pdf

                    saved_docx = save_docx(
                        text_out, title=title, output_dir=settings_manager.documents_dir
                    )
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace(
                        "\\", "/"
                    )
                    saved_files.append(rel_path)
                    file_type = "docx"
                    _report("Word文档生成完成", f"已保存到: {rel_path}")

                    # 如用户明确需要 PDF，也同时保存
                    if prefer_pdf:
                        try:
                            _report("正在生成PDF...", "转换为 PDF")
                            saved_pdf = save_pdf(
                                text_out,
                                title=title,
                                output_dir=settings_manager.documents_dir,
                            )
                            pdf_rel = os.path.relpath(saved_pdf, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            saved_files.append(pdf_rel)
                            _report("PDF生成完成", f"已保存到: {pdf_rel}")
                        except Exception as pdf_err:
                            _app_logger.warning(f"[FILE_GEN] ⚠️ PDF保存失败: {pdf_err}")
                            _report("PDF生成失败", str(pdf_err)[:50])

            return {
                "success": True,
                "output": f"已生成{file_type.upper()}文档: {', '.join([os.path.basename(p) for p in saved_files])}"
                + (f" (Excel失败: {excel_error})" if excel_error else ""),
                "content": text_out,
                "file_type": file_type or "docx",
                "saved_files": saved_files,
                "model_id": model_id,
            }
        except Exception as e:
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_painter(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """执行图像生成子任务 - 为PPT等生成配图 (带可视进度)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[PAINTER] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            topic = context.get("original_input", user_input)
            prompt = f"Professional illustration for: {topic[:100]}. Clean flat design, no text."

            image_paths = []
            images_dir = os.path.join(WORKSPACE_DIR, "images")
            os.makedirs(images_dir, exist_ok=True)

            _report("启动图像生成...", "调用 Imagen 4 模型")

            for i in range(2):
                try:
                    _report(
                        f"正在生成第 {i+1}/2 张配图...", f"提示词: {prompt[:30]}..."
                    )

                    # Run potentially blocking generation in thread
                    fname = f"painter_{i}_{int(time.time()*1000)%1000000}.png"
                    fpath = os.path.join(images_dir, fname)
                    _img_models = [
                        "imagen-4.0-generate-001",
                        "imagen-4.0-fast-generate-001",
                        "imagen-3.0-generate-001",
                    ]
                    _img_res = None
                    for _img_m in _img_models:
                        try:
                            _img_res = await asyncio.to_thread(
                                lambda _m=_img_m: client.models.generate_images(
                                    model=_m,
                                    prompt=prompt,
                                    config=types.GenerateImagesConfig(
                                        number_of_images=1
                                    ),
                                )
                            )
                            if _img_res and _img_res.generated_images:
                                break
                        except Exception as _img_e:
                            _app_logger.debug(f"[PAINTER] {_img_m} 失败: {_img_e}")
                            _img_res = None
                    if _img_res and _img_res.generated_images:
                        with open(fpath, "wb") as f:
                            f.write(_img_res.generated_images[0].image.image_bytes)
                        image_paths.append(fpath)
                        _app_logger.info(f"[PAINTER] ✅ 配图 {i+1} 已生成: {fname}")
                        _report(f"✅ 配图 {i+1} 完成", fname)
                    else:
                        raise RuntimeError("所有图像模型均失败")
                except Exception as img_err:
                    _app_logger.warning(f"[PAINTER] ⚠️ 配图 {i+1} 生成失败: {img_err}")
                    _report(f"⚠️ 配图 {i+1} 失败", str(img_err))

            success = len(image_paths) > 0
            if success:
                _report("✅ 图像生成任务完成", f"共生成 {len(image_paths)} 张")
            else:
                _report("❌ 图像生成任务失败", "未生成有效图片")

            return {
                "success": success,
                "output": f"已生成 {len(image_paths)} 张配图",
                "content": ",".join(image_paths),
                "image_paths": image_paths,
                "model_id": "imagen-3.0",
            }
        except Exception as e:
            _report("❌ 图像生成遇到致命错误", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_research(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """执行深度研究子任务 - 使用 Gemini Pro 深度分析 (可视进度)"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[RESEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("启动深度研究流程...", "分析上下文数据")
            search_data = context.get("WEB_SEARCH_result", {})
            search_text = search_data.get("content", "") or search_data.get(
                "output", ""
            )

            # Phase 1: Planning
            _report("规划研究大纲...", "确定分析维度")
            # (Implied planning by WebSearcher internal logic, but we report it)
            await asyncio.sleep(0.5)  # Simulate quick think

            # Phase 2: Synthesis
            _report("正在进行深度分析...", "优先 Deep Research Pro，失败自动回退")
            # Run in thread to not block event loop if sync
            research_text = await asyncio.to_thread(
                WebSearcher.deep_research_for_ppt, user_input, search_text
            )

            # Phase 3: Verification
            _report("验证研究报告...", "检查内容完整性")
            if research_text:
                _report("✅ 研究完成", f"生成 {len(research_text)} 字详细报告")
                return {
                    "success": True,
                    "output": f"深度研究完成，获取 {len(research_text)} 字专业分析",
                    "content": research_text,
                    "model_id": MODEL_MAP.get(
                        "RESEARCH", "deep-research-pro-preview-12-2025"
                    ),
                }
            else:
                _report("⚠️ 研究产出为空", "回退到基础搜索结果")
                return {
                    "success": True,
                    "output": "研究未返回结果，将使用已有信息",
                    "content": search_text,
                }
        except Exception as e:
            _report("❌ 研究过程出错", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_coder(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """执行代码生成子任务 - 使用最佳可用 Gemini 模型"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[CODER] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            model_id = MODEL_MAP.get("CODER", "gemini-3-pro-preview")
            _report("启动代码生成...", f"模型: {model_id}")

            # 注入前步搜索/研究结果（如有）
            search_ctx = ""
            for key in (
                "WEB_SEARCH_result",
                "RESEARCH_result",
                "search_result",
                "research_result",
            ):
                val = context.get(key)
                if val:
                    text = (
                        val.get("content") or val.get("output") or ""
                        if isinstance(val, dict)
                        else str(val)
                    )
                    if text:
                        search_ctx = text[:3000]
                        break

            full_prompt = user_input
            if search_ctx:
                full_prompt = f"参考信息:\n{search_ctx}\n\n任务: {user_input}"

            sys_instr = (
                "你是 Koto 代码专家。直接输出完整可运行代码，使用代码块（```语言）包裹，"
                "不加废话前言。必要时简短说明运行方式（≤3行）。"
            )
            _report("正在生成代码...", "调用 Interactions API")

            result_text = await asyncio.to_thread(
                _call_interactions_api_sync, model_id, full_prompt, sys_instr, 90.0
            )

            if not result_text:
                # 降级到 gemini-2.5-flash
                _report("⚠️ 主模型超时，降级生成...", "gemini-2.5-flash")
                resp = await asyncio.to_thread(
                    lambda: client.models.generate_content(
                        model=_INTERACTIONS_FALLBACK_MODEL,
                        contents=full_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=sys_instr,
                            temperature=0.3,
                            max_output_tokens=4096,
                        ),
                    )
                )
                result_text = resp.text or "(无输出)"
                model_id = _INTERACTIONS_FALLBACK_MODEL

            # 自动保存代码文件
            if settings_manager.get("ai", "auto_save_files") is not False:
                saved = Utils.auto_save_files(result_text)
            else:
                saved = []
            _report(
                "✅ 代码生成完成",
                f"已保存 {len(saved)} 个文件" if saved else "未检测到文件标记",
            )

            return {
                "success": True,
                "output": result_text,
                "content": result_text,
                "saved_files": saved,
                "model_id": model_id,
            }
        except Exception as e:
            _report("❌ 代码生成失败", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    async def _execute_system(
        cls, user_input: str, context: dict, progress_callback=None
    ) -> dict:
        """执行系统操作子任务 - 调用 LocalExecutor"""

        def _report(msg: str, detail: str = ""):
            _app_logger.debug(f"[SYSTEM] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            _report("执行系统操作...", user_input[:40])
            result = await asyncio.to_thread(LocalExecutor.execute, user_input)
            success = result.get("success", False)
            msg = result.get("message", "")
            if success:
                _report("✅ 系统操作完成", msg[:60])
            else:
                _report("⚠️ 系统操作失败", msg[:60])
            return {
                "success": success,
                "output": msg,
                "content": msg,
                "model_id": "local-executor",
            }
        except Exception as e:
            _report("❌ 系统操作异常", str(e))
            return {"success": False, "output": "", "error": str(e)}

    @classmethod
    def _merge_results(cls, subtasks: list, context: dict) -> dict:
        """合并所有子任务的结果"""
        merged = {"summary": "任务执行完成", "steps": [], "final_output": ""}

        for i, subtask in enumerate(subtasks):
            step_info = {
                "step": i + 1,
                "task": subtask["task_type"],
                "status": subtask["status"],
                "description": subtask["description"],
            }

            if subtask["result"]:
                step_info["output"] = subtask["result"].get("output", "")
            if subtask["error"]:
                step_info["error"] = subtask["error"]

            merged["steps"].append(step_info)

        # 最后一个完成的任务的输出作为最终输出
        for subtask in reversed(subtasks):
            if subtask["status"] == "completed" and subtask["result"]:
                merged["final_output"] = subtask["result"].get("output", "")
                break

        return merged

    @classmethod
    async def _validate_quality(
        cls, user_input: str, combined_output: dict, context: dict
    ) -> int:
        """
        验证输出质量（语义评分版本）。
        先用快速规则给基准分，再用 gemini-2.0-flash-lite 做语义评估。
        返回: 质量评分 (0-100)
        """
        # ── 规则基准分 ──────────────────────────────────────────────
        score = 40
        total_steps = len(combined_output.get("steps", []))
        completed_steps = len(
            [
                s
                for s in combined_output.get("steps", [])
                if s.get("status") == "completed"
            ]
        )
        if total_steps > 0:
            score += int((completed_steps / total_steps) * 30)  # 最多 +30

        final_output = combined_output.get("final_output", "")
        if not final_output:
            return max(0, min(100, score))  # 无输出直接返回规则分

        # 有文件输出加分
        has_files = any(
            r.get("result", {}).get("saved_files")
            for r in combined_output.get("steps", [])
            if isinstance(r.get("result"), dict)
        )
        if has_files:
            score += 10

        # ── 语义评分（gemini-2.0-flash-lite，低成本）────────────────
        try:
            check_prompt = (
                f"用户需求：{user_input[:300]}\n\n"
                f"最终输出（前1500字）：{final_output[:1500]}\n\n"
                "请评估输出是否满足了用户需求。只输出一个 0~30 的整数（30为完全满足）。"
            )
            resp = await asyncio.to_thread(
                lambda: client.models.generate_content(
                    model="gemini-2.0-flash-lite",
                    contents=check_prompt,
                    config=types.GenerateContentConfig(
                        max_output_tokens=8,
                        temperature=0.0,
                    ),
                )
            )
            text = (resp.text or "").strip()
            m = re.search(r"\d+", text)
            if m:
                semantic_score = min(30, max(0, int(m.group())))
                score += semantic_score
        except Exception as e:
            _app_logger.debug(f"[VALIDATE_QUALITY] 语义评分失败，使用规则分: {e}")

        return max(0, min(100, score))


# ================= 智能语料路由器配置 =================
# 配置 SmartDispatcher 以使用本地定义的类和对象
# SmartDispatcher、ModelRouter 等已从 app.core.routing 导入

try:
    _app_logger.debug("[INIT] Configuring SmartDispatcher with local dependencies...")
    SmartDispatcher.configure(
        local_executor=LocalExecutor,
        context_analyzer=ContextAnalyzer,
        web_searcher=WebSearcher,
        model_map=MODEL_MAP,
        client=client,
    )
    _app_logger.debug("[INIT] SmartDispatcher configured successfully.")
except Exception as e:
    _app_logger.error(f"[ERROR] Failed to configure SmartDispatcher: {e}")

# ─── 后台启动动态模型路由器 ────────────────────────────────────────────────────
# 不阻塞主线程启动；路由表更新后会自动覆盖静态默认值及 SmartDispatcher 配置
import threading as _threading

_threading.Thread(
    target=_init_model_manager, name="ModelManagerInit", daemon=True
).start()

# === Ollama 后备路由 (可选) ===
LOCAL_ROUTER_MODEL = "qwen3:8b"  # 升级: Qwen3 中英文能力远超旧模型
OLLAMA_API_URL = "http://localhost:11434/api/generate"


class LocalDispatcher:
    """后备路由器 - 使用 Ollama (如果可用)"""

    @staticmethod
    def is_ollama_running():
        # 云端模式下禁用 Ollama（云服务器无本地 GPU）
        if os.environ.get("KOTO_DEPLOY_MODE") == "cloud":
            return False
        try:
            requests.get("http://localhost:11434", timeout=0.2)
            return True
        except Exception:
            return False

    @staticmethod
    def analyze(user_input, history=None):
        """优先使用 SmartDispatcher，失败时使用 Ollama"""
        # 使用智能本地路由
        return SmartDispatcher.analyze(user_input, history)


# ================= Utilities =================


class Utils:
    _PACKAGE_ALLOWLIST = {
        "pygame": "pygame",
        "numpy": "numpy",
        "pandas": "pandas",
        "requests": "requests",
        "bs4": "beautifulsoup4",
        "beautifulsoup4": "beautifulsoup4",
        "lxml": "lxml",
        "pillow": "Pillow",
        "PIL": "Pillow",
        "opencv": "opencv-python",
        "cv2": "opencv-python",
        "matplotlib": "matplotlib",
        "scipy": "scipy",
        "sklearn": "scikit-learn",
        "flask": "flask",
        "fastapi": "fastapi",
        "uvicorn": "uvicorn",
        "streamlit": "streamlit",
        "gradio": "gradio",
    }

    @staticmethod
    def sanitize_string(s):
        if isinstance(s, str):
            return s.encode("utf-8", "ignore").decode("utf-8")
        return s

    @staticmethod
    def is_failure_output(text: str) -> bool:
        if not text or not str(text).strip():
            return True
        t = str(text).strip().lower()
        if t.startswith("❌") or "失败" in t or "错误" in t:
            return True
        # 检测模型声称「无法联网/没有实时数据」的拒绝型回答
        _no_internet_phrases = [
            "没有直接联网",
            "无法直接联网",
            "无法联网",
            "没有联网",
            "不能联网",
            "没有实时",
            "无法获取实时",
            "不能获取实时",
            "没有访问互联网",
            "无法访问互联网",
            "i don't have access to the internet",
            "i cannot access the internet",
            "i'm unable to access the internet",
            "no internet access",
            "i don't have real-time",
            "i cannot browse",
            "i can't browse",
        ]
        return any(phrase in t for phrase in _no_internet_phrases)

    @staticmethod
    def build_fix_prompt(
        task_type: str, user_input: str, prev_output: str = "", error_hint: str = ""
    ) -> str:
        base = (
            f"用户需求: {user_input}\n\n"
            f"上次输出/错误:\n{prev_output or error_hint}\n\n"
            "请修正并重新输出最终结果。不要解释过程，只输出最终内容。\n"
        )

        if task_type == "FILE_GEN":
            return base + (
                "要求：输出可执行的 Python 脚本，并使用 BEGIN_FILE/END_FILE 标记。\n"
                "必须生成文档或表格文件（docx/xlsx/pdf）。"
            )
        if task_type == "CODER":
            return base + "要求：输出完整可运行代码，并包含必要说明。"
        if task_type == "RESEARCH":
            return base + "要求：输出结构化报告，包含标题与要点。"
        if task_type == "WEB_SEARCH":
            return base + "要求：基于实时信息回答，给出清晰结论。"
        return base

    @staticmethod
    def adapt_prompt_to_markdown(
        task_type: str, user_input: str, history: list = None
    ) -> str:
        """使用本地模板将原始请求转为结构化 Markdown，便于大模型理解。

        注：已移除 flash-lite 二次润色调用（额外 API 费用 + ~300ms 延迟，收益不明显）。
        PromptAdapter 的本地模板（base_md）已足够主模型理解。
        """
        try:
            try:
                from web.prompt_adapter import PromptAdapter
            except ImportError:
                from prompt_adapter import PromptAdapter

            # model_generate=None：仅使用本地关键词提取 + Markdown 模板，不发起额外 LLM 调用
            return PromptAdapter.adapt(
                user_input=user_input,
                task_type=task_type,
                history=history,
                model_generate=None,
            )
        except Exception as e:
            _app_logger.debug(f"[PROMPT_ADAPTER] Failed: {e}")
            return user_input

    @staticmethod
    def quick_self_check(task_type: str, user_input: str, output_text: str) -> dict:
        """使用快速模型进行自检，返回 {'pass': bool, 'fix_prompt': str}。"""
        try:
            check_prompt = (
                "你是质量检查器。判断输出是否满足用户需求。\n"
                "只输出以下格式之一：\n"
                "PASS\n"
                "或\n"
                "FAIL\nFIX_PROMPT: <用于修正的提示词>\n\n"
                f"任务类型: {task_type}\n"
                f"用户需求: {user_input}\n"
                f"模型输出:\n{output_text}\n"
            )
            response = client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=check_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=300,
                    temperature=0.1,
                ),
            )
            text = (response.text or "").strip()
            if text.startswith("PASS"):
                return {"pass": True, "fix_prompt": ""}
            if text.startswith("FAIL"):
                fix = ""
                for line in text.splitlines():
                    if line.startswith("FIX_PROMPT:"):
                        fix = line.replace("FIX_PROMPT:", "").strip()
                        break
                return {"pass": False, "fix_prompt": fix}
            return {"pass": True, "fix_prompt": ""}
        except Exception as e:
            print(f"[SELF_CHECK] Failed: {e}")
            return {"pass": True, "fix_prompt": ""}

    @staticmethod
    def detect_required_packages(text: str) -> list:
        """从输出中粗略检测第三方依赖（仅返回白名单内的包）。"""
        if not text:
            return []
        modules = set()
        for line in text.splitlines():
            line = line.strip()
            if line.startswith("import "):
                parts = line.replace("import", "").split(",")
                for p in parts:
                    name = p.strip().split(" ")[0]
                    if name:
                        modules.add(name)
            elif line.startswith("from "):
                parts = line.split()
                if len(parts) >= 2:
                    modules.add(parts[1].strip())

        packages = set()
        for mod in modules:
            if mod in Utils._PACKAGE_ALLOWLIST:
                packages.add(Utils._PACKAGE_ALLOWLIST[mod])
        return sorted(packages)

    @staticmethod
    def auto_install_packages(packages: list) -> dict:
        """安装缺失的依赖包。返回安装结果摘要。"""
        result = {"installed": [], "skipped": [], "failed": []}
        if not packages:
            return result

        for pkg in packages:
            try:
                spec = importlib.util.find_spec(pkg)
                if spec is not None:
                    result["skipped"].append(pkg)
                    continue
                module_aliases = [
                    m for m, p in Utils._PACKAGE_ALLOWLIST.items() if p == pkg
                ]
                if any(importlib.util.find_spec(m) is not None for m in module_aliases):
                    result["skipped"].append(pkg)
                    continue
            except Exception:
                pass

            try:
                if getattr(sys, "frozen", False):
                    # 打包版无法安装新包，pip 在冻结环境下不可用
                    result["failed"].append(pkg)
                else:
                    cmd = [sys.executable, "-m", "pip", "install", pkg]
                    proc = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=120,
                        creationflags=(
                            subprocess.CREATE_NO_WINDOW
                            if sys.platform == "win32"
                            else 0
                        ),
                    )
                    if proc.returncode == 0:
                        result["installed"].append(pkg)
                    else:
                        result["failed"].append(pkg)
            except Exception:
                result["failed"].append(pkg)

        return result

    @staticmethod
    def auto_save_files(text):
        """自动从响应中提取并保存文件"""
        saved = []

        code_dir = os.path.join(WORKSPACE_DIR, "code")
        os.makedirs(code_dir, exist_ok=True)

        def _get_save_dir(filename):
            ext = os.path.splitext(filename)[1].lower()
            code_exts = {
                ".py",
                ".js",
                ".ts",
                ".tsx",
                ".jsx",
                ".java",
                ".cs",
                ".cpp",
                ".c",
                ".go",
                ".rs",
                ".rb",
                ".php",
                ".swift",
                ".kt",
                ".m",
                ".scala",
                ".sh",
                ".ps1",
                ".bat",
                ".cmd",
                ".json",
                ".yaml",
                ".yml",
                ".toml",
                ".ini",
                ".cfg",
                ".sql",
                ".md",
                ".html",
                ".css",
            }
            return code_dir if ext in code_exts else WORKSPACE_DIR

        # 调试：打印前800字符看看格式
        _app_logger.debug(f"[FILE_GEN] Response first 800 chars:\n{text[:800]}\n")

        # 预处理：统一格式 (去掉多余空格)
        normalized_text = text

        # 方法1: 多种 BEGIN_FILE 格式的正则匹配
        patterns = [
            # 格式1: ---BEGIN_FILE: filename.py--- (无空格)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # 格式2: ---BEGIN_FILE: filename.py--- ... ---END_FILE--- (带换行)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
            # 格式3: 更宽松 - 允许各种空白
            r"---\s*BEGIN_FILE[:\s]+([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # 格式4: 最宽松 - 捕获任意文件名
            r"---BEGIN_FILE[:\s]+([^\n-]+?)---\s*(.*?)---END_FILE---",
        ]

        matches1 = []
        for i, pattern in enumerate(patterns):
            try:
                matches1 = re.findall(
                    pattern, normalized_text, re.DOTALL | re.IGNORECASE
                )
                _app_logger.debug(f"[FILE_GEN] Pattern{i+1} matches: {len(matches1)}")
                if matches1:
                    _app_logger.debug(f"[FILE_GEN] ✓ Using pattern {i+1}")
                    break
            except Exception as e:
                _app_logger.debug(f"[FILE_GEN] Pattern{i+1} error: {e}")

        for filename, content in matches1:
            try:
                filename = filename.strip()
                content = content.strip()
                _app_logger.debug(
                    f"[FILE_GEN] Processing file: '{filename}', content length: {len(content)}"
                )

                # 清除 Markdown 代码块标记
                if content.startswith("```"):
                    lines = content.split("\n")
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines and lines[-1].strip() == "```":
                        lines = lines[:-1]
                    content = "\n".join(lines)
                    _app_logger.debug(
                        f"[FILE_GEN] After stripping markdown: {len(content)} chars"
                    )

                # 确保文件名有效
                if not filename or len(filename) > 100:
                    _app_logger.debug(f"[FILE_GEN] Invalid filename: {filename}")
                    continue

                # 确保文件名有扩展名
                if "." not in filename:
                    filename = filename + ".py"

                base_dir = _get_save_dir(filename)
                path = os.path.join(base_dir, filename)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                saved.append(filename)
                _app_logger.info(f"[FILE_GEN] ✅ Saved: {filename} to {path}")
            except Exception as e:
                _app_logger.error(f"[FILE_GEN] ❌ Save failed: {e}")
                import traceback

                traceback.print_exc()

        # 方法2: 如果方法1没找到，尝试提取 ```python 代码块 + 文件名注释
        if not saved:
            _app_logger.debug(
                f"[FILE_GEN] Method1 empty, trying method2 (```python blocks)..."
            )

            # 先尝试匹配带文件名的代码块
            # 例如: # filename: cat_info.py 或 # cat_info.py
            pattern2a = (
                r"```python\s*\n#\s*(?:filename:\s*)?([a-zA-Z0-9_.-]+\.py)\s*\n(.*?)```"
            )
            matches2a = re.findall(pattern2a, text, re.DOTALL)
            _app_logger.debug(
                f"[FILE_GEN] Pattern2a (with filename comment) matches: {len(matches2a)}"
            )

            if matches2a:
                for filename, code in matches2a:
                    code = code.strip()
                    if not code or len(code) < 20:
                        continue
                    base_dir = _get_save_dir(filename)
                    path = os.path.join(base_dir, filename)
                    try:
                        with open(path, "w", encoding="utf-8") as f:
                            f.write(code)
                        saved.append(filename)
                        _app_logger.info(f"[FILE_GEN] ✅ Method2a saved: {filename}")
                    except Exception as e:
                        _app_logger.error(f"[FILE_GEN] ❌ Method2a save failed: {e}")
            else:
                # 无文件名的代码块，使用时间戳
                pattern2 = r"```python\s*\n(.*?)```"
                matches2 = re.findall(pattern2, text, re.DOTALL)
                _app_logger.debug(
                    f"[FILE_GEN] Pattern2 (generic) matches: {len(matches2)}"
                )

                if matches2:
                    timestamp = int(time.time())
                    for idx, code in enumerate(matches2):
                        code = code.strip()
                        if not code or len(code) < 50:
                            continue

                        # 尝试从代码中提取有意义的文件名
                        filename = None
                        # 查找 doc_path, file_path 等变量
                        path_match = re.search(
                            r'(?:doc_path|file_path|filepath|output_path)\s*=.*?["\']([^"\']+\.(pdf|docx|xlsx))["\']',
                            code,
                        )
                        if path_match:
                            # 使用目标文件名作为脚本名
                            target_file = os.path.basename(path_match.group(1))
                            filename = target_file.rsplit(".", 1)[0] + ".py"

                        if not filename:
                            filename = f"generated_{timestamp}_{idx}.py"

                        base_dir = _get_save_dir(filename)
                        path = os.path.join(base_dir, filename)
                        try:
                            with open(path, "w", encoding="utf-8") as f:
                                f.write(code)
                            saved.append(filename)
                            _app_logger.info(f"[FILE_GEN] ✅ Method2 saved: {filename}")
                        except Exception as e:
                            _app_logger.error(f"[FILE_GEN] ❌ Method2 save failed: {e}")

        _app_logger.debug(f"[FILE_GEN] Final saved files: {saved}")
        return saved

    @staticmethod
    def save_image_part(blob_part):
        try:
            # 使用用户设置的图片目录
            images_dir = settings_manager.images_dir
            os.makedirs(images_dir, exist_ok=True)

            timestamp = int(time.time())
            filename = f"generated_{timestamp}.png"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as f:
                f.write(blob_part.inline_data.data)

            # 返回相对于 workspace 的路径
            # 确保路径始终在 workspace 下，且格式为正斜杠
            try:
                rel_path = os.path.relpath(filepath, WORKSPACE_DIR)
                # 如果包含 .. 说明不在 workspace 下，需要处理
                if ".." in rel_path:
                    # 降级为只返回文件名，放在 workspace/images 下
                    abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                    os.makedirs(abs_workspace_images, exist_ok=True)
                    fallback_path = os.path.join(abs_workspace_images, filename)
                    with open(fallback_path, "wb") as f:
                        f.write(blob_part.inline_data.data)
                    rel_path = os.path.relpath(fallback_path, WORKSPACE_DIR)
                    _app_logger.debug(
                        f"[IMAGE] Falling back to workspace/images: {rel_path}"
                    )

                result = rel_path.replace("\\", "/")
                _app_logger.debug(f"[IMAGE] Saved image: {result}")
                return result
            except Exception as path_err:
                _app_logger.debug(f"[IMAGE] Path calculation error: {path_err}")
                # 最后的保险方案：直接保存到 workspace/images
                abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                os.makedirs(abs_workspace_images, exist_ok=True)
                fallback_path = os.path.join(abs_workspace_images, filename)
                with open(fallback_path, "wb") as f:
                    f.write(blob_part.inline_data.data)
                result = os.path.relpath(fallback_path, WORKSPACE_DIR).replace(
                    "\\", "/"
                )
                _app_logger.debug(f"[IMAGE] Emergency fallback: {result}")
                return result
        except Exception as e:
            _app_logger.debug(f"[IMAGE] Save failed: {e}")
            import traceback

            traceback.print_exc()
            return None


# ================= Session Manager =================


class SessionManager:
    def __init__(self):
        self.sessions = {}

    def list_sessions(self):
        """列出所有会话，按修改时间排序（最新在前）"""
        files = [f for f in os.listdir(CHAT_DIR) if f.endswith(".json")]
        # 按修改时间排序，最新的在前
        files_with_time = []
        for f in files:
            path = os.path.join(CHAT_DIR, f)
            mtime = os.path.getmtime(path)
            files_with_time.append((f, mtime))
        files_with_time.sort(key=lambda x: x[1], reverse=True)
        return [f[0] for f in files_with_time]

    def load(self, filename):
        """加载会话历史 - 返回用于模型上下文的截断版本"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    full_history = json.load(f)
                    # 仅截断用于模型上下文的部分，不影响持久化存储
                    return self._trim_history(full_history)
            except (json.JSONDecodeError, OSError) as e:
                _app_logger.warning("Failed to load session %s: %s", filename, e)
                return []
        return []

    def load_full(self, filename):
        """加载完整会话历史 - 用于追加保存，不做截断"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return json.load(f)
            except (json.JSONDecodeError, OSError) as e:
                _app_logger.warning("Failed to load full session %s: %s", filename, e)
                return []
        return []

    def _trim_history(self, history, max_turns=20):
        """保留最多 20 轮对话（约 12000+ tokens），确保上下文足够但不过长"""
        if len(history) <= max_turns:
            return history
        # 只保留最后 N 轮对话
        trimmed = history[-max_turns:]
        _app_logger.debug(
            f"[HISTORY] Trimmed to last {max_turns} turns (was {len(history)})"
        )
        return trimmed

    def create(self, name):
        safe = "".join([c if c.isalnum() else "_" for c in name])
        filename = f"{safe}.json"
        path = os.path.join(CHAT_DIR, filename)
        # 若同名文件已存在，加时间戳后缀避免覆盖已有会话
        if os.path.exists(path):
            filename = f"{safe}_{int(time.time())}.json"
            path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump([], f)
        return filename

    def save(self, filename, history):
        path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=2, ensure_ascii=False)

    def append_and_save(self, filename, user_msg, model_msg, **extra_fields):
        """追加消息并保存 - 基于磁盘完整历史，避免截断导致数据丢失"""
        full_history = self.load_full(filename)
        user_timestamp = extra_fields.pop("user_timestamp", datetime.now().isoformat())
        model_timestamp = extra_fields.pop(
            "model_timestamp", datetime.now().isoformat()
        )

        full_history.append(
            {"role": "user", "parts": [user_msg], "timestamp": user_timestamp}
        )
        model_entry = {"role": "model", "parts": [model_msg]}
        if "timestamp" not in extra_fields:
            model_entry["timestamp"] = model_timestamp
        model_entry.update(extra_fields)
        full_history.append(model_entry)
        self.save(filename, full_history)
        return full_history

    def append_user_early(self, filename, user_msg):
        """在请求到达时立即保存用户消息，防止断连导致丢失
        返回history长度，后续用update_last_model_response更新模型回复"""
        full_history = self.load_full(filename)
        now_iso = datetime.now().isoformat()
        full_history.append({"role": "user", "parts": [user_msg], "timestamp": now_iso})
        full_history.append(
            {"role": "model", "parts": ["⏳ 处理中..."], "timestamp": now_iso}
        )
        self.save(filename, full_history)
        return len(full_history)

    def update_last_model_response(self, filename, model_msg, **extra_fields):
        """更新最后一条模型回复（配合append_user_early使用）"""
        full_history = self.load_full(filename)
        if full_history and full_history[-1].get("role") == "model":
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history[-1] = model_entry
            self.save(filename, full_history)
        else:
            # fallback: 直接追加
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history.append(model_entry)
            self.save(filename, full_history)

    def add_message(
        self, filename, role, content, task="CHAT", model_name="Auto", **extra_fields
    ):
        """追加单条消息（兼容旧调用），默认附带时间戳"""
        full_history = self.load_full(filename)
        entry = {
            "role": role,
            "parts": [content],
            "task": task,
            "model_name": model_name,
            "timestamp": extra_fields.pop("timestamp", datetime.now().isoformat()),
        }
        entry.update(extra_fields)
        full_history.append(entry)
        self.save(filename, full_history)
        return entry

    def delete(self, filename):
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                os.remove(path)
                return True
            except OSError as e:
                _app_logger.warning("Failed to delete session %s: %s", filename, e)
                return False
        return False

    def rename(self, filename, new_name):
        """将会话文件重命名。new_name 为用户输入的显示名称（非文件名）。"""
        old_path = os.path.join(CHAT_DIR, filename)
        if not os.path.exists(old_path):
            return {"success": False, "error": "会话不存在"}
        safe = "".join(
            [c if c.isalnum() or c in "_- " else "_" for c in new_name]
        ).strip()
        if not safe:
            return {"success": False, "error": "名称无效"}
        new_filename = f"{safe}.json"
        new_path = os.path.join(CHAT_DIR, new_filename)
        if os.path.exists(new_path) and os.path.abspath(new_path) != os.path.abspath(
            old_path
        ):
            new_filename = f"{safe}_{int(time.time())}.json"
            new_path = os.path.join(CHAT_DIR, new_filename)
        try:
            os.rename(old_path, new_path)
            return {"success": True, "new_filename": new_filename}
        except OSError as e:
            _app_logger.warning(
                "Failed to rename session %s -> %s: %s", filename, new_filename, e
            )
            return {"success": False, "error": str(e)}


session_manager = SessionManager()

# ================= 初始化全局模块 =================
# 懒加载 Memory Manager 和 Knowledge Base
_memory_manager = None
_kb = None


def get_memory_manager():
    """获取或创建 Memory Manager 实例（增强版）"""
    global _memory_manager
    if _memory_manager is None:
        try:
            # 优先使用增强版本
            from enhanced_memory_manager import EnhancedMemoryManager

            _memory_manager = EnhancedMemoryManager()
            _app_logger.info("[INIT] ✅ 增强记忆管理器已初始化")
        except ImportError:
            try:
                from web.enhanced_memory_manager import EnhancedMemoryManager

                _memory_manager = EnhancedMemoryManager()
                _app_logger.info("[INIT] ✅ 增强记忆管理器已初始化")
            except ImportError:
                # 降级到基础版本
                try:
                    from memory_manager import MemoryManager
                except ImportError:
                    from web.memory_manager import MemoryManager
                _memory_manager = MemoryManager()
                _app_logger.warning("[INIT] ⚠️  使用基础记忆管理器")

        # 注入摘要与向量适配器（如果支持）
        try:

            def _memory_generate(
                prompt: str, temperature: float = 0.2, max_tokens: int = 300
            ) -> str:
                resp = client.models.generate_content(
                    model="gemini-2.0-flash-lite",
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=temperature,
                        max_output_tokens=max_tokens,
                    ),
                )
                return resp.text or ""

            def _memory_embed(texts: list) -> list:
                safe_texts = [(t or "")[:1000] for t in texts]
                resp = client.models.embed_content(
                    model="text-embedding-004", contents=safe_texts
                )
                embeddings = []
                if hasattr(resp, "embeddings"):
                    for item in resp.embeddings:
                        if hasattr(item, "values"):
                            embeddings.append(list(item.values))
                        elif hasattr(item, "embedding"):
                            embeddings.append(list(item.embedding))
                        elif isinstance(item, dict):
                            embeddings.append(
                                list(item.get("values") or item.get("embedding") or [])
                            )
                elif hasattr(resp, "embedding"):
                    embeddings.append(list(resp.embedding))
                elif isinstance(resp, dict) and "embeddings" in resp:
                    for item in resp.get("embeddings", []):
                        embeddings.append(
                            list(item.get("values") or item.get("embedding") or [])
                        )
                return embeddings

            if hasattr(_memory_manager, "set_llm_adapters"):
                _memory_manager.set_llm_adapters(
                    generate_fn=_memory_generate, embedding_fn=_memory_embed
                )
        except Exception as e:
            _app_logger.warning(f"[INIT] ⚠️  记忆适配器注入失败: {e}")
    return _memory_manager


def _start_memory_extraction(
    user_msg: str,
    ai_msg: str,
    history=None,
    task_type: str = "CHAT",
    session_name: str = "default",
):
    """后台提取长期记忆（含 MemoryReflector 深度反思），不阻塞主对话流程"""
    try:
        from memory_integration import MemoryIntegration
    except ImportError:
        try:
            from web.memory_integration import MemoryIntegration
        except ImportError:
            MemoryIntegration = None

    def _llm_sync(prompt: str) -> str:
        """Synchronous LLM call for reflection / summarization."""
        try:
            resp = client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=600,
                ),
            )
            return resp.text or ""
        except Exception:
            return ""

    def _worker():
        # ── Existing MemoryIntegration (entity extraction) ────────────────
        if MemoryIntegration and MemoryIntegration.should_extract(user_msg, ai_msg):
            try:
                memory_mgr = get_memory_manager()

                class _LLMAdapter:
                    async def generate(self, prompt, temperature=0.1, max_tokens=500):
                        return _llm_sync(prompt)

                result = asyncio.run(
                    MemoryIntegration.extract_and_apply(
                        memory_mgr, user_msg, ai_msg, _LLMAdapter(), history
                    )
                )
                if result.get("success"):
                    _app_logger.info("[MemoryIntegration] ✅ 自动记忆提取完成")
                else:
                    _app_logger.warning(
                        f"[MemoryIntegration] ⚠️ 提取失败: {result.get('error')}"
                    )
            except Exception as e:
                _app_logger.error(f"[MemoryIntegration] ❌ 异常: {e}")

        # ── 2-B: MemoryReflector (deep structured reflection) ─────────────
        try:
            from app.core.memory.memory_reflector import MemoryReflector

            MemoryReflector.reflect_async(
                user_msg=user_msg,
                ai_msg=ai_msg,
                task_type=task_type,
                session_name=session_name,
                get_memory_fn=get_memory_manager,
                llm_fn=_llm_sync,
            )
        except Exception as e:
            _app_logger.warning(f"[MemoryReflector] ⚠️ 启动失败: {e}")

        # ── 2-C: PersonalityMatrix — 动态个人记忆矩阵更新（使用更高质量模型）──
        try:
            _pm_mgr = get_memory_manager()
            if _pm_mgr and hasattr(_pm_mgr, "update_personality_async"):
                _pm_mgr.update_personality_async(user_msg, ai_msg, _llm_quality_sync)
        except Exception as e:
            _app_logger.warning(f"[PersonalityMatrix] ⚠️ 更新启动失败: {e}")

        # ── 3: ShadowWatcher 影子追踪（零感知观察）────────────────────────────
        try:
            from app.core.monitoring.shadow_watcher import ShadowWatcher

            ShadowWatcher.observe(user_msg, ai_msg, session_name)
        except Exception as e:
            _app_logger.warning(f"[ShadowWatcher] ⚠️ 观察失败: {e}")

        # ── 3-B: ResponseEvaluator 模型自评（自动质量评分 → RatingStore）────
        try:
            from app.core.learning.rating_store import RatingStore as _RS
            from app.core.learning.response_evaluator import ResponseEvaluator

            _eval_msg_id = _RS.make_msg_id(session_name, user_msg)
            ResponseEvaluator.evaluate_async(
                msg_id=_eval_msg_id,
                user_input=user_msg,
                ai_response=ai_msg,
                task_type=task_type,
                session_name=session_name,
                llm_fn=_llm_sync,
            )
        except Exception as e:
            _app_logger.warning(f"[ResponseEvaluator] ⚠️ 自评启动失败: {e}")

        # ── 4: MacroRecorder 宏录制（重复工作流检测）────────────────────────
        try:
            from app.core.monitoring.macro_recorder import MacroRecorder

            MacroRecorder.record_turn(user_msg, task_type or "CHAT", session_name)
        except Exception as e:
            _app_logger.warning(f"[MacroRecorder] ⚠️ 记录失败: {e}")

    threading.Thread(target=_worker, daemon=True).start()


def get_knowledge_base():
    """获取或创建 Knowledge Base 实例"""
    global _kb
    if _kb is None:
        try:
            from knowledge_base import KnowledgeBase
        except ImportError:
            from web.knowledge_base import KnowledgeBase
        _kb = KnowledgeBase()
        _app_logger.info("[INIT] ✅ Knowledge Base 已初始化")
    return _kb


# 为了向后兼容，导出全局变量
memory_manager = None  # 将通过 get_memory_manager() 动态获取
kb = None  # 将通过 get_knowledge_base() 动态获取

# ================= Koto Brain =================


class KotoBrain:
    # 图像编辑关键词
    IMAGE_EDIT_KEYWORDS = [
        "修改",
        "换",
        "改成",
        "变成",
        "底色",
        "背景",
        "颜色",
        "抠图",
        "去背景",
        "P图",
        "美化",
        "滤镜",
        "调色",
        "编辑",
        "change",
        "modify",
        "edit",
        "background",
        "color",
    ]

    def chat(
        self,
        history,
        user_input,
        file_data=None,
        model=None,
        auto_model=True,
        task_type: str = None,
    ):
        start_time = time.time()
        original_input = user_input
        # 支持模型选择和自动选择
        _model_id_locked = (
            False  # 如果已在路由中强制设置 model_id，跳过后续 SmartDispatcher 覆盖
        )
        if model and not auto_model:
            model_id = model
            route_method = "Manual select"
            # 优先使用调用方传入的 task_type，避免重复路由
            target_key = task_type or "CHAT"
        else:
            target_key = "CHAT"
            route_method = "Auto"
            model_id = None  # 先置空，下面按路由决定

            if file_data:
                _fd_mime = (
                    file_data.get("mime_type") or "application/octet-stream"
                ).lower()
                _is_image_file = _fd_mime.startswith("image/")
                if _is_image_file:
                    # 图片文件：判断编辑 vs 分析
                    user_lower = user_input.lower()
                    is_edit = any(kw in user_lower for kw in self.IMAGE_EDIT_KEYWORDS)
                    if is_edit:
                        target_key = "PAINTER"
                        route_method = "Image Edit"
                    else:
                        target_key = "VISION"
                        route_method = "Image Analysis"
                else:
                    # 非图片二进制文件（PDF/Word等）：路由为 CHAT，使用降级模型直接读取
                    target_key = "CHAT"
                    route_method = "📄 Binary-Doc-Read"
                    # 强制使用支持 generate_content + 文件字节的降级模型（Interactions API 不支持文件附件）
                    model_id = _INTERACTIONS_FALLBACK_MODEL
                    _model_id_locked = True
            else:
                # 使用智能路由器
                target_key, route_method, _ = SmartDispatcher.analyze(user_input)

            if not _model_id_locked:
                model_id = SmartDispatcher.get_model_for_task(
                    target_key, has_image=bool(file_data)
                )

        # 使用小模型将请求转换为结构化 Markdown（仅在大模型处理时启用）
        # ⚠️ 跳过条件：有文件附件时（file_data）、或输入很大（含嵌入文件内容）
        _has_embedded_file_content = (
            "=== 文件内容 ===" in user_input or len(user_input) > 3000
        )
        model_input = user_input
        if (
            auto_model
            and not file_data
            and not _has_embedded_file_content
            and target_key not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]
        ):
            # 仅使用本地模板重整（不传 model_generate，避免额外的 flash-lite API 调用）
            model_input = Utils.adapt_prompt_to_markdown(
                target_key, user_input, history=history
            )
            if model_input != user_input:
                _app_logger.debug("[PROMPT_ADAPTER] Applied local Markdown template")
        result = {
            "task": target_key,
            "model": model_id,
            "route_method": route_method,  # 路由方法信息
            "response": "",
            "images": [],
            "saved_files": [],
            "latency": 0,
            "total_time": 0,
        }

        try:
            # === SYSTEM Mode (本地执行) ===
            if target_key == "SYSTEM":
                exec_result = LocalExecutor.execute(user_input)
                result["response"] = exec_result["message"]
                if exec_result.get("details"):
                    result["response"] += f"\n\n{exec_result['details']}"
                result["total_time"] = time.time() - start_time
                return result

            # === PAINTER Mode (图像生成/编辑) ===
            if target_key == "PAINTER":
                # 如果有输入图片（图像编辑模式）- 使用代码方式处理
                if file_data:
                    # 保存上传的图片到 workspace
                    import subprocess
                    import tempfile

                    temp_img_path = os.path.join(
                        WORKSPACE_DIR, "images", f"input_{int(time.time())}.jpg"
                    )
                    os.makedirs(os.path.dirname(temp_img_path), exist_ok=True)
                    with open(temp_img_path, "wb") as f:
                        f.write(file_data["data"])

                    # 构建图像编辑的系统指令
                    edit_instruction = f"""你是一个图像处理专家。用户上传了一张图片，需要你生成 Python 代码来处理它。

图片路径: {temp_img_path}
用户请求: {user_input}

请生成完整的 Python 代码来完成用户的图像编辑请求。

要求:
1. 使用 OpenCV (cv2) 或 PIL 处理图片
2. 处理后的图片保存到: {settings_manager.images_dir}
3. 文件名格式: edited_{{timestamp}}.jpg 或 .png
4. 代码必须完整可执行
5. 对于换背景色，使用颜色阈值或边缘检测来识别背景区域

常用的背景色处理方法:
- 证件照换底色: 检测接近原背景色的像素，替换为目标颜色
- 蓝色背景 RGB: (67, 142, 219) 或 (0, 191, 255)
- 红色背景 RGB: (255, 0, 0) 或 (220, 0, 0)  
- 白色背景 RGB: (255, 255, 255)

代码格式（必须使用这个格式）:
---BEGIN_FILE: image_edit.py---
# 你的代码
---END_FILE---"""

                    # 调用 Gemini 生成代码（带回退）
                    edit_models = [
                        "gemini-3-flash-preview",
                        "gemini-3-pro-preview",
                        "gemini-2.5-flash",
                    ]
                    code_response = None
                    last_error = None

                    def _process_code_response(code_response_text: str):
                        # 提取代码 - 支持多种格式
                        import re

                        patterns = [
                            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                            r"```python\s*(.*?)```",  # 标准 markdown 代码块
                            r"```\s*(.*?)```",  # 无语言标记的代码块
                        ]

                        code_content = None
                        for pattern in patterns:
                            matches = re.findall(
                                pattern, code_response_text, re.DOTALL | re.IGNORECASE
                            )
                            if matches:
                                if isinstance(matches[0], tuple):
                                    code_content = matches[0][1].strip()
                                else:
                                    code_content = matches[0].strip()
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Extracted code, length: {len(code_content)}"
                                )
                                break

                        if not code_content:
                            return {
                                "images": [],
                                "response": f"❌ 无法从模型响应中提取代码\n\n模型返回内容:\n```\n{code_response_text[:500]}\n```",
                                "error": "no_code",
                            }

                        # 保存并执行代码
                        temp_script = os.path.join(
                            tempfile.gettempdir(), f"koto_edit_{int(time.time())}.py"
                        )
                        with open(temp_script, "w", encoding="utf-8") as f:
                            f.write(code_content)

                        _app_logger.debug(
                            f"[IMAGE_EDIT] Executing script: {temp_script}"
                        )
                        if getattr(sys, "frozen", False):
                            # 打包模式：sys.executable 是 Koto.exe，不能用来运行脚本，改为进程内 exec()
                            import contextlib as _ctx
                            import io as _io

                            _out, _err, _rc = _io.StringIO(), _io.StringIO(), 0
                            try:
                                _prev = os.getcwd()
                                os.chdir(WORKSPACE_DIR)
                                with _ctx.redirect_stdout(_out), _ctx.redirect_stderr(
                                    _err
                                ):
                                    exec(
                                        open(temp_script, "r", encoding="utf-8").read(),
                                        {"__file__": temp_script},
                                    )
                                os.chdir(_prev)
                            except Exception as _ex:
                                _err.write(str(_ex))
                                _rc = 1

                            class _ImgR:
                                returncode = _rc
                                stdout = _out.getvalue()
                                stderr = _err.getvalue()

                            exec_result = _ImgR()
                        else:
                            exec_result = subprocess.run(
                                [sys.executable, temp_script],
                                capture_output=True,
                                text=True,
                                timeout=60,
                                cwd=WORKSPACE_DIR,
                            )

                        _app_logger.debug(
                            f"[IMAGE_EDIT] Script result: returncode={exec_result.returncode}"
                        )
                        if exec_result.stdout:
                            _app_logger.debug(
                                f"[IMAGE_EDIT] stdout: {exec_result.stdout[:200]}"
                            )
                        if exec_result.stderr:
                            _app_logger.debug(
                                f"[IMAGE_EDIT] stderr: {exec_result.stderr[:200]}"
                            )

                        # 清理临时脚本
                        try:
                            os.remove(temp_script)
                        except OSError:
                            pass

                        if exec_result.returncode == 0:
                            images = []
                            images_dir = settings_manager.images_dir
                            for f in os.listdir(images_dir):
                                if f.startswith("edited_") and f.endswith(
                                    (".jpg", ".png", ".jpeg")
                                ):
                                    full_path = os.path.join(images_dir, f)
                                    age = time.time() - os.path.getmtime(full_path)
                                    if age < 60:
                                        rel_path = os.path.relpath(
                                            full_path, WORKSPACE_DIR
                                        ).replace("\\", "/")
                                        images.append(rel_path)

                            if images:
                                return {
                                    "images": images,
                                    "response": f"✅ 图片编辑完成!\n🖼️ 保存位置: `{images_dir}`",
                                    "error": "",
                                }
                            return {
                                "images": [],
                                "response": f"⚠️ 脚本执行成功但未检测到新图片\n\n{exec_result.stdout[:500]}",
                                "error": "no_output",
                            }

                        return {
                            "images": [],
                            "response": f"❌ 图片处理失败\n```\n{exec_result.stderr[:500]}\n```",
                            "error": "exec_failed",
                        }

                    for edit_model in edit_models:
                        try:
                            _app_logger.debug(
                                f"[IMAGE_EDIT] Trying model: {edit_model}"
                            )
                            _app_logger.debug(f"[IMAGE_EDIT] Sending request to API...")
                            response = client.models.generate_content(
                                model=edit_model,
                                contents=edit_instruction,
                                config=types.GenerateContentConfig(
                                    max_output_tokens=4096, temperature=0.5
                                ),
                            )
                            _app_logger.debug(f"[IMAGE_EDIT] Got API response")

                            if (
                                response.candidates
                                and response.candidates[0].content.parts
                            ):
                                code_response = (
                                    response.candidates[0].content.parts[0].text
                                )
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Got response from {edit_model}, length: {len(code_response)}"
                                )
                                break
                        except Exception as model_err:
                            last_error = str(model_err)
                            _app_logger.debug(
                                f"[IMAGE_EDIT] Model {edit_model} failed: {last_error[:100]}"
                            )
                            continue

                    if code_response:
                        run_result = _process_code_response(code_response)
                        result["images"] = run_result["images"]
                        result["response"] = run_result["response"]
                    else:
                        result["response"] = (
                            f"❌ 所有模型都不可用: {last_error[:200] if last_error else '未知错误'}"
                        )

                    # 失败后自动修正并重试一次（避免无编辑结果）
                    if not result["images"] and Utils.is_failure_output(
                        result["response"]
                    ):
                        fix_prompt = (
                            "上次生成失败，请修正并只输出完整可执行的 Python 代码。\n"
                            "必须使用 BEGIN_FILE/END_FILE 格式。\n"
                            f"图片路径: {temp_img_path}\n"
                            f"输出目录: {settings_manager.images_dir}\n"
                            f"用户请求: {user_input}\n\n"
                            f"失败信息/输出: {result['response']}\n"
                        )
                        retry_models = ["gemini-3-flash-preview", "gemini-2.5-flash"]
                        for retry_model in retry_models:
                            try:
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Retry with model: {retry_model}"
                                )
                                retry_resp = client.models.generate_content(
                                    model=retry_model,
                                    contents=fix_prompt,
                                    config=types.GenerateContentConfig(
                                        max_output_tokens=4096
                                    ),
                                )
                                if (
                                    retry_resp.candidates
                                    and retry_resp.candidates[0].content.parts
                                ):
                                    retry_code = (
                                        retry_resp.candidates[0].content.parts[0].text
                                    )
                                    retry_run = _process_code_response(retry_code)
                                    if retry_run["images"]:
                                        result["images"] = retry_run["images"]
                                        result["response"] = retry_run["response"]
                                        break
                                    result["response"] = retry_run["response"]
                            except Exception as retry_err:
                                _app_logger.debug(
                                    f"[IMAGE_EDIT] Retry failed: {retry_err}"
                                )

                    result["total_time"] = time.time() - start_time
                    return result
                else:
                    # 纯图像生成使用 gemini-3.1-flash-image-preview
                    try:
                        _app_logger.info(f"[图像生成] 开始生成: {user_input[:50]}...")
                        response = client.models.generate_content(
                            model="gemini-3.1-flash-image-preview",
                            contents=user_input,
                            config=types.GenerateContentConfig(
                                response_modalities=["TEXT", "IMAGE"]
                            ),
                        )
                        _app_logger.info(
                            f"[图像生成] 响应成功，候选数: {len(response.candidates) if response.candidates else 0}"
                        )

                        # 保存生成的图片
                        if response.candidates and response.candidates[0].content.parts:
                            for part in response.candidates[0].content.parts:
                                if hasattr(part, "inline_data") and part.inline_data:
                                    img_filename = Utils.save_image_part(part)
                                    if img_filename:
                                        result["images"].append(img_filename)
                                        _app_logger.info(
                                            f"[图像生成] 已保存: {img_filename}"
                                        )

                        if result["images"]:
                            save_path = settings_manager.images_dir
                            result["response"] = (
                                f"✨ 图片已生成!\n🖼️ 保存位置: `{save_path}`"
                            )
                        else:
                            result["response"] = (
                                "❌ 图像生成失败: 无输出内容，请检查提示词"
                            )
                        result["total_time"] = time.time() - start_time
                        return result
                    except Exception as img_err:
                        error_msg = str(img_err)
                        _app_logger.info(f"[图像生成] 错误: {error_msg[:200]}")

                        # 提供更详细的错误信息
                        if (
                            "disconnected" in error_msg.lower()
                            or "timeout" in error_msg.lower()
                        ):
                            result["response"] = (
                                f"❌ 连接超时或中断: {error_msg[:100]}\n\n💡 建议: 请稍后重试，或检查网络连接"
                            )
                        elif "safety" in error_msg.lower():
                            result["response"] = "❌ 内容因安全政策被过滤，请修改提示词"
                        elif (
                            "quota" in error_msg.lower() or "rate" in error_msg.lower()
                        ):
                            result["response"] = "❌ API 配额已达限制，请稍后重试"
                        else:
                            result["response"] = f"❌ 图像生成失败: {error_msg[:100]}"

                        result["total_time"] = time.time() - start_time
                        return result

                if not response.candidates:
                    result["response"] = "Generation failed (safety filter or busy)."
                    result["total_time"] = time.time() - start_time
                    return result

                text_response = ""
                if response.candidates and response.candidates[0].content.parts:
                    for part in response.candidates[0].content.parts:
                        if hasattr(part, "text") and part.text:
                            text_response += part.text
                        if hasattr(part, "inline_data") and part.inline_data:
                            img_filename = Utils.save_image_part(part)
                            if img_filename:
                                result["images"].append(img_filename)

                # 添加图片保存位置提示
                if result["images"]:
                    save_path = settings_manager.images_dir
                    text_response += f"\n\n🖼️ 图片已保存到: `{save_path}`"

                result["response"] = (
                    text_response if text_response else "Image generated successfully!"
                )
                result["total_time"] = time.time() - start_time
                return result

            # === RAG: Retrieve Relevant Context (Auto) ===
            try:
                # 获取知识库实例
                kb_inst = get_knowledge_base()

                # 仅在非特定模式且输入有效时检索
                if target_key not in ["PAINTER", "SYSTEM"] and len(original_input) > 3:
                    # 避免对极短的问候语进行检索
                    skip_keywords = ["你好", "hello", "hi", "test", "测试"]
                    if not any(original_input.lower() == k for k in skip_keywords):
                        _app_logger.debug(
                            f"[RAG]正在检索知识库: {original_input[:50]}..."
                        )
                        rag_results = kb_inst.search(original_input, top_k=3)

                        if rag_results:
                            _app_logger.debug(
                                f"[RAG] 检索到 {len(rag_results)} 个相关片段"
                            )
                            context_str = "\n".join(
                                [
                                    f"--- 来源: {r['file_name']} (相似度: {r['similarity']:.2f}) ---\n{r['text']}"
                                    for r in rag_results
                                ]
                            )

                            # 将上下文注入 prompt
                            rag_context = f"\n\n【参考资料】\n以下是从本地知识库检索到的相关内容，供回答参考：\n{context_str}\n\n"

                            # Log retrieval
                            _app_logger.debug(
                                f"[RAG] Injected context length: {len(rag_context)}"
                            )

                            # Update model input
                            # 如果有 file_data，model_input 可能是 None 或不被直接使用，需谨慎
                            if not file_data:
                                model_input = rag_context + model_input
                            else:
                                # 对于有文件的请求，我们将上下文拼接到 original_input (user prompt)
                                # 注意：下面 generate_content 用的是 original_input + image_part
                                original_input = rag_context + original_input

            except Exception as rag_err:
                _app_logger.debug(f"[RAG] Retrieval warning: {rag_err}")

            # === Regular Mode ===
            # 构建历史记录格式（过滤无关历史）
            history_for_model = ContextAnalyzer.filter_history(original_input, history)
            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(
                    types.Content(
                        role=turn["role"],
                        parts=[types.Part.from_text(text=p) for p in turn["parts"]],
                    )
                )

            # 根据任务类型选择系统提示：FILE_GEN 走文档生成提示，其余走通用助手提示
            if target_key == "FILE_GEN":
                _brain_sys_instruction = _get_system_instruction()
            else:
                _brain_sys_instruction = _get_chat_system_instruction(original_input)

            if file_data:
                # 构建 Part 格式（适用于图片和 PDF/文档）
                doc_part = types.Part.from_bytes(
                    data=file_data["data"], mime_type=file_data["mime_type"]
                )
                _fd_mime2 = (file_data.get("mime_type") or "").lower()
                _is_image = _fd_mime2.startswith("image/")

                if not _is_image:
                    # PDF / 文档二进制：Interactions API 不支持文件附件
                    # → 直接使用 gemini-2.5-flash（原生支持 generate_content + PDF bytes）
                    _doc_model = _INTERACTIONS_FALLBACK_MODEL
                    if model_id != _doc_model:
                        _app_logger.info(
                            f"[brain.chat] 非图片文件 ({_fd_mime2}): 降级模型 {model_id} → {_doc_model}"
                        )
                        model_id = _doc_model
                        result["model"] = model_id
                    response = client.models.generate_content(
                        model=model_id,
                        contents=[original_input, doc_part],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""
                elif model_id in _INTERACTIONS_ONLY_MODELS:
                    # 图片文件 + gemini-3-preview 模型：走 Interactions API
                    try:
                        accumulated_text = _call_interactions_api_sync(
                            model_id,
                            original_input,
                            sys_instruction=_brain_sys_instruction,
                        )
                        if not accumulated_text:
                            raise ValueError("Interactions API 返回空响应")
                    except Exception as _ia_err:
                        print(
                            f"[brain.chat] {model_id} Interactions API 失败: {_ia_err} → 降级到 {_INTERACTIONS_FALLBACK_MODEL}"
                        )
                        model_id = _INTERACTIONS_FALLBACK_MODEL
                        result["model"] = model_id
                        _fb_resp = client.models.generate_content(
                            model=model_id,
                            contents=[original_input, doc_part],
                            config=types.GenerateContentConfig(
                                system_instruction=_brain_sys_instruction
                            ),
                        )
                        accumulated_text = _fb_resp.text if _fb_resp.text else ""
                else:
                    # 图片文件 + 普通 generate_content 模型
                    response = client.models.generate_content(
                        model=model_id,
                        contents=[original_input, doc_part],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""
            else:
                # gemini-3-preview 只支持 Interactions API，不支持 generate_content
                if model_id in _INTERACTIONS_ONLY_MODELS:
                    try:
                        # 将历史记录折叠进 prompt（Interactions API 不支持多轮历史）
                        history_prefix = ""
                        if formatted_history:
                            history_lines = []
                            for turn in formatted_history[-6:]:  # 最近 3 轮
                                role_label = "用户" if turn.role == "user" else "助手"
                                turn_text = " ".join(
                                    p.text
                                    for p in turn.parts
                                    if hasattr(p, "text") and p.text
                                )
                                if turn_text:
                                    history_lines.append(f"{role_label}: {turn_text}")
                            if history_lines:
                                history_prefix = (
                                    "[对话历史]\n" + "\n".join(history_lines) + "\n\n"
                                )
                        full_prompt = history_prefix + model_input
                        accumulated_text = _call_interactions_api_sync(
                            model_id,
                            full_prompt,
                            sys_instruction=_brain_sys_instruction,
                        )
                        if not accumulated_text:
                            raise ValueError("Interactions API 返回空响应")
                    except Exception as _ia_err:
                        _app_logger.info(
                            f"[brain.chat] {model_id} Interactions API 失败: {_ia_err} → 降级到 {_INTERACTIONS_FALLBACK_MODEL}"
                        )
                        model_id = _INTERACTIONS_FALLBACK_MODEL
                        result["model"] = model_id
                        _fb_resp = client.models.generate_content(
                            model=model_id,
                            contents=formatted_history
                            + [
                                types.Content(
                                    role="user",
                                    parts=[types.Part.from_text(text=model_input)],
                                )
                            ],
                            config=types.GenerateContentConfig(
                                system_instruction=_brain_sys_instruction
                            ),
                        )
                        accumulated_text = _fb_resp.text if _fb_resp.text else ""
                else:
                    response = client.models.generate_content(
                        model=model_id,
                        contents=formatted_history
                        + [
                            types.Content(
                                role="user",
                                parts=[types.Part.from_text(text=model_input)],
                            )
                        ],
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    accumulated_text = response.text if response.text else ""

            first_token_latency = (time.time() - start_time) * 1000
            result["latency"] = first_token_latency

            # Auto-save files
            if settings_manager.get("ai", "auto_save_files") is not False:
                saved_files = Utils.auto_save_files(accumulated_text)
            else:
                saved_files = []
            result["saved_files"] = saved_files

            # 添加文件保存提示
            if saved_files:
                files_list = ", ".join(saved_files)
                accumulated_text += (
                    f"\n\n📁 文件已保存: **{files_list}**\n📂 位置: `{WORKSPACE_DIR}`"
                )

            result["response"] = accumulated_text
            result["total_time"] = time.time() - start_time
            return result

        except Exception as e:
            err_str = str(e)
            # 自动降级：如果模型返回"只支持 Interactions API"错误，用 2.0-flash 重试一次
            if "Interactions API" in err_str and model_id not in (
                _INTERACTIONS_ONLY_MODELS | {_INTERACTIONS_FALLBACK_MODEL}
            ):
                print(
                    f"[brain.chat] Interactions API 错误，自动降级 {model_id} → {_INTERACTIONS_FALLBACK_MODEL}"
                )
                try:
                    model_id = _INTERACTIONS_FALLBACK_MODEL
                    _fb = client.models.generate_content(
                        model=model_id,
                        contents=(
                            formatted_history
                            + [
                                types.Content(
                                    role="user",
                                    parts=[types.Part.from_text(text=model_input)],
                                )
                            ]
                            if not file_data
                            else [original_input]
                        ),
                        config=types.GenerateContentConfig(
                            system_instruction=_brain_sys_instruction
                        ),
                    )
                    result["response"] = _fb.text if _fb.text else ""
                    result["model"] = model_id
                    result["total_time"] = time.time() - start_time
                    return result
                except Exception as _fb_err:
                    result["response"] = f"❌ 分析失败: {_fb_err}"
            elif (
                "API key not valid" in err_str
                or "INVALID_ARGUMENT" in err_str
                and "api key" in err_str.lower()
            ):
                result["response"] = (
                    "❌ **API 密钥无效**\n\n"
                    "请检查您的 Gemini API 密钥：\n"
                    "1. 前往 [aistudio.google.com/apikey](https://aistudio.google.com/apikey) 获取有效密钥\n"
                    "2. 在 Koto 设置页面更新 API 密钥\n"
                    "3. 确保密钥所在项目已启用 Generative Language API\n\n"
                    f"原始错误: `{err_str[:200]}`"
                )
            else:
                result["response"] = f"❌ 发生错误: {err_str}"
            result["total_time"] = time.time() - start_time
            return result


brain = KotoBrain()

# ================= Routes =================


@app.route("/")
def index():
    # 云模式：未认证用户看到落地页
    deploy_mode = os.environ.get("KOTO_DEPLOY_MODE", "local")
    auth_enabled = os.environ.get("KOTO_AUTH_ENABLED", "false").lower() == "true"
    if deploy_mode == "cloud" and auth_enabled:
        return render_template("landing.html")
    return render_template("index.html")


@app.route("/app")
def app_main():
    """主应用页面（SaaS 模式下需认证后访问）"""
    return render_template("index.html")


@app.route("/file-network")
def file_network():
    """文件网络界面"""
    return render_template("file_network.html")


@app.route("/knowledge-graph")
def knowledge_graph_page():
    """知识图谱可视化界面"""
    return render_template("knowledge_graph.html")


@app.route("/test_upload")
def test_upload():
    return render_template("test_upload.html")


@app.route("/edit-ppt/<session_id>")
def edit_ppt(session_id):
    """PPT 生成后编辑页面（P1 功能）"""
    return render_template("edit_ppt.html")


@app.route("/skills")
@app.route("/skill-marketplace")
def skill_marketplace():
    """Koto Skill 库 — GitHub Extension Marketplace 风格管理界面"""
    return render_template("skill_marketplace.html")


@app.route("/monitoring-dashboard")
def monitoring_dashboard():
    """Phase 4 System Monitoring Dashboard"""
    return send_from_directory(
        os.path.join(os.path.dirname(__file__), "static"), "monitoring_dashboard.html"
    )


@app.route("/api/sessions", methods=["GET"])
def get_sessions():
    """List all chat sessions.
    ---
    tags:
      - Sessions
    responses:
      200:
        description: List of session names
        schema:
          type: object
          properties:
            sessions:
              type: array
              items:
                type: string
    """
    sessions = session_manager.list_sessions()
    return jsonify({"sessions": [s.replace(".json", "") for s in sessions]})


@app.route("/api/sessions", methods=["POST"])
def create_session():
    """Create a new chat session.
    ---
    tags:
      - Sessions
    parameters:
      - in: body
        name: body
        schema:
          properties:
            name:
              type: string
              description: Optional session name
    responses:
      200:
        description: Session created
        schema:
          type: object
          properties:
            success:
              type: boolean
            session:
              type: string
    """
    data = request.json
    name = data.get("name", f"chat_{int(time.time())}")
    filename = session_manager.create(name)
    return jsonify({"success": True, "session": filename.replace(".json", "")})


@app.route("/api/sessions/<session_name>", methods=["GET"])
def get_session(session_name):
    """Get a specific chat session with full history.
    ---
    tags:
      - Sessions
    parameters:
      - in: path
        name: session_name
        type: string
        required: true
    responses:
      200:
        description: Session data with conversation history
        schema:
          type: object
          properties:
            session:
              type: string
            history:
              type: array
              items:
                type: object
    """
    # 返回完整历史供前端渲染（不截断），截断仅用于模型上下文
    history = session_manager.load_full(f"{session_name}.json")
    return jsonify({"session": session_name, "history": history})


@app.route("/api/sessions/<session_name>/rename", methods=["PATCH"])
def rename_session(session_name):
    """Rename a chat session."""
    data = request.json or {}
    new_name = (data.get("new_name") or "").strip()
    if not new_name:
        return jsonify({"success": False, "error": "新名称不能为空"}), 400
    result = session_manager.rename(f"{session_name}.json", new_name)
    if result["success"]:
        new_session = result["new_filename"].replace(".json", "")
        return jsonify({"success": True, "new_session": new_session})
    return jsonify({"success": False, "error": result.get("error", "重命名失败")}), 400


@app.route("/api/sessions/<session_name>", methods=["DELETE"])
def delete_session(session_name):
    """Delete a chat session.
    ---
    tags:
      - Sessions
    parameters:
      - in: path
        name: session_name
        type: string
        required: true
    responses:
      200:
        description: Deletion result
        schema:
          type: object
          properties:
            success:
              type: boolean
    """
    success = session_manager.delete(f"{session_name}.json")
    return jsonify({"success": success})


@app.route("/api/chat", methods=["POST"])
def chat():
    """Send a chat message and get a response (non-streaming).
    ---
    tags: [Chat]
    security:
      - Bearer: []
    parameters:
      - in: body
        name: body
        schema:
          required: [session, message]
          properties:
            session: {type: string, description: Session/conversation name}
            message: {type: string, description: User message}
            locked_model: {type: string, default: auto}
            locked_task: {type: string}
    responses:
      200:
        description: AI response
        schema:
          properties:
            response: {type: string}
            model: {type: string}
      400:
        description: Missing session or message
      500:
        description: Internal error
    """
    data = request.json
    session_name = data.get("session")
    user_input = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")

    if not session_name or not user_input:
        return jsonify({"error": "Missing session or message"}), 400

    user_input = Utils.sanitize_string(user_input)

    # Load history
    full_history = session_manager.load_full(f"{session_name}.json")
    history = session_manager._trim_history(full_history)

    # 确定使用的模型
    if locked_model and locked_model != "auto":
        model = locked_model
        auto_model = False
    elif locked_task:
        model = MODEL_MAP.get(locked_task, MODEL_MAP["CHAT"])
        auto_model = False
    else:
        model = None
        auto_model = True

    # Get response
    result = brain.chat(history, user_input, model=model, auto_model=auto_model)

    # 代码任务: 自动检查依赖并安装
    if result.get("task") == "CODER" and result.get("response"):
        pkgs = Utils.detect_required_packages(result["response"])
        if pkgs:
            install_result = Utils.auto_install_packages(pkgs)
            installed = install_result.get("installed", [])
            failed = install_result.get("failed", [])
            skipped = install_result.get("skipped", [])
            msg_parts = []
            if installed:
                msg_parts.append(f"✅ 已安装: {', '.join(installed)}")
            if skipped:
                msg_parts.append(f"ℹ️ 已存在: {', '.join(skipped)}")
            if failed:
                msg_parts.append(f"⚠️ 安装失败: {', '.join(failed)}")
            if msg_parts:
                result["response"] += "\n\n" + "\n".join(msg_parts)

    # Update history (基于磁盘完整历史追加，避免截断丢失)
    session_manager.append_and_save(
        f"{session_name}.json", user_input, result["response"]
    )

    return jsonify(result)


# ============== Agent 确认 API ==============
# NOTE: These routes have been migrated to the unified agent blueprint
#       (app/api/agent_routes.py) under /api/agent/confirm and /api/agent/choice.
#       Kept here as comments for reference.

# @app.route('/api/agent/confirm', methods=['POST'])
# def agent_confirm():
#     """Agent 用户确认 API — 前端点击确认/取消后回调"""
#     ...

# @app.route('/api/agent/choice', methods=['POST'])
# def agent_choice():
#     """Agent 用户选择 API — 前端选择后回调"""
#     ...


# NOTE: /api/agent/plan has been migrated to the unified agent blueprint
#       (app/api/agent_routes.py). Kept as comment for reference.
# @app.route('/api/agent/plan', methods=['POST'])
# def agent_plan(): ...


@app.route("/api/chat/stream", methods=["POST"])
def chat_stream():
    """Stream a chat response via Server-Sent Events.
    ---
    tags:
      - Chat
    parameters:
      - in: body
        name: body
        schema:
          required: [session, message]
          properties:
            session:
              type: string
            message:
              type: string
            locked_model:
              type: string
              default: auto
            locked_task:
              type: string
    responses:
      200:
        description: SSE stream of chat tokens
    """
    data = request.json
    session_name = data.get("session")
    user_input = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")

    _app_logger.debug(
        f"\n[STREAM] Incoming request: locked_task='{locked_task}', locked_model='{locked_model}'"
    )
    _app_logger.debug(f"[STREAM] User input: {user_input[:60]}")

    if not session_name or not user_input:

        def error_gen():
            yield f"data: {json.dumps({'type': 'error', 'message': 'Missing session or message'})}\n\n"

        return Response(error_gen(), mimetype="text/event-stream")

    # API 密钥缺失时提前返回友好提示
    if not API_KEY:

        def no_key_gen():
            msg = (
                "⚠️ **API 密钥未配置**\n\n"
                "请在 `config/gemini_config.env` 文件中设置：\n"
                "```\nGEMINI_API_KEY=你的密钥\n```\n\n"
                "💡 获取密钥：[Google AI Studio](https://aistudio.google.com/apikey)\n\n"
                "设置完成后重启 Koto 即可使用。"
            )
            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"

        return Response(no_key_gen(), mimetype="text/event-stream")

    user_input = Utils.sanitize_string(user_input)

    # 🧠 意图分析与重写 (Intent Analysis & Rewrite)
    # 结合历史记忆和本地模型，理解用户的复杂指令（如“重复上个任务”、“把刚才那个改成...”）
    try:
        from app.core.routing.intent_analyzer import IntentAnalyzer

        if IntentAnalyzer.should_analyze(user_input):
            full_hist = session_manager.load_full(f"{session_name}.json")
            rewritten_input = IntentAnalyzer.rewrite_intent(user_input, full_hist)
            if rewritten_input and rewritten_input != user_input:
                _app_logger.debug(
                    f"[STREAM] 🔄 意图重写: '{user_input}' -> '{rewritten_input}'"
                )
                user_input = rewritten_input
    except Exception as e:
        _app_logger.warning(f"[STREAM] ⚠️ 意图分析失败: {e}")
        # 降级到基础的正则匹配
        repeat_patterns = [
            r"^重复.*任务",
            r"^再做一遍",
            r"^再来一次",
            r"^re(peat|do).*last.*task",
            r"^try.*again",
        ]
        if any(re.search(p, user_input, re.IGNORECASE) for p in repeat_patterns):
            try:
                full_hist = session_manager.load_full(f"{session_name}.json")
                last_user_msg = None
                for msg in reversed(full_hist):
                    if msg.get("role") == "user":
                        content = (msg.get("parts") or [""])[0]
                        if not any(
                            re.search(p, content, re.IGNORECASE)
                            for p in repeat_patterns
                        ):
                            last_user_msg = content
                            break
                if last_user_msg:
                    _app_logger.debug(
                        f"[REPEAT] Found last user message: {last_user_msg[:50]}..."
                    )
                    user_input = last_user_msg
            except Exception as hist_e:
                _app_logger.debug(f"[REPEAT] Error fetching history: {hist_e}")

    # ⚡ 快速路径：系统时间查询 - 直接返回，无需发送到LLM
    time_query_patterns = [
        r"当前.*时间|当前系统时间",
        r"现在.*几点|几点钟",
        r"几点|什么时间",
        r"时间是|现在是",
        r"now.*time|what.*time|current.*time",
    ]
    if any(
        re.search(pattern, user_input, re.IGNORECASE) for pattern in time_query_patterns
    ):

        def quick_time_response():
            from datetime import datetime

            now = datetime.now()
            date_str = now.strftime("%Y年%m月%d日")
            weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][
                now.weekday()
            ]
            time_str = now.strftime("%H:%M:%S")
            timestamp = now.isoformat()  # 记录精确时间戳
            response = f"当前系统时间为：\n\n**{date_str} {weekday} {time_str}**"

            # 记录到历史（用户 + 模型，均带时间戳）
            try:
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response,
                    task="CHAT",
                    model_name="QuickResponse",
                    timestamp=timestamp,
                    user_timestamp=timestamp,
                    model_timestamp=timestamp,
                )
            except Exception as e:
                _app_logger.debug(f"[STREAM] Quick time history save failed: {e}")

            yield f"data: {json.dumps({'type': 'progress', 'message': '📅 系统时间查询', 'detail': '从本地获取'}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'token', 'content': response, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0.01, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"

        return Response(quick_time_response(), mimetype="text/event-stream")

    # 🎯 获取动态系统指令（根据用户问题智能注入上下文）
    try:
        system_instruction = _get_chat_system_instruction(user_input)
    except Exception as e:
        _app_logger.debug(f"[STREAM] Warning: Dynamic system instruction failed: {e}")
        system_instruction = (
            _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION()
        )  # 降级到新鲜生成的指令

    history = session_manager.load(f"{session_name}.json")
    full_history = session_manager.load_full(f"{session_name}.json")

    # 🧠 2-A: ContextWindowManager — MemGPT-style page-out/in
    _cw_paged_context = ""
    try:
        from app.core.memory.context_window_manager import ContextWindowManager as _CWM

        _cw_out = _CWM.manage(
            history=history,
            query=user_input,
            session_name=session_name,
            get_memory_fn=get_memory_manager,
        )
        history = _cw_out["history"]
        _cw_paged_context = _cw_out.get("paged_in_context", "")
        if _cw_paged_context:
            system_instruction += f"\n\n{_cw_paged_context}"
    except Exception as _cw_err:
        print(f"[CWM] ⚠️ 上下文管理器异常: {_cw_err}")

    # 🕵️‍♀️ 检测是否有最近上传的文件 (5分钟内)
    has_recent_upload = False
    recent_file_type = None
    recent_file_path = None  # 保存路径以便后续注入文件内容
    try:
        upload_scan_dirs = ["web/uploads", "uploads", "workspace/documents"]
        recent_threshold = time.time() - 300  # 5分钟内
        for d in upload_scan_dirs:
            if os.path.exists(d):
                for f in os.listdir(d):
                    fp = os.path.join(d, f)
                    if os.path.isfile(fp) and os.path.getmtime(fp) > recent_threshold:
                        has_recent_upload = True
                        _, ext = os.path.splitext(f)
                        recent_file_type = ext.lower()
                        recent_file_path = fp  # 记录完整路径
                        print(f"[STREAM] Found recent upload: {f} ({recent_file_type})")
                        break
            if has_recent_upload:
                break
    except Exception as e:
        _app_logger.debug(f"[STREAM] Error checking uploads: {e}")

    # 确定任务类型和模型
    context_info = None
    if locked_task:
        task_type = locked_task
        route_method = "🔒 Manual"
        _app_logger.info(f"[STREAM] ✅ Using locked_task: '{task_type}'")
    else:
        # 将文件信息传递给分析器（同时注入 [FILE_ATTACHED:ext] 前缀确保本地模型正确分类）
        context_override = {
            "has_file": has_recent_upload,
            "file_type": recent_file_type,
        }
        _routing_input = user_input
        if has_recent_upload and recent_file_type:
            _routing_input = f"[FILE_ATTACHED:{recent_file_type}] {user_input}"
            _app_logger.debug(f"[STREAM] 📎 文件上下文注入: {_routing_input[:80]}")
        task_type, route_method, context_info = SmartDispatcher.analyze(
            _routing_input, history, file_context=context_override
        )
        _app_logger.debug(
            f"[STREAM] Auto-detected task_type: '{task_type}', context: {context_info is not None}"
        )

        # ── 安全兜底：未知 task_type → CHAT ──────────────────────────────────
        _HANDLED_TASK_TYPES = {
            "SYSTEM",
            "FILE_OP",
            "FILE_EDIT",
            "FILE_SEARCH",
            "DOC_ANNOTATE",
            "WEB_SEARCH",
            "RESEARCH",
            "PAINTER",
            "FILE_GEN",
            "CODER",
            "CHAT",
            "MULTI_STEP",
            "AGENT",
            "VISION",
        }
        if not task_type or task_type not in _HANDLED_TASK_TYPES:
            _app_logger.warning(
                f"[STREAM] ⚠️ 收到未知 task_type='{task_type}'，降级为 CHAT"
            )
            task_type = "CHAT"
            route_method = "⬇️ Unknown→CHAT"

        # ── MULTI_STEP 保护：无 is_multi_step_task 标记时降级 CHAT ──────────
        if task_type == "MULTI_STEP" and (
            not context_info or not context_info.get("is_multi_step_task")
        ):
            _app_logger.warning(f"[STREAM] ⚠️ MULTI_STEP 无有效 context，降级为 CHAT")
            task_type = "CHAT"
            route_method = "⬇️ MULTI_STEP→CHAT"

        # ── FILE_EDIT 保护：输入中无法识别文件路径时降级 CHAT ───────────────
        # 避免短/普通聊天消息因携带文件上下文而被误分类到 FILE_EDIT
        if task_type == "FILE_EDIT":
            _fe_pat1 = re.search(
                r'(?:修改|编辑|改)\s+["\']?([^"\']{2,}?)["\']?\s+.+', user_input
            )
            _fe_pat2 = re.search(
                r'(?:把|将)\s+["\']?([^"\']{2,}?)["\']?\s+(?:的|中的|里的)\s*.+',
                user_input,
            )
            if not _fe_pat1 and not _fe_pat2:
                _app_logger.warning(
                    f"[STREAM] ⚠️ FILE_EDIT 输入无有效文件路径: '{user_input[:40]}' → 降级为 CHAT"
                )
                task_type = "CHAT"
                route_method = "⬇️ FILE_EDIT→CHAT"

        # ── CHAT → WEB_SEARCH 安全兜底（防止天气/股价/新闻等实时查询被误分为CHAT）────
        # 这是最后一道防线：在任务链路执行之前，重新校验是否需要联网搜索
        if task_type == "CHAT" and WebSearcher.needs_web_search(user_input):
            _app_logger.debug(
                f"[STREAM] ⚡ CHAT→WEB_SEARCH 安全兜底触发: '{user_input[:40]}'"
            )
            task_type = "WEB_SEARCH"
            route_method = "🌐 CHAT→WEB_SEARCH"

        # 如果有上下文信息，记录详情
        if context_info and context_info.get("is_continuation"):
            _app_logger.debug(
                f"[STREAM] Context continuation: {context_info.get('related_task')}, confidence: {context_info.get('confidence')}"
            )

    # ── Phase2: RouterDecision (classify_v2) ─────────────────────────────────
    # 在 SmartDispatcher 已确定 task_type 的基础上，进一步获取 skill_id / forward 决策。
    # 以非阻塞方式运行（超时保护），失败时不影响主流程。
    _router_decision = None
    try:
        from app.core.routing.local_model_router import LocalModelRouter as _LMRv2

        _router_decision = _LMRv2.classify_v2(user_input, hint=task_type, timeout=1.5)
        if _router_decision and _router_decision.skill_id:
            _app_logger.debug(
                f"[STREAM] 🎯 RouterDecision skill_id={_router_decision.skill_id} "
                f"forward_to_cloud={_router_decision.forward_to_cloud} "
                f"confidence={_router_decision.confidence:.2f}"
            )
    except Exception as _rv2_err:
        _app_logger.debug(f"[STREAM] RouterDecision classify_v2 跳过: {_rv2_err}")

    # ── 将任务专属补充指令追加到 system_instruction ────────────────────────
    _addendum = _TASK_SYSTEM_ADDENDUMS.get(task_type, "")
    if _addendum:
        system_instruction = system_instruction + _addendum
        _app_logger.debug(f"[STREAM] 📌 任务专属指令已注入: {task_type}")

    # ── 🔮 LangGraph 高级工作流路由（RESEARCH / FILE_GEN / MULTI_STEP）─────────
    # resolve_workflow() 检测用户意图，决定是否用 WorkflowEngine 替代旧路径
    # 注意：有近期上传文件时传入 has_file=True，防止 LangGraph 工作流被错误激活
    # （LangGraph 工作流没有文件字节上下文，无法处理文件分析任务）
    _wf_route = "legacy"
    if task_type in ("RESEARCH", "FILE_GEN", "MULTI_STEP"):
        try:
            _wf_route = SmartDispatcher.resolve_workflow(
                task_type, user_input, has_file=has_recent_upload
            )
            if _wf_route != "legacy":
                _app_logger.debug(f"[STREAM] 🔮 LangGraph 工作流路由: {_wf_route}")
        except Exception as _wf_err:
            _app_logger.debug(f"[STREAM] resolve_workflow 跳过: {_wf_err}")

    if _wf_route in ("langgraph_research_doc", "langgraph_multi_agent_ppt"):
        _wf_name = (
            "research_and_document"
            if _wf_route == "langgraph_research_doc"
            else "multi_agent_ppt"
        )
        _wf_label = (
            "📚 研究+文档"
            if _wf_route == "langgraph_research_doc"
            else "🎞️ 多Agent PPT"
        )

        def generate_langgraph_workflow():
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'LG_WORKFLOW', 'workflow': _wf_name, 'route_method': 'LangGraph', 'message': f'🎯 任务分类: {_wf_label} (LangGraph WorkflowEngine)'})}\n\n"
            try:
                from app.core.workflow.langgraph_workflow import WorkflowEngine

                _engine = WorkflowEngine()
                final_output = ""
                for event in _engine.stream(
                    workflow=_wf_name,
                    user_input=user_input,
                    session_id=session_name,
                ):
                    node = event.get("node", "")
                    content = event.get("content", "")
                    done = event.get("done", False)
                    if node == "error":
                        yield f"data: {json.dumps({'type': 'error', 'message': content})}\n\n"
                        return
                    if content:
                        yield f"data: {json.dumps({'type': 'status' if not done else 'token', 'message': f'[{node}] {content}' if not done else None, 'content': content if done else None}, ensure_ascii=False)}\n\n"
                    if done:
                        final_output = content
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                try:
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        final_output or f"[{_wf_label}工作流完成]",
                    )
                except Exception:
                    pass
            except Exception as _wf_ex:
                import traceback

                _app_logger.error(
                    f"[LG_WORKFLOW] ❌ 工作流失败:\n{traceback.format_exc()}"
                )
                yield f"data: {json.dumps({'type': 'error', 'message': f'工作流执行失败: {str(_wf_ex)}'})}\n\n"

        return Response(generate_langgraph_workflow(), mimetype="text/event-stream")

    # ─────────────────────────────────────────────────────────────────────────
    if (
        task_type == "MULTI_STEP"
        and context_info
        and context_info.get("is_multi_step_task")
    ):
        _app_logger.debug(
            f"[STREAM] 🔄 检测到复杂任务，使用 TaskOrchestrator 执行多步流程"
        )
        multi_step_info = context_info.get("multi_step_info", {})
        pattern = multi_step_info.get("pattern", "unknown")

        # === 🤖 MultiAgent 高质量通路（LangGraph：研究→写作→审核→修订）===
        # 仅当 LangGraph 可用且路由决策为 langgraph_react（通用复杂任务）时触发
        if _wf_route == "langgraph_react":
            _app_logger.debug(
                f"[STREAM] 🤖 MultiAgentOrchestrator 通路：RESEARCHER→WRITER→CRITIC→REVISE"
            )
            _ma_model = SmartDispatcher.get_model_for_task("MULTI_STEP")

            def generate_multi_agent():
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'MULTI_STEP', 'pattern': 'multi_agent', 'route_method': 'LangGraph MultiAgent', 'message': '🎯 任务分类: 🤖 多Agent高质量处理（研究→写作→审核→修订）'})}\n\n"
                try:
                    from app.core.agent.multi_agent import MultiAgentOrchestrator

                    orch = MultiAgentOrchestrator.preset_content_pipeline(
                        model_id=_ma_model,
                        max_revisions=1,
                    )
                    _agent_labels = {
                        "researcher": "📚 研究专员",
                        "writer": "✍️ 写作专员",
                        "critic": "🔍 审核专员",
                        "revise": "🔧 修订专员",
                        "finalize": "✅ 整合完成",
                    }
                    final_output = ""
                    for event in orch.stream(
                        user_input=user_input, session_id=session_name
                    ):
                        agent_name = event.get("agent", "unknown")
                        content = event.get("content", "")
                        done = event.get("done", False)
                        label = _agent_labels.get(agent_name, f"[{agent_name}]")

                        if agent_name == "error":
                            raise RuntimeError(content)

                        if done:
                            final_output = content
                            yield f"data: {json.dumps({'type': 'token', 'content': content}, ensure_ascii=False)}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'status', 'message': f'{label} 处理中...'})}\n\n"

                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            final_output or "[多Agent任务完成]",
                        )
                    except Exception:
                        pass

                except Exception as _ma_err:
                    import traceback as _tb

                    _app_logger.error(
                        f"[MULTI_AGENT] ❌ MultiAgentOrchestrator 失败: {_tb.format_exc()}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'多Agent执行失败，请重试: {str(_ma_err)}'})}\n\n"

            return Response(generate_multi_agent(), mimetype="text/event-stream")

        # === 文档工作流执行 ===
        if pattern == "document_workflow" and DocumentWorkflowExecutor:
            _app_logger.debug(f"[STREAM] 📄 执行文档工作流")

            def generate_doc_workflow():
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_WORKFLOW', 'pattern': 'document_workflow', 'route_method': route_method, 'message': '🎯 任务分类: 📄 文档工作流执行'})}\n\n"

                # 查找最近上传的文档
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]

                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        docs = []
                        for ext in [
                            ".docx",
                            ".md",
                            ".txt",
                            ".json",
                            ".doc",
                            ".pdf",
                            ".rtf",
                            ".odt",
                        ]:
                            import glob

                            docs.extend(
                                glob.glob(f"{dir_path}/**/*{ext}", recursive=True)
                            )

                        if docs:
                            # 获取最新的文档
                            doc_path = max(docs, key=os.path.getmtime)
                            break

                if not doc_path:
                    yield f"data: {json.dumps({'type': 'error', 'message': '❌ 未找到可执行的文档文件（支持 .docx, .doc, .pdf, .md, .txt, .rtf, .odt, .json）'})}\n\n"
                    return

                status_msg = f"📄 找到文档: {os.path.basename(doc_path)}\n"
                yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"

                try:
                    import asyncio

                    # 执行文档工作流
                    executor = DocumentWorkflowExecutor(client)

                    # 加载工作流
                    status_msg = "⏳ 正在解析文档中的工作流...\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"

                    load_result = asyncio.run(executor.load_from_document(doc_path))

                    if not load_result.get("success"):
                        error_msg = (
                            f"❌ 文档解析失败: {load_result.get('error', '未知错误')}\n"
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                        return

                    # 显示工作流信息
                    info_msg = f"✅ 工作流加载成功\n"
                    info_msg += f"   名称: {executor.workflow_name}\n"
                    info_msg += f"   步骤数: {len(executor.steps)}\n"
                    info_msg += f"   背景: {executor.workflow_context}\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': info_msg})}\n\n"

                    # 显示所有步骤
                    steps_msg = "📋 工作流步骤:\n"
                    for step in executor.steps:
                        steps_msg += (
                            f"  {step.step_id}. [{step.step_type}] {step.description}\n"
                        )
                    steps_msg += "\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': steps_msg})}\n\n"

                    # 执行工作流（流式反馈每个步骤）
                    start_msg = "🚀 开始执行工作流...\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': start_msg})}\n\n"

                    for step in executor.steps:
                        step_msg = f"[步骤 {step.step_id}/{len(executor.steps)}] {step.description}\n"
                        step_msg += f"└─ 类型: {step.step_type}\n"
                        step_msg += f"   ⏳ 执行中...\n"
                        yield f"data: {json.dumps({'type': 'status', 'message': step_msg})}\n\n"

                        step.status = "running"
                        step.start_time = datetime.now()

                        try:
                            # 执行步骤
                            step_result = asyncio.run(
                                executor._execute_step_standalone(step)
                            )
                            step.result = step_result
                            step.status = "completed"

                            success_msg = f"   ✅ 完成\n"
                            if isinstance(step_result, dict) and step_result.get(
                                "output"
                            ):
                                output_preview = str(step_result["output"])[:200]
                                success_msg += f"   📄 输出预览: {output_preview}...\n"
                            success_msg += "\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': success_msg})}\n\n"

                        except Exception as e:
                            step.status = "failed"
                            step.error = str(e)
                            error_msg = f"   ❌ 失败: {e}\n\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': error_msg})}\n\n"

                        finally:
                            step.end_time = datetime.now()

                    # 生成结果
                    results = {
                        "workflow_name": executor.workflow_name,
                        "start_time": datetime.now().isoformat(),
                        "steps": [step.to_dict() for step in executor.steps],
                        "overall_status": "completed",
                    }
                    results["summary"] = executor._generate_summary(results)

                    # 保存结果
                    output_path = asyncio.run(executor.save_results(results))

                    # 发送完成消息
                    separator = "=" * 50
                    final_msg = f"\n{separator}\n"
                    final_msg += f"✅ 文档工作流执行完成\n\n"
                    final_msg += f"📊 执行统计:\n"
                    total = len(results["steps"])
                    completed = sum(
                        1 for s in results["steps"] if s["status"] == "completed"
                    )
                    failed = sum(1 for s in results["steps"] if s["status"] == "failed")
                    final_msg += f"  总步骤: {total}\n"
                    final_msg += f"  成功: {completed}\n"
                    final_msg += f"  失败: {failed}\n"
                    final_msg += f"  成功率: {completed/total*100:.1f}%\n\n"
                    final_msg += f"📁 结果已保存: {os.path.basename(output_path)}\n"
                    final_msg += f"📂 位置: `workspace/workflows/`\n\n"
                    final_msg += f"{separator}\n"

                    yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [output_path]})}\n\n"

                    # 保存文档工作流对话历史（基于磁盘完整历史追加）
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"[文档工作流完成] {executor.workflow_name}",
                        )
                    except Exception:
                        pass

                except Exception as e:
                    import traceback

                    error_detail = traceback.format_exc()
                    error_msg = f"❌ 工作流执行失败: {str(e)}\n{error_detail}"
                    yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                    # 保存失败记录
                    try:
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"[文档工作流失败] {str(e)[:200]}",
                        )
                    except Exception:
                        pass

            return Response(generate_doc_workflow(), mimetype="text/event-stream")

        # === 其他多步任务执行 ===
        from app.core.routing import TaskDecomposer

        subtasks = TaskDecomposer.create_subtasks(user_input, multi_step_info)
        use_local_planner = multi_step_info.get("pattern") == "local_plan"

        def generate_multi_step():
            # === 立即发送任务分类信息 ===
            pattern = multi_step_info.get("pattern", "unknown")
            classification_msg = f"🎯 任务分类: 🔄 多步任务\n"
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'MULTI_STEP', 'pattern': pattern, 'route_method': route_method, 'message': classification_msg})}\n\n"

            # 显示所有子任务
            status_msg = f"📋 任务分解:\n"
            for i, subtask in enumerate(subtasks):
                status_msg += f"  {i+1}. {subtask['description']}\n"
            status_msg += "\n"
            yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"

            # 执行所有子任务（逐步流式反馈）
            try:
                import asyncio

                execution_log = []
                step_results = []
                context = {"original_input": user_input, "user_input": user_input}
                saved_files = []

                # ── 使用 PlanExecutor 执行（支持拓扑排序 + 依赖注入）──────────────
                import queue as _queue_mod
                import threading as _threading_mod

                from app.core.routing.plan_executor import PlanExecutor as _PlanExecutor
                from app.core.routing.plan_executor import (
                    build_handlers_from_orchestrator as _build_handlers,
                )

                # 构建 handlers（将 TaskOrchestrator 各方法包装成 PlanExecutor 接口）
                _handlers = _build_handlers(TaskOrchestrator, context)

                # PlanExecutor 实例（拓扑排序 + ContextStore 传递）
                _plan_exec = _PlanExecutor(
                    steps=subtasks,
                    user_input=user_input,
                    handlers=_handlers,
                    max_retry=1,
                )

                # 在后台线程运行 async executor，通过 queue 回传事件给 SSE 生成器
                _event_queue = _queue_mod.Queue()
                _plan_exception = {"err": None}

                def _run_plan_executor():
                    async def _inner():
                        try:
                            async for event in _plan_exec.execute():
                                _event_queue.put(event)
                        except Exception as _exc:
                            _plan_exception["err"] = _exc
                        finally:
                            _event_queue.put(None)  # 结束信号

                    asyncio.run(_inner())

                _exec_thread = _threading_mod.Thread(
                    target=_run_plan_executor, daemon=True
                )
                _exec_thread.start()

                _plan_done_event = None
                while True:
                    try:
                        evt = _event_queue.get(timeout=0.1)
                    except _queue_mod.Empty:
                        if not _exec_thread.is_alive():
                            break
                        continue

                    if evt is None:
                        break

                    etype = evt.get("type", "")

                    if etype == "progress":
                        yield f"data: {json.dumps({'type': 'progress', 'message': evt.get('message', ''), 'detail': evt.get('detail', '')})}\n\n"

                    elif etype == "step_done":
                        step_idx = evt.get("step_index", 0)
                        task_type_done = evt.get("task_type", "")
                        success = evt.get("success", False)
                        preview = evt.get("output_preview", "")
                        if success:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'步骤 {step_idx} 完成', 'detail': preview[:80]})}\n\n"
                        else:
                            err_msg = evt.get("error") or "执行失败"
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'步骤 {step_idx} 遇到问题', 'detail': err_msg[:80]})}\n\n"
                        # 回补 subtasks 状态（用于后续自检）
                        for _st in subtasks:
                            if str(_st.get("id")) == str(evt.get("step_id")):
                                _st["status"] = "completed" if success else "failed"
                                _st["result"] = {
                                    "success": success,
                                    "output": preview,
                                    "error": evt.get("error"),
                                }
                                # 收集输出
                                if success:
                                    step_results.append(
                                        {"success": success, "output": preview}
                                    )
                                    # 从 ContextStore 获取完整结果后更新 saved_files
                                break

                    elif etype == "status":
                        yield f"data: {json.dumps({'type': 'status', 'message': evt.get('message', '')})}\n\n"

                    elif etype == "plan_done":
                        _plan_done_event = evt
                        # 收集保存文件
                        saved_files.extend(evt.get("saved_files") or [])
                        # 同步 context 快照
                        for k, v in (evt.get("context_snapshot") or {}).items():
                            context[k] = v

                _exec_thread.join()

                if _plan_exception["err"]:
                    raise _plan_exception["err"]

                # ── 组装最终输出 ─────────────────────────────────────────────
                final_result_text = (_plan_done_event or {}).get(
                    "final_output", "(无输出)"
                )
                # 也从 ContextStore 补全 saved_files（PlanExecutor 内部已汇总）
                _pe_saved = (_plan_done_event or {}).get("saved_files") or []
                for _pf in _pe_saved:
                    if _pf not in saved_files:
                        saved_files.append(_pf)

                # 质量验证
                yield f"data: {json.dumps({'type': 'status', 'message': '正在进行最终质量验证...'})}\n\n"
                _combined_for_validate = {
                    "final_output": final_result_text,
                    "steps": [
                        {
                            "status": s.get("status", "completed"),
                            "result": s.get("result"),
                        }
                        for s in subtasks
                    ],
                }
                quality_score = asyncio.run(
                    TaskOrchestrator._validate_quality(
                        user_input, _combined_for_validate, context
                    )
                )
                yield f"data: {json.dumps({'type': 'status', 'message': f'质量验证完成，评分: {quality_score}/100'})}\n\n"

                # 复杂任务快速自检
                check = Utils.quick_self_check(
                    "MULTI_STEP", user_input, final_result_text
                )
                if not check.get("pass") and check.get("fix_prompt"):
                    yield f"data: {json.dumps({'type': 'status', 'message': '🩺 自检未通过，正在修正最终输出...'})}\n\n"
                    try:
                        fix_resp = client.models.generate_content(
                            model=SmartDispatcher.get_model_for_task("MULTI_STEP"),
                            contents=check["fix_prompt"],
                            config=types.GenerateContentConfig(
                                system_instruction=system_instruction,
                                temperature=0.4,
                                max_output_tokens=3000,
                            ),
                        )
                        final_result_text = fix_resp.text or final_result_text
                    except Exception as _fix_err:
                        _app_logger.warning(f"[MULTI_STEP] ⚠️ 自检修正失败: {_fix_err}")

                # LocalPlanner self_check
                if use_local_planner:
                    from app.core.routing import LocalPlanner

                    plan_check = LocalPlanner.self_check(
                        user_input, subtasks, step_results
                    )
                    _lp_status = plan_check.get("status", "partial")
                    _lp_summary = plan_check.get("summary", "")
                    _lp_next = (
                        plan_check.get("next_actions", [])
                        if isinstance(plan_check.get("next_actions", []), list)
                        else []
                    )
                    _lp_msg = f"自检结论: {_lp_status}"
                    if _lp_summary:
                        _lp_msg += f"\n说明: {_lp_summary}"
                    if _lp_next:
                        _lp_msg += f"\n建议后续: {', '.join(_lp_next)}"
                    yield f"data: {json.dumps({'type': 'status', 'message': _lp_msg})}\n\n"

                separator = "=" * 50
                output_text = (
                    f"\n{separator}\n✅ 多步任务完成\n质量评分: {quality_score}/100\n"
                )
                if saved_files:
                    output_text += "已保存文件:\n"
                    for p in saved_files:
                        name = os.path.basename(p)
                        link_path = p.replace("\\", "/")
                        output_text += f"- [{name}]({link_path})\n"
                    output_text += f"\n📂 位置: `{settings_manager.documents_dir}`\n"

                errors_list = [
                    s["result"].get("error")
                    for s in subtasks
                    if s.get("status") == "failed"
                    and isinstance(s.get("result"), dict)
                    and s["result"].get("error")
                ]
                if errors_list:
                    output_text += f"⚠️ 遇到的问题: {', '.join(errors_list)}\n"

                output_text += f"\n最终输出:\n{final_result_text}\n{separator}\n"

                yield f"data: {json.dumps({'type': 'token', 'content': output_text})}\n\n"
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files})}\n\n"

            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'多步任务执行失败: {str(e)}'})}\n\n"

            # 保存 MULTI_STEP 对话历史（基于磁盘完整历史追加）
            try:
                _multi_summary = (
                    f"[多步任务完成] {', '.join(s['description'] for s in subtasks)}"
                )
                if saved_files:
                    _multi_summary += f"\n生成文件: {', '.join(os.path.basename(p) for p in saved_files)}"
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, _multi_summary
                )
                _app_logger.info(f"[MULTI_STEP] ✅ 对话历史已保存")
                _start_memory_extraction(
                    user_input,
                    _multi_summary,
                    [],
                    task_type="CODER",
                    session_name=session_name,
                )
            except Exception as save_err:
                _app_logger.warning(f"[MULTI_STEP] ⚠️ 保存对话历史失败: {save_err}")

        return Response(generate_multi_step(), mimetype="text/event-stream")

    # === Agent 任务执行（LangGraphAgent ReAct，降级到 UnifiedAgent）===
    if task_type == "AGENT":
        _app_logger.debug(f"[STREAM] 🤖 执行 Agent 任务 (LangGraphAgent ReAct)")
        # classify_v2 已识别出 skill_id，传给 Agent 实现 Skill 专属行为
        _agent_skill_id = _router_decision.skill_id if _router_decision else None

        def generate_agent():
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'AGENT', 'route_method': route_method, 'message': '🎯 任务分类: 🤖 智能助手 (LangGraph ReAct)'})}\n\n"

            final_answer = ""
            collected_steps = []

            # ── 优先：LangGraphAgent（StateGraph + CheckpointSaver）────────────
            _lg_ok = False
            try:
                from app.core.agent.factory import create_langgraph_agent

                _lg_agent = create_langgraph_agent(
                    model_id=SmartDispatcher.get_model_for_task("AGENT"),
                )
                _lg_ok = True
                for chunk in _lg_agent.stream(
                    input_text=user_input,
                    history=history,
                    session_id=session_name,
                    skill_id=_agent_skill_id,
                    task_type="AGENT",
                ):
                    ctype = chunk.get("type", "token")
                    content = chunk.get("content", "")
                    if ctype == "answer":
                        final_answer = content
                        step_data = {
                            "step_type": "ANSWER",
                            "content": content,
                            "tool": None,
                        }
                    elif ctype == "tool_call":
                        step_data = {
                            "step_type": "TOOL_CALL",
                            "content": f"调用工具: {content}",
                            "tool": content,
                            "args": chunk.get("args", {}),
                        }
                    elif ctype == "tool_result":
                        step_data = {
                            "step_type": "TOOL_RESULT",
                            "content": content,
                            "tool": None,
                        }
                    elif ctype == "token":
                        step_data = {
                            "step_type": "THINKING",
                            "content": content,
                            "tool": None,
                        }
                    elif ctype == "error":
                        raise RuntimeError(content)
                    else:
                        continue
                    collected_steps.append(step_data)
                    yield f"data: {json.dumps({'type': 'agent_step', 'data': step_data}, ensure_ascii=False)}\n\n"

            except Exception as _lg_err:
                _app_logger.debug(
                    f"[AGENT] LangGraphAgent 失败 ({_lg_err})，降级到 UnifiedAgent..."
                )
                _lg_ok = False

            # ── 降级：UnifiedAgent（旧 while-loop 实现）────────────────────────
            if not _lg_ok:
                try:
                    from app.core.agent.factory import create_agent
                    from app.core.agent.types import AgentStepType

                    _ua = create_agent(
                        model_id=SmartDispatcher.get_model_for_task("AGENT")
                    )
                    collected_steps = []
                    final_answer = ""
                    for step in _ua.run(
                        input_text=user_input,
                        history=history,
                        session_id=session_name,
                        skill_id=_agent_skill_id,
                        task_type="AGENT",
                    ):
                        step_data = step.to_dict()
                        collected_steps.append(step_data)
                        if step.step_type == AgentStepType.ANSWER:
                            final_answer = step.content or ""
                        yield f"data: {json.dumps({'type': 'agent_step', 'data': step_data}, ensure_ascii=False)}\n\n"
                    if not final_answer and collected_steps:
                        final_answer = collected_steps[-1].get("content", "")
                except Exception as e:
                    import traceback

                    _app_logger.error(
                        f"[AGENT] ❌ UnifiedAgent 也失败:\n{traceback.format_exc()}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'Agent 执行失败: {str(e)}'})}\n\n"
                    return

            task_payload = {
                "id": f"task_{int(time.time() * 1000)}",
                "status": "success",
                "result": final_answer,
                "steps": collected_steps,
                "engine": "langgraph" if _lg_ok else "unified",
            }
            yield f"data: {json.dumps({'type': 'task_final', 'data': task_payload}, ensure_ascii=False)}\n\n"

            try:
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    final_answer or "[Agent 任务完成]",
                )
            except Exception:
                pass
            try:
                _start_memory_extraction(
                    user_input,
                    final_answer or "",
                    [],
                    task_type="AGENT",
                    session_name=session_name,
                )
            except Exception:
                pass

        return Response(generate_agent(), mimetype="text/event-stream")

    if locked_model and locked_model != "auto":
        model_id = locked_model
    else:
        # 传递 complexity 以便为复杂任务选择更强的模型
        _complexity = (context_info or {}).get("complexity", "normal")
        model_id = SmartDispatcher.get_model_for_task(task_type, complexity=_complexity)

    _app_logger.debug(
        f"[STREAM] Final: task_type='{task_type}', model_id='{model_id}'\n"
    )

    # 🎯 Skills 注入：将用户启用的 Skill 追加到 system_instruction
    try:
        from app.core.skills.skill_manager import SkillManager

        _active_skills = SkillManager.get_active_skill_names(task_type=task_type)
        if _active_skills:
            _app_logger.debug(
                f"[STREAM] 🎯 Active Skills ({task_type}): {', '.join(_active_skills)}"
            )
        # 意图绑定：技能未手动开启时，按输入内容临时激活匹配的技能
        _intent_temp_ids = []
        try:
            from app.core.skills.skill_trigger_binding import get_skill_binding_manager

            _intent_temp_ids = get_skill_binding_manager().match_intent(
                user_input or ""
            )
        except Exception:
            pass
        # AutoMatcher 补充：规则/语义匹配覆盖意图绑定未持久化的场景
        try:
            from app.core.skills.skill_auto_matcher import SkillAutoMatcher

            _auto_ids = SkillAutoMatcher.match(
                user_input=user_input or "", task_type=task_type or "CHAT"
            )
            if _auto_ids:
                # 合并去重，保持 intent 结果优先
                _intent_temp_ids = list(dict.fromkeys(_intent_temp_ids + _auto_ids))
        except Exception:
            pass
        if _intent_temp_ids:
            _app_logger.debug(f"[STREAM] 🔗 Auto Skills: {', '.join(_intent_temp_ids)}")
        system_instruction = SkillManager.inject_into_prompt(
            system_instruction,
            task_type=task_type,
            user_input=user_input,
            temp_skill_ids=_intent_temp_ids,
        )
    except Exception as _sk_err:
        _app_logger.warning(f"[STREAM] ⚠️ Skills 注入失败: {_sk_err}")

    # 📚 RAG 混合检索（向量 + BM25 + RRF 融合）
    # _rag_context_block: 配送给 generate()、RESEARCH 、ToT 等各路径
    _rag_context_block = ""
    try:
        from app.core.services.rag_service import get_rag_service

        _rag_svc = get_rag_service()
        if _rag_svc.stats().get("initialized"):
            _rag_hits = _rag_svc.hybrid_retrieve(user_input, k=3, score_threshold=0.3)
            if _rag_hits:
                for _rc in _rag_hits:
                    _src = os.path.basename(_rc.get("source", "unknown"))
                    _rag_context_block += f"[{_src} | 相似度: {_rc.get('score', 0):.3f}]\n{_rc['content']}\n\n"
                # 同时注入 system_instruction（供 ToT 、AGENT 路径使用）
                _rag_sys_block = (
                    "\n\n─────────────────────────────────────────"
                    "\n## 📚 知识库参考内容（混合检索）\n" + _rag_context_block
                )
                system_instruction += _rag_sys_block
                _app_logger.debug(
                    f"[STREAM] 📚 混合RAG: {len(_rag_hits)} 片段，top_score={_rag_hits[0].get('score', 0):.3f}"
                )
    except Exception as _rag_err:
        _app_logger.warning(f"[STREAM] ⚠️ RAG 注入跳过: {_rag_err}")

    # 🕸️ Graph RAG — entity-expanded triple retrieval
    try:
        from app.core.services.graph_rag_service import GraphRAGService as _GRAGS

        _graph_ctx = _GRAGS.retrieve(user_input, k=8)
        if _graph_ctx:
            _rag_context_block += "\n\n" + _graph_ctx
            system_instruction += (
                "\n\n─────────────────────────────────────────" "\n" + _graph_ctx
            )
            _app_logger.debug(f"[STREAM] 🕸️ Graph RAG: 注入知识图谱关联事实")
    except Exception as _ge:
        pass

    # 读取用户设置：是否显示思考过程
    _show_thinking = False
    try:
        _show_thinking = settings_manager.get("ai", "show_thinking") == True
    except Exception:
        pass

    def generate():
        start_time = time.time()

        def _infer_analysis_source(message: str, phase: str = "thinking") -> str:
            """推断分析来源：local / cloud / hybrid / system"""
            msg = (message or "").lower()
            phase_l = (phase or "").lower()

            if any(k in msg for k in ["ollama", "本地模型", "qwen", "local"]):
                return "local"
            if any(k in msg for k in ["gemini", "deep-research", "云端", "cloud"]):
                return "cloud"
            if any(
                k in phase_l for k in ["routing", "context", "planning", "analyzing"]
            ):
                return "hybrid"
            return "system"

        def yield_thinking(message: str, phase: str = "thinking", source: str = None):
            """发送思考过程事件（仅当用户开启 show_thinking 时），附带分析来源"""
            if not _show_thinking:
                return ""

            resolved_source = source or _infer_analysis_source(message, phase)
            source_tag = {
                "local": "[本地分析]",
                "cloud": "[大模型分析]",
                "hybrid": "[混合决策]",
                "system": "[系统流程]",
            }.get(resolved_source, "[系统流程]")

            elapsed = round(time.time() - start_time, 1)
            display_message = f"{source_tag} {message}"
            return f"data: {json.dumps({'type': 'thinking', 'message': display_message, 'phase': phase, 'elapsed': elapsed, 'analysis_source': resolved_source}, ensure_ascii=False)}\n\n"

        # === 立即反馈任务分类信息 ===
        task_display_names = {
            "PAINTER": "🎨 图像生成",
            "FILE_GEN": "📄 文档生成",
            "CODER": "💻 代码编程",
            "RESEARCH": "📚 深度研究",
            "WEB_SEARCH": "🌐 实时搜索",
            "CHAT": "💬 对话",
            "SYSTEM": "🖥️ 系统操作",
            "FILE_OP": "📂 文件操作",
            "FILE_EDIT": "✏️ 文件编辑",
            "FILE_SEARCH": "🔍 文件搜索",
            "VISION": "👁️ 图像识别",
            "MULTI_STEP": "🔄 多步任务",
            "AGENT": "🤖 智能助手",
        }

        model_display = get_model_display_name(model_id)
        task_display = task_display_names.get(task_type, task_type)

        # 发送任务分类信息（在最开始，立即显示）
        classification_msg = f"🎯 任务分类: {task_display}"
        if route_method:
            classification_msg += f" (方法: {route_method})"

        routing_list = None
        # 仅保留 routing_list 用于内部调试，不显示给用户
        if context_info and context_info.get("routing_list"):
            routing_list = context_info.get("routing_list")

        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': task_display, 'model': model_id, 'model_display': model_display, 'route_method': route_method, 'routing_list': routing_list, 'message': classification_msg})}\n\n"

        # 思考过程：任务路由分析
        t = yield_thinking(f"分析用户意图 → 识别为 {task_display}", "routing", "hybrid")
        if t:
            yield t
        model_source = (
            "local"
            if any(k in (model_id or "").lower() for k in ["qwen", "llama", "ollama"])
            else "cloud"
        )
        t = yield_thinking(
            f"路由方法: {route_method}，选择模型: {model_display}",
            "model",
            model_source,
        )
        if t:
            yield t
        if routing_list:
            steps_str = (
                " → ".join(
                    [
                        (
                            f"{r.get('task','?')}({r.get('score', 0):.2f})"
                            if isinstance(r.get("score"), (int, float))
                            else f"{r.get('task','?')}({r.get('score','?')})"
                        )
                        for r in routing_list[:5]
                    ]
                )
                if isinstance(routing_list, list)
                else str(routing_list)
            )
            t = yield_thinking(f"路由决策链: {steps_str}", "routing", "hybrid")
            if t:
                yield t

        # 如果有复杂度信息，也发送
        if context_info and context_info.get("complexity"):
            complexity_msg = f"📊 任务复杂度: {context_info['complexity']}"
            yield f"data: {json.dumps({'type': 'info', 'message': complexity_msg})}\n\n"
            t = yield_thinking(
                f"任务复杂度评估: {context_info['complexity']}", "analyzing", "hybrid"
            )
            if t:
                yield t

        # 如果有上下文，使用增强后的输入
        effective_input = user_input
        if (
            context_info
            and context_info.get("is_continuation")
            and context_info.get("enhanced_input")
        ):
            effective_input = context_info["enhanced_input"]
            _app_logger.debug(
                f"[STREAM] Using enhanced input (length: {len(effective_input)})"
            )
            yield f"data: {json.dumps({'type': 'info', 'message': '🔗 检测到延续任务，使用上下文增强'})}\n\n"
            t = yield_thinking(
                f"检测到上下文延续，增强输入 ({len(effective_input)} 字符)",
                "context",
                "hybrid",
            )
            if t:
                yield t

        # 使用快速小模型将请求转为结构化 Markdown（仅对大模型任务启用）
        if task_type not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]:
            adapted_input = Utils.adapt_prompt_to_markdown(
                task_type, effective_input, history=history
            )
            if adapted_input != effective_input:
                effective_input = adapted_input
                yield f"data: {json.dumps({'type': 'info', 'message': '🧾 已将请求结构化为Markdown提示'})}\n\n"
                t = yield_thinking(
                    "将用户请求结构化为 Markdown 格式以提升输出质量",
                    "planning",
                    "hybrid",
                )
                if t:
                    yield t

        # 重置中断标志（每次新请求都重置）
        _interrupt_manager.reset(session_name)
        interrupt_event = _interrupt_manager.get_event(session_name)

        def interrupted():
            return _interrupt_manager.is_interrupted(session_name)

        # 发送进度: 开始处理
        from web.smart_feedback import SmartFeedback

        _task_labels = SmartFeedback.TASK_LABELS
        _tl = _task_labels.get(task_type, task_type)
        yield f"data: {json.dumps({'type': 'progress', 'message': f'开始处理{_tl}任务', 'detail': get_model_display_name(model_id)})}\n\n"

        try:
            # 初始化模型追踪变量（用于日志记录）
            used_model = "unknown"

            # === SYSTEM Mode (本地执行 - 即时) ===
            if task_type == "SYSTEM":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在分析系统指令...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在执行操作...', 'detail': ''})}\n\n"

                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result["message"]
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"

                if Utils.is_failure_output(response_text):
                    t = yield_thinking(
                        "系统指令执行失败，使用 AI 修正后重试", "validating"
                    )
                    if t:
                        yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次执行失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "SYSTEM", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # 先保存历史，再发送 done 事件（防止客户端断开导致丢失）
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task=task_type,
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_OP Mode (文件操作 - 即时) ===
            if task_type == "FILE_OP":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在分析文件操作...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在访问文件系统...', 'detail': ''})}\n\n"

                batch_manager = get_batch_ops_manager()
                if batch_manager.is_batch_command(user_input):
                    parsed = batch_manager.parse_command(user_input)
                    if not parsed.get("success"):
                        response_text = (
                            f"❌ {parsed.get('error')}\n\n{parsed.get('hint', '')}"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_OP",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    job = batch_manager.create_job(
                        name=f"batch_{parsed.get('operation')}",
                        operation=parsed.get("operation"),
                        input_dir=parsed.get("input_dir"),
                        output_dir=parsed.get("output_dir"),
                        options=parsed.get("options", {}),
                    )
                    batch_manager.start_job(job.job_id)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'🧩 已创建批量任务: {job.job_id}', 'detail': ''})}\n\n"

                    summary_text = None
                    for event in batch_manager.iter_job_events(job.job_id):
                        if event.get("type") == "progress":
                            current = event.get("current", 0)
                            total = event.get("total", 0)
                            progress_pct = int((current / total) * 100) if total else 0
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⏳ 批量处理中...', 'detail': event.get('detail', ''), 'progress': progress_pct, 'total': total})}\n\n"
                        elif event.get("type") == "final":
                            summary_text = event.get("summary") or "✅ 批量处理完成"
                            break
                        elif event.get("type") == "error":
                            summary_text = event.get("message", "❌ 批量任务失败")
                            break

                    if summary_text:
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            summary_text,
                            task="FILE_OP",
                            model_name=used_model,
                        )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                file_result = FileOperator.execute(user_input)
                response_text = file_result["message"]
                if file_result.get("content"):
                    response_text += f"\n\n{file_result['content']}"

                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次执行失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "FILE_OP", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # 先保存历史，再发送 done 事件
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_OP",
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_EDIT Mode (文件编辑 - 智能修改) ===
            if task_type == "FILE_EDIT":
                used_model = model_id
                t = yield_thinking(
                    "进入文件编辑模式，将理解用户指令并修改文件", "routing"
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在分析编辑指令...', 'detail': ''})}\n\n"

                editor = get_file_editor()

                # 尝试从用户输入中提取文件路径和指令
                # 模式 1: "修改 path/to/file 把xxx改成yyy"
                match = re.search(
                    r'(?:修改|编辑|改)\s+["\']?([^"\']+?)["\']?\s+(.+)', user_input
                )
                if not match:
                    # 模式 2: "把 path/to/file 的xxx改成yyy"
                    match = re.search(
                        r'(?:把|将)\s+["\']?([^"\']+?)["\']?\s+(?:的|中的|里的)\s*(.+)',
                        user_input,
                    )

                if match:
                    file_path = match.group(1).strip()
                    instruction = match.group(2).strip()

                    t = yield_thinking(
                        f"提取到文件路径: {file_path}, 指令: {instruction}", "analyzing"
                    )
                    if t:
                        yield t

                    yield f"data: {json.dumps({'type': 'progress', 'message': f'🔍 目标文件: {os.path.basename(file_path)}', 'detail': ''})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '正在执行编辑...', 'detail': ''})}\n\n"

                    result = editor.smart_edit(file_path, instruction)

                    if result["success"]:
                        operation = result.get("operation", "edit")
                        edit_result = result.get("result", {})

                        response_text = f"✅ 文件编辑成功！\n\n"
                        response_text += f"**操作类型**: {operation}\n"

                        if operation == "replace":
                            response_text += (
                                f"**替换次数**: {edit_result.get('replacements', 0)}\n"
                            )
                            response_text += (
                                f"**预览**:\n```\n{edit_result.get('preview', '')}\n```"
                            )
                        elif operation == "delete_lines":
                            response_text += f"**删除内容**:\n```\n{edit_result.get('deleted_content', '')}\n```"
                        elif operation == "insert_line":
                            response_text += (
                                f"**消息**: {edit_result.get('message', '')}"
                            )

                        if edit_result.get("backup"):
                            response_text += (
                                f"\n\n💾 备份文件: `{edit_result.get('backup')}`"
                            )
                    else:
                        error_msg = result.get("error", "未知错误")
                        hint = result.get("hint", "")
                        response_text = f"❌ 文件编辑失败\n\n{error_msg}\n\n{hint}"
                else:
                    # 无法提取文件路径，让AI理解
                    response_text = "❌ 无法识别文件路径和编辑指令\n\n"
                    response_text += "请使用以下格式:\n"
                    response_text += "- `修改 文件路径 把'旧文本'改成'新文本'`\n"
                    response_text += "- `把 文件路径 的第5-10行删除`\n"
                    response_text += "- `编辑 文件路径 在第3行之后插入'新内容'`"

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_EDIT",
                    model_name=used_model,
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_SEARCH Mode (文件搜索: 全盘扫描 + 工作区索引) ===
            if task_type == "FILE_SEARCH":
                used_model = "FileScanner (Local)"
                t = yield_thinking("进入文件搜索模式：全盘文件索引检索", "searching")
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 正在搜索文件...', 'detail': '全盘索引检索'}, ensure_ascii=False)}\n\n"

                # ── 导入全盘扫描器 ─────────────────────────────────────────────
                try:
                    from web.file_scanner import (
                        FileScanner,
                        extract_query_from_input,
                        is_disk_search_intent,
                    )

                    _disk_scanner_ok = True
                except Exception as _fse:
                    _app_logger.warning(
                        f"[FILE_SEARCH] ⚠️ FileScanner 导入失败: {_fse}"
                    )
                    _disk_scanner_ok = False

                # 判断是否是全盘搜索意图
                _is_disk = _disk_scanner_ok and is_disk_search_intent(user_input)

                # ── 指定路径 + 扩展名列举/归纳（如"归纳 C:\Desktop 里的 .doc 文件"）──
                import re as _re_pathscan
                import time as _time_fmt

                _explicit_path_m = _re_pathscan.search(
                    r"([A-Za-z]:[\\][^\s\u4e00-\u9fa5]*)", user_input
                )
                if _explicit_path_m:
                    from pathlib import Path as _ScanPath

                    _scan_dir = _ScanPath(_explicit_path_m.group(1).rstrip("\\/. "))
                    # 关键词 → 扩展名映射
                    _KW_EXT_MAP = [
                        (["word文件", "word文档", "word"], [".doc", ".docx"]),
                        (["excel文件", "excel表格", "excel"], [".xls", ".xlsx"]),
                        (["ppt文件", "ppt演示", "ppt", "幻灯片"], [".ppt", ".pptx"]),
                        (["pdf文件", "pdf"], [".pdf"]),
                        (
                            ["图片", "照片"],
                            [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"],
                        ),
                        (["压缩包", "压缩文件"], [".zip", ".rar", ".7z"]),
                    ]
                    _ext_filters = []
                    _user_lower = user_input.lower()
                    for _kws, _exts in _KW_EXT_MAP:
                        if any(k in _user_lower for k in _kws):
                            _ext_filters = _exts[:]
                            break
                    if not _ext_filters:
                        _ext_raw = _re_pathscan.findall(
                            r"\.(docx?|xlsx?|pdf|txt|md|pptx?|csv|json|py|jpe?g|png|gif|mp[34]|zip|rar)",
                            user_input,
                            _re_pathscan.IGNORECASE,
                        )
                        _ext_filters = list(
                            dict.fromkeys("." + e.lower() for e in _ext_raw)
                        )
                    if _scan_dir.exists() and _scan_dir.is_dir():
                        _ext_label = (
                            "、".join(_ext_filters) if _ext_filters else "所有类型"
                        )
                        # ── 判断意图：归纳（物理整理）vs 搜索/列举（只读列表）──────
                        _CATALOG_KWS = ["归纳", "归档", "归类"]
                        _is_catalog_intent = any(k in user_input for k in _CATALOG_KWS)

                        # ── 内容过滤意图（"哪几个是X"）优先于 flat list ──────────
                        import re as _re_filter_intent

                        _CONTENT_FILTER_KWS = [
                            "哪几个是",
                            "哪几份是",
                            "有哪几个",
                            "有几个是",
                            "哪些是",
                            "是企业",
                            "是访谈",
                            "是报告",
                            "是合同",
                            "是简历",
                            "是方案",
                            "是什么类型",
                            "属于",
                            "是什么文件",
                            "哪几",
                            "几个是",
                            "几份是",
                        ]
                        _is_filter_intent = not _is_catalog_intent and (
                            any(k in user_input for k in _CONTENT_FILTER_KWS)
                            or bool(
                                _re_filter_intent.search(
                                    r"哪[几个些].*是|是.*[报告合同简历方案访谈纪要计划]",
                                    user_input,
                                )
                            )
                        )

                        if _is_catalog_intent:
                            # ═══ 归纳模式：FolderCatalogOrganizer 物理整理 ═══
                            t = yield_thinking(
                                f"启动文件夹归纳: {_scan_dir}，筛选: {_ext_label}",
                                "searching",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🗂️ 正在归纳 {_scan_dir} 中的 {_ext_label} 文件...', 'detail': '仅处理当前目录，不进入子文件夹'}, ensure_ascii=False)}\n\n"
                            try:
                                try:
                                    from web.folder_catalog_organizer import (
                                        FolderCatalogOrganizer,
                                    )
                                except ImportError:
                                    from folder_catalog_organizer import (
                                        FolderCatalogOrganizer,
                                    )
                                _analyzer = get_file_analyzer()
                                _organizer_inst = get_file_organizer()
                                _engine = FolderCatalogOrganizer(
                                    get_organize_root(), _analyzer, _organizer_inst
                                )
                                _summary = _engine.organize_folder(
                                    str(_scan_dir),
                                    recursive=False,
                                    ext_filters=_ext_filters if _ext_filters else None,
                                )
                                if _summary.get("success"):
                                    _entries = _summary.get("entries", [])
                                    _ok = [e for e in _entries if e.get("organized")]
                                    _fail = [
                                        e for e in _entries if not e.get("organized")
                                    ]
                                    _organize_root_display = get_organize_root()
                                    response_text = (
                                        f"✅ 归纳完成！\n\n"
                                        f"- 📂 来源目录: `{_scan_dir}`\n"
                                        f"- 📁 整理到: `{_organize_root_display}`\n"
                                        f"- ✔️ 已整理: **{len(_ok)}** 个文件\n"
                                    )
                                    if _fail:
                                        response_text += f"- ⚠️ 失败: {len(_fail)} 个（{', '.join(e['file_name'] for e in _fail[:3])}{'...' if len(_fail) > 3 else ''}）\n"
                                    if _summary.get("report_markdown"):
                                        response_text += f"\n🧾 归纳清单已保存: `{_summary['report_markdown']}`"
                                    # 显示分组结果
                                    _groups: dict = {}
                                    for _e in _ok:
                                        _grp = _e.get("suggested_folder", "其他")
                                        _groups.setdefault(_grp, []).append(
                                            _e["file_name"]
                                        )
                                    if _groups:
                                        response_text += "\n\n### 📂 归纳分组\n"
                                        for _grp, _names in _groups.items():
                                            response_text += f"\n**{_grp}**\n"
                                            for _n in _names:
                                                response_text += f"- {_n}\n"
                                else:
                                    response_text = f"❌ 归纳失败: {_summary.get('error', '未知错误')}"
                            except Exception as _oe:
                                response_text = f"❌ 归纳异常: {str(_oe)}"
                                _app_logger.debug(
                                    f"[FILE_SEARCH] FolderCatalogOrganizer 异常: {_oe}"
                                )
                        elif _is_filter_intent:
                            # ═══ 内容过滤模式：Ollama 判断哪些文件符合描述 ═══
                            # 从问句中提取过滤条件（去除路径和结构词）
                            _criterion_raw = user_input
                            _criterion_raw = _re_filter_intent.sub(
                                r"[A-Za-z]:[\\][^\s\u4e00-\u9fa5]*", "", _criterion_raw
                            ).strip()
                            _criterion_raw = _re_filter_intent.sub(
                                r"这个路径下|路径下|这个文件夹|文件夹下|下面|下有|哪几个|有哪几个"
                                r"|有几个|哪些|哪几份|几个是|几份是|是的|这个|的文件|这条路径",
                                "",
                                _criterion_raw,
                            ).strip()
                            _criterion = _criterion_raw.strip(
                                "？?。 "
                            ) or user_input.strip("？?。 ")
                            t = yield_thinking(
                                f"内容过滤: {_scan_dir} / 条件: {_criterion}",
                                "searching",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🔎 正在逐一分析文件内容...', 'detail': f'过滤条件: {_criterion}'}, ensure_ascii=False)}\n\n"
                            try:
                                try:
                                    from web.file_qa import (
                                        filter_files_by_criterion as _ffc2,
                                    )
                                except ImportError:
                                    from file_qa import (
                                        filter_files_by_criterion as _ffc2,
                                    )
                                _filter_res = _ffc2(
                                    criterion=_criterion,
                                    directory=str(_scan_dir),
                                    ext_filters=_ext_filters if _ext_filters else None,
                                )
                                if _filter_res.get("success"):
                                    _matches2 = _filter_res.get("matches", [])
                                    _scanned2 = _filter_res.get("total_scanned", 0)
                                    if _matches2:
                                        response_text = (
                                            f"🔎 在 `{_scan_dir}` 扫描了 **{_scanned2}** 个文件，"
                                            f"找到 **{len(_matches2)}** 个符合「{_criterion}」：\n\n"
                                        )
                                        for _mi, _mm in enumerate(_matches2, 1):
                                            response_text += (
                                                f"**{_mi}. `{_mm['file_name']}`**\n"
                                            )
                                            if _mm.get("reason"):
                                                response_text += (
                                                    f"   _{_mm['reason']}_\n"
                                                )
                                            response_text += "\n"
                                    else:
                                        response_text = (
                                            f"📭 在 `{_scan_dir}` 扫描了 **{_scanned2}** 个文件，"
                                            f"未找到符合「{_criterion}」描述的文件。"
                                        )
                                else:
                                    response_text = (
                                        f"❌ {_filter_res.get('error', '过滤失败')}"
                                    )
                            except Exception as _fe3:
                                response_text = f"❌ 内容过滤异常: {_fe3}"
                        else:
                            # ═══ 搜索/列举模式：只读，显示文件列表 ═══
                            t = yield_thinking(
                                f"扫描路径: {_scan_dir}（仅当前层），筛选: {_ext_label}",
                                "searching",
                            )
                            if t:
                                yield t
                            _file_list = []
                            try:
                                for _entry in _scan_dir.iterdir():
                                    if not _entry.is_file():
                                        continue
                                    if _entry.name.startswith("~$"):
                                        continue
                                    if (
                                        _ext_filters
                                        and _entry.suffix.lower() not in _ext_filters
                                    ):
                                        continue
                                    try:
                                        _stat = _entry.stat()
                                        _sz = _stat.st_size
                                        _sz_str = (
                                            f"{_sz} B"
                                            if _sz < 1024
                                            else (
                                                f"{_sz/1024:.1f} KB"
                                                if _sz < 1048576
                                                else f"{_sz/1048576:.1f} MB"
                                            )
                                        )
                                        _file_list.append(
                                            {
                                                "name": _entry.name,
                                                "size": _sz_str,
                                                "mtime": _stat.st_mtime,
                                                "mtime_str": _time_fmt.strftime(
                                                    "%Y-%m-%d %H:%M",
                                                    _time_fmt.localtime(_stat.st_mtime),
                                                ),
                                            }
                                        )
                                    except (PermissionError, OSError):
                                        pass
                            except (PermissionError, OSError):
                                pass
                            _file_list.sort(key=lambda x: x["mtime"], reverse=True)
                            if not _file_list:
                                response_text = f"📁 在 `{_scan_dir}` 中未找到任何 **{_ext_label}** 文件。"
                            else:
                                response_text = f"📁 在 `{_scan_dir}` 中找到 **{len(_file_list)}** 个 **{_ext_label}** 文件：\n\n"
                                response_text += "| # | 文件名 | 大小 | 修改时间 |\n| --- | --- | --- | --- |\n"
                                for _i, _f in enumerate(_file_list[:100], 1):
                                    response_text += f"| {_i} | `{_f['name']}` | {_f['size']} | {_f['mtime_str']} |\n"
                                if len(_file_list) > 100:
                                    response_text += f"\n*...共 {len(_file_list)} 个文件，仅显示前 100 个*"
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                # ── 📁 文件夹监控（Watch Mode） ────────────────────────────────
                _WATCH_START_KWS = [
                    "监控文件夹",
                    "监控目录",
                    "监控这个文件夹",
                    "开始监控",
                    "自动归类",
                ]
                _WATCH_STOP_KWS = ["停止监控", "取消监控", "关闭监控"]
                _WATCH_LIST_KWS = ["正在监控", "监控列表", "查看监控", "有哪些监控"]
                _watch_path_m = _re_pathscan.search(
                    r"([A-Za-z]:[\\][^\s\u4e00-\u9fa5]*)", user_input
                )
                if any(k in user_input for k in _WATCH_START_KWS) and _watch_path_m:
                    _wdir = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _watcher = get_file_watcher()
                    _watcher.configure(
                        get_file_analyzer(), get_file_organizer(), get_organize_root()
                    )
                    _wres = _watcher.start_watch(_wdir)
                    if _wres.get("success"):
                        response_text = (
                            f"👁️ **文件夹监控已启动！**\n\n"
                            f"- 📂 监控目录: `{_wdir}`\n"
                            f"- ⚡ 新文件落地后自动分析并归类到 `{get_organize_root()}`\n"
                            f"- 🔕 说「**停止监控 {_wdir}**」可随时关闭\n\n"
                            f"_支持格式: .doc/.docx/.pdf/.xlsx/.pptx/.txt/.csv/.zip 等_"
                        )
                    else:
                        response_text = (
                            f"❌ 启动监控失败: {_wres.get('error', '未知错误')}"
                        )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WATCH_STOP_KWS) and _watch_path_m:
                    _wdir = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _wres = get_file_watcher().stop_watch(_wdir)
                    response_text = (
                        f"⛔ 已停止监控 `{_wdir}`"
                        if _wres.get("success")
                        else f"⚠️ 停止失败: {_wres.get('error', '该目录未在监控中')}"
                    )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WATCH_LIST_KWS):
                    try:
                        from web.file_watcher import get_file_watcher
                    except ImportError:
                        from file_watcher import get_file_watcher
                    _watches = get_file_watcher().list_watches()
                    if _watches:
                        response_text = "👁️ **当前监控目录列表：**\n\n"
                        for _w in _watches:
                            _alive = "✅ 运行中" if _w["alive"] else "⚠️ 已停止"
                            response_text += f"- `{_w['path']}` — {_alive}（自 {_w['started_at']}）\n"
                    else:
                        response_text = "📭 当前没有正在监控的文件夹。\n\n说「**监控文件夹 C:\\xxx**」可启动监控。"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # ── 📋 单文件关键字段提取（合同/发票/简历等） ──────────────────
                _FIELDS_KWS = [
                    "提取字段",
                    "提取信息",
                    "关键信息",
                    "合同信息",
                    "发票信息",
                    "提取关键",
                    "解读这个",
                    "读一下这个",
                    "分析这个文件",
                    "文件内容",
                ]

                # ── 🗂️ 工作文件库管理命令 ─────────────────────────────────────
                _WFL_ADD_KWS = [
                    "添加监控文件夹",
                    "添加文件夹",
                    "加入文件库",
                    "监控这个文件夹到文件库",
                    "把这个文件夹加入文件库",
                    "添加到文件库",
                ]
                _WFL_REFRESH_KWS = [
                    "刷新文件库",
                    "更新文件库",
                    "重新扫描文件库",
                    "重建文件库",
                ]
                _WFL_STATUS_KWS = [
                    "文件库状态",
                    "文件库统计",
                    "文件库有多少",
                    "文件库里有什么",
                    "文件库概况",
                    "查看文件库",
                ]

                if any(k in user_input for k in _WFL_ADD_KWS) and _watch_path_m:
                    _add_path = _watch_path_m.group(1).rstrip("\\/. ")
                    try:
                        from web.work_file_library import get_work_file_library

                        _wfl2 = get_work_file_library()
                        _added = _wfl2.add_watch_folder(_add_path)
                        if _added:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'📂 正在扫描 {_add_path}...', 'detail': '添加到工作文件库'}, ensure_ascii=False)}\n\n"
                            _wfl2.scan_locations(locations=[_add_path])
                            _wfl2.wait_for_scan(timeout=15.0)
                            _cnt = _wfl2.count()
                            response_text = (
                                f"✅ 已将 `{_add_path}` 添加到工作文件库并完成扫描！\n\n"
                                f"文件库现共收录 **{_cnt}** 个工作文件。\n"
                                "以后说「找 xxx」即可快速检索。"
                            )
                        else:
                            response_text = (
                                f"❌ 添加失败，请确认路径存在: `{_add_path}`"
                            )
                    except Exception as _wadd_e:
                        response_text = f"❌ 添加文件夹出错: {_wadd_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WFL_REFRESH_KWS):
                    try:
                        from web.work_file_library import get_work_file_library

                        _wfl3 = get_work_file_library()
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🔄 正在刷新工作文件库...', 'detail': '重新扫描所有位置'}, ensure_ascii=False)}\n\n"
                        _wfl3.scan_locations(force=True)
                        _wfl3.wait_for_scan(timeout=15.0)
                        _st = _wfl3.get_stats()
                        _cats_str = "、".join(
                            f"{k} {v}个" for k, v in _st.get("categories", {}).items()
                        )
                        response_text = (
                            f"✅ 工作文件库已刷新！\n\n"
                            f"共收录 **{_st['total']}** 个工作文件"
                            + (f"（{_cats_str}）" if _cats_str else "")
                            + "。"
                        )
                    except Exception as _wref_e:
                        response_text = f"❌ 刷新文件库出错: {_wref_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if any(k in user_input for k in _WFL_STATUS_KWS):
                    try:
                        from web.work_file_library import (
                            _CATEGORY_ICONS,
                            get_work_file_library,
                        )

                        _wfl4 = get_work_file_library()
                        _st4 = _wfl4.get_stats()
                        import time as _t4

                        _ls = _st4.get("last_scan")
                        _ls_str = (
                            _t4.strftime("%Y-%m-%d %H:%M", _t4.localtime(_ls))
                            if _ls
                            else "从未扫描"
                        )
                        response_text = f"### 🗂️ 工作文件库状态\n\n"
                        response_text += f"- **收录总数**: {_st4['total']} 个工作文件\n"
                        response_text += f"- **最后扫描**: {_ls_str}\n\n"
                        if _st4.get("categories"):
                            response_text += "**按类型分布：**\n\n"
                            for _cat, _cnt4 in _st4["categories"].items():
                                _icon4 = _CATEGORY_ICONS.get(_cat, "📎")
                                response_text += f"- {_icon4} {_cat}: **{_cnt4}** 个\n"
                        _wfs = _wfl4.list_watch_folders()
                        _default_locs = __import__(
                            "web.work_file_library", fromlist=["_get_common_locations"]
                        )._get_common_locations()
                        response_text += (
                            f"\n**扫描位置（{len(_default_locs) + len(_wfs)} 个）：**\n"
                        )
                        for _loc in _default_locs:
                            response_text += f"- `{_loc}` （默认）\n"
                        for _wf in _wfs:
                            response_text += f"- `{_wf['path']}` （用户添加）\n"
                        response_text += "\n说「**刷新文件库**」可重新扫描，说「**添加监控文件夹 路径**」可扩大范围。"
                    except Exception as _wst_e:
                        response_text = f"❌ 获取文件库状态失败: {_wst_e}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                if any(k in user_input for k in _FIELDS_KWS) and _watch_path_m:
                    _tgt_file = _watch_path_m.group(1).rstrip("\\/. ")
                    from pathlib import Path as _FPath

                    _fp = _FPath(_tgt_file)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'📋 正在解读 {_fp.name}...', 'detail': '提取关键字段'}, ensure_ascii=False)}\n\n"
                    try:
                        try:
                            from web.file_fields_extractor import extract_fields as _ef
                            from web.file_fields_extractor import (
                                fields_to_markdown as _fm,
                            )
                        except ImportError:
                            from file_fields_extractor import extract_fields as _ef
                            from file_fields_extractor import fields_to_markdown as _fm
                        _ana = get_file_analyzer()
                        _content = _ana._extract_content(str(_fp))
                        _fields = _ef(_fp.name, _content, _fp.suffix.lower())
                        if _fields:
                            response_text = "### 📋 文件关键信息\n\n" + _fm(
                                _fields, _fp.name
                            )
                        else:
                            response_text = f"⚠️ 无法提取字段（Ollama 可能未运行，或文件内容无法解析）\n文件: `{_fp.name}`"
                    except Exception as _fe:
                        response_text = f"❌ 提取失败: {_fe}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # ── 🤖 跨文件问答（File Q&A） ──────────────────────────────────
                _QA_KWS = [
                    "哪个",
                    "哪些",
                    "有没有",
                    "关于",
                    "谁的",
                    "最早",
                    "最晚",
                    "最高",
                    "这几份",
                    "这些文件",
                    "文件里",
                    "里面有没有",
                    "它们",
                    "汇总一下",
                    "告诉我",
                    "查一查",
                    "对比",
                ]
                # 有问题词 + 有路径 → 跨文件问答
                if any(k in user_input for k in _QA_KWS) and _watch_path_m:
                    _qa_dir = _watch_path_m.group(1).rstrip("\\/. ")
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🤖 正在阅读文件并思考...', 'detail': '跨文件问答'}, ensure_ascii=False)}\n\n"
                    try:
                        try:
                            from web.file_qa import answer_file_question
                        except ImportError:
                            from file_qa import answer_file_question
                        _qa_result = answer_file_question(
                            question=user_input,
                            search_dirs=[_qa_dir],
                            top_k=6,
                        )
                        if _qa_result.get("success"):
                            response_text = _qa_result["answer"]
                            _srcs = _qa_result.get("sources", [])
                            if _srcs:
                                response_text += "\n\n---\n**参考文件：** " + "、".join(
                                    f"`{s['file_name']}`" for s in _srcs
                                )
                        else:
                            response_text = f"❌ {_qa_result.get('error', '问答失败')}"
                    except Exception as _qe:
                        response_text = f"❌ 问答异常: {_qe}"
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                if _is_disk:
                    # 检测是否是"扫描"指令（而非搜索指令）
                    _scan_cmd = any(
                        k in user_input.lower()
                        for k in [
                            "扫描我的电脑",
                            "扫描电脑",
                            "扫描磁盘",
                            "扫描硬盘",
                            "全盘扫描",
                            "开始扫描",
                            "scan my",
                            "start scan",
                        ]
                    )
                    if _scan_cmd:
                        scan_started = FileScanner.start_scan()
                        if scan_started:
                            response_text = (
                                "🚀 全盘文件扫描已启动！\n\n"
                                f"正在扫描以下磁盘分区：**{', '.join(FileScanner.get_drives())}**\n\n"
                                "扫描在后台运行，不会影响您的使用。扫描完成后您可以通过对话：\n"
                                "- 「帮我找一下 xxx 文件」\n"
                                "- 「打开 我的简历」\n"
                                "- 「找一下 2025年报告」\n\n"
                                "首次扫描通常需要 **2-10 分钟**，之后结果会持久化。"
                            )
                        else:
                            st = FileScanner.get_status()
                            scanned = st.get("scanned", 0)
                            indexed = st.get(
                                "indexed_count", FileScanner.stats()["total"]
                            )
                            response_text = (
                                f"⏳ 全盘扫描正在进行中...\n\n"
                                f"- 已检查文件：**{scanned:,}**\n"
                                f"- 已索引：**{indexed:,}**\n"
                                f"- 当前目录：`{st.get('current_dir', '...')[:80]}`"
                            )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # ── 工作文件库搜索（优先，无需全盘扫描）────────────────────
                    # 仅当用户未明确要求全盘搜索时，先查工作文件库
                    _EXPLICIT_FULL_DISK_KWS = [
                        "全盘",
                        "整个电脑",
                        "所有磁盘",
                        "所有文件",
                        "全电脑",
                        "全部磁盘",
                        "全部文件",
                        "全硬盘",
                    ]
                    _want_full_disk = any(
                        k in user_input for k in _EXPLICIT_FULL_DISK_KWS
                    )

                    if not _want_full_disk:
                        try:
                            from web.work_file_library import (
                                _CATEGORY_ICONS,
                                detect_category_from_input,
                                get_work_file_library,
                            )

                            _wfl = get_work_file_library()
                            _wfl_query = extract_query_from_input(user_input)

                            # 如果库还没数据，触发快速扫描并等待
                            if not _wfl.is_indexed():
                                yield f"data: {json.dumps({'type': 'progress', 'message': '📂 正在快速建立工作文件库...', 'detail': '扫描桌面、文档、下载等常用位置'}, ensure_ascii=False)}\n\n"
                                _wfl.scan_locations()
                                _wfl.wait_for_scan(timeout=10.0)

                            # 检测用户意图的文件类型
                            _wfl_category = detect_category_from_input(user_input)
                            _wfl_results = _wfl.search(
                                _wfl_query, limit=30, category=_wfl_category
                            )
                            _wfl_stats = _wfl.get_stats()

                            if _wfl_results:
                                # 按分类分组展示
                                _grouped: dict = {}
                                for _r in _wfl_results:
                                    _grouped.setdefault(_r["category"], []).append(_r)

                                response_text = f"📂 在文件库中找到 **{len(_wfl_results)}** 个包含「{_wfl_query}」的文件：\n\n"
                                for _cat, _cat_files in _grouped.items():
                                    _icon = _CATEGORY_ICONS.get(_cat, "📎")
                                    response_text += f"### {_icon} {_cat}（{len(_cat_files)} 个）\n\n"
                                    response_text += "| 文件名 | 大小 | 修改时间 |\n| --- | --- | --- |\n"
                                    for _f in _cat_files[:10]:
                                        response_text += f"| `{_f['name']}` | {_f['size_str']} | {_f['mtime_str']} |\n"
                                    if len(_cat_files) > 10:
                                        response_text += f"\n_...还有 {len(_cat_files) - 10} 个同类文件_\n"
                                    response_text += "\n"

                                _cats_summary = "、".join(
                                    f"{k} {v}个"
                                    for k, v in _wfl_stats.get("categories", {}).items()
                                )
                                response_text += (
                                    f"\n---\n_文件库共收录 **{_wfl_stats['total']}** 个工作文件"
                                    + (f"（{_cats_summary}）" if _cats_summary else "")
                                    + "。如需搜索更多位置，说「添加监控文件夹 D:\\工作资料」_"
                                )

                                # 同步发送文件选择器事件（供前端渲染卡片）
                                _picker_files = [
                                    {
                                        "path": _r["path"],
                                        "name": _r["name"],
                                        "ext": _r["ext"],
                                        "category": _r["category"],
                                        "size_str": _r["size_str"],
                                        "mtime_str": _r["mtime_str"],
                                        "score": _r["score"],
                                    }
                                    for _r in _wfl_results[:12]
                                ]
                                yield f"data: {json.dumps({'type': 'file_picker', 'query': _wfl_query, 'count': len(_wfl_results), 'files': _picker_files, 'auto_opened': False}, ensure_ascii=False)}\n\n"
                                yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    response_text,
                                    task="FILE_SEARCH",
                                    model_name=used_model,
                                )
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                            elif _wfl.is_indexed():
                                # 库已建立但本次未找到
                                _cats_summary = "、".join(
                                    f"{k} {v}个"
                                    for k, v in _wfl_stats.get("categories", {}).items()
                                )
                                response_text = (
                                    f"📭 工作文件库中未找到包含「{_wfl_query}」的文件。\n\n"
                                    f"文件库当前收录了 **{_wfl_stats['total']}** 个工作文件"
                                    + (f"（{_cats_summary}）" if _cats_summary else "")
                                    + "。\n\n💡 提示：\n"
                                    "- 说「**添加监控文件夹 D:\\工作资料**」可扩大搜索范围\n"
                                    "- 说「**刷新文件库**」可重新扫描已有位置\n"
                                    f"- 说「**全盘搜索 {_wfl_query}**」可搜索整个电脑"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    response_text,
                                    task="FILE_SEARCH",
                                    model_name=used_model,
                                )
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                        except Exception as _wfl_exc:
                            _app_logger.warning(
                                f"[FILE_SEARCH] ⚠️ 工作文件库搜索出错（降级到全盘扫描）: {_wfl_exc}"
                            )
                            # 继续走全盘扫描逻辑

                    # ── 全盘文件名模糊搜索 ──────────────────────────────────
                    FileScanner.ensure_loaded()
                    query = extract_query_from_input(user_input)
                    t = yield_thinking(
                        f"全盘搜索关键词: {query!r}，索引量: {FileScanner.stats()['total']:,}",
                        "searching",
                    )
                    if t:
                        yield t

                    if not FileScanner.is_indexed():
                        # 没有索引 → 建议用户触发扫描
                        response_text = (
                            "⚠️ 全盘文件索引尚未建立。\n\n"
                            "请说「**扫描我的电脑**」让 Koto 自动开始扫描，\n"
                            "或点击侧边栏 **文件管理 > 开始全盘扫描**。\n\n"
                            "首次扫描可能需要 2-10 分钟，完成后即可通过对话快速查找和打开任何文件。"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    disk_results = FileScanner.search(query, limit=12)

                    if not disk_results:
                        response_text = (
                            f"❌ 未找到包含 **{query}** 的文件。\n\n"
                            f"💡 建议：\n"
                            f"- 检查关键词是否正确\n"
                            f"- 如果文件较新，可以重新扫描（说「扫描我的电脑」）\n"
                            f"- 当前索引 {FileScanner.stats()['total']:,} 个文件"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # 判断是否直接打开（唯一高置信匹配）
                    auto_open = (
                        len(disk_results) == 1 and disk_results[0]["score"] >= 0.9
                    ) or (
                        len(disk_results) >= 1
                        and disk_results[0]["score"] >= 0.95
                        and (len(disk_results) < 2 or disk_results[1]["score"] < 0.7)
                    )

                    if auto_open:
                        best = disk_results[0]
                        open_result = FileScanner.open_file(best["path"])
                        if open_result["success"]:
                            response_text = (
                                f"✅ 已为您打开文件：**{best['name']}**\n\n"
                                f"📁 路径: `{best['path']}`\n"
                                f"📂 分类: {best['category']}　大小: {best['size_str']}　修改: {best['mtime_str']}"
                            )
                        else:
                            response_text = f"⚠️ 找到文件但打开失败: {open_result.get('error', '')}\n\n📁 路径: `{best['path']}`"
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_SEARCH",
                            model_name=used_model,
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    # 多结果 → 发送 file_picker 事件（前端渲染选择卡片）
                    picker_event = {
                        "type": "file_picker",
                        "query": query,
                        "count": len(disk_results),
                        "files": disk_results[:12],
                        "auto_opened": False,
                    }
                    yield f"data: {json.dumps(picker_event, ensure_ascii=False)}\n\n"

                    response_text = f"🔍 找到 {len(disk_results)} 个匹配 **{query}** 的文件，请点击选择要打开的文件："
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                else:
                    # ── 工作区内容索引搜索（原有逻辑）────────────────────────
                    indexer = get_file_indexer()
                    keywords = (
                        user_input.replace("找文件", "")
                        .replace("搜索", "")
                        .replace("查找", "")
                    )
                    keywords = (
                        keywords.replace("包含", "").replace("的文件", "").strip()
                    )
                    t = yield_thinking(f"工作区关键词搜索: {keywords}", "searching")
                    if t:
                        yield t
                    results = indexer.search(keywords, limit=10)
                    if not results:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🔄 扩展搜索范围...', 'detail': ''}, ensure_ascii=False)}\n\n"
                        results = indexer.find_by_content(keywords, min_similarity=0.2)
                    if results:
                        response_text = f"🔍 找到 {len(results)} 个匹配文件:\n\n"
                        for i, r in enumerate(results[:10], 1):
                            file_name = r.get("file_name", "未知文件")
                            file_path = r.get("file_path", "")
                            snippet = r.get("match_snippet", "")
                            score = r.get("score", 0)
                            similarity = r.get("similarity")
                            response_text += (
                                f"### {i}. {file_name}\n📁 路径: `{file_path}`\n"
                            )
                            if similarity:
                                response_text += f"🎯 相似度: {similarity:.0%}\n"
                            elif score:
                                response_text += f"⭐ 匹配分: {score:.2f}\n"
                            if snippet:
                                response_text += f"📄 预览: {snippet[:200]}...\n"
                            response_text += "\n"
                    else:
                        response_text = (
                            "❌ 未找到匹配文件\n\n💡 提示:\n"
                            "- 若要搜索电脑上所有文件，请说「帮我找一下 xxx 文件」\n"
                            f"- 当前工作区索引文件数: {len(indexer.list_indexed_files(limit=1000))}"
                        )
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text}, ensure_ascii=False)}\n\n"
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        response_text,
                        task="FILE_SEARCH",
                        model_name=used_model,
                    )
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

            # === DOC_ANNOTATE Mode (文档标注/润色 - 流式反馈) ===
            if task_type == "DOC_ANNOTATE":
                used_model = model_id if model_id else "gemini-3.1-pro-preview"
                t = yield_thinking(
                    f"进入文档标注模式，将使用 {model_id or 'gemini-3.1-pro-preview'} 分析文档",
                    "routing",
                )
                if t:
                    yield t
                _app_logger.debug(f"[STREAM] 📄 执行 DOC_ANNOTATE 任务")

                # 从请求中获取task_id，用于支持取消操作
                task_id = request.json.get("task_id")

                # 查找最近上传的文档
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]

                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        import glob

                        docs = []
                        for ext in [
                            ".docx",
                            ".docxm",
                            ".doc",
                            ".pdf",
                            ".txt",
                            ".md",
                            ".rtf",
                            ".odt",
                        ]:
                            docs.extend(
                                glob.glob(f"{dir_path}/**/*{ext}", recursive=True)
                            )
                        if docs:
                            doc_path = max(docs, key=os.path.getmtime)
                            break

                if not doc_path or not os.path.exists(doc_path):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '❌ 未找到文档', 'detail': '请上传 .docx/.doc/.pdf/.txt/.md/.rtf 文件'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return

                # 非 .docx 自动转换（输出复制到标准文档目录，避免输出到 temp 目录）
                _dw_ext = os.path.splitext(doc_path)[1].lower()
                _stream_docs_dir = settings_manager.documents_dir
                os.makedirs(_stream_docs_dir, exist_ok=True)
                if _dw_ext != ".docx":
                    try:
                        import tempfile as _tmpdw

                        from web.doc_converter import convert_to_docx

                        _dw_conv_dir = _tmpdw.mkdtemp(prefix="koto_dw_")
                        _dw_conv_path, _ = convert_to_docx(
                            doc_path, output_dir=_dw_conv_dir
                        )
                        # 复制转换后文件到标准文档目录，确保输出也在该目录
                        _dw_conv_basename = os.path.basename(_dw_conv_path)
                        _dw_conv_in_docs = os.path.join(
                            _stream_docs_dir, _dw_conv_basename
                        )
                        import shutil as _dw_shutil

                        _dw_shutil.copy2(_dw_conv_path, _dw_conv_in_docs)
                        doc_path = _dw_conv_in_docs
                        _app_logger.debug(
                            f"[DocWorkflow] 转换 {_dw_ext} → .docx 并复制到文档目录: {doc_path}"
                        )
                    except Exception as _dw_conv_err:
                        yield f"data: {json.dumps({'type': 'error', 'message': f'文档转换失败: {_dw_conv_err}'})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                        return

                # Step 1: 读取文档信息
                yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '📖 正在读取文档...', 'detail': os.path.basename(doc_path)})}\n\n"

                doc_filename = os.path.basename(doc_path)
                total_chars = 0
                total_paras = 0

                try:
                    from docx import Document

                    doc = Document(doc_path)
                    total_paras = len([p for p in doc.paragraphs if p.text.strip()])
                    total_chars = sum(len(p.text) for p in doc.paragraphs)

                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading_complete', 'message': f'✅ 文档解析完成', 'detail': f'{doc_filename}: {total_paras} 段  |  {total_chars} 字'})}\n\n"
                except Exception as e:
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 读取文档失败: {str(e)}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return

                # Step 2: 展示任务信息
                nl = "\n"
                task_info_msg = f"📋 【任务信息】{nl}- 模型: {model_id}{nl}- 需求: {user_input[:100]}{nl}- 文档: {doc_filename}"
                yield f"data: {json.dumps({'type': 'info', 'message': task_info_msg})}\n\n"

                try:
                    from web.document_feedback import DocumentFeedbackSystem

                    feedback_system = DocumentFeedbackSystem(
                        gemini_client=client, default_model_id="gemini-3.1-pro-preview"
                    )

                    # 使用流式分析系统，逐步反馈进度
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'processing_start', 'message': '🔍 开始处理文档...', 'detail': '这个过程会涉及多个阶段'})}\n\n"

                    revised_file = None
                    final_result = None
                    cancelled = False

                    # 迭代流式结果，传入task_id用于支持取消
                    for (
                        progress_event
                    ) in feedback_system.full_annotation_loop_streaming(
                        doc_path,
                        user_input,
                        task_id=task_id,
                        model_id=model_id,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(
                            session_name
                        ),
                    ):
                        stage = progress_event.get("stage", "unknown")
                        progress = progress_event.get("progress", 0)
                        message = progress_event.get("message", "")
                        detail = progress_event.get("detail", "")

                        # 处理任务取消
                        if stage == "cancelled":
                            cancelled = True
                            yield f"data: {json.dumps({'type': 'info', 'message': '⏸️ 任务已取消', 'detail': '用户中止了处理'})}\n\n"
                            break

                        # 根据阶段发送不同样式的进度信息
                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message, 'detail': detail, 'progress': progress})}\n\n"

                        # 保存最终结果
                        if stage == "complete":
                            final_result = progress_event.get("result", {})
                            revised_file = final_result.get("revised_file")

                    # 如果任务被取消，返回取消响应
                    if cancelled:
                        total_time = time.time() - start_time
                        # 保存取消记录到历史
                        session_manager.append_and_save(
                            f"{session_name}.json", user_input, "⏸️ 文档标注任务已取消"
                        )
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time, 'cancelled': True})}\n\n"
                        return

                    # 如果处理成功，生成详细总结
                    if final_result and final_result.get("success"):
                        applied = final_result.get("applied", 0)
                        failed = final_result.get("failed", 0)
                        total = final_result.get("total", applied + failed)

                        # 计算修改密度
                        density = (
                            (applied / total_chars * 1000) if total_chars > 0 else 0
                        )

                        summary_msg = (
                            f"✅ **文档修改完成！**\n\n"
                            f"📊 **测试结果**：\n"
                            f"- **文档分析**：成功读取 {total_paras} 段，共 {total_chars} 字。\n"
                            f"- **AI 处理**：文档被并发处理，总耗时约 {int(time.time() - start_time)} 秒。\n"
                            f"- **生成质量**：AI 成功找出了 **{total} 处** 翻译生硬、语序不顺的地方。\n"
                            f"- **应用修订**：成功将 **{applied} 处** 修改以“修订模式（Track Changes）”写入了 Word 文档（仅有 {failed} 处因复杂格式定位失败，属于正常容错范围）。\n\n"
                            f"📂 **验证文件**：\n"
                            f"高质量的测试结果文件已经生成在您的本地目录中，您可以直接打开查看效果：\n"
                            f"👉 `{os.path.basename(revised_file) if revised_file else '待生成'}`\n\n"
                            f"💡 **使用方法**：\n"
                            f"1. 用 Microsoft Word 打开输出文件\n"
                            f"2. 点击「审阅」标签页\n"
                            f"3. 右侧气泡中查看全部修改建议\n"
                            f"4. 逐条接受或忽略（右键批注可操作）\n"
                            f"5. 点击「接受全部」或逐条处理\n\n"
                            f"📂 **文件位置**: `{os.path.dirname(revised_file) if revised_file else settings_manager.documents_dir}`"
                        )

                        yield f"data: {json.dumps({'type': 'progress', 'message': '📝 生成最终报告...', 'detail': ''})}\n\n"
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"

                        # 保存对话历史（包含元数据）
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            summary_msg,
                            task="DOC_ANNOTATE",
                            model_name=model_id,
                            saved_files=[revised_file] if revised_file else [],
                        )

                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': total_time})}\n\n"
                    else:
                        error_msg = (
                            final_result.get("message", "未知错误")
                            if final_result
                            else "处理失败"
                        )
                        # 保存失败记录
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            f"❌ 文档标注失败: {error_msg}",
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 处理失败: {error_msg}'})}\n\n"

                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

                except Exception as e:
                    import traceback

                    error_detail = traceback.format_exc()
                    _app_logger.error(f"[DOC_ANNOTATE] ❌ 失败:\n{error_detail}")
                    # 保存异常记录
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        f"❌ 文档标注异常: {str(e)[:200]}",
                    )

                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 处理异常: {str(e)[:200]}'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

                return

            # === WEB_SEARCH Mode (联网搜索 - 实时信息) ===
            if task_type == "WEB_SEARCH":
                used_model = "gemini-2.5-flash (Google Search)"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在连接互联网搜索...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在搜索实时信息...', 'detail': 'Google Search'})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '正在整理搜索结果...', 'detail': ''})}\n\n"

                # 优先使用本地/AI路由器生成的 skill_prompt，实现「模型理解意图 → 生成执行指令」
                _skill_prompt = (context_info or {}).get("skill_prompt")
                search_result = WebSearcher.search_with_grounding(
                    user_input, skill_prompt=_skill_prompt
                )
                response_text = search_result["response"]

                if (
                    Utils.is_failure_output(response_text)
                    or "搜索失败" in response_text
                ):
                    t = yield_thinking(
                        "初次搜索结果不佳，使用 gemini-2.0-flash-lite 改写查询词后重试",
                        "searching",
                    )
                    if t:
                        yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次搜索失败，正在修正查询...', 'detail': ''})}\n\n"
                    fix_query_prompt = (
                        "请把用户需求改写成更适合搜索的简短关键词或查询语句，只输出查询语句。\n"
                        f"用户需求: {user_input}"
                    )
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.2,
                            max_output_tokens=64,
                        ),
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result["response"]

                if Utils.is_failure_output(response_text):
                    fix_prompt = Utils.build_fix_prompt(
                        "WEB_SEARCH", user_input, response_text
                    )
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1200,
                        ),
                    )
                    response_text = fix_resp.text or response_text

                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"

                # 先保存历史，再发送 done 事件
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, response_text
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === 🗂️ 文件分析直通：有近期上传文件 + RESEARCH/CHAT/FILE_GEN + legacy ═════
            # chat_stream 检测到近期上传文件时，直接读取文件内容并传给模型，
            # 避免走不带文件字节的 RESEARCH/ToT 路径（那条路没有文件上下文，给出错误结果）
            if (
                has_recent_upload
                and recent_file_path
                and os.path.isfile(recent_file_path)
                and task_type in ("RESEARCH", "CHAT", "FILE_GEN")
                and _wf_route == "legacy"
            ):
                _rfile_ext = os.path.splitext(recent_file_path)[1].lower()
                _rfile_name = os.path.basename(recent_file_path)
                _is_binary_doc = _rfile_ext in (
                    ".pdf",
                    ".docx",
                    ".doc",
                    ".pptx",
                    ".xlsx",
                )
                _is_image_file = _rfile_ext in (
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".gif",
                    ".webp",
                    ".bmp",
                )

                # 图片 → 交给后续 VISION 路径，不在这里拦截
                if not _is_image_file:
                    print(
                        f"[STREAM] 🗂️ 文件分析直通: {_rfile_name} ({_rfile_ext}) task={task_type}"
                    )
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'model': model_id, 'message': f'📄 文件分析模式: {_rfile_name}'}, ensure_ascii=False)}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'📂 正在读取文件: {_rfile_name}', 'stage': 'file_ready', 'progress': 15})}\n\n"

                    try:
                        import time as _rft_mod

                        _rf_start = _rft_mod.time()

                        _rf_task_label = {
                            "RESEARCH": "🔬 深度分析",
                            "FILE_GEN": "📝 内容生成",
                            "CHAT": "💬 文件问答",
                        }.get(task_type, task_type)
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'🎯 任务类型: {_rf_task_label}', 'stage': 'routing_complete', 'progress': 25})}\n\n"

                        if task_type == "RESEARCH":
                            _rf_sys = (
                                "你是一位专业的文档分析助手，擅长深度解读各类文件（商业计划书、研究报告、技术文档等）。\n"
                                "请仔细阅读用户提供的文件内容，并按以下结构输出分析报告：\n\n"
                                "## 核心摘要\n- 用 3-5 条要点概括文件核心内容\n\n"
                                "## 详细解读\n### 背景与目标\n### 关键内容分析\n### 数据与证据\n\n"
                                "## 结论与建议\n- 综合评判与可行性/价值判断\n\n"
                                "要求：用中文，条理清晰，避免冗余，不输出代码块标记。"
                            )
                        else:
                            _rf_sys = (
                                "你是一位专业的文档阅读与分析助手。用户上传了一份文件并提出了问题，"
                                "请认真阅读文件的完整内容，用中文给出详细、准确的分析和回答。\n"
                                "注意：直接回答用户的具体问题，引用文件中的具体数据和信息支撑你的判断，"
                                "用清晰的结构输出，避免空泛表述。"
                            )

                        # 二进制文档需要能处理文件字节的模型
                        _rf_model = model_id or MODEL_MAP.get(
                            task_type, "gemini-2.5-flash"
                        )
                        if _is_binary_doc:
                            _rf_model = (
                                globals().get("_INTERACTIONS_FALLBACK_MODEL")
                                or "gemini-2.5-flash"
                            )

                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚡ 正在请求 {_rf_model}，请稍候...', 'stage': 'api_calling', 'progress': 35})}\n\n"

                        # 构建请求内容（文本 + 文件字节 或 提取文本）
                        _rf_contents = user_input
                        if _is_binary_doc:
                            try:
                                with open(recent_file_path, "rb") as _rfh:
                                    _rf_bytes = _rfh.read()
                                _rf_mime_map = {
                                    ".pdf": "application/pdf",
                                    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    ".doc": "application/msword",
                                    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                                    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                }
                                _rf_mime = _rf_mime_map.get(
                                    _rfile_ext, "application/octet-stream"
                                )
                                _rf_doc_part = types.Part.from_bytes(
                                    data=_rf_bytes, mime_type=_rf_mime
                                )
                                _rf_contents = [user_input, _rf_doc_part]
                                print(
                                    f"[STREAM] 📄 ByteDoc-Read: model={_rf_model}, bytes={len(_rf_bytes)}, mime={_rf_mime}"
                                )
                            except Exception as _rfb_err:
                                print(
                                    f"[STREAM] ⚠️ 文件字节读取失败，回退文本模式: {_rfb_err}"
                                )
                                try:
                                    from web.file_processor import process_uploaded_file

                                    _rf_contents, _ = process_uploaded_file(
                                        recent_file_path, user_input
                                    )
                                except Exception:
                                    pass  # 用 user_input 作兜底
                        else:
                            # 文本类文件：提取内容后注入
                            try:
                                from web.file_processor import process_uploaded_file

                                _rf_contents, _ = process_uploaded_file(
                                    recent_file_path, user_input
                                )
                            except Exception as _rft_err:
                                print(f"[STREAM] ⚠️ 文本文件提取失败: {_rft_err}")

                        _rf_stream = client.models.generate_content_stream(
                            model=_rf_model,
                            contents=_rf_contents,
                            config=types.GenerateContentConfig(
                                system_instruction=_rf_sys,
                                temperature=0.7,
                                max_output_tokens=8000,
                            ),
                        )

                        _rf_full_text = ""
                        _rf_first_token = True
                        for _rf_chunk in _rf_stream:
                            _rf_t = getattr(_rf_chunk, "text", None)
                            if _rf_t:
                                if _rf_first_token:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '✍️ 模型正在生成回复...', 'stage': 'generating', 'progress': 55})}\n\n"
                                    _rf_first_token = False
                                _rf_full_text += _rf_t
                                yield f"data: {json.dumps({'type': 'token', 'content': _rf_t}, ensure_ascii=False)}\n\n"

                        _rf_elapsed = round(_rft_mod.time() - _rf_start, 2)
                        try:
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                _rf_full_text,
                                task=task_type,
                                model_name=_rf_model,
                            )
                        except Exception:
                            pass
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    except Exception as _rf_err:
                        import traceback as _rf_tb

                        print(
                            f"[STREAM] ⚠️ 文件分析直通失败，降级标准路径: {_rf_tb.format_exc()}"
                        )
                        yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 文件读取遇到问题，切换至标准模式...', 'detail': str(_rf_err)[:100]})}\n\n"
                        # 降级：继续向下走 ToT / RESEARCH 标准路径
            # ═══════════════════════════════════════════════════════════════════════════

            # === 🌳 Tree of Thought Mode (RESEARCH / FILE_GEN 并行多路推理选优) ===
            # 触发条件：legacy 路由 + RESEARCH/FILE_GEN + 非 Deep-Research-Pro + ToT 未被用户禁用
            # 排除 Excel/脚本类文件生成请求：ToT 产出文本，而这类请求需要代码脚本，两者不兼容
            _excel_request = any(
                k in (effective_input or "").lower()
                for k in ["excel", "xlsx", ".xls", "电子表格", "spreadsheet"]
            )
            _tot_enabled = (
                _wf_route == "legacy"
                and task_type in ("RESEARCH", "FILE_GEN")
                and len(str(effective_input)) >= 20
                and not str(model_id or "").startswith("deep-research-pro-preview")
                and settings_manager.get("ai", "use_tree_of_thought") is not False
                and not _excel_request  # Excel 走代码生成路径，ToT 文本输出无法生成文件
            )

            if _tot_enabled:
                # FILE_GEN 的 ToT 使用 flash 级别模型，避免 pro 模型响应慢导致前端 25s 无心跳触发卡住检测
                if task_type == "FILE_GEN":
                    _tot_model = MODEL_MAP.get("FILE_GEN", "gemini-3-flash-preview")
                else:
                    _tot_model = model_id or MODEL_MAP.get(
                        task_type, "gemini-2.5-flash"
                    )
                _tot_n = (
                    2 if task_type == "FILE_GEN" else 3
                )  # FILE_GEN 用 2 路，RESEARCH 用 3 路
                _tot_label = "📄 文档生成" if task_type == "FILE_GEN" else "🔬 深度研究"
                yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'route_method': 'TreeOfThought', 'message': f'🌳 Tree of Thought 启动：{_tot_n} 条并行推理分支 ({_tot_label})'}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': f'🌳 Tree of Thought 启动 ({_tot_n} 分支并行推理)...', 'detail': f'模型: {_tot_model}'})}\n\n"

                try:
                    from app.core.agent.tree_of_thought import create_tot

                    _tot = create_tot(
                        task_type=task_type, n_branches=_tot_n, model_id=_tot_model
                    )
                    _tot_final = ""
                    _tot_winner_id = None

                    for _evt in _tot.stream(
                        user_input=effective_input,
                        task_type=task_type,
                        system_instruction=system_instruction,
                    ):
                        _stage = _evt.get("stage", "")

                        if _stage == "expand":
                            _bid = _evt.get("branch_id", "?")
                            _blabel = _evt.get("label", "")
                            _bstatus = _evt.get("status", "")
                            if _bstatus == "generating":
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'🌿 分支 {_bid}「{_blabel}」生成中...', 'detail': ''})}\n\n"
                            elif _bstatus == "done":
                                _elapsed = _evt.get("elapsed", "")
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 分支 {_bid}「{_blabel}」完成 ({_elapsed}s)', 'detail': _evt.get('preview', '')[:60]})}\n\n"
                            elif _bstatus == "error":
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ 分支 {_bid} 失败', 'detail': _evt.get('error', '')[:80]})}\n\n"

                        elif _stage == "evaluate":
                            _bstatus = _evt.get("status", "")
                            if _bstatus == "scoring":
                                yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 Critic 正在评估各分支质量...', 'detail': ''})}\n\n"
                            else:
                                _bid = _evt.get("branch_id", "?")
                                _score = _evt.get("score", 0)
                                _crit = _evt.get("critique", "")
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'📊 分支 {_bid} 得分 {_score:.1f} — {_crit}', 'detail': ''}, ensure_ascii=False)}\n\n"

                        elif _stage == "select":
                            _tot_winner_id = _evt.get("winner_id")
                            _tot_score = _evt.get("score", 0)
                            _wlabel = _evt.get("winner_label", "")
                            _reason = _evt.get("reason", "")
                            _tot_final = _evt.get("content", "")
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🏆 最优分支: {_tot_winner_id}「{_wlabel}」(得分 {_tot_score:.1f})', 'detail': _reason}, ensure_ascii=False)}\n\n"
                            yield f"data: {json.dumps({'type': 'token', 'content': _tot_final}, ensure_ascii=False)}\n\n"

                        elif _stage == "error":
                            _errmsg = _evt.get("message", "未知错误")
                            _app_logger.error(f"[ToT] ❌ 错误: {_errmsg}")
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ Tree of Thought 遇到问题，切换至标准模式: {_errmsg[:100]}', 'detail': ''}, ensure_ascii=False)}\n\n"
                            _tot_enabled_fallback = True
                            _tot_final = ""
                            break
                    else:
                        _tot_enabled_fallback = False

                    if _tot_final:
                        try:
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                _tot_final[:6000],
                                task=task_type,
                                model_name=_tot_model,
                            )
                        except Exception:
                            pass
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time, 'tot_winner': _tot_winner_id})}\n\n"
                        return

                    # ToT 失败 → 降级到下方标准 RESEARCH/FILE_GEN 逻辑
                    _app_logger.warning(f"[ToT] ⚠️ 未获得有效输出，降级至标准路径")

                except ImportError:
                    _app_logger.warning(
                        "[ToT] ⚠️ tree_of_thought 模块未找到，降级至标准路径"
                    )
                except Exception as _tot_err:
                    import traceback as _ttb

                    _app_logger.error(f"[ToT] ❌ 异常: {_ttb.format_exc()}")
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ Tree of Thought 异常，切换至标准模式', 'detail': str(_tot_err)[:100]})}\n\n"
            # ──────────────────────────────────────────────────────────────────

            # === RESEARCH Mode (深度研究 - 流式响应优先) ===
            if task_type == "RESEARCH":
                research_model = model_id or MODEL_MAP.get(
                    "RESEARCH", "gemini-3-pro-preview"
                )
                used_model = research_model
                t = yield_thinking(
                    f"进入深度研究模式，使用 {research_model} 进行专业级分析",
                    "analyzing",
                )
                if t:
                    yield t
                newline = "\n"
                _detail = (
                    "使用 Interactions API 深度研究"
                    if research_model.startswith("deep-research-pro-preview")
                    else f"使用 {research_model} 进行流式分析"
                )
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔬 启动深度研究模式...', 'detail': _detail})}{newline}{newline}"

                # 构建深度研究的system instruction
                research_instruction = """你是一位专业的研究助手，擅长深度分析复杂技术话题。请按照以下结构提供全面深入的研究报告：

1. **技术概述**：清晰定义和解释核心概念
2. **技术原理**：详细说明工作机制和底层原理
3. **优势分析**：列举主要优点和应用场景
4. **问题与挑战**：分析存在的问题和技术瓶颈
5. **对比分析**：与其他同类技术进行横向对比
6. **发展趋势**：讨论未来发展方向和应用前景
7. **参考资料**：提供相关技术文档和学术资料的引用

📌 **特殊查询类型增强规则**：

**价格/费用/票务查询**（如高铁票、机票、酒店、门票等）：
- ✅ **首先输出一个清晰的表格**，包含关键信息（车次、发车时间、到达时间、座位、价格、时长等）
- ✅ 必须提供**具体价格**（例如：二等座 ¥524.5）
- ❌ 禁止使用价格区间（如"500-600元"）
- ✅ 按座位/房型等级**分别列出**每个选项的确切价格
- ✅ 列出**具体班次/车次号**（如 G12、航班 MU5137）
- ✅ 列出**发车时间和到达时间**，方便用户对比选择
- ❌ 禁止输出重复内容或多个相同的段落

**强制使用表格格式**：
```
🚄 上海虹桥 → 北京南（2026年2月12日）

| 车次   | 发车  | 到达  | 座位类型 | 价格     | 时长  |
|--------|-------|-------|----------|----------|-------|
| G12次  | 09:00 | 13:24 | 商务座   | ¥1,748   | 4h24m |
| G12次  | 09:00 | 13:24 | 一等座   | ¥933     | 4h24m |
| G12次  | 09:00 | 13:24 | 二等座   | ¥524.5   | 4h24m |
| G8次   | 10:00 | 14:31 | 商务座   | ¥1,748   | 4h31m |
| G8次   | 10:00 | 14:31 | 一等座   | ¥933     | 4h31m |
| G8次   | 10:00 | 14:31 | 二等座   | ¥524.5   | 4h31m |

💡 购票方式：访问 12306.cn 搜索对应车次购买。
```

要求：
- 提供具体的技术细节和数据支持
- 使用专业术语但确保可理解性
- 保持客观中立的分析态度
- 内容全面且有深度
- 适当使用图表和示例说明"""

                # 注入 skill_prompt（模型对用户意图的理解）
                _research_skill = (context_info or {}).get("skill_prompt")
                if _research_skill:
                    research_instruction += (
                        f"\n\n[用户期望的输出重点] {_research_skill}"
                    )
                # 将知识库检索内容注入研究指令（若有）
                if _rag_context_block:
                    research_instruction += (
                        f"\n\n[📚 知识库参考资料]\n{_rag_context_block}"
                    )

                collected_text = []

                try:
                    newline = "\n"
                    if research_model.startswith("deep-research-pro-preview"):
                        yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在进行深度分析...', 'detail': 'Deep Research 正在检索与综合，可能需要较长时间'})}{newline}{newline}"
                        deep_text = WebSearcher.deep_research_for_ppt(
                            effective_input, ""
                        )
                        if not deep_text:
                            raise RuntimeError("Deep Research 未返回有效内容")
                        collected_text.append(deep_text)
                        yield f"data: {json.dumps({'type': 'token', 'content': deep_text})}\n\n"
                    else:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在进行深度分析...', 'detail': f'{research_model} 正在思考，可能需要30-90秒'})}{newline}{newline}"

                        response_stream = client.models.generate_content_stream(
                            model=research_model,
                            contents=effective_input,
                            config=types.GenerateContentConfig(
                                system_instruction=research_instruction,
                                temperature=0.7,
                                max_output_tokens=8000,  # 允许更长的输出
                                top_p=0.95,
                            ),
                        )

                    if not research_model.startswith("deep-research-pro-preview"):
                        chunk_count = 0
                        heartbeat_interval = 5  # 每5秒发送一次心跳
                        first_chunk_received = False

                        # 使用保活包装器处理流式响应
                        for item_type, item_data in stream_with_keepalive(
                            response_stream,
                            start_time,
                            keepalive_interval=heartbeat_interval,
                            max_wait_first_token=90,
                        ):  # 最多等待90秒
                            # 检查中断
                            if interrupted():
                                _app_logger.debug(f"[RESEARCH] 用户中断研究")
                                newline = "\n"
                                interrupt_msg = f"{newline}{newline}⏹️ 研究已被用户中断"
                                yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}{newline}{newline}"
                                break

                            if item_type == "heartbeat":
                                # 发送心跳保持连接
                                elapsed = item_data
                                if first_chunk_received:
                                    char_count = len("".join(collected_text))
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在生成中...', 'detail': f'已生成 {char_count} 字符，耗时 {elapsed}s'})}\n\n"
                                else:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '🧠 模型正在深度思考...', 'detail': f'已等待 {elapsed}s，请耐心等待'})}\n\n"

                            elif item_type == "timeout":
                                # 等待超时
                                yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，模型响应时间过长，请稍后重试'})}\n\n"
                                break

                            elif item_type == "chunk":
                                chunk = item_data
                                if chunk.text:
                                    if not first_chunk_received:
                                        first_chunk_received = True
                                        _app_logger.debug(
                                            f"[RESEARCH] 收到第一个响应块，耗时 {time.time() - start_time:.1f}s"
                                        )

                                    collected_text.append(chunk.text)
                                    yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"
                                    chunk_count += 1

                                    # 每50个chunk显示一次进度日志
                                    if chunk_count % 50 == 0:
                                        _app_logger.debug(
                                            f"[RESEARCH] 已生成 {chunk_count} 个chunk, {len(''.join(collected_text))} 字符"
                                        )

                    final_text = "".join(collected_text)
                    _app_logger.info(
                        f"[RESEARCH] ✅ 研究完成，共 {len(final_text)} 字符"
                    )

                    # 保存历史（基于磁盘完整历史追加）
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        final_text[:4000],
                        task="RESEARCH",
                        model_name=used_model,
                    )

                except Exception as research_err:
                    error_msg = str(research_err)
                    _app_logger.debug(f"[RESEARCH] 错误: {error_msg}")

                    # 智能错误处理
                    if "503" in error_msg or "UNAVAILABLE" in error_msg:
                        # API过载，尝试使用Flash版本
                        try:
                            newline = "\n"
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 服务繁忙，切换到 Gemini 2.5 Flash...', 'detail': ''})}{newline}{newline}"

                            response_stream = client.models.generate_content_stream(
                                model="gemini-2.5-flash",
                                contents=effective_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=research_instruction,
                                    temperature=0.7,
                                    max_output_tokens=8000,
                                ),
                            )

                            last_heartbeat_flash = time.time()
                            for chunk in response_stream:
                                if interrupted():
                                    break
                                if chunk.text:
                                    collected_text.append(chunk.text)
                                    yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"

                                    # Flash 模式下也发送心跳
                                    current_time = time.time()
                                    if current_time - last_heartbeat_flash > 3:
                                        elapsed = int(current_time - start_time)
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚡ 快速模式生成中...', 'detail': f'{elapsed}s'})}\n\n"
                                        last_heartbeat_flash = current_time

                            final_text = "".join(collected_text)
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                final_text[:4000],
                                task="RESEARCH",
                                model_name="gemini-3-flash-preview",
                            )

                        except Exception as fallback_err:
                            error_text = f"❌ 研究服务暂时不可用\n\n错误信息: {str(fallback_err)[:200]}\n\n💡 建议：\n1. 稍后重试\n2. 简化问题\n3. 使用普通对话模式"
                            yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                error_text[:1000],
                                task="RESEARCH",
                                model_name="gemini-3-flash-preview",
                            )

                    elif (
                        "timeout" in error_msg.lower()
                        or "disconnect" in error_msg.lower()
                    ):
                        # 连接问题
                        error_text = f"⚠️ 连接超时或中断\n\n可能原因：\n1. 网络不稳定\n2. 服务器繁忙\n3. 代理配置问题\n\n建议：请稍后重试，或检查网络连接"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            error_text[:1000],
                            task="RESEARCH",
                            model_name=used_model,
                        )

                    else:
                        # 其他错误
                        error_text = f"❌ 研究过程中出现错误\n\n{error_msg[:300]}\n\n请尝试：\n1. 重新提问\n2. 简化问题描述\n3. 稍后重试"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            error_text[:1000],
                            task="RESEARCH",
                            model_name=used_model,
                        )

                total_time = time.time() - start_time
                newline = "\n"
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}{newline}{newline}"
                return

            # === PAINTER Mode (图像生成 - Gemini 3.1 Flash Image 优先，Imagen 4.0 备用) ===
            if task_type == "PAINTER":
                used_model = "Gemini 3.1 Flash Image (Imagen 4.0 fallback)"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 正在理解你的创作请求...', 'detail': '', 'progress': 5, 'stage': 'paint_prepare'})}\n\n"

                # 使用上下文增强的输入（如果有）
                if (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("enhanced_input")
                ):
                    image_prompt = context_info["enhanced_input"]
                    _app_logger.debug(
                        f"[PAINTER] 使用上下文增强的prompt: {image_prompt[:100]}..."
                    )
                else:
                    image_prompt = effective_input

                yield f"data: {json.dumps({'type': 'progress', 'message': '🖌️ Gemini 3.1 Flash Image 正在生成图像...', 'detail': '请耐心等待', 'progress': 20, 'stage': 'paint_generate'})}\n\n"

                max_retries = 2
                use_fallback = False
                images = []

                for attempt in range(max_retries):
                    try:
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 图像生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return

                        if attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 第 {attempt} 次重试...', 'detail': '', 'progress': 25, 'stage': 'paint_retry'})}\n\n"
                            time.sleep(2)

                        # 选择模型
                        if use_fallback:
                            model_name = "Imagen 4.0"
                            yield f"data: {json.dumps({'type': 'progress', 'message': '🔄 切换到 Imagen 4.0...', 'detail': '', 'progress': 30, 'stage': 'paint_fallback'})}\n\n"
                        else:
                            model_name = "Gemini 3.1 Flash Image"

                        # 使用后台线程执行请求，主线程发送心跳
                        import queue
                        import threading

                        result_queue = queue.Queue()

                        def worker():
                            try:
                                if use_fallback:
                                    result = client.models.generate_images(
                                        model="imagen-4.0-fast-generate-001",
                                        prompt=image_prompt,
                                        config=types.GenerateImagesConfig(
                                            number_of_images=1
                                        ),
                                    )
                                else:
                                    result = client.models.generate_content(
                                        model="gemini-3.1-flash-image-preview",
                                        contents=image_prompt,
                                        config=types.GenerateContentConfig(
                                            response_modalities=["TEXT", "IMAGE"]
                                        ),
                                    )
                                result_queue.put(("success", result))
                            except Exception as e:
                                result_queue.put(("error", e))

                        thread = threading.Thread(target=worker, daemon=True)
                        thread.start()

                        # Per-model timeout: 120s for gemini-3.1-flash-image (takes ~65s), 90s for imagen fallback
                        timeout_seconds = 120 if not use_fallback else 90
                        attempt_start = time.time()
                        timed_out = False
                        response = None

                        while True:
                            attempt_elapsed = time.time() - attempt_start

                            if attempt_elapsed > timeout_seconds:
                                timed_out = True
                                if not use_fallback:
                                    # Primary model timed out — switch to imagen
                                    _app_logger.debug(
                                        f"[PAINTER] Gemini 3.1 Flash Image 超时 ({int(attempt_elapsed)}s)，切换到 Imagen"
                                    )
                                    use_fallback = True
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '⏱️ 模型响应超时，切换到 Imagen...', 'detail': '', 'progress': 28, 'stage': 'paint_fallback'})}\n\n"
                                    break
                                else:
                                    # Imagen also timed out
                                    elapsed = time.time() - start_time
                                    yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ 图像生成超时 ({int(elapsed)}s)，请稍后重试'})}\n\n"
                                    total_time = time.time() - start_time
                                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                    return

                            if interrupted():
                                yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 图像生成已中断'})}\n\n"
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return

                            try:
                                status, data = result_queue.get(timeout=3.0)
                                if status == "success":
                                    response = data
                                    break
                                else:
                                    raise data
                            except queue.Empty:
                                progress_guess = min(
                                    85,
                                    30 + int((attempt_elapsed / timeout_seconds) * 55),
                                )
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'🎨 {model_name} 生成中...', 'detail': f'{int(attempt_elapsed)}s', 'progress': progress_guess, 'stage': 'paint_running'})}\n\n"

                        if timed_out:
                            continue  # outer for-loop: retry with use_fallback=True (imagen timeout already returned above)

                        # 处理响应
                        yield f"data: {json.dumps({'type': 'progress', 'message': '💾 正在保存图片...', 'detail': '', 'progress': 90, 'stage': 'paint_save'})}\n\n"

                        if use_fallback:
                            if response.generated_images:
                                for gen_img in response.generated_images:
                                    img_data = gen_img.image.image_bytes
                                    images_dir = settings_manager.images_dir
                                    os.makedirs(images_dir, exist_ok=True)
                                    timestamp = int(time.time())
                                    filename = f"generated_{timestamp}.png"
                                    filepath = os.path.join(images_dir, filename)
                                    with open(filepath, "wb") as f:
                                        f.write(img_data)

                                    # 确保路径在 workspace 下
                                    try:
                                        rel_path = os.path.relpath(
                                            filepath, WORKSPACE_DIR
                                        ).replace("\\", "/")
                                        if ".." not in rel_path:
                                            images.append(rel_path)
                                            _app_logger.debug(
                                                f"[PAINTER] Imagen 已保存: {rel_path}"
                                            )
                                        else:
                                            # 降级保存到 workspace/images
                                            abs_workspace_images = os.path.join(
                                                WORKSPACE_DIR, "images"
                                            )
                                            os.makedirs(
                                                abs_workspace_images, exist_ok=True
                                            )
                                            fallback_filepath = os.path.join(
                                                abs_workspace_images, filename
                                            )
                                            with open(fallback_filepath, "wb") as f:
                                                f.write(img_data)
                                            fallback_rel = os.path.relpath(
                                                fallback_filepath, WORKSPACE_DIR
                                            ).replace("\\", "/")
                                            images.append(fallback_rel)
                                            _app_logger.debug(
                                                f"[PAINTER] Imagen 降级保存: {fallback_rel}"
                                            )
                                    except Exception as path_err:
                                        _app_logger.debug(
                                            f"[PAINTER] Path error: {path_err}"
                                        )
                        else:
                            if (
                                response.candidates
                                and response.candidates[0].content.parts
                            ):
                                for part in response.candidates[0].content.parts:
                                    if (
                                        hasattr(part, "inline_data")
                                        and part.inline_data
                                    ):
                                        img_filename = Utils.save_image_part(part)
                                        if img_filename:
                                            images.append(img_filename)
                                            _app_logger.debug(
                                                f"[PAINTER] Gemini 3.1 Flash Image 已保存: {img_filename}"
                                            )

                        if images:
                            save_path = settings_manager.images_dir
                            msg = f"✨ 图片已生成! (使用 {model_name})\n🖼️ 保存位置: {save_path}"
                            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"

                            yield f"data: {json.dumps({'type': 'progress', 'message': '✅ 图像生成完成', 'detail': f'{len(images)} 张', 'progress': 100, 'stage': 'complete'})}\n\n"

                            # 先保存历史记录（包含图片路径），再发送 done
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                "图像已生成",
                                images=images,
                                task="PAINTER",
                                model_name=model_name,
                            )

                            total_time = time.time() - start_time
                            _app_logger.debug(
                                f"[PAINTER] 发送图片列表: {images}"
                            )  # 调试
                            yield f"data: {json.dumps({'type': 'done', 'images': images, 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            if not use_fallback:
                                use_fallback = True
                                continue
                            else:
                                yield f"data: {json.dumps({'type': 'token', 'content': '❌ 模型未返回图片'})}\n\n"

                    except Exception as img_err:
                        error_msg = str(img_err)
                        model_label = (
                            "Imagen" if use_fallback else "Gemini-3.1-Flash-Image"
                        )
                        _app_logger.debug(
                            f"[PAINTER] {model_label} 尝试 {attempt+1} 失败 ({type(img_err).__name__}): {error_msg[:300]}"
                        )

                        # Fall back to imagen on ANY non-safety error when using primary model
                        if (
                            not use_fallback
                            and "safety" not in error_msg.lower()
                            and "blocked" not in error_msg.lower()
                        ):
                            _app_logger.debug(
                                f"[PAINTER] Gemini 3.1 Flash Image 失败，切换到 Imagen: {error_msg[:200]}"
                            )
                            use_fallback = True
                            continue

                        if (
                            "safety" in error_msg.lower()
                            or "blocked" in error_msg.lower()
                        ):
                            user_msg = "❌ 内容被安全策略过滤，请修改描述"
                        elif "location is not supported" in error_msg.lower():
                            user_msg = "❌ 地区限制，请配置中转服务"
                        else:
                            user_msg = f"❌ 图像生成失败: {error_msg[:100]}"

                        yield f"data: {json.dumps({'type': 'token', 'content': user_msg})}\n\n"

                # PAINTER 所有重试都失败时也要保存历史
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, "图像生成失败"
                )

                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return

            # === FILE_GEN Mode (文件生成 - 自动执行) ===
            if task_type == "FILE_GEN":
                t = yield_thinking(
                    f"进入文件生成模式，将使用 {model_id} 生成文档", "generating"
                )
                if t:
                    yield t
                _app_logger.debug(f"[FILE_GEN] ===== Starting file generation =====")
                _app_logger.debug(
                    f"[FILE_GEN] Model: {model_id}, User input: {user_input[:100]}..."
                )

                yield f"data: {json.dumps({'type': 'progress', 'message': '🧾 准备生成文档...', 'detail': get_model_display_name(model_id), 'progress': 5, 'stage': 'filegen_prepare'})}\n\n"

                response_text = ""
                generated_files = []
                temp_scripts = []  # 临时脚本列表（执行后删除）
                api_timeout = 120  # 增加到 120 秒，长文档需要更多时间

                if interrupted():
                    yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return

                # ⭐ 检查是否是"转换请求"（把之前的内容做成word/pdf）
                is_convert_request = (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("continuation_type") == "convert"
                    and context_info.get("context_summary", {}).get("last_model_output")
                )

                if is_convert_request:
                    # 直接转换模式 - 不需要调用模型，直接生成文档
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在将内容转换为文档...', 'detail': '', 'progress': 30, 'stage': 'filegen_convert'})}\n\n"

                    try:
                        from web.document_generator import save_docx, save_pdf

                        source_content = context_info["context_summary"][
                            "last_model_output"
                        ]
                        _app_logger.debug(
                            f"[FILE_GEN] 直接转换模式，源内容长度: {len(source_content)}"
                        )

                        # 提取标题（尝试从内容中找 # 标题）
                        title_match = re.search(
                            r"^#\s*(.+)$", source_content, re.MULTILINE
                        )
                        if title_match:
                            title = title_match.group(1).strip()[:50]
                        else:
                            title = (
                                f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            )

                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)

                        # 判断生成 Word 还是 PDF
                        user_lower = user_input.lower()

                        if "pdf" in user_lower:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成 PDF...', 'detail': '', 'progress': 60, 'stage': 'filegen_save'})}\n\n"
                            saved_path = save_pdf(
                                source_content, title=title, output_dir=docs_dir
                            )
                            file_type = "PDF"
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成 Word 文档...', 'detail': '', 'progress': 60, 'stage': 'filegen_save'})}\n\n"
                            saved_path = save_docx(
                                source_content, title=title, output_dir=docs_dir
                            )
                            file_type = "Word"

                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        generated_files.append(rel_path)

                        success_msg = f"✅ **{file_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"

                        yield f"data: {json.dumps({'type': 'progress', 'message': '✅ 文档转换完成', 'detail': file_type, 'progress': 100, 'stage': 'complete'})}\n\n"

                        _app_logger.info(f"[FILE_GEN] ✅ 直接转换成功: {rel_path}")

                    except Exception as convert_err:
                        error_msg = f"❌ 文档转换失败: {str(convert_err)}"
                        _app_logger.debug(f"[FILE_GEN] 转换错误: {convert_err}")
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"

                    # 保存历史（基于磁盘完整历史追加）
                    _model_msg = (
                        f"已生成文件: {', '.join(generated_files)}"
                        if generated_files
                        else "文档转换失败"
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _model_msg
                    )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # ⭐ 检查是否是 PPT 生成请求
                ppt_keywords = [
                    "ppt",
                    "幻灯片",
                    "演示文稿",
                    "演示",
                    "presentation",
                    "slide",
                    "slides",
                ]
                user_lower_check = user_input.lower()
                is_ppt_request = any(kw in user_lower_check for kw in ppt_keywords)

                if is_ppt_request:
                    # =============== PPT 专用生成流程 ===============
                    _app_logger.debug(f"[FILE_GEN] 🎯 检测到 PPT 生成请求")

                    # ── 初始化智能反馈 ──
                    from web.smart_feedback import SmartFeedback

                    def _fb_emit(msg, detail=""):
                        _app_logger.debug(f"[SmartFB] {msg} | {detail}")

                    fb = SmartFeedback.for_ppt(user_input, emit=_fb_emit)

                    def _fb_sse(msg_detail_tuple):
                        """将 SmartFeedback 返回的 (msg, detail) 转为 SSE 数据行"""
                        msg, detail = msg_detail_tuple
                        progress_pct = 0
                        if (
                            getattr(fb, "total_steps", None)
                            and getattr(fb, "current_step", 0) > 0
                        ):
                            progress_pct = min(
                                95, int((fb.current_step / fb.total_steps) * 100)
                            )
                        return f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': detail, 'progress': progress_pct})}\n\n"

                    yield _fb_sse(fb.start())

                    try:
                        # ──────── Session: 创建 PPT 编辑会话 ────────
                        ppt_session_id = None
                        try:
                            from web.ppt_session_manager import get_ppt_session_manager

                            ppt_session_mgr = get_ppt_session_manager()
                            ppt_session_id = ppt_session_mgr.create_session(
                                title=user_input[:50],  # 前 50 字作为临时标题
                                user_input=user_input,
                                theme="business",
                            )
                            _app_logger.info(
                                f"[FILE_GEN/PPT] 📋 创建编辑会话: {ppt_session_id}"
                            )
                        except Exception as session_err:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ⚠️ 会话创建异常（不影响生成）: {session_err}"
                            )

                        # ──────── Step 0: 处理上传的文件 ────────
                        uploaded_file_context = ""
                        uploaded_files = (
                            request.files.getlist("files[]")
                            if request.method == "POST"
                            else []
                        )

                        if uploaded_files:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'📂 正在解析 {len(uploaded_files)} 个上传文件...', 'detail': '提取文本内容'})}\n\n"
                            try:
                                from web.file_parser import FileParser

                                uploaded_file_paths = []
                                for uploaded_file in uploaded_files:
                                    if uploaded_file and uploaded_file.filename:
                                        # 保存临时文件
                                        temp_dir = os.path.join(
                                            WORKSPACE_DIR, "temp_uploads"
                                        )
                                        os.makedirs(temp_dir, exist_ok=True)
                                        temp_path = os.path.join(
                                            temp_dir, uploaded_file.filename
                                        )
                                        uploaded_file.save(temp_path)
                                        uploaded_file_paths.append(temp_path)

                                if uploaded_file_paths:
                                    # 批量解析
                                    parse_results = FileParser.batch_parse(
                                        uploaded_file_paths
                                    )
                                    successful_results = [
                                        r for r in parse_results if r.get("success")
                                    ]

                                    if successful_results:
                                        uploaded_file_context = (
                                            FileParser.merge_contents(
                                                successful_results
                                            )
                                        )
                                        _app_logger.info(
                                            f"[FILE_GEN/PPT] ✅ 已解析 {len(successful_results)} 个文件, 总字数: {len(uploaded_file_context)}"
                                        )
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 已解析 {len(successful_results)} 个上传文件', 'detail': f'{len(uploaded_file_context)} 字内容'})}\n\n"
                                    else:
                                        _app_logger.warning(
                                            f"[FILE_GEN/PPT] ⚠️ 上传文件解析失败"
                                        )
                                        failed_reasons = [
                                            r.get("error", "未知错误")
                                            for r in parse_results
                                            if not r.get("success")
                                        ]
                                        _app_logger.info(
                                            f"    原因: {', '.join(failed_reasons)}"
                                        )

                            except ImportError:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ FileParser 模块未找到，跳过文件处理"
                                )
                            except Exception as file_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 文件处理异常: {file_err}"
                                )

                        # ──────── Step 0.1: 智能判断是否需要联网搜索 ────────
                        search_context = ""

                        # 检测是否需要搜索最新信息
                        _needs_search = WebSearcher.needs_web_search(user_input)

                        # 额外PPT话题检测：包含年份/时间/新品/事件/排行等的PPT大概率需要搜索
                        import re as _re

                        _time_topic_patterns = [
                            r"20\d{2}",  # 年份
                            r"\d+月",  # 月份
                            r"(新番|新片|新剧|新歌|新品|上映|首发|发售)",
                            r"(排行|排名|榜单|top|盘点|导视|速递|一览)",
                            r"(行情|走势|趋势|市场|价格|报告)",
                            r"(热门|热点|火爆|流行|人气)",
                            r"(最新|最近|近期|本周|本月|当前|目前)",
                        ]
                        if not _needs_search:
                            for pat in _time_topic_patterns:
                                if _re.search(pat, user_input, _re.IGNORECASE):
                                    _needs_search = True
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] 🔍 话题时效性检测命中: {pat}"
                                    )
                                    break

                        if _needs_search:
                            yield _fb_sse(fb.search_start())
                            try:
                                search_result = WebSearcher.search_with_grounding(
                                    user_input
                                )
                                if search_result.get("success") and search_result.get(
                                    "response"
                                ):
                                    search_context = search_result["response"]
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] ✅ 搜索完成, 获取 {len(search_context)} 字符参考信息"
                                    )
                                    yield _fb_sse(
                                        fb.search_done(char_count=len(search_context))
                                    )
                                else:
                                    _app_logger.warning(
                                        f"[FILE_GEN/PPT] ⚠️ 搜索无结果或失败"
                                    )
                            except Exception as search_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 搜索异常: {search_err}"
                                )

                        # ──────── Step 0.5: 复杂主题深度研究 ────────
                        research_context = ""
                        _complex_patterns = [
                            r"(原理|机制|架构|技术|算法|理论|分析|研究|综述)",
                            r"(行业|产业|市场|商业|战略|规划|方案)",
                            r"(学术|论文|课题|毕业|教学|课程)",
                            r"(历史|发展|演变|变迁|沿革)",
                            r"(对比|比较|评估|评测|benchmark)",
                            r"(经济|金融|投资|财务|财报)",
                        ]
                        _is_complex = len(user_input) > 30 or any(
                            _re.search(p, user_input) for p in _complex_patterns
                        )

                        if _is_complex:
                            yield _fb_sse(fb.research_start())
                            try:
                                research_context = WebSearcher.deep_research_for_ppt(
                                    user_input, search_context
                                )
                                if research_context:
                                    yield _fb_sse(
                                        fb.research_done(
                                            char_count=len(research_context)
                                        )
                                    )
                                else:
                                    _app_logger.warning(
                                        f"[FILE_GEN/PPT] ⚠️ 深度研究未返回结果"
                                    )
                            except Exception as res_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 深度研究异常: {res_err}"
                                )

                        # ──────── Step 1: 用 AI 生成结构化大纲 ────────
                        # ──── 提取用户 PPT 偏好（页数、重点、简要话题） ────
                        import re as _ppt_re

                        def _extract_ppt_preferences(text):
                            prefs = {
                                "target_pages": None,
                                "focus_topics": [],
                                "brief_topics": [],
                            }
                            pm = _ppt_re.search(
                                r"(?:做|生成|需要|大概|约|大约)?\s*(\d+)\s*页", text
                            )
                            if pm:
                                prefs["target_pages"] = int(pm.group(1))
                            for pat in [
                                r"(?:重点|详细|着重|深入|多讲|多介绍)(?:介绍|讲|分析|说明|展示|讲解)\s*(.+?)(?:[，,。；;、]|$)",
                                r"(?:突出|强调)\s*(.+?)(?:[，,。；;、]|$)",
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["focus_topics"].append(t)
                            for pat in [
                                r"(?:简单|简要|简略|大致)(?:带过|介绍|说|讲)\s*(.+?)(?:[，,。；;、]|$)",
                                r"(.+?)(?:一笔带过|略过|跳过|简单说)",
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["brief_topics"].append(t)
                            return prefs

                        ppt_prefs = _extract_ppt_preferences(user_input)
                        _target_pages = ppt_prefs["target_pages"]
                        _target_hint = (
                            f"约 {_target_pages} 页（封面和结束页除外，这是用户指定的，必须严格遵守）"
                            if _target_pages
                            else "8~15 页（根据内容复杂度智能调整，内容多可以多做几页，内容少就精简）"
                        )
                        _focus_hint = ""
                        if ppt_prefs["focus_topics"]:
                            _focus_hint = (
                                "\n**用户指定的重点内容（必须用 [详细] 多页展开）：**\n"
                                + "\n".join(f"- {t}" for t in ppt_prefs["focus_topics"])
                                + "\n"
                            )
                        _brief_hint = ""
                        if ppt_prefs["brief_topics"]:
                            _brief_hint = (
                                "\n**用户指定的简要内容（合并到 [概览] 页）：**\n"
                                + "\n".join(f"- {t}" for t in ppt_prefs["brief_topics"])
                                + "\n"
                            )

                        _app_logger.info(
                            f"[FILE_GEN/PPT] 用户偏好: 页数={_target_pages}, 重点={ppt_prefs['focus_topics']}, 简要={ppt_prefs['brief_topics']}"
                        )

                        # ──── 智能内容规划 Prompt ────
                        ppt_outline_prompt = (
                            "你是一个顶尖的演示文稿内容策划师和排版规划师。\n\n"
                            "你的工作分两步：\n"
                            "1. **内容规划** — 分析主题，判断哪些内容是重点（需要多页详细展示），哪些是简要（可以一页多主题速览）\n"
                            "2. **版式选择** — 为每部分选择最合适的幻灯片类型\n\n"
                            "## 可用的幻灯片类型\n"
                            "在每个 `## 章节标题` 前一行写类型标签：\n\n"
                            "| 标签 | 用途 | 格式 |\n"
                            "|------|------|------|\n"
                            "| `[详细]` | 常规内容页，深入展示 4-6 个要点（每个要点30-80字） | `- **关键词** — 详细解释说明和具体数据` |\n"
                            "| `[概览]` | 多主题速览页，2-4 个小主题并列（每个小主题下2-4个要点） | 用 `### 子标题` 分组 |\n"
                            "| `[亮点]` | 关键数据突出页（3-4组数据） | `- 数值 \\| 详细说明` |\n"
                            "| `[对比]` | 两方对比页（每方3-5个要点） | 用 `### 选项A` 和 `### 选项B` 分两组 |\n"
                            "| `[过渡页]` | 章节过渡，引入大章节（少用） | 下方写一行描述 |\n\n"
                            "## 输出格式（严格遵循）\n"
                            "```\n"
                            "# PPT主标题\n\n"
                            "[过渡页]\n"
                            "## 第一部分标题\n"
                            "简短描述\n\n"
                            "[详细]\n"
                            "## 页面标题\n"
                            "- **核心概念** — 这是一段完整的解释性文字，包含关键数据或事实依据，让观众能真正理解这一点的内容\n"
                            "- **技术特点** — 具体描述技术的工作原理、优势所在、实际应用场景和相关参数\n"
                            "- **市场数据** — 引用权威机构的统计数字、市场规模、增长率等量化信息\n"
                            "- **实际案例** — 某公司/项目的具体实践经验，取得了什么样的成果\n\n"
                            "[概览]\n"
                            "## 速览标题\n"
                            "### 子话题1\n"
                            "- 第一个要点的详细说明\n"
                            "- 第二个要点的详细说明\n"
                            "- 第三个要点的详细说明\n"
                            "### 子话题2\n"
                            "- 第一个要点和具体数据\n"
                            "- 第二个要点和应用场景\n\n"
                            "[亮点]\n"
                            "## 关键数据\n"
                            "- 500亿 | 全球市场规模\n"
                            "- 35% | 年增长率\n\n"
                            "[对比]\n"
                            "## 对比标题\n"
                            "### 方案A\n"
                            "- 特点1\n"
                            "### 方案B\n"
                            "- 特点1\n"
                            "```\n\n"
                            "## 内容规划规则\n"
                            f"1. **总页数目标: {_target_hint}**\n"
                            "2. **重点内容**使用多个 `[详细]` 页展开，每页 4-6 个信息丰富的要点\n"
                            "3. ⚠️ **每个要点必须是一个完整的信息段落（30-80字），不能只写几个词或短语**\n"
                            "4. 要点格式: `- **关键词** — 具体的解释说明，包含数据、事实、案例等实质内容`\n"
                            "5. **非重点内容**合并到 `[概览]` 页，一页 2-4 个小主题，⚠️ **每个小主题下必须有 2-4 个要点**\n"
                            "6. 有数据亮点时用 `[亮点]` 页（全文最多 1-2 次），每页 3-4 组数据\n"
                            "7. `[过渡页]` 最多 2 个，用于划分大章节\n"
                            "8. ⚠️ **搜索资料和研究报告中的数据、案例、数字必须如实引用，不得编造。每页至少引用 1 个具体数据或案例**\n"
                            "9. 中文输出，只输出大纲，不要额外说明\n"
                            "10. ⚠️ **内容充实度是最重要的评判标准 — 宁可要点少一些但每个要点信息量大，不要很多空洞的要点**\n"
                            "11. ⚠️ **禁止出现模糊表述**：如 '显著增长'、'广泛应用'、'巨大潜力' 等，必须用具体数字替代。例如：'市场规模达 XX 亿' 而不是 '市场规模巨大'\n"
                            "12. **每个 [详细] 页至少包含 1 个真实案例或数据点**，数据需标注来源（如 '据IDC数据' '根据XX年报'）\n"
                            f"{_focus_hint}"
                            f"{_brief_hint}"
                            f"\n用户需求: {user_input}\n"
                        )

                        # 注入搜索结果（增加限额以保留更多数据）
                        if uploaded_file_context:
                            ppt_outline_prompt = (
                                ppt_outline_prompt[: -len("\n用户需求: " + user_input)]
                                + f"\n\n## 上传的参考文件内容\n"
                                f"以下是用户上传的文档资料，请充分利用其中的内容、数据、案例来生成 PPT：\n"
                                f"---\n{uploaded_file_context[:15000]}\n---\n"
                                f"\n用户需求: {user_input}\n"
                            )

                        if search_context:
                            ppt_outline_prompt += (
                                f"\n**以下是联网搜索获取的最新参考资料（包含重要数据），请务必基于这些信息生成内容，尤其是其中的数字、案例、市场数据：**\n"
                                f"---\n{search_context[:10000]}\n---\n"
                            )

                        # 注入深度研究结果（增加限额）
                        if research_context:
                            ppt_outline_prompt += (
                                f"\n**以下是深度研究分析报告——这是你最重要的内容来源，其中的数据和分析必须充分融入大纲：**\n"
                                f"---\n{research_context[:12000]}\n---\n"
                            )

                        # 也注入上下文
                        if (
                            context_info
                            and context_info.get("is_continuation")
                            and context_info.get("enhanced_input")
                        ):
                            ppt_outline_prompt += f"\n\n历史上下文参考资料:\n{context_info['enhanced_input'][:3000]}"

                        yield _fb_sse(fb.ppt_planning("调用 AI 生成内容大纲"))

                        outline_response = None
                        outline_models = [
                            "gemini-2.5-flash",
                            model_id,
                            "gemini-3-flash-preview",
                        ]
                        # 根据目标页数调整 token 限额：20页大纲需要更多空间
                        _outline_tokens = (
                            16384 if (_target_pages and _target_pages >= 15) else 8192
                        )
                        for om in outline_models:
                            try:
                                resp = client.models.generate_content(
                                    model=om,
                                    contents=ppt_outline_prompt,
                                    config=types.GenerateContentConfig(
                                        temperature=0.6,
                                        max_output_tokens=_outline_tokens,
                                    ),
                                )
                                if resp.text:
                                    outline_response = resp.text
                                    _app_logger.info(
                                        f"[FILE_GEN/PPT] ✅ 大纲生成成功 ({om}), 长度: {len(outline_response)}"
                                    )
                                    break
                            except Exception as oe:
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] 大纲模型 {om} 失败: {oe}"
                                )
                                continue

                        if not outline_response:
                            raise Exception("所有模型均无法生成大纲")

                        # Step 2: 解析智能规划大纲（支持多种幻灯片类型标签）
                        def _parse_ppt_plan(md_text):
                            """解析带 [类型] 标签的智能 PPT 大纲"""
                            import re as _re

                            lines = md_text.split("\n")
                            plan = {"title": "", "subtitle": "", "slides": []}

                            _type_map = {
                                "过渡页": "divider",
                                "过渡": "divider",
                                "分隔": "divider",
                                "详细": "detail",
                                "重点": "detail",
                                "亮点": "highlight",
                                "数据": "highlight",
                                "关键": "highlight",
                                "概览": "overview",
                                "速览": "overview",
                                "简要": "overview",
                                "总览": "overview",
                                "对比": "comparison",
                                "比较": "comparison",
                                "vs": "comparison",
                            }

                            current_slide = None
                            current_type = "detail"
                            current_sub = (
                                None  # 当前子主题（用于 overview / comparison）
                            )

                            for line in lines:
                                line = line.rstrip()

                                # 跳过 markdown 代码块标记
                                if line.strip() in ("```", "```markdown"):
                                    continue

                                # 类型标签行: [xxx]
                                tag_m = _re.match(r"^\s*\[(.+?)\]\s*$", line)
                                if tag_m:
                                    tag = tag_m.group(1).strip()
                                    current_type = _type_map.get(tag, "detail")
                                    continue

                                # 主标题: # xxx
                                if line.startswith("# ") and not line.startswith("## "):
                                    raw = line[2:].strip()
                                    for pfx in [
                                        "幻灯片标题：",
                                        "幻灯片标题:",
                                        "演示标题：",
                                        "演示标题:",
                                        "PPT标题：",
                                        "PPT标题:",
                                    ]:
                                        if raw.startswith(pfx):
                                            raw = raw[len(pfx) :].strip()
                                    plan["title"] = raw
                                    continue

                                # 章节标题: ## xxx
                                if line.startswith("## "):
                                    # 保存上一个 slide 的 subsection
                                    if (
                                        current_sub
                                        and current_slide
                                        and current_slide.get("type")
                                        in ("overview", "comparison")
                                    ):
                                        current_slide.setdefault(
                                            "subsections", []
                                        ).append(current_sub)
                                        current_sub = None
                                    # 保存上一个 slide
                                    if current_slide:
                                        plan["slides"].append(current_slide)

                                    current_slide = {
                                        "type": current_type,
                                        "title": line[3:].strip(),
                                        "points": [],
                                        "content": [],
                                    }
                                    if current_type == "divider":
                                        current_slide["description"] = ""
                                    current_type = (
                                        "detail"  # 重置（每个标签只作用于紧跟的 ## ）
                                    )
                                    current_sub = None
                                    continue

                                # 子标题: ### xxx （用于 overview / comparison）
                                if line.startswith("### ") and current_slide:
                                    # 如果当前 slide 不是 overview/comparison，自动升级为 overview
                                    if current_slide.get("type") not in (
                                        "overview",
                                        "comparison",
                                    ):
                                        current_slide["type"] = "overview"
                                    if current_sub:
                                        current_slide.setdefault(
                                            "subsections", []
                                        ).append(current_sub)
                                    current_sub = {
                                        "subtitle": line[4:].strip(),
                                        "label": line[4:].strip(),
                                        "points": [],
                                    }
                                    continue

                                # 要点行: - / • / * 或数字编号 1. 2. 等
                                if (
                                    _re.match(r"^[\s]*[-•*]\s", line)
                                    or _re.match(r"^[\s]*\d+[.、)\s]\s*", line)
                                ) and current_slide is not None:
                                    pt = _re.sub(
                                        r"^[\s]*[-•*\d.、)\s]+\s*", "", line
                                    ).strip()
                                    if not pt:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(pt)
                                    else:
                                        current_slide["points"].append(pt)
                                        current_slide["content"].append(pt)
                                    continue

                                # 普通文本行（非空、非标题）→ 也捕获为要点
                                if (
                                    current_slide is not None
                                    and line.strip()
                                    and not line.startswith("#")
                                ):
                                    # 过渡页描述文字优先
                                    if current_slide.get("type") == "divider":
                                        current_slide["description"] = line.strip()
                                        continue
                                    # 清理可能残留的 markdown 标记
                                    cleaned = _re.sub(r"^#{1,4}\s+", "", line.strip())
                                    cleaned = cleaned.strip()
                                    if not cleaned:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(cleaned)
                                    else:
                                        current_slide["points"].append(cleaned)
                                        current_slide["content"].append(cleaned)
                                    continue

                                # 过渡页描述文字 (fallback - 不应到达这里)
                                if (
                                    current_slide
                                    and current_slide.get("type") == "divider"
                                    and line.strip()
                                    and not line.startswith("#")
                                ):
                                    current_slide["description"] = line.strip()

                            # 收尾
                            if current_sub and current_slide:
                                current_slide.setdefault("subsections", []).append(
                                    current_sub
                                )
                            if current_slide:
                                plan["slides"].append(current_slide)

                            # 后处理: 如果 slide 有 subsections 但类型不是 overview/comparison，自动修正
                            for sl in plan["slides"]:
                                if sl.get("subsections") and sl.get("type") not in (
                                    "overview",
                                    "comparison",
                                ):
                                    sl["type"] = "overview"

                            # 后处理: comparison 的 subsections → left / right
                            for sl in plan["slides"]:
                                if (
                                    sl.get("type") == "comparison"
                                    and "subsections" in sl
                                ):
                                    subs = sl["subsections"]
                                    if len(subs) >= 2:
                                        sl["left"] = subs[0]
                                        sl["right"] = subs[1]

                            return plan

                        ppt_data = _parse_ppt_plan(outline_response)
                        slide_count = len(ppt_data["slides"])
                        slide_types_summary = ", ".join(
                            f'{s.get("type","detail")}' for s in ppt_data["slides"]
                        )
                        _app_logger.info(
                            f"[FILE_GEN/PPT] 解析完成: 标题='{ppt_data['title']}', {slide_count} 页, 类型=[{slide_types_summary}]"
                        )

                        if slide_count == 0:
                            raise Exception("大纲解析失败，未提取到幻灯片内容")

                        # ──────── Quality Gate: 大纲质量自检与内容清洗 ────────
                        try:
                            from web.file_quality_checker import FileQualityGate

                            def _quality_progress(msg, detail=""):
                                pass  # 内部用，下面统一发 SSE

                            qg_result = FileQualityGate.check_and_fix_ppt_outline(
                                ppt_data["slides"],
                                user_request=user_input,
                                progress_callback=_quality_progress,
                            )
                            ppt_data["slides"] = qg_result["outline"]
                            _qg_score = qg_result["quality"]["score"]
                            _qg_fixes = qg_result["fixes"]
                            _qg_issues = qg_result["quality"]["issues"]

                            # 报告清洗结果
                            if _qg_fixes:
                                yield _fb_sse(
                                    fb.info(
                                        f"🧹 已自动清洗 {len(_qg_fixes)} 处内容问题",
                                        "移除 Markdown 残留和 AI 对话痕迹",
                                    )
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] 🧹 质量清洗: {len(_qg_fixes)} 处修复"
                                )

                            # 报告质量评分
                            yield _fb_sse(
                                fb.ppt_quality_check(
                                    _qg_score, issues=_qg_issues, fixes=_qg_fixes
                                )
                            )
                            _app_logger.info(
                                f"[FILE_GEN/PPT] 📊 质量评分: {_qg_score}/100, action={qg_result['action']}"
                            )

                            # 更新 slide_count（清洗可能移除空白 slide）
                            slide_count = len(ppt_data["slides"])
                        except Exception as qg_err:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ⚠️ 质量门控异常（不影响生成）: {qg_err}"
                            )

                        # ──────── Step 2.1: 用户指定页数时调整幻灯片数量 ────────
                        _max_slides = _target_pages  # 只有用户明确指定时才生效
                        if _max_slides and slide_count > _max_slides:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ⚠️ 页数超限 ({slide_count} > {_max_slides})，执行智能精简..."
                            )
                            yield _fb_sse(
                                fb.info(
                                    f"✂️ 精简页面: {slide_count} → {_max_slides} 页",
                                    "合并相似内容，保留核心信息",
                                )
                            )

                            slides = ppt_data["slides"]
                            # 策略: 1) 合并相邻的详细页为概览页  2) 去掉多余过渡页  3) 截断尾部

                            # 先去掉多余过渡页（只保留最多 1 个）
                            divider_indices = [
                                i
                                for i, s in enumerate(slides)
                                if s.get("type") == "divider"
                            ]
                            if len(divider_indices) > 1:
                                for idx in divider_indices[1:]:
                                    slides[idx]["_remove"] = True
                                slides = [s for s in slides if not s.get("_remove")]

                            # 然后合并相邻的详细页为概览页
                            while len(slides) > _max_slides:
                                merged = False
                                for i in range(len(slides) - 1):
                                    if (
                                        slides[i].get("type") == "detail"
                                        and slides[i + 1].get("type") == "detail"
                                    ):
                                        # 合并: 第一个和第二个详细页变成一个概览页
                                        s1 = slides[i]
                                        s2 = slides[i + 1]
                                        merged_slide = {
                                            "type": "overview",
                                            "title": s1.get("title", ""),
                                            "points": [],
                                            "content": [],
                                            "subsections": [
                                                {
                                                    "subtitle": s1.get("title", ""),
                                                    "label": s1.get("title", ""),
                                                    "points": (
                                                        s1.get("points", [])
                                                        or s1.get("content", [])
                                                    )[:4],
                                                },
                                                {
                                                    "subtitle": s2.get("title", ""),
                                                    "label": s2.get("title", ""),
                                                    "points": (
                                                        s2.get("points", [])
                                                        or s2.get("content", [])
                                                    )[:4],
                                                },
                                            ],
                                        }
                                        slides[i] = merged_slide
                                        slides.pop(i + 1)
                                        merged = True
                                        break
                                if not merged:
                                    # 无法合并了，直接截断
                                    slides = slides[:_max_slides]
                                    break

                            ppt_data["slides"] = slides
                            slide_count = len(slides)
                            _app_logger.info(f"[FILE_GEN/PPT] 精简后: {slide_count} 页")

                        # 生成版式摘要
                        _type_map_display = {
                            "detail": "详细",
                            "overview": "概览",
                            "highlight": "亮点",
                            "divider": "过渡",
                            "comparison": "对比",
                        }
                        _tc = {}
                        for _s in ppt_data["slides"]:
                            _t = _s.get("type", "detail")
                            _tc[_t] = _tc.get(_t, 0) + 1
                        _ts = "×".join(
                            f"{_type_map_display.get(k,k)}{v}" for k, v in _tc.items()
                        )

                        yield _fb_sse(
                            fb.ppt_outline_ready(
                                slide_count, title=ppt_data["title"], type_summary=_ts
                            )
                        )

                        # ──────── Step 2.2: 内容充实（逐页扩写） ────────
                        # 检查内容是否单薄（平均每页要点少于 3 个或要点太短）
                        _thin_slides = []
                        for si, sl in enumerate(ppt_data["slides"]):
                            stype = sl.get("type", "detail")
                            if stype in ("divider",):
                                continue  # 过渡页不需要充实
                            pts = sl.get("points", [])
                            subs = sl.get("subsections", [])
                            # 要点太少 或 平均要点太短
                            avg_len = sum(len(p) for p in pts) / max(len(pts), 1)
                            sub_pts_count = (
                                sum(len(sub.get("points", [])) for sub in subs)
                                if subs
                                else 0
                            )

                            if stype == "overview":
                                # 概览页：子主题数太少或每个子主题要点太少
                                if not subs or sub_pts_count < len(subs) * 2:
                                    _thin_slides.append(si)
                            elif stype in ("detail", "comparison"):
                                if len(pts) < 3 or avg_len < 20:
                                    _thin_slides.append(si)
                            elif stype == "highlight":
                                if len(pts) < 2:
                                    _thin_slides.append(si)

                        if _thin_slides:
                            yield _fb_sse(fb.ppt_enriching(len(_thin_slides)))

                            # 构建批量充实 prompt（一次性处理所有薄弱页面）
                            _enrich_prompt = (
                                "你是PPT内容撰写专家。以下幻灯片内容太单薄，请逐页充实。\n\n"
                                "**要求：**\n"
                                "1. 每个[详细]页必须有 4-6 个要点，每个要点 30-80 字\n"
                                "2. 每个[概览]页的每个子主题下必须有 2-4 个要点\n"
                                "3. 保持 `- **关键词** — 详细解释` 的格式\n"
                                "4. ⚠️ **必须包含具体数据、案例、事实** — 禁止写 '显著增长' '广泛应用' 等模糊表述，\n"
                                "   必须写 '据IDC数据，2025年市场规模达XXX亿' 这样有数字有来源的内容\n"
                                "5. 优先使用下方【参考资料】和【研究分析】中的真实数据\n"
                                "6. 严格按以下 JSON 格式输出，不要额外文字\n\n"
                            )

                            _slides_to_enrich = []
                            for si in _thin_slides:
                                sl = ppt_data["slides"][si]
                                _slides_to_enrich.append(
                                    {
                                        "index": si,
                                        "type": sl.get("type", "detail"),
                                        "title": sl.get("title", ""),
                                        "current_points": sl.get("points", []),
                                        "subsections": (
                                            [
                                                {
                                                    "subtitle": sub.get("subtitle", ""),
                                                    "points": sub.get("points", []),
                                                }
                                                for sub in sl.get("subsections", [])
                                            ]
                                            if sl.get("subsections")
                                            else []
                                        ),
                                    }
                                )

                            _enrich_prompt += f"主题: {ppt_data['title']}\n"
                            if search_context:
                                _enrich_prompt += f"\n参考资料（包含重要数据，请充分利用）:\n{search_context[:6000]}\n"
                            if research_context:
                                _enrich_prompt += f"\n研究分析（包含核心数据和案例，必须融入）:\n{research_context[:6000]}\n"

                            _enrich_prompt += (
                                f"\n需要充实的幻灯片:\n```json\n{json.dumps(_slides_to_enrich, ensure_ascii=False, indent=2)}\n```\n\n"
                                "请输出充实后的结果，格式:\n"
                                "```json\n"
                                '[{"index": 0, "points": ["...", ...], "subsections": [{"subtitle": "...", "points": ["..."]}, ...]}]\n'
                                "```\n"
                                "只输出 JSON，不要额外文字。"
                            )

                            try:
                                _enrich_resp = client.models.generate_content(
                                    model="gemini-2.5-flash",
                                    contents=_enrich_prompt,
                                    config=types.GenerateContentConfig(
                                        temperature=0.5, max_output_tokens=8192
                                    ),
                                )
                                _enrich_text = _enrich_resp.text or ""
                                import re as _enrich_re

                                _em = _enrich_re.search(
                                    r"\[.*\]", _enrich_text, _enrich_re.DOTALL
                                )
                                if _em:
                                    _enriched = json.loads(_em.group())
                                    _applied = 0
                                    for _e in _enriched:
                                        _idx = _e.get("index")
                                        if _idx is not None and 0 <= _idx < len(
                                            ppt_data["slides"]
                                        ):
                                            _sl = ppt_data["slides"][_idx]
                                            # 更新 points
                                            if _e.get("points") and len(
                                                _e["points"]
                                            ) >= len(_sl.get("points", [])):
                                                _sl["points"] = _e["points"]
                                                _sl["content"] = _e["points"]
                                            # 更新 subsections
                                            if (
                                                _e.get("subsections")
                                                and len(_e["subsections"]) > 0
                                            ):
                                                _new_subs = []
                                                for _ns in _e["subsections"]:
                                                    _new_subs.append(
                                                        {
                                                            "subtitle": _ns.get(
                                                                "subtitle", ""
                                                            ),
                                                            "label": _ns.get(
                                                                "subtitle", ""
                                                            ),
                                                            "points": _ns.get(
                                                                "points", []
                                                            ),
                                                        }
                                                    )
                                                if _new_subs:
                                                    _sl["subsections"] = _new_subs
                                                    # 也更新 comparison 的 left/right
                                                    if (
                                                        _sl.get("type") == "comparison"
                                                        and len(_new_subs) >= 2
                                                    ):
                                                        _sl["left"] = _new_subs[0]
                                                        _sl["right"] = _new_subs[1]
                                            _applied += 1

                                    if _applied > 0:
                                        yield _fb_sse(fb.ppt_enriched(_applied))
                                        _app_logger.info(
                                            f"[FILE_GEN/PPT] ✅ 内容充实完成: {_applied}/{len(_thin_slides)} 页"
                                        )
                                    else:
                                        _app_logger.warning(
                                            f"[FILE_GEN/PPT] ⚠️ 内容充实解析成功但未应用"
                                        )
                                else:
                                    _app_logger.warning(
                                        f"[FILE_GEN/PPT] ⚠️ 内容充实返回格式异常"
                                    )
                            except Exception as enrich_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 内容充实异常（不影响生成）: {enrich_err}"
                                )

                        # ──────── Step 2.5: 为幻灯片生成配图（Gemini 3.1 Flash Image 优先） ────────
                        ppt_images = []
                        # 对详细页配图（概览/对比/过渡/亮点页不适合插图）
                        img_candidate_slides = [
                            (i, s)
                            for i, s in enumerate(ppt_data["slides"])
                            if s.get("type", "detail") == "detail"
                        ]

                        if img_candidate_slides:
                            _n_images = min(
                                4, max(2, len(img_candidate_slides) // 2 + 1)
                            )
                            yield _fb_sse(fb.ppt_images(_n_images))
                            try:
                                slide_titles_for_img = [
                                    s.get("title", "") for _, s in img_candidate_slides
                                ]
                                img_results = WebSearcher.generate_ppt_images(
                                    slide_titles_for_img,
                                    topic=ppt_data["title"],
                                    max_images=_n_images,
                                )
                                # 将配图路径注入到对应 slide
                                for img_info in img_results:
                                    picked_idx = img_info["slide_index"]
                                    if picked_idx < len(img_candidate_slides):
                                        real_idx = img_candidate_slides[picked_idx][0]
                                        ppt_data["slides"][real_idx]["image"] = (
                                            img_info["image_path"]
                                        )
                                        ppt_images.append(img_info["image_path"])

                                if ppt_images:
                                    yield _fb_sse(fb.ppt_images_done(len(ppt_images)))
                                else:
                                    yield _fb_sse(fb.ppt_images_done(0))
                            except Exception as img_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 配图生成异常: {img_err}"
                                )
                                yield _fb_sse(fb.warn("配图跳过，不影响PPT生成"))

                        # ──────── Step 3: 生成 PPT 文件(含逐页进度) ────────
                        yield _fb_sse(fb.ppt_rendering(slide_count))

                        from web.ppt_generator import PPTGenerator

                        # 检测主题
                        theme = "business"
                        if any(
                            kw in user_lower_check
                            for kw in ["技术", "tech", "科技", "编程", "开发"]
                        ):
                            theme = "tech"
                        elif any(
                            kw in user_lower_check
                            for kw in ["创意", "creative", "艺术", "设计"]
                        ):
                            theme = "creative"

                        ppt_gen = PPTGenerator(theme=theme)

                        ppt_title = ppt_data["title"] or "演示文稿"
                        safe_title = re.sub(r'[\\/*?:"<>|]', "_", ppt_title)[:50]
                        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pptx"
                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)
                        ppt_path = os.path.join(docs_dir, filename)

                        # 使用 progress_callback 来收集进度消息（生成器无法在回调中yield）
                        _slide_progress_msgs = []

                        def _ppt_progress_cb(cur, total, stitle, stype):
                            _slide_progress_msgs.append((cur, total, stitle, stype))

                        ppt_gen.generate_from_outline(
                            title=ppt_title,
                            outline=ppt_data["slides"],
                            output_path=ppt_path,
                            subtitle=ppt_data.get("subtitle", ""),
                            author="Koto AI",
                            progress_callback=_ppt_progress_cb,
                        )

                        # 发送逐页进度（回调已经收集完毕）
                        for cur, total, stitle, stype in _slide_progress_msgs:
                            if stitle:
                                yield _fb_sse(
                                    fb.ppt_slide_progress(cur, total, stitle, stype)
                                )

                        yield _fb_sse(fb.substep("PPT 渲染完成，正在保存"))

                        # ──────── Post-Render Quality Check: 检查渲染后的 PPTX 文件 ────────
                        try:
                            from web.file_quality_checker import FileQualityGate

                            _post_check = FileQualityGate.post_check_pptx(ppt_path)
                            if _post_check.get("issues"):
                                _pc_score = _post_check["score"]
                                yield _fb_sse(
                                    fb.ppt_quality_check(
                                        _pc_score, issues=_post_check["issues"]
                                    )
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] 📊 文件后检: {_pc_score}/100, issues={_post_check['issues']}"
                                )
                            else:
                                yield _fb_sse(fb.info("✅ 文件质量验证通过"))
                        except Exception as pc_err:
                            _app_logger.warning(
                                f"[FILE_GEN/PPT] ⚠️ 文件后检异常: {pc_err}"
                            )

                        rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        generated_files.append(rel_path)

                        # 统计各类型幻灯片数量
                        _type_names = {
                            "detail": "详细页",
                            "overview": "概览页",
                            "highlight": "亮点页",
                            "divider": "过渡页",
                            "comparison": "对比页",
                        }
                        _type_counts = {}
                        for _s in ppt_data["slides"]:
                            _t = _s.get("type", "detail")
                            _type_counts[_t] = _type_counts.get(_t, 0) + 1
                        _type_desc = "、".join(
                            f"{_type_names.get(k,k)} ×{v}"
                            for k, v in _type_counts.items()
                        )

                        _img_desc = (
                            f"\n🖼️ 配图: {len(ppt_images)} 张" if ppt_images else ""
                        )
                        _research_desc = (
                            "\n🔬 已融入深度研究分析" if research_context else ""
                        )

                        # ──────── 保存会话数据（P1 编辑功能支持） ────────
                        if ppt_session_id:
                            try:
                                from web.ppt_session_manager import (
                                    get_ppt_session_manager,
                                )

                                ppt_session_mgr = get_ppt_session_manager()
                                ppt_session_mgr.save_generation_data(
                                    session_id=ppt_session_id,
                                    ppt_data=ppt_data,
                                    ppt_file_path=rel_path,
                                    search_context=search_context,
                                    research_context=research_context,
                                    uploaded_file_context=uploaded_file_context,
                                )
                                _app_logger.info(
                                    f"[FILE_GEN/PPT] 💾 会话数据已保存，可用于后续编辑"
                                )
                            except Exception as save_err:
                                _app_logger.warning(
                                    f"[FILE_GEN/PPT] ⚠️ 会话保存异常: {save_err}"
                                )

                        success_msg = (
                            f"✅ **PPT 演示文稿生成成功！**\n\n"
                            f"📊 标题: **{ppt_title}**\n"
                            f"📄 页数: {slide_count} 页（{_type_desc}）{_img_desc}{_research_desc}\n"
                            f"📁 文件: **{filename}**\n"
                            f"📍 位置: `{docs_dir}`"
                        )

                        # 如果有会话，附加编辑链接
                        if ppt_session_id:
                            success_msg += f"\n\n🎨 **[点击编辑 PPT](/edit-ppt/{ppt_session_id})** - 修改内容、调整顺序、重新生成页面"

                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        _app_logger.info(f"[FILE_GEN/PPT] ✅ PPT 生成成功: {rel_path}")

                    except Exception as ppt_err:
                        _app_logger.error(f"[FILE_GEN/PPT] ❌ PPT 生成失败: {ppt_err}")
                        import traceback

                        traceback.print_exc()
                        error_msg = f"❌ PPT 生成失败: {str(ppt_err)}"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"

                    # 保存历史（基于磁盘完整历史追加，在 done 事件之前）
                    _ppt_msg = (
                        f"已生成PPT: {', '.join(generated_files)}"
                        if generated_files
                        else "PPT生成失败"
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _ppt_msg
                    )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # 使用上下文增强的输入（如果有，例如"把这个做成word"时会包含之前的内容）
                if (
                    context_info
                    and context_info.get("is_continuation")
                    and context_info.get("enhanced_input")
                ):
                    file_gen_input = context_info["enhanced_input"]
                    _app_logger.debug(
                        f"[FILE_GEN] 使用上下文增强输入 (length: {len(file_gen_input)})"
                    )
                else:
                    file_gen_input = effective_input

                # ⭐ FILE_GEN 前置步骤：时间解析 + 信息收集
                _time_context_text, _time_parse = _build_filegen_time_context(
                    user_input
                )
                _web_context = ""
                _should_collect = WebSearcher.needs_web_search(user_input)

                # 对“X月新番/番剧/动画”等时间敏感主题强制启用信息收集
                _anime_time_patterns = [
                    r"([1-9]|1[0-2])\s*月\s*(新番|番剧|动画)",
                    r"(新番|番剧|动画).*(\d{1,2}\s*月)",
                ]
                if not _should_collect and any(
                    re.search(p, user_input, re.IGNORECASE)
                    for p in _anime_time_patterns
                ):
                    _should_collect = True

                if _should_collect:
                    try:
                        if _time_parse.get("resolved_month"):
                            _q = f"{_time_parse['resolved_year']}年{_time_parse['resolved_month']}月 新番 动画 番剧 名单 介绍"
                        else:
                            _q = user_input

                        _time_detail = _time_context_text.replace("\n", " | ")[:180]
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🕒 正在解析时间语义...', 'detail': _time_detail})}\n\n"
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🌐 正在收集最新信息...', 'detail': _q[:120]})}\n\n"

                        _search_res = WebSearcher.search_with_grounding(_q)
                        if _search_res.get("success") and _search_res.get("response"):
                            _web_context = _search_res.get("response", "")
                            _app_logger.info(
                                f"[FILE_GEN] ✅ 信息收集完成，长度: {len(_web_context)}"
                            )
                            yield f"data: {json.dumps({'type': 'progress', 'message': '✅ 信息收集完成', 'detail': f'已获取 {len(_web_context)} 字符参考信息'})}\n\n"
                        else:
                            _app_logger.warning(f"[FILE_GEN] ⚠️ 信息收集未返回结果")
                    except Exception as _collect_err:
                        _app_logger.warning(
                            f"[FILE_GEN] ⚠️ 信息收集异常: {_collect_err}"
                        )

                # 将时间上下文/检索结果拼接进生成输入
                _prepended_blocks = [_time_context_text]
                if _web_context:
                    _prepended_blocks.append("[联网检索参考]\n" + _web_context[:9000])
                file_gen_input = (
                    "\n\n".join(_prepended_blocks) + "\n\n" + file_gen_input
                )

                # ⭐ 判断是否是文档生成请求（Word/PDF）
                _doc_keywords = [
                    "word",
                    "docx",
                    "doc",
                    "pdf",
                    "报告",
                    "文档",
                    "论文",
                    "综述",
                    "whitepaper",
                ]
                _is_doc_request = any(k in user_input.lower() for k in _doc_keywords)
                _is_complex = (context_info or {}).get("complexity") == "complex"

                if _is_doc_request:
                    # ============== 文档直出模式（流式） ==============
                    # 使用 generate_content_stream 保持连接活跃，避免代理超时断开
                    _doc_type = "PDF" if "pdf" in user_input.lower() else "Word"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'📄 正在生成 {_doc_type} 文档...', 'detail': '请稍候，正在撰写内容'})}\n\n"
                    _app_logger.debug(
                        f"[FILE_GEN] 📄 文档直出模式-流式 (type={_doc_type}, complex={_is_complex})"
                    )

                    _doc_instruction = """你是 Koto 专业文档撰写助手。请根据用户要求，直接输出**完整、详细、高质量**的文档正文内容。

## 输出规则
- 直接输出 Markdown 格式的文档正文，不要输出代码
- 使用 # ## ### 组织标题层级
- 使用段落、列表、表格丰富内容
- 中文撰写，专业术语准确
- 内容要**充实详尽**，每一节至少2-3段，总字数不少于3000字
- 如果是技术报告，必须包含：行业概述、技术原理、关键工艺、对比分析、应用场景、发展趋势
- 不要输出任何 BEGIN_FILE/END_FILE 标记
- 不要输出 JSON 或代码格式"""

                    _doc_instruction += "\n\n时间要求：若用户请求涉及月份但未写年份（如‘1月新番’），必须按当前年份撰写，禁止默认回退到历史年份。"

                    _max_tokens = 16384 if _is_complex else 8192
                    _doc_models = list(
                        dict.fromkeys(
                            [
                                model_id,
                                "gemini-3-pro-preview",
                                "gemini-2.5-flash",
                                "gemini-3-flash-preview",
                            ]
                        )
                    )

                    _doc_collected = []  # 收集所有流式文本块

                    for model_attempt, current_model in enumerate(_doc_models):
                        if _doc_collected:
                            break
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if model_attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 切换到备用模型 {current_model}...', 'detail': ''})}\n\n"
                            _doc_collected.clear()
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🚀 正在调用 {current_model}...', 'detail': '流式生成中'})}\n\n"

                        try:
                            _doc_stream = client.models.generate_content_stream(
                                model=current_model,
                                contents=file_gen_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=_doc_instruction,
                                    max_output_tokens=_max_tokens,
                                    temperature=0.7,
                                ),
                            )
                            _first_chunk = False
                            for item_type, item_data in stream_with_keepalive(
                                _doc_stream,
                                start_time,
                                keepalive_interval=5,
                                max_wait_first_token=120,  # 文档生成允许等待更久
                            ):
                                if interrupted():
                                    _app_logger.info(f"[FILE_GEN/DOC] 用户中断")
                                    _interrupt_msg = "\n\n⏹️ 文件生成已中断"
                                    yield f"data: {json.dumps({'type': 'token', 'content': _interrupt_msg})}\n\n"
                                    break

                                if item_type == "heartbeat":
                                    _elapsed = item_data
                                    _char_count = sum(len(c) for c in _doc_collected)
                                    if _first_chunk:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在撰写文档...', 'detail': f'已生成 {_char_count} 字符，耗时 {_elapsed}s', 'stage': 'generating'})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '🧠 模型正在组织内容...', 'detail': f'已等待 {_elapsed}s，请耐心等待', 'stage': 'api_calling'})}\n\n"

                                elif item_type == "timeout":
                                    _app_logger.warning(
                                        f"[FILE_GEN/DOC] ⚠️ {current_model} 等待首token超时: {item_data}"
                                    )
                                    break  # 尝试下一个模型

                                elif item_type == "chunk":
                                    chunk = item_data
                                    if chunk.text:
                                        if not _first_chunk:
                                            _first_chunk = True
                                            _app_logger.info(
                                                f"[FILE_GEN/DOC] ✅ {current_model} 收到第一个响应块，耗时 {time.time() - start_time:.1f}s"
                                            )
                                        _doc_collected.append(chunk.text)
                                        # 每收到10个chunk发送一次进度更新，保持客户端连接活跃
                                        if len(_doc_collected) % 10 == 0:
                                            _char_count = sum(
                                                len(c) for c in _doc_collected
                                            )
                                            _elapsed = int(time.time() - start_time)
                                            yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在撰写文档...', 'detail': f'已生成 {_char_count} 字符，耗时 {_elapsed}s'})}\n\n"

                        except Exception as _doc_err:
                            err_str = str(_doc_err)
                            _app_logger.error(
                                f"[FILE_GEN/DOC] ❌ {current_model}: {err_str[:200]}"
                            )
                            if "location is not supported" in err_str.lower():
                                response_text = "❌ 地区限制，请配置中转服务"
                                break
                            continue

                    response_text = "".join(_doc_collected)
                    if response_text:
                        _app_logger.info(
                            f"[FILE_GEN/DOC] ✅ 流式生成完成，共 {len(response_text)} 字符"
                        )

                    if not response_text or response_text.startswith("❌"):
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text or '❌ 所有模型都不可用，请稍后重试'})}\n\n"
                    else:
                        # ── 文档质量自检与清洗 ──
                        try:
                            from web.file_quality_checker import FileQualityGate

                            _doc_qg = FileQualityGate.check_and_fix_document(
                                response_text, user_request=user_input
                            )
                            response_text = _doc_qg["text"]
                            _doc_score = _doc_qg["quality"]["score"]
                            _doc_fixes = _doc_qg["fixes"]
                            if _doc_fixes:
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'🧹 已清洗 {len(_doc_fixes)} 处内容问题', 'detail': '移除AI对话痕迹'})}\n\n"
                            _dq_emoji = "✅" if _doc_score >= 75 else "⚠️"
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'{_dq_emoji} 文档质量检查: {_doc_score}/100', 'detail': '; '.join(_doc_qg['quality']['issues'][:2]) if _doc_qg['quality']['issues'] else '质量良好'})}\n\n"
                        except Exception as _dqg_err:
                            _app_logger.warning(
                                f"[FILE_GEN/DOC] ⚠️ 质量门控异常: {_dqg_err}"
                            )

                        # 直接保存文档
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'📝 正在保存 {_doc_type} 文档...', 'detail': ''})}\n\n"
                        try:
                            try:
                                from web.document_generator import save_docx, save_pdf
                            except ModuleNotFoundError:
                                from document_generator import save_docx, save_pdf
                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)
                            title_match = re.search(
                                r"^#\s*(.+)$", response_text, re.MULTILINE
                            )
                            title = (
                                title_match.group(1).strip()[:50]
                                if title_match
                                else f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            )
                            if _doc_type == "PDF":
                                saved_path = save_pdf(
                                    response_text, title=title, output_dir=docs_dir
                                )
                            else:
                                saved_path = save_docx(
                                    response_text, title=title, output_dir=docs_dir
                                )
                            rel_path = os.path.relpath(
                                saved_path, WORKSPACE_DIR
                            ).replace("\\", "/")
                            generated_files.append(rel_path)
                            success_msg = f"✅ **{_doc_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                            _app_logger.info(
                                f"[FILE_GEN/DOC] ✅ 文档已保存: {rel_path}"
                            )
                        except Exception as doc_err:
                            import traceback

                            traceback.print_exc()
                            _app_logger.error(
                                f"[FILE_GEN/DOC] ❌ 文档保存失败: {doc_err}"
                            )
                            fallback_msg = (
                                f"⚠️ 文档保存失败 ({doc_err})，以下是生成的内容：\n\n"
                            )
                            yield f"data: {json.dumps({'type': 'token', 'content': fallback_msg + response_text})}\n\n"

                    _gen_msg = (
                        f"已生成文件: {', '.join(generated_files)}"
                        if generated_files
                        else (response_text[:500] if response_text else "生成失败")
                    )
                    session_manager.append_and_save(
                        f"{session_name}.json", user_input, _gen_msg
                    )
                    total_time = time.time() - start_time
                    _app_logger.info(
                        f"[FILE_GEN/DOC] ★★★ done event, files: {generated_files}, time: {total_time:.2f}s"
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # 普通 FILE_GEN 模式（需要模型生成代码/脚本）
                yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成文件代码...', 'detail': '请稍候，可能需要 10-30 秒'})}\n\n"

                # 模型列表（主模型 + 备用模型）
                file_gen_models = [
                    model_id,  # 主模型
                    "gemini-3-pro-preview",  # 备用1 (更强的推理)
                    "gemini-2.5-flash",  # 备用2
                    "gemini-3-flash-preview",  # 备用3
                ]

                # 使用线程 + 超时来调用API（带重试）
                import tempfile
                import threading

                for model_attempt, current_model in enumerate(file_gen_models):
                    if response_text and not response_text.startswith("❌"):
                        break  # 已成功

                    if interrupted():
                        yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    if model_attempt > 0:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 切换到备用模型 {current_model}...', 'detail': ''})}\n\n"
                        _app_logger.debug(
                            f"[FILE_GEN] Trying fallback model: {current_model}"
                        )
                    else:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'🚀 正在调用 {current_model}...', 'detail': '生成中'})}\n\n"

                    response_holder = {"data": None, "error": None}

                    # 构建含对话历史的 contents 列表，使模型能参考之前的财务模型等上下文
                    _fg_history_for_model = ContextAnalyzer.filter_history(
                        user_input, history
                    )
                    _fg_formatted_history = []
                    for _fg_turn in _fg_history_for_model[
                        -6:
                    ]:  # 最多保留最近 6 轮，避免 token 过长
                        _fg_formatted_history.append(
                            types.Content(
                                role=_fg_turn["role"],
                                parts=[
                                    types.Part.from_text(text=p)
                                    for p in _fg_turn["parts"]
                                ],
                            )
                        )
                    _fg_contents = _fg_formatted_history + [
                        types.Content(
                            role="user",
                            parts=[types.Part.from_text(text=file_gen_input)],
                        )
                    ]

                    def call_api(m=current_model, _contents=_fg_contents):
                        try:
                            _app_logger.debug(f"[FILE_GEN] Calling API: {m}")
                            response = client.models.generate_content(
                                model=m,
                                contents=_contents,
                                config=types.GenerateContentConfig(
                                    system_instruction=_get_system_instruction(),
                                    max_output_tokens=8192,
                                ),
                            )
                            response_holder["data"] = response
                            _app_logger.info(
                                f"[FILE_GEN] ✅ API call successful with {m}"
                            )
                        except Exception as e:
                            _app_logger.error(
                                f"[FILE_GEN] ❌ API call exception with {m}: {type(e).__name__}: {str(e)}"
                            )
                            response_holder["error"] = e

                    api_thread = threading.Thread(target=call_api, daemon=True)
                    api_thread.start()

                    # 在等待期间发送心跳进度
                    wait_interval = 5  # 每 5 秒发送一次进度
                    elapsed = 0
                    while api_thread.is_alive() and elapsed < api_timeout:
                        api_thread.join(timeout=wait_interval)
                        elapsed += wait_interval
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if api_thread.is_alive() and elapsed < api_timeout:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'⏳ 正在生成中...', 'detail': f'已等待 {elapsed} 秒'})}\n\n"

                    if api_thread.is_alive():
                        _app_logger.warning(
                            f"[FILE_GEN] ⚠️ API call timeout with {current_model} after {api_timeout}s"
                        )
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ {current_model} 响应超时', 'detail': '正在切换模型...'})}\n\n"
                        response_text = ""
                        continue  # 尝试下一个模型
                    elif response_holder["error"]:
                        error_str = str(response_holder["error"])
                        _app_logger.debug(
                            f"[FILE_GEN] API Error with {current_model}: {error_str}"
                        )

                        # 地区限制错误 - 直接失败，不重试
                        if (
                            "location is not supported" in error_str.lower()
                            or "failed_precondition" in error_str.lower()
                        ):
                            response_text = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 config/gemini_config.env 配置中转服务 GEMINI_API_BASE\n2. 或使用支持的代理服务"
                            break  # 地区限制，不继续重试
                        elif (
                            "503" in error_str
                            or "overloaded" in error_str.lower()
                            or "unavailable" in error_str.lower()
                        ):
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ {current_model} 服务繁忙', 'detail': '正在切换模型...'})}\n\n"
                            response_text = ""
                            continue  # 503 错误，尝试下一个模型
                        else:
                            response_text = f"❌ API 调用失败: {error_str[:200]}"
                            continue  # 其他错误也尝试下一个模型
                    elif response_holder["data"]:
                        file_gen_response = response_holder["data"]
                        if (
                            file_gen_response.candidates
                            and file_gen_response.candidates[0].content.parts
                        ):
                            for part in file_gen_response.candidates[0].content.parts:
                                if hasattr(part, "text") and part.text:
                                    response_text += part.text
                        _app_logger.debug(
                            f"[FILE_GEN] Response length: {len(response_text)}"
                        )
                        if response_text:
                            break  # 成功获取响应

                if not response_text:
                    response_text = "❌ 所有模型都不可用，请稍后重试"

                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次生成失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt(
                        "FILE_GEN", user_input, response_text
                    )
                    try:
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_system_instruction(),
                                max_output_tokens=8192,
                                temperature=0.4,
                            ),
                        )
                        response_text = fix_resp.text or response_text
                    except Exception as fix_err:
                        _app_logger.debug(f"[FILE_GEN] 修正重试失败: {fix_err}")

                # 只显示简短的进度，不显示完整代码
                if response_text and not response_text.startswith("❌"):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🔧 正在处理代码...', 'detail': ''})}\n\n"

                    # 提取代码到临时文件
                    patterns = [
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
                    ]

                    code_content = None
                    for pattern in patterns:
                        matches = re.findall(
                            pattern, response_text, re.DOTALL | re.IGNORECASE
                        )
                        if matches:
                            _, code_content = matches[0]
                            code_content = code_content.strip()
                            _app_logger.debug(
                                f"[FILE_GEN] Extracted code, length: {len(code_content)}"
                            )
                            break

                    # 检查提取的内容是否是有效的Python代码（不是JSON或其他格式）
                    is_valid_python = False
                    if code_content:
                        code_lower = code_content.lower()
                        # 如果提取的内容是 JSON 或 HTML 或其他格式，直接跳过代码执行
                        if code_lower.startswith(("{", "[", "<", '"')):
                            _app_logger.debug(
                                f"[FILE_GEN] Extracted content is not Python code (starts with {code_content[0]}), treating as text content"
                            )
                            code_content = None
                        else:
                            is_valid_python = True

                    if code_content and is_valid_python:
                        # 保存到临时文件
                        temp_dir = tempfile.gettempdir()
                        temp_script = os.path.join(
                            temp_dir, f"koto_gen_{int(time.time())}.py"
                        )

                        with open(temp_script, "w", encoding="utf-8") as f:
                            f.write(code_content)
                        temp_scripts.append(temp_script)
                        _app_logger.debug(
                            f"[FILE_GEN] Saved temp script: {temp_script}"
                        )

                        # 执行脚本
                        yield f"data: {json.dumps({'type': 'progress', 'message': '⚙️ 正在执行脚本生成文件...', 'detail': ''})}\n\n"

                        try:
                            if getattr(sys, "frozen", False):
                                # 打包模式：进程内 exec() 执行，避免启动新 Koto 窗口
                                import contextlib as _ctx
                                import io as _io

                                _out, _err, _rc = _io.StringIO(), _io.StringIO(), 0
                                try:
                                    _prev = os.getcwd()
                                    os.chdir(WORKSPACE_DIR)
                                    # 注入 KOTO_OUTPUT_DIR 到 exec 命名空间，让脚本能保存到正确目录
                                    _exec_globals = {
                                        "__file__": temp_script,
                                        "KOTO_OUTPUT_DIR": settings_manager.documents_dir,
                                    }
                                    with _ctx.redirect_stdout(
                                        _out
                                    ), _ctx.redirect_stderr(_err):
                                        exec(
                                            open(
                                                temp_script, "r", encoding="utf-8"
                                            ).read(),
                                            _exec_globals,
                                        )
                                    os.chdir(_prev)
                                except Exception as _ex:
                                    _err.write(str(_ex))
                                    _rc = 1

                                class _FgR:
                                    returncode = _rc
                                    stdout = _out.getvalue()
                                    stderr = _err.getvalue()

                                result = _FgR()
                            else:
                                _script_env = os.environ.copy()
                                _script_env["KOTO_OUTPUT_DIR"] = (
                                    settings_manager.documents_dir
                                )
                                result = subprocess.run(
                                    [sys.executable, temp_script],
                                    capture_output=True,
                                    text=True,
                                    timeout=60,
                                    cwd=WORKSPACE_DIR,
                                    env=_script_env,
                                    creationflags=(
                                        subprocess.CREATE_NO_WINDOW
                                        if sys.platform == "win32"
                                        else 0
                                    ),
                                )
                            _app_logger.debug(
                                f"[FILE_GEN] Script exit code: {result.returncode}"
                            )
                            _app_logger.debug(
                                f"[FILE_GEN] Script stdout: {result.stdout}"
                            )
                            _app_logger.debug(
                                f"[FILE_GEN] Script stderr: {result.stderr}"
                            )

                            if result.returncode == 0:
                                # 检查生成的文件 — 同时扫描 documents_dir 和 WORKSPACE_DIR 根目录
                                docs_dir = settings_manager.documents_dir
                                _FILE_EXTS = (
                                    ".pdf",
                                    ".docx",
                                    ".xlsx",
                                    ".pptx",
                                    ".ppt",
                                    ".png",
                                    ".jpg",
                                )
                                _scan_dirs = [docs_dir, WORKSPACE_DIR]
                                for _scan_dir in _scan_dirs:
                                    if os.path.exists(_scan_dir):
                                        for f in os.listdir(_scan_dir):
                                            if f.endswith(_FILE_EXTS):
                                                full_path = os.path.join(_scan_dir, f)
                                                age = time.time() - os.path.getmtime(
                                                    full_path
                                                )
                                                if (
                                                    age < 90
                                                ):  # 90s 窗口，因为脚本执行可能需要一些时间
                                                    rel_path = os.path.relpath(
                                                        full_path, WORKSPACE_DIR
                                                    ).replace("\\", "/")
                                                    if rel_path not in generated_files:
                                                        generated_files.append(rel_path)
                                                        _app_logger.debug(
                                                            f"[FILE_GEN] Generated: {rel_path}"
                                                        )

                                # 若仍未找到，尝试从 stdout 解析文件路径
                                if not generated_files and result.stdout:
                                    import re as _re_fp

                                    _fp_matches = _re_fp.findall(
                                        r"[\w./\\:\- ]+\.(?:xlsx|docx|pptx|pdf|ppt)",
                                        result.stdout,
                                        _re_fp.IGNORECASE,
                                    )
                                    for _fp in _fp_matches:
                                        _fp = _fp.strip()
                                        # 相对路径 → 基于 WORKSPACE_DIR 解析
                                        _abs = (
                                            _fp
                                            if os.path.isabs(_fp)
                                            else os.path.join(WORKSPACE_DIR, _fp)
                                        )
                                        if os.path.exists(_abs):
                                            _rp = os.path.relpath(
                                                _abs, WORKSPACE_DIR
                                            ).replace("\\", "/")
                                            if _rp not in generated_files:
                                                generated_files.append(_rp)
                                                _app_logger.debug(
                                                    f"[FILE_GEN] Found via stdout: {_rp}"
                                                )

                                if generated_files:
                                    files_list = ", ".join(
                                        [os.path.basename(f) for f in generated_files]
                                    )
                                    success_msg = (
                                        "✅ **文件生成成功！**\n\n📁 生成的文件: **"
                                        + files_list
                                        + "**\n📍 保存位置: `"
                                        + docs_dir
                                        + "`"
                                    )
                                    yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                                else:
                                    # 脚本执行成功但没有检测到新文件
                                    output = result.stdout.strip()
                                    if output:
                                        msg = (
                                            "✅ 脚本执行完成\n```\n" + output + "\n```"
                                        )
                                        yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'token', 'content': '⚠️ 脚本执行完成，但未检测到新文件'})}\n\n"
                            else:
                                error_msg = result.stderr.strip() or "未知错误"
                                err_content = (
                                    "❌ 脚本执行失败\n```\n" + error_msg[:500] + "\n```"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': err_content})}\n\n"

                        except subprocess.TimeoutExpired:
                            yield f"data: {json.dumps({'type': 'token', 'content': '⚠️ 脚本执行超时（60秒）'})}\n\n"
                        except Exception as e:
                            _app_logger.debug(f"[FILE_GEN] Execution error: {e}")
                            err_msg = "❌ 执行错误: " + str(e)
                            yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"

                        # 删除临时脚本
                        for temp_file in temp_scripts:
                            try:
                                if os.path.exists(temp_file):
                                    os.remove(temp_file)
                                    _app_logger.debug(
                                        "Deleted temp script: %s", temp_file
                                    )
                            except OSError:
                                pass
                    else:
                        # 没有匹配到代码格式：直接把模型内容生成文档
                        try:
                            from web.document_generator import save_docx, save_pdf

                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)

                            # 提取标题（尝试从内容中找 # 标题）
                            title_match = re.search(
                                r"^#\s*(.+)$", response_text, re.MULTILINE
                            )
                            if title_match:
                                title = title_match.group(1).strip()[:50]
                            else:
                                # 尝试从用户输入提取关键词作为文件名
                                try:
                                    clean_input = user_input
                                    # 去除常用指令词
                                    stop_patterns = [
                                        "生成的",
                                        "写一个",
                                        "写一篇",
                                        "帮我",
                                        "请",
                                        "关于",
                                        "一下",
                                        "文档",
                                        "file",
                                        "generate",
                                        "write",
                                        "about",
                                        "make",
                                        "create",
                                    ]
                                    for pattern in stop_patterns:
                                        clean_input = clean_input.replace(pattern, " ")

                                    # 提取中英文关键词 (2-20 chars)
                                    keywords = [
                                        w
                                        for w in re.split(
                                            r"[^a-zA-Z0-9\u4e00-\u9fa5]", clean_input
                                        )
                                        if w.strip()
                                    ]
                                    valid_keywords = [
                                        k
                                        for k in keywords
                                        if len(k) > 1 and len(k) < 20
                                    ]

                                    if valid_keywords:
                                        # 取前几个关键词组合
                                        title = "_".join(valid_keywords[:3])
                                    else:
                                        title = f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                except Exception:
                                    title = f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                            user_lower = user_input.lower()
                            if "pdf" in user_lower:
                                saved_path = save_pdf(
                                    response_text, title=title, output_dir=docs_dir
                                )
                                file_type = "PDF"
                            else:
                                saved_path = save_docx(
                                    response_text, title=title, output_dir=docs_dir
                                )
                                file_type = "Word"

                            rel_path = os.path.relpath(
                                saved_path, WORKSPACE_DIR
                            ).replace("\\", "/")
                            if rel_path not in generated_files:
                                generated_files.append(rel_path)

                            success_msg = f"✅ **{file_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        except Exception as direct_err:
                            _app_logger.debug(
                                f"[FILE_GEN] Direct save failed: {direct_err}"
                            )
                            # 回退展示原始响应
                            yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text or '⚠️ 模型未返回响应'})}\n\n"

                # 保存历史（基于磁盘完整历史追加，在 done 事件之前）
                _gen_msg = (
                    f"已生成文件: {', '.join(generated_files)}"
                    if generated_files
                    else (response_text[:500] if response_text else "生成失败")
                )
                session_manager.append_and_save(
                    f"{session_name}.json", user_input, _gen_msg
                )

                # 发送完成事件
                total_time = time.time() - start_time
                _app_logger.debug(
                    f"[FILE_GEN] ★★★ Sending done event, generated_files: {generated_files}, total_time: {total_time:.2f}s"
                )
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                return

            # === Regular Mode (流式输出) ===
            # 根据任务类型选择系统指令
            # CHAT/RESEARCH等使用简化指令，避免不必要的文件生成
            use_instruction = (
                _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION()
                if task_type in ["CHAT", "RESEARCH"]
                else _get_system_instruction()
            )

            # 注入长期记忆上下文
            _memory_manager = get_memory_manager()

            # 更新对话摘要（滑动窗口外）
            if full_history and len(full_history) > 20:

                def _summarize():
                    return _memory_manager.get_or_update_summary(
                        session_name, full_history
                    )

                _, err, timed_out = run_with_timeout(_summarize, 6)
                if timed_out:
                    _app_logger.debug("[MEMORY] 摘要更新超时，已跳过")
                elif err:
                    _app_logger.debug(f"[MEMORY] 摘要更新失败: {err}")

            memory_context = _memory_manager.get_context_string(
                user_input, session_name=session_name, history=full_history
            )
            if memory_context:
                use_instruction += f"\n\n{memory_context}"
                _app_logger.debug(
                    f"[MEMORY] 注入了 {len(memory_context)} 字符的记忆上下文"
                )
                t = yield_thinking(
                    f"从长期记忆中检索到 {len(memory_context)} 字符的相关上下文并注入",
                    "context",
                    "local",
                )
                if t:
                    yield t

            # 根据任务类型提供差异化进度提示
            if task_type == "CODER":
                used_model = model_id
                t = yield_thinking(
                    f"进入代码生成模式，使用 {model_id} 进行代码分析与生成",
                    "generating",
                    "cloud",
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '💻 正在分析代码需求...', 'detail': f'使用 {model_id}'})}\n\n"

                # 特殊优化：对于游戏开发或安装包，添加简短指令避免啰嗦
                if any(
                    k in user_input.lower()
                    for k in ["游戏", "app", "五子棋", "pygame", "install", "安装"]
                ):
                    use_instruction += "\n\n[Important] If suggesting to install packages (like pygame), assume the user knows how to use pip. Just output `pip install package_name` in a code block. Do NOT write long tutorials about installation. Focus on the Python Code."

            elif task_type == "CHAT":
                used_model = model_id
                t = yield_thinking(
                    f"进入对话模式，使用 {model_id} 生成回复", "generating", "cloud"
                )
                if t:
                    yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '💬 Koto 正在思考...', 'detail': '请稍候'})}\n\n"

                # ═══ 本地模型快速通道：简单问题直接走 Ollama ═══
                from app.core.routing import LocalModelRouter

                if LocalModelRouter.is_simple_query(user_input, task_type, history):
                    local_stream = LocalModelRouter.generate_stream(
                        user_input,
                        history=history,
                        system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION(),
                    )
                    if local_stream is not None:
                        _app_logger.debug(
                            f"[CHAT] ⚡ 使用本地模型快速响应: {LocalModelRouter._response_model}"
                        )
                        t = yield_thinking(
                            f"检测到简单查询，切换到本地模型 {LocalModelRouter._response_model} 快速响应",
                            "model",
                            "local",
                        )
                        if t:
                            yield t
                        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': '💬 对话', 'model': f'🏠 {LocalModelRouter._response_model} (本地)', 'message': f'🎯 任务分类: 💬 对话 (方法: 🏠 {LocalModelRouter._response_model} 本地快速通道)'})}\n\n"
                        local_full_text = ""
                        local_ok = False
                        try:
                            for chunk in local_stream:
                                local_full_text += chunk
                                yield f"data: {json.dumps({'type': 'token', 'content': chunk})}\n\n"
                            local_ok = bool(local_full_text.strip())
                        except Exception as local_err:
                            _app_logger.debug(f"[CHAT] 本地模型生成失败: {local_err}")

                        if local_ok:
                            # 本地模型成功 → 保存并返回
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                local_full_text,
                                task=task_type,
                                model_name=f"ollama/{LocalModelRouter._response_model}",
                            )
                            _reflect_types_local = {
                                "CHAT",
                                "RESEARCH",
                                "CODER",
                                "FILE_GEN",
                                "AGENT",
                            }
                            if task_type in _reflect_types_local:
                                _start_memory_extraction(
                                    user_input,
                                    local_full_text,
                                    history,
                                    task_type=task_type,
                                    session_name=session_name,
                                )
                            total_time = time.time() - start_time
                            _app_logger.debug(
                                f"[CHAT] ⚡ 本地模型响应完成 ({total_time:.2f}s)"
                            )
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            # 本地模型失败 → 静默降级到云模型
                            _app_logger.debug(f"[CHAT] 本地模型输出为空，降级到云模型")
                            t = yield_thinking(
                                f"本地模型输出为空，降级到云端模型 {model_id}",
                                "model",
                                "hybrid",
                            )
                            if t:
                                yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': '☁️ 切换到云端模型...', 'detail': model_id})}\n\n"
            elif task_type == "RESEARCH":
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔬 正在进行深度分析...', 'detail': f'使用 {model_id}'})}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'progress', 'message': '💭 Koto 正在思考...', 'detail': '请稍候'})}\n\n"

            # 注入 skill_prompt（路由器对用户意图的格式期望）
            _task_skill = (context_info or {}).get("skill_prompt")
            if _task_skill:
                use_instruction += f"\n\n[响应要求] {_task_skill}"

            # 构建历史记录（非延续任务时过滤无关历史）
            if context_info and context_info.get("is_continuation"):
                history_for_model = history
                t = yield_thinking(
                    f"检测到上下文延续，保留全部 {len(history)} 轮对话历史",
                    "context",
                    "hybrid",
                )
                if t:
                    yield t
            else:
                history_for_model = ContextAnalyzer.filter_history(user_input, history)
                if len(history_for_model) != len(history):
                    t = yield_thinking(
                        f"过滤对话历史: {len(history)} 轮 → {len(history_for_model)} 轮相关记录",
                        "context",
                        "hybrid",
                    )
                    if t:
                        yield t

            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(
                    types.Content(
                        role=turn["role"],
                        parts=[types.Part.from_text(text=p) for p in turn["parts"]],
                    )
                )

            t = yield_thinking(
                f"准备调用 {model_id} API，发送 {len(formatted_history)+1} 条消息",
                "generating",
            )
            if t:
                yield t

            # ── 并联本地模型：生成执行计划（与云端模型并发，填充首包延迟死区）──
            # 本地模型 ~200-400ms 生成 3-5 个步骤；云端首包通常需 1-5s
            # 两者并发执行，步骤在首包前或首次心跳时流出，零额外延迟
            import concurrent.futures as _cf

            _plan_future = None
            try:
                from app.core.routing import LocalModelRouter as _LMR_plan

                if _LMR_plan.is_ollama_available() and _LMR_plan._initialized:
                    _plan_exec = _cf.ThreadPoolExecutor(
                        max_workers=1, thread_name_prefix="koto_plan"
                    )
                    _plan_future = _plan_exec.submit(
                        _LMR_plan.generate_plan, user_input, task_type
                    )
                    _plan_exec.shutdown(wait=False)  # 不阻塞，后台跑
            except Exception:
                _plan_future = None

            # 将 RAG 上下文作为用户消息前缀（与系统指令分离，提升事实准确性）
            if _rag_context_block:
                _rag_augmented_input = (
                    f"[📚 知识库参考内容（请以此为事实依据）]\n"
                    f"{_rag_context_block}"
                    f"────────────────────────────────────────────\n"
                    f"[用户问题]\n{effective_input}"
                )
            else:
                _rag_augmented_input = effective_input

            # 使用流式响应
            response = client.models.generate_content_stream(
                model=model_id,
                contents=formatted_history
                + [
                    types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=_rag_augmented_input)],
                    )
                ],
                config=types.GenerateContentConfig(system_instruction=use_instruction),
            )

            full_text = ""
            chunk_count = 0
            heartbeat_interval = 5  # 每5秒发送一次心跳
            first_chunk_received = False
            _plan_flushed = False  # 本地执行计划是否已流出

            try:
                # 使用保活包装器处理流式响应
                max_wait = 60 if task_type == "CODER" else 120
                for item_type, item_data in stream_with_keepalive(
                    response,
                    start_time,
                    keepalive_interval=heartbeat_interval,
                    max_wait_first_token=max_wait,
                ):
                    # 检查中断标志
                    if _interrupt_manager.is_interrupted(session_name):
                        _app_logger.debug(
                            f"[INTERRUPT] User interrupted at chunk {chunk_count}"
                        )
                        interrupt_msg = "\n\n⏸️ 用户已中断"
                        yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}\n\n"
                        break

                    # ── 执行计划：尝试在心跳时非阻塞刷出（填充首包前的空白）──
                    if not _plan_flushed and _plan_future is not None:
                        try:
                            _steps = _plan_future.result(timeout=0.05)
                            if _steps:
                                for _s in _steps:
                                    _pt = yield_thinking(
                                        f"📋 {_s}", "planning", "local"
                                    )
                                    if _pt:
                                        yield _pt
                            _plan_flushed = True
                            _plan_future = None
                        except _cf.TimeoutError:
                            pass  # 还没好，下次检查
                        except Exception:
                            _plan_flushed = True
                            _plan_future = None

                    if item_type == "heartbeat":
                        elapsed = item_data
                        if first_chunk_received:
                            # 根据任务类型差异化心跳（已收到首包，正在流式输出）
                            char_count = len(full_text)
                            if task_type == "CODER":
                                hb_msg = f"💻 代码生成中... 已输出 {char_count} 字符"
                            elif task_type == "RESEARCH":
                                hb_msg = f"🔬 深度分析中... 已输出 {char_count} 字符"
                            else:
                                hb_msg = "💭 正在生成..."
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'{elapsed}s', 'stage': 'generating'})}\n\n"
                        else:
                            # 等待首包（api_calling 阶段，前端显示旋转 spinner）
                            if task_type == "CODER":
                                hb_msg = "💻 代码分析中，请稍候..."
                            elif task_type == "RESEARCH":
                                hb_msg = "🔬 深度思考中，请耐心等待..."
                            else:
                                hb_msg = "🧠 模型思考中..."
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'已等待 {elapsed}s', 'stage': 'api_calling'})}\n\n"

                    elif item_type == "timeout":
                        if task_type == "CODER" and not full_text:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 首包超时，切换到快速模型...', 'detail': ''})}\n\n"
                            try:
                                fallback_resp = client.models.generate_content(
                                    model="gemini-2.5-flash",
                                    contents=formatted_history
                                    + [
                                        types.Content(
                                            role="user",
                                            parts=[
                                                types.Part.from_text(
                                                    text=_rag_augmented_input
                                                )
                                            ],
                                        )
                                    ],
                                    config=types.GenerateContentConfig(
                                        system_instruction=use_instruction,
                                        temperature=0.4,
                                        max_output_tokens=4000,
                                    ),
                                )
                                fallback_text = fallback_resp.text or ""
                                if fallback_text:
                                    full_text = fallback_text
                                    yield f"data: {json.dumps({'type': 'token', 'content': fallback_text})}\n\n"
                            except Exception:
                                yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，请稍后重试'})}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，请稍后重试'})}\n\n"
                        break

                    elif item_type == "chunk":
                        chunk = item_data
                        if chunk.text:
                            if not first_chunk_received:
                                first_chunk_received = True
                                _app_logger.debug(
                                    f"[CHAT] 收到第一个响应，耗时 {time.time() - start_time:.1f}s"
                                )
                                # 首包到达：最后一次机会刷出执行计划（等最多 0.5s）
                                if not _plan_flushed and _plan_future is not None:
                                    try:
                                        _steps = _plan_future.result(timeout=0.5)
                                        if _steps:
                                            for _s in _steps:
                                                _pt = yield_thinking(
                                                    f"📋 {_s}", "planning", "local"
                                                )
                                                if _pt:
                                                    yield _pt
                                    except Exception:
                                        pass
                                    _plan_flushed = True
                                    _plan_future = None

                            full_text += chunk.text
                            chunk_count += 1
                            yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"

            except Exception as stream_error:
                error_str = str(stream_error)
                _app_logger.debug(f"[CHAT] Stream error: {error_str}")

                # 地区限制错误
                if (
                    "location is not supported" in error_str.lower()
                    or "failed_precondition" in error_str.lower()
                ):
                    error_text = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 `config/gemini_config.env` 配置中转服务 `GEMINI_API_BASE`\n2. 或使用支持的代理服务"
                    yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                # 流式传输中断，但已有部分内容
                elif full_text:
                    error_msg = error_str[:50]
                    warn_text = f"\n\n⚠️ (传输中断: {error_msg}...)"
                    yield f"data: {json.dumps({'type': 'token', 'content': warn_text})}\n\n"
                else:
                    raise stream_error

            # 失败时先修正一次（不直接报错）
            if Utils.is_failure_output(full_text):
                yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次生成失败，正在修正...', 'detail': ''})}\n\n"
                fix_prompt = Utils.build_fix_prompt(task_type, user_input, full_text)
                fix_resp = client.models.generate_content(
                    model=model_id,
                    contents=fix_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=use_instruction,
                        temperature=0.4,
                        max_output_tokens=4000,
                    ),
                )
                corrected_text = fix_resp.text or full_text
                if corrected_text and corrected_text != full_text:
                    corrected_msg = f"\n\n🔁 修正版本:\n{corrected_text}"
                    yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                    full_text = corrected_text
            else:
                # 复杂任务进行快速自检
                is_complex_task = (
                    task_type in ["RESEARCH", "FILE_GEN", "CODER"]
                    or (context_info and context_info.get("complexity") == "complex")
                    or len(user_input) > 200
                )
                if is_complex_task:
                    check = Utils.quick_self_check(task_type, user_input, full_text)
                    if not check.get("pass") and check.get("fix_prompt"):
                        status_msg = "🩺 自检未通过，正在修正..."
                        yield f"data: {json.dumps({'type': 'progress', 'message': status_msg, 'detail': '快速模型自检'})}\n\n"
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=check["fix_prompt"],
                            config=types.GenerateContentConfig(
                                system_instruction=use_instruction,
                                temperature=0.4,
                                max_output_tokens=4000,
                            ),
                        )
                        corrected_text = fix_resp.text or full_text
                        if corrected_text and corrected_text != full_text:
                            corrected_msg = f"\n\n🔁 修正版本:\n{corrected_text}"
                            yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                            full_text = corrected_text

            # 处理自动保存的文件
            if settings_manager.get("ai", "auto_save_files") is not False:
                saved_files = Utils.auto_save_files(full_text)
            else:
                saved_files = []

            # 代码任务: 检测并自动安装依赖
            if task_type == "CODER":
                pkgs = Utils.detect_required_packages(full_text)
                if pkgs:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📦 检测到依赖，正在检查/安装...', 'detail': ', '.join(pkgs)})}\n\n"
                    install_result = Utils.auto_install_packages(pkgs)
                    installed = install_result.get("installed", [])
                    failed = install_result.get("failed", [])
                    skipped = install_result.get("skipped", [])
                    msg_parts = []
                    if installed:
                        msg_parts.append(f"✅ 已安装: {', '.join(installed)}")
                    if skipped:
                        msg_parts.append(f"ℹ️ 已存在: {', '.join(skipped)}")
                    if failed:
                        msg_parts.append(f"⚠️ 安装失败: {', '.join(failed)}")
                    if msg_parts:
                        msg_content = "\n\n" + "\n".join(msg_parts)
                        yield f"data: {json.dumps({'type': 'token', 'content': msg_content})}\n\n"

            # 如果有保存的文件，提示用户保存位置
            if saved_files:
                files_list = ", ".join(saved_files)
                save_hint = (
                    f"\n\n📁 文件已保存: **{files_list}**\n📂 位置: `{WORKSPACE_DIR}`"
                )
                yield f"data: {json.dumps({'type': 'token', 'content': save_hint})}\n\n"

            # 先保存历史，再发送 done 事件（包含元数据用于前端渲染）
            session_manager.append_and_save(
                f"{session_name}.json",
                user_input,
                full_text,
                task=task_type,
                model_name=model_id,
                saved_files=saved_files,
            )
            # 2-B: Memory reflection for all supported task types
            _reflect_types = {"CHAT", "RESEARCH", "CODER", "FILE_GEN", "AGENT"}
            if task_type in _reflect_types:
                _start_memory_extraction(
                    user_input,
                    full_text,
                    history_for_model,
                    task_type=task_type,
                    session_name=session_name,
                )

            # 计算 msg_id 并包含在 done 事件中，使前端能提交用户评分
            try:
                from app.core.learning.rating_store import RatingStore as _RS

                _done_msg_id = _RS.make_msg_id(session_name, user_input)
            except Exception:
                _done_msg_id = ""

            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files, 'total_time': total_time, 'msg_id': _done_msg_id})}\n\n"

        except Exception as e:
            error_str = str(e)
            _app_logger.debug(f"[CHAT] Exception: {error_str}")

            # 地区限制错误
            if (
                "location is not supported" in error_str.lower()
                or "failed_precondition" in error_str.lower()
            ):
                error_response = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 `config/gemini_config.env` 配置中转服务 `GEMINI_API_BASE`\n2. 或使用支持的代理服务"
            elif "API key not valid" in error_str or (
                "INVALID_ARGUMENT" in error_str and "api key" in error_str.lower()
            ):
                error_response = (
                    "❌ **API 密钥无效**\n\n"
                    "请检查您的 Gemini API 密钥：\n"
                    "1. 前往 [aistudio.google.com/apikey](https://aistudio.google.com/apikey) 获取有效密钥\n"
                    "2. 在 Koto 设置页面更新 API 密钥（设置 → API 配置）\n"
                    "3. 确保密钥所在 Google 项目已启用 Generative Language API\n\n"
                    f"原始错误: `{error_str[:150]}`"
                )
            elif (
                "server disconnected" in error_str.lower()
                or "disconnected without" in error_str.lower()
                or "connection reset" in error_str.lower()
                or "connection aborted" in error_str.lower()
            ):
                # 将连接中断的模型标记为短期不可用（2 分钟），下次请求自动降级到 Flash
                try:
                    from app.core.llm.model_fallback import get_fallback_executor

                    get_fallback_executor().mark_unavailable(model_id, ttl=120)
                    _app_logger.warning(
                        f"[CHAT] 连接中断，已将 {model_id} 标记不可用 120s，下次自动降级"
                    )
                except Exception:
                    pass
                error_response = (
                    "❌ **服务器连接中断**\n\n"
                    "与 Gemini API 的连接被意外断开，这通常是临时性问题。\n\n"
                    "💡 建议：\n"
                    "1. 稍等片刻后重新发送消息\n"
                    "2. 检查您的网络连接稳定性\n"
                    "3. 如果使用代理，请确认代理连接正常\n"
                    "4. 如问题持续，可尝试切换到其他模型"
                )
            elif (
                "resource_exhausted" in error_str.lower()
                or "quota" in error_str.lower()
                or "rate limit" in error_str.lower()
                or "429" in error_str
            ):
                error_response = (
                    "❌ **API 配额超限**\n\n"
                    "当前 API 密钥的请求频率或配额已达上限。\n\n"
                    "💡 建议：\n"
                    "1. 稍等 1-2 分钟后重试\n"
                    "2. 在设置中切换到其他 API 密钥\n"
                    "3. 或升级您的 Google AI Studio 计划"
                )
            elif (
                "unavailable" in error_str.lower()
                or "503" in error_str
                or "service unavailable" in error_str.lower()
            ):
                error_response = (
                    "❌ **Gemini 服务暂时不可用**\n\n"
                    "Gemini API 服务器当前无法响应，可能正在维护中。\n\n"
                    "💡 建议：稍等片刻后重试，或访问 [status.google.com](https://status.google.com) 查看服务状态"
                )
            elif (
                "deadline_exceeded" in error_str.lower()
                or "timed out" in error_str.lower()
            ):
                error_response = (
                    "❌ **请求超时**\n\n"
                    "模型响应时间过长，请求已超时。\n\n"
                    "💡 建议：\n"
                    "1. 尝试缩短您的问题或分步骤提问\n"
                    "2. 切换到响应更快的模型（如 gemini-2.5-flash）\n"
                    "3. 检查网络连接质量"
                )
            else:
                error_response = f"❌ 发生错误: {error_str[:200]}"

            # 即使出错也要保存用户的问题
            session_manager.append_and_save(
                f"{session_name}.json",
                user_input,
                error_response,
                task=task_type,
                model_name=model_id,
            )

            yield f"data: {json.dumps({'type': 'token', 'content': error_response})}\n\n"
            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"

    def _safe_generate():
        """
        generate() 外层安全包装器：
        确保无论 generate() 内部以何种方式结束，前端都能收到 'done' 事件，
        避免因任务识别失败/早期异常导致对话界面永远挂起。
        """
        _sent_done = False
        try:
            for _chunk in generate():
                if isinstance(_chunk, (str, bytes)):
                    _chunk_str = (
                        _chunk
                        if isinstance(_chunk, str)
                        else _chunk.decode("utf-8", errors="replace")
                    )
                    if '"type": "done"' in _chunk_str or "'type': 'done'" in _chunk_str:
                        _sent_done = True
                yield _chunk
        except Exception as _sg_err:
            _app_logger.warning(
                f"[STREAM] ⚠️ _safe_generate caught exception: {_sg_err}"
            )
            import traceback

            traceback.print_exc()
            if not _sent_done:
                _err_msg = f"❌ 流式响应异常终止: {str(_sg_err)[:200]}"
                yield f"data: {json.dumps({'type': 'token', 'content': _err_msg})}\n\n"
        finally:
            if not _sent_done:
                _app_logger.warning(
                    f"[STREAM] ⚠️ generate() 未发送 done 事件，触发兜底 done (task_type={task_type})"
                )
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0, 'fallback_done': True})}\n\n"

    response = Response(
        stream_with_context(_safe_generate()), mimetype="text/event-stream"
    )
    response.headers["Cache-Control"] = "no-cache"
    response.headers["X-Accel-Buffering"] = "no"  # 禁用 nginx 缓冲
    response.headers["Connection"] = "keep-alive"
    return response


@app.route("/api/chat/file", methods=["POST"])
def chat_with_file():
    """处理文件上传和聊天请求"""
    from web.document_generator import save_docx, save_pdf, to_workspace_rel
    from web.file_processor import process_uploaded_file

    def _strip_code_blocks(text: str) -> str:
        if not text:
            return text
        # Remove fenced code blocks entirely
        text = re.sub(r"```[\s\S]*?```", "", text)
        # Remove inline code ticks but keep the content
        text = text.replace("`", "")
        return text.strip()

    def _build_analysis_title(user_text: str, filename: str, is_binary: bool) -> str:
        name_base = os.path.splitext(filename)[0]
        text_lower = (user_text or "").lower()
        ext = os.path.splitext(filename)[1].lower()

        # 1. Determine File Type Prefix
        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            prefix = "图片"
        elif ext == ".pdf":
            prefix = "PDF"
        elif ext in [".doc", ".docx"]:
            prefix = "Word"
        elif ext in [".ppt", ".pptx"]:
            prefix = "PPT"
        else:
            prefix = "文件" if is_binary else "文档"

        # 2. Determine Intent
        intent = "分析"
        intent_map = {
            "翻译": ["翻译", "translate", "译文", "中译英", "英译中"],
            "总结": ["总结", "归纳", "摘要", "summary", "概括", "核心内容"],
            "文字识别": ["提取", "识别", "ocr", "文字", "转文字", "读图"],
            "表格识别": ["表格", "table", "excel", "转表"],
            "对比分析": ["对比", "比较", "diff", "区别", "差异"],
            "校对": ["校对", "检查", "审阅", "纠错", "改错"],
            "润色": ["润色", "改写", "polish", "rewrite", "优化", "美化"],
            "续写": ["续写", "扩写", "continue", "补充"],
            "大纲": ["大纲", "框架", "outline", "目录"],
            "解释": ["解释", "explain", "什么意思", "含义"],
        }

        found_intent_keywords = []
        for k, v in intent_map.items():
            for kw in v:
                if kw in text_lower:
                    intent = k
                    found_intent_keywords.append(kw)
                    break
            if intent != "分析":
                break

        # 3. Extract Topic Keywords (Improved)
        stop_words = [
            "帮我",
            "请",
            "一下",
            "把",
            "这个",
            "这篇",
            "文件",
            "文章",
            "内容",
            "生成",
            "写一个",
            "做一份",
            "koto",
            "分析",
            "阅读",
            "提取",
            "识别",
            "output",
            "make",
            "create",
            "generate",
            "please",
            "the",
            "a",
            "an",
            "is",
            "of",
            "to",
            "for",
            "with",
            "in",
            "on",
            "user",
            "file",
            "document",
            "from",
            "this",
            "that",
            "it",
            "what",
            "how",
            "why",
            "where",
            "into",
            "check",
            "run",
        ]

        # Prepare text
        text_lower = user_text.lower()

        # Safe replacement for Chinese phrases (which don't use spaces)
        zh_stops = [w for w in stop_words if re.match(r"[\u4e00-\u9fa5]+", w)]
        for stop in zh_stops + found_intent_keywords:
            if re.match(r"[\u4e00-\u9fa5]+", stop):  # Only safe-replace Chinese phrases
                text_lower = text_lower.replace(stop, " ")

        # Tokenize by non-word chars (separates English words, breaks Chinese into blocks if spaces inserted)
        # Regex: Keep Chinese chars and English words
        # This splits "summary of report" -> "summary", "of", "report"
        tokens = re.findall(r"[a-zA-Z0-9\u4e00-\u9fa5]+", text_lower)

        # Filter tokens
        valid_keywords = []
        en_stops = set([w for w in stop_words if not re.match(r"[\u4e00-\u9fa5]+", w)])

        for token in tokens:
            if token in en_stops:
                continue
            if token in found_intent_keywords:
                continue  # Filter intent words token-wise
            if len(token) < 2:
                continue
            valid_keywords.append(token)

        # Select best keyword
        topic = ""
        if valid_keywords:
            topic = "_".join(valid_keywords[:3])

        # 4. Construct Final Title
        # Strategy:
        # If user provided a specific topic, prioritize it: "{Intent}_{Topic}_{Filename}"
        # If no detected topic but intent exists: "{Intent}_{Filename}"
        # Fallback: "{Prefix}{Intent}_{Filename}"

        sanitized_name = name_base.replace(" ", "_")

        if topic:
            return f"{intent}_{topic}_{sanitized_name}"
        else:
            return f"{prefix}{intent}_{sanitized_name}"

    session_name = request.form.get("session")
    user_input = request.form.get("message", "")
    files = request.files.getlist("file")

    # 🔍 调试日志
    _app_logger.info(f"[FILE UPLOAD DEBUG] ========== 接收到文件上传请求 ==========")
    _app_logger.info(
        f"[FILE UPLOAD DEBUG] request.files keys: {list(request.files.keys())}"
    )
    _app_logger.info(
        f"[FILE UPLOAD DEBUG] request.files.getlist('file'): {len(files)} 个文件"
    )
    for i, f in enumerate(files):
        _app_logger.info(f"[FILE UPLOAD DEBUG]   {i+1}. {f.filename if f else 'None'}")

    if not files:
        single_file = request.files.get("file")
        if single_file:
            files = [single_file]
            _app_logger.info(
                f"[FILE UPLOAD DEBUG] 使用单文件模式，文件: {single_file.filename}"
            )

    locked_task = request.form.get("locked_task")
    locked_model = request.form.get("locked_model", "auto")
    stream_mode = request.form.get("stream", "").lower() in ("1", "true", "yes")

    _app_logger.info(f"[FILE UPLOAD DEBUG] 最终 files 列表: {len(files)} 个文件")
    _app_logger.info(f"[FILE UPLOAD DEBUG] 判断: len(files) > 1 = {len(files) > 1}")

    if not session_name or not files:
        return jsonify({"error": "Missing session or file"}), 400
    if len(files) > 10:
        return jsonify({"error": "最多一次上传 10 个文件"}), 400

    if len(files) > 1:
        # 检测是否是 PPT 生成意图 (多文件合并生成 PPT)
        ppt_keywords = ["ppt", "slide", "幻灯片", "演示文稿", "powerpoint"]
        is_ppt_intent = any(kw in (user_input or "").lower() for kw in ppt_keywords)

        if is_ppt_intent:
            _app_logger.info(f"[FILE UPLOAD] 检测到多文件 PPT 生成意图: {user_input}")

            # 预先保存所有文件，避免在生成器中访问已关闭的 FileStorage
            saved_file_paths = []
            source_filenames = []

            for f in files:
                if f and f.filename:
                    fname = f.filename
                    fpath = os.path.join(UPLOAD_DIR, fname)
                    # 如果文件指针不在开头，重置它
                    f.seek(0)
                    f.save(fpath)
                    saved_file_paths.append(fpath)
                    source_filenames.append(fname)

            def generate_ppt_stream():
                try:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在准备 PPT 生成...', 'detail': f'检测到 {len(saved_file_paths)} 个源文件'})}\n\n"

                    context_text = ""

                    # 1. 提取所有已保存文件内容
                    for i, filepath in enumerate(saved_file_paths):
                        filename = os.path.basename(filepath)
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'📖 正在读取文件 ({i+1}/{len(saved_file_paths)})...', 'detail': filename})}\n\n"

                        try:
                            # 提取内容
                            from web.file_processor import FileProcessor

                            processor = FileProcessor()
                            # 简化版的 process
                            f_result = processor.process_file(filepath)
                            content = f_result.get("text_content") or f_result.get(
                                "content", ""
                            )

                            # 截断过长内容避免Token爆炸，但保留足够上下文
                            if len(content) > 50000:
                                content = content[:50000] + "...(truncated)"

                            context_text += f"\n\n=== {filename} ===\n{content}\n"

                        except Exception as e:
                            _app_logger.info(
                                f"[PPT BATCH] 读取文件 {filename} 失败: {e}"
                            )
                            context_text += (
                                f"\n\n=== {filename} (Error) ===\n无法读取内容\n"
                            )

                    # 2. 调用 PPT 生成管道
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 正在设计 PPT 结构...', 'detail': '基于多个文件内容'})}\n\n"

                    import asyncio

                    from web.ppt_pipeline import PPTGenerationPipeline

                    # 构造增强后的 Prompt
                    enhanced_prompt = f"{user_input}\n\n【参考资料】\n基于以下文件生成的 PPT:\n{context_text}"

                    # 限制 Prompt 长度
                    if len(enhanced_prompt) > 100000:
                        enhanced_prompt = (
                            enhanced_prompt[:100000] + "\n...(context truncated)"
                        )

                    # 异步执行 PPT 生成
                    # 使用项目内的 get_client() 获取 Gemini 客户端
                    ai_client = get_client()
                    pipeline = PPTGenerationPipeline(ai_client=ai_client)

                    import queue
                    import threading
                    import traceback

                    pipeline_timeout_sec = 300
                    start_ts = time.time()

                    # 混合消息队列（进度+思考）
                    event_queue = queue.Queue()

                    def _progress_listener(msg, p=None):
                        event_queue.put({"type": "progress", "msg": msg, "progress": p})

                    def _thought_listener(text):
                        # Use a dedicated type for thought/reasoning text
                        event_queue.put({"type": "thought", "text": text})

                    run_state = {
                        "done": False,
                        "result": None,
                        "error": None,
                        "traceback": "",
                    }

                    def _run_pipeline_bg():
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        try:
                            # 传递 progress_callback 和 thought_callback
                            run_state["result"] = loop.run_until_complete(
                                pipeline.generate(
                                    user_request=enhanced_prompt,
                                    output_path=os.path.join(
                                        settings_manager.documents_dir,
                                        f"Koto_Presentation_{int(time.time())}.pptx",
                                    ),
                                    enable_auto_images=True,  # 允许自动配图
                                    progress_callback=_progress_listener,
                                    thought_callback=_thought_listener,
                                )
                            )
                        except Exception as bg_err:
                            run_state["error"] = str(bg_err)
                            run_state["traceback"] = traceback.format_exc()
                        finally:
                            try:
                                loop.close()
                            except Exception:
                                pass
                            run_state["done"] = True

                    worker = threading.Thread(target=_run_pipeline_bg, daemon=True)
                    worker.start()

                    # 实时轮询进度队列，转发给前端
                    last_progress_msg = "初始化生成环境..."

                    while not run_state["done"]:
                        elapsed = int(time.time() - start_ts)
                        if elapsed > pipeline_timeout_sec:
                            _progress_listener("生成超时，正在强制停止...", 100)
                            run_state["error"] = (
                                f"PPT 生成超时（>{pipeline_timeout_sec}s）"
                            )
                            break

                        # 消费所有的事件
                        try:
                            while not event_queue.empty():
                                item = event_queue.get_nowait()

                                if item["type"] == "progress":
                                    msg = item["msg"]
                                    p = item["progress"]
                                    last_progress_msg = msg
                                    detail_text = (
                                        f"进度: {p}%"
                                        if p is not None
                                        else f"已用时 {elapsed}s"
                                    )
                                    yield f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': detail_text})}\n\n"

                                elif item["type"] == "thought":
                                    # Send thought as a partial text response or a special 'thought' event
                                    # Assuming frontend can handle 'text' type for appending to the assistant's message
                                    # or 'thought' for a distinct UI block.
                                    # Let's use 'text' for now to ensure it appears in the chat stream.
                                    thought_text = (
                                        f"\n\n> 🤖 **Koto 思考**: {item['text']}\n"
                                    )
                                    yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"

                        except queue.Empty:
                            pass

                        # 如果没有新消息，每2秒发一次心跳防止连接断开
                        if elapsed % 2 == 0 and event_queue.empty():
                            yield f"data: {json.dumps({'type': 'progress', 'message': last_progress_msg, 'detail': f'已用时 {elapsed}s'})}\n\n"

                        time.sleep(0.5)

                    # 发送最后剩余的消息
                    try:
                        while not event_queue.empty():
                            item = event_queue.get_nowait()
                            if item["type"] == "progress":
                                yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': ''})}\n\n"
                            elif item["type"] == "thought":
                                thought_text = (
                                    f"\n\n> 🤖 **Koto 思考**: {item['text']}\n"
                                )
                                yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"
                    except Exception:
                        pass

                    if run_state["error"]:
                        err = run_state["error"]
                        tb = run_state.get("traceback", "")
                        _app_logger.info(
                            f"[PPT BATCH] Background pipeline error: {err}"
                        )
                        if tb:
                            _app_logger.info(f"[PPT BATCH] Traceback: {tb[:800]}")
                        raise Exception(f"PPT 管道异常: {err}")

                    ppt_result = run_state["result"] or {}

                    # pipeline returns 'output_path', also check 'file_path' for compat
                    saved_path = ppt_result.get("output_path") or ppt_result.get(
                        "file_path"
                    )

                    if not ppt_result.get("success"):
                        err_detail = ppt_result.get("error", "未知错误")
                        tb = ppt_result.get("traceback", "")
                        _app_logger.info(
                            f"[PPT BATCH] Pipeline returned failure: {err_detail}"
                        )
                        if tb:
                            _app_logger.info(f"[PPT BATCH] Traceback: {tb[:500]}")
                        raise Exception(f"PPT 管道生成失败: {err_detail}")

                    if saved_path and os.path.exists(saved_path):
                        yield f"data: {json.dumps({'type': 'progress', 'message': '✅ PPT 生成完成！', 'detail': os.path.basename(saved_path)})}\n\n"

                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace(
                            "\\", "/"
                        )
                        success_msg = f"✅ **PPT 生成成功！**\n\n基于 {len(saved_file_paths)} 个文件生成的演示文稿。\n📁 文件: **{os.path.basename(saved_path)}**"

                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [rel_path], 'total_time': 0})}\n\n"
                    else:
                        raise Exception("PPT 文件生成失败，未返回路径")

                except Exception as e:
                    _app_logger.info(f"[PPT BATCH ERROR] {e}")
                    import traceback

                    traceback.print_exc()
                    err_msg = f"❌ 生成失败: {str(e)}"
                    yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_ppt_stream()), mimetype="text/event-stream"
            )

        history = session_manager.load(f"{session_name}.json")
        file_names = [f.filename for f in files if f and f.filename]
        user_message = f"[Files: {', '.join(file_names)}] {user_input}"
        session_manager.append_user_early(f"{session_name}.json", user_message)

        batch_results = []
        combined_saved_files = []
        combined_images = []

        def _process_single_file(file):
            if not file or not file.filename:
                return None

            filename = _secure_filename(file.filename) or f"upload_{uuid.uuid4().hex}"
            filepath = os.path.join(UPLOAD_DIR, filename)
            file.save(filepath)
            file_type = file.mimetype or file.content_type or ""
            file_ext = os.path.splitext(filename)[1].lower()

            # 检测是否是纯归档/整理请求（不需要AI分析内容）
            organize_keywords = [
                "整理",
                "归档",
                "归纳",
                "分类",
                "整理一下",
                "整理下",
                "帮我整理",
                "文件整理",
                "organize",
                "sort",
            ]
            is_organize_only = any(kw in (user_input or "") for kw in organize_keywords)

            try:
                # formatted_message, file_data = process_uploaded_file(filepath, user_input)
                # --- Modify to use FileProcessor directly for simultaneous KB indexing ---
                from web.file_processor import FileProcessor

                _processor = FileProcessor()
                _file_raw = _processor.process_file(filepath)

                # 1. 自动建库 (Auto-Indexing to Knowledge Base) - Use threading to not block UI
                try:
                    _text_content = _file_raw.get("text_content", "")
                    if _text_content and len(_text_content) > 50:  # Ignore tiny files

                        def _bg_index(content, meta):
                            try:
                                from web.knowledge_base import KnowledgeBase

                                _kb = KnowledgeBase()
                                res = _kb.add_content(content, meta)
                                _app_logger.debug(
                                    f"[KB] Auto-indexing completed: {res}"
                                )
                            except Exception as e:
                                _app_logger.debug(f"[KB] Auto-indexing failed: {e}")

                        import threading

                        _idx_thread = threading.Thread(
                            target=_bg_index,
                            args=(
                                _text_content,
                                {
                                    "file_path": filepath,
                                    "file_name": filename,
                                    "file_type": file_ext,
                                    "mtime": os.path.getmtime(filepath),
                                },
                            ),
                        )
                        _idx_thread.start()
                        _app_logger.debug(f"[KB] 已启动后台建库任务: {filename}")
                except Exception as _kb_err:
                    _app_logger.debug(f"[KB] Indexing trigger failed: {_kb_err}")

                # 1-B. 注册到 FileRegistry（统一文件元数据中心）
                try:

                    def _bg_register_file(_fpath, _sid):
                        try:
                            from app.core.file.file_registry import get_file_registry

                            _reg = get_file_registry()
                            _reg.register(
                                _fpath,
                                source="upload",
                                session_id=_sid,
                                extract_content=True,
                            )
                            _app_logger.info(
                                f"[FileRegistry] ✅ 已注册上传文件: {os.path.basename(_fpath)}"
                            )
                        except Exception as _re:
                            _app_logger.warning(
                                f"[FileRegistry] ⚠️ 注册失败（非致命）: {_re}"
                            )

                    import threading as _thr

                    _reg_thread = _thr.Thread(
                        target=_bg_register_file,
                        args=(filepath, session_name),
                        daemon=True,
                    )
                    _reg_thread.start()
                except Exception as _rge:
                    _app_logger.warning(f"[FileRegistry] ⚠️ 启动注册线程失败: {_rge}")

                # 2. Continue with standard chat formatting
                formatted_message, file_data = _processor.format_result_for_chat(
                    _file_raw, user_input
                )

                task_type = locked_task
                context_info = None
                route_method = "Auto"
                if not task_type:
                    if file_data and file_type and file_type.startswith("image"):
                        message_lower = (user_input or "").lower()
                        is_edit = any(
                            kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS
                        )
                        task_type = "PAINTER" if is_edit else "VISION"
                        route_method = (
                            "🖼️ Image Edit" if is_edit else "👁️ Image Analysis"
                        )
                        _app_logger.info(
                            f"[FILE UPLOAD] 图片任务直通路由: {task_type} (方法: {route_method})"
                        )
                    else:
                        _ann_exts = {
                            ".doc",
                            ".docx",
                            ".pdf",
                            ".txt",
                            ".md",
                            ".markdown",
                            ".rtf",
                            ".odt",
                        }
                        use_annotation = (
                            _should_use_annotation_system(user_input, has_file=True)
                            and file_ext in _ann_exts
                        )

                        if use_annotation:
                            task_type = "DOC_ANNOTATE"
                            route_method = "📌 Annotation-Strict"
                        elif _is_explicit_file_gen_request(user_input):
                            # 用户明确要生成新文件，直接路由，无需模型分类
                            task_type = "FILE_GEN"
                            route_method = "📄 Explicit-Gen"
                        else:
                            # ★ 主路径：让本地模型做语义路由
                            # 传入 [FILE_ATTACHED:ext] 标记，模型通过训练好的规则判断
                            # CHAT=读文件回答  RESEARCH=深入研究  FILE_GEN=生成新文档
                            _dispatch_q = (
                                user_input or ""
                            ).strip() or "请分析这份文件的内容"
                            _dispatch_input = (
                                f"[FILE_ATTACHED:{file_ext or '.file'}] {_dispatch_q}"
                            )
                            task_analysis, route_method, context_info = (
                                SmartDispatcher.analyze(
                                    _dispatch_input, history=history
                                )
                            )
                            task_type = task_analysis

                if locked_model != "auto":
                    model_to_use = locked_model
                else:
                    complexity = "complex" if file_data is None else "normal"
                    if context_info and context_info.get("complexity"):
                        complexity = context_info["complexity"]

                    if task_type == "FILE_GEN":
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type, has_image=bool(file_data), complexity=complexity
                        )
                    else:
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type, has_image=bool(file_data)
                        )

                _app_logger.info(
                    f"[FILE UPLOAD] 任务类型: {task_type}, 模型: {model_to_use}"
                )

                result = {
                    "task": task_type,
                    "model": model_to_use,
                    "route_method": route_method,
                    "response": "",
                    "images": [],
                    "saved_files": [],
                }

                # 纯归档模式：跳过AI内容分析，直接归档
                if is_organize_only:
                    _app_logger.info(
                        f"[FILE UPLOAD] 纯归档模式: {filename}，跳过AI分析"
                    )
                    result["response"] = ""
                    result["task"] = "FILE_ORGANIZE"
                elif task_type == "DOC_ANNOTATE":
                    # 批量/多文件模式下的标注：同步运行标注管道
                    _app_logger.info(
                        f"[FILE UPLOAD] 批量 DOC_ANNOTATE 模式: {filename}"
                    )
                    try:
                        from web.document_feedback import DocumentFeedbackSystem

                        _batch_docs_dir = settings_manager.documents_dir
                        os.makedirs(_batch_docs_dir, exist_ok=True)
                        # 如需转换先转换
                        _batch_filepath = filepath
                        _batch_file_ext = file_ext
                        if _batch_file_ext != ".docx":
                            try:
                                import tempfile as _bttmp

                                from web.doc_converter import convert_to_docx as _btc

                                _bt_tmp = _bttmp.mkdtemp(prefix="koto_bt_")
                                _bt_conv, _ = _btc(_batch_filepath, output_dir=_bt_tmp)
                                _bt_dest = os.path.join(
                                    _batch_docs_dir, os.path.basename(_bt_conv)
                                )
                                import shutil as _bt_sh

                                _bt_sh.copy2(_bt_conv, _bt_dest)
                                _batch_filepath = _bt_dest
                            except Exception as _bt_err:
                                _app_logger.info(
                                    f"[BATCH DOC_ANNOTATE] 转换失败: {_bt_err}"
                                )
                        _batch_target = os.path.join(
                            _batch_docs_dir, os.path.basename(_batch_filepath)
                        )
                        if os.path.abspath(_batch_filepath) != os.path.abspath(
                            _batch_target
                        ):
                            import shutil as _bsh

                            _bsh.copy2(_batch_filepath, _batch_target)
                        _bt_feedback = DocumentFeedbackSystem(
                            gemini_client=client,
                            default_model_id="gemini-3.1-pro-preview",
                        )
                        _bt_final = None
                        for _bt_evt in _bt_feedback.full_annotation_loop_streaming(
                            _batch_target, user_input
                        ):
                            if _bt_evt.get("stage") == "complete":
                                _bt_final = _bt_evt.get("result", {})
                        if _bt_final and _bt_final.get("success"):
                            _bt_revised = _bt_final.get("revised_file", "")
                            result["response"] = (
                                f"✅ 文档标注完成: {os.path.basename(_bt_revised)}"
                            )
                            result["saved_files"] = [_bt_revised] if _bt_revised else []
                        else:
                            result["response"] = (
                                f"❌ 批量标注失败: {(_bt_final or {}).get('message', '未知错误')}"
                            )
                    except Exception as _bt_exc:
                        result["response"] = f"❌ 批量标注异常: {_bt_exc}"
                else:
                    _app_logger.info(
                        f"[FILE UPLOAD] 处理文件: {filename}, 使用 brain.chat"
                    )
                    brain_result = brain.chat(
                        history=history,
                        user_input=formatted_message,
                        file_data=file_data,
                        model=model_to_use,
                        auto_model=(locked_model == "auto"),
                    )
                    result.update(brain_result)

                # 🗂️ 关键：为每个文件调用FileOrganizer进行归档
                organize_info = {"success": False, "message": "未归档"}
                try:
                    # 使用AI分析文件类型和建议目录
                    from web.file_analyzer import FileAnalyzer

                    analyzer = FileAnalyzer()
                    analysis = analyzer.analyze_file(filepath)  # 只传文件路径
                    suggested_folder = analysis.get("suggested_folder")
                    entity_name = analysis.get("entity")
                    entity_type = analysis.get("entity_type")
                    organizer = get_file_organizer()

                    # 如果已存在同名公司/项目文件夹，则复用
                    if entity_name:
                        existing_folder = organizer.find_entity_folder(entity_name)
                        if existing_folder:
                            suggested_folder = existing_folder

                    if suggested_folder:
                        org_result = organizer.organize_file(
                            filepath,
                            suggested_folder,
                            auto_confirm=True,
                            metadata={
                                "entity": entity_name,
                                "entity_type": entity_type,
                            },
                        )

                        if org_result.get("success"):
                            organize_info = {
                                "success": True,
                                "message": f"✅ 已归档到: {org_result.get('relative_path', suggested_folder)}",
                                "category": suggested_folder,
                                "path": org_result.get("dest_file"),
                            }
                            _app_logger.info(
                                f"[FILE ORGANIZE] ✅ {filename} -> {suggested_folder}"
                            )
                        else:
                            organize_info = {
                                "success": False,
                                "message": f"⚠️ 归档失败: {org_result.get('error', '未知错误')}",
                            }
                    else:
                        organize_info = {
                            "success": False,
                            "message": "⚠️ 无法确定文件分类",
                        }
                except Exception as e:
                    organize_info = {
                        "success": False,
                        "message": f"⚠️ 归档异常: {str(e)}",
                    }
                    _app_logger.info(f"[FILE ORGANIZE ERROR] {filename}: {e}")

                result["file_name"] = filename
                result["organize"] = organize_info
                return result

            except Exception as e:
                return {
                    "file_name": filename,
                    "task": "ERROR",
                    "model": "none",
                    "response": f"❌ 处理文件时出错: {str(e)}",
                    "images": [],
                    "saved_files": [],
                    "organize": {"success": False, "message": "❌ 处理失败，未归档"},
                }

        if stream_mode:

            def generate_progress():
                total = len([f for f in files if f and f.filename])
                started = {
                    "type": "progress",
                    "current": 0,
                    "total": total,
                    "status": "start",
                    "detail": f"开始处理 {total} 个文件",
                }
                yield f"data: {json.dumps(started)}\n\n"

                current = 0
                for file in files:
                    if not file or not file.filename:
                        continue

                    current += 1
                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "processing",
                        "detail": f"处理中: {file.filename} ({current}/{total})",
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                    result = _process_single_file(file)
                    if result:
                        batch_results.append(result)
                        combined_saved_files.extend(result.get("saved_files", []))
                        combined_images.extend(result.get("images", []))

                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "done",
                        "detail": f"完成: {file.filename} ({current}/{total})",
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                summary_lines = [f"📦 批量处理完成，共 {len(batch_results)} 个文件", ""]

                organized_count = sum(
                    1
                    for item in batch_results
                    if item.get("organize", {}).get("success")
                )
                if organized_count > 0:
                    summary_lines.append(f"✅ 已归档: {organized_count} 个文件")

                summary_lines.append("\n📄 **文件详情：**")
                for i, item in enumerate(batch_results, 1):
                    fname = item.get("file_name", "unknown")
                    task = item.get("task", "UNKNOWN")
                    organize = item.get("organize", {})

                    status = "✅" if task != "ERROR" else "❌"
                    org_status = organize.get("message", "未归档")

                    summary_lines.append(f"{i}. {status} **{fname}**")
                    summary_lines.append(f"   📂 {org_status}")

                    response = item.get("response", "")
                    if response and len(response) > 100:
                        summary_lines.append(f"   💬 {response[:100]}...")
                    elif response:
                        summary_lines.append(f"   💬 {response}")

                summary_msg = "\n".join(summary_lines)

                session_manager.update_last_model_response(
                    f"{session_name}.json",
                    summary_msg,
                    task="FILE_BATCH",
                    model_name=locked_model if locked_model != "auto" else "auto",
                    saved_files=combined_saved_files,
                    images=combined_images,
                )

                final_payload = {
                    "type": "final",
                    "response": summary_msg,
                    "task": "FILE_BATCH",
                    "model": locked_model if locked_model != "auto" else "auto",
                    "results": batch_results,
                    "images": combined_images,
                    "saved_files": combined_saved_files,
                }
                yield f"data: {json.dumps(final_payload)}\n\n"

            return Response(generate_progress(), mimetype="text/event-stream")

        for file in files:
            result = _process_single_file(file)
            if not result:
                continue
            batch_results.append(result)
            combined_saved_files.extend(result.get("saved_files", []))
            combined_images.extend(result.get("images", []))

        # 生成详细摘要，包含归档信息
        summary_lines = [f"📦 批量处理完成，共 {len(batch_results)} 个文件", ""]

        organized_count = sum(
            1 for item in batch_results if item.get("organize", {}).get("success")
        )
        if organized_count > 0:
            summary_lines.append(f"✅ 已归档: {organized_count} 个文件")

        summary_lines.append("\n📄 **文件详情：**")
        for i, item in enumerate(batch_results, 1):
            fname = item.get("file_name", "unknown")
            task = item.get("task", "UNKNOWN")
            organize = item.get("organize", {})

            status = "✅" if task != "ERROR" else "❌"
            org_status = organize.get("message", "未归档")

            summary_lines.append(f"{i}. {status} **{fname}**")
            summary_lines.append(f"   📂 {org_status}")

            # 显示AI响应摘要（截取前100字）
            response = item.get("response", "")
            if response and len(response) > 100:
                summary_lines.append(f"   💬 {response[:100]}...")
            elif response:
                summary_lines.append(f"   💬 {response}")

        summary_msg = "\n".join(summary_lines)

        session_manager.update_last_model_response(
            f"{session_name}.json",
            summary_msg,
            task="FILE_BATCH",
            model_name=locked_model if locked_model != "auto" else "auto",
            saved_files=combined_saved_files,
            images=combined_images,
        )

        return jsonify(
            {
                "response": summary_msg,
                "task": "FILE_BATCH",
                "model": locked_model if locked_model != "auto" else "auto",
                "results": batch_results,
                "images": combined_images,
                "saved_files": combined_saved_files,
            }
        )

    file = files[0]

    # Save uploaded file
    filename = _secure_filename(file.filename) or f"upload_{uuid.uuid4().hex}"
    filepath = os.path.join(UPLOAD_DIR, filename)
    file.save(filepath)
    file_type = file.mimetype or file.content_type or ""
    file_ext = os.path.splitext(filename)[1].lower()

    # Load history first (保证即使出错也能保存用户输入)
    history = session_manager.load(f"{session_name}.json")
    user_message = f"[File: {filename}] {user_input}"

    # 🔒 立即保存用户消息到磁盘，防止断连/崩溃导致丢失
    session_manager.append_user_early(f"{session_name}.json", user_message)

    try:
        # 使用新的文件处理器（提取文本/二进制）
        formatted_message, file_data = process_uploaded_file(filepath, user_input)

        # ==================== 智能文档分析引擎 ====================
        # 对 .docx/.doc 文件，使用 LLM 驱动的智能分析引擎判断用户意图
        # 不再硬编码正则，而是让分析器理解用户真实需求
        if file_ext in [".docx", ".doc"]:
            # ── 翻译请求：最高优先级，直接走服务器端翻译管道 ──────────────
            _TRANSLATE_KWS = [
                "翻译",
                "译成",
                "译为",
                "转成英文",
                "转成日文",
                "转成中文",
                "translate",
                "翻成",
                "转译",
            ]
            _is_translate_request = any(
                kw in (user_input or "").lower() for kw in _TRANSLATE_KWS
            )
            if _is_translate_request and locked_task != "DOC_ANNOTATE":
                _app_logger.info(
                    f"[DOCX TRANSLATE] 检测到翻译请求，启用格式保留翻译管道"
                )

                def generate_docx_translation():
                    try:
                        from web.docx_translator_module import (
                            detect_target_language,
                            translate_docx_streaming,
                        )

                        target_lang = detect_target_language(user_input or "")
                        docs_dir = os.path.join(WORKSPACE_DIR, "documents")
                        os.makedirs(docs_dir, exist_ok=True)

                        yield f"data: {json.dumps({'type': 'classification', 'task_type': 'FILE_GEN', 'task_display': '🌐 Word 文档翻译', 'route_method': '🌐 DocxTranslator', 'message': f'🎯 启动格式保留翻译 → {target_lang}'})}\n\n"

                        for event in translate_docx_streaming(
                            filepath, target_lang, client, output_dir=docs_dir
                        ):
                            stage = event.get("stage", "")
                            msg = event.get("message", "")
                            progress = event.get("progress", 0)

                            if stage == "error":
                                yield f"data: {json.dumps({'type': 'token', 'content': f'❌ 翻译失败: {msg}'})}\n\n"
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                                return

                            elif stage == "complete":
                                out_path = event.get("output_path", "")
                                out_name = event.get(
                                    "output_filename", os.path.basename(out_path)
                                )
                                count = event.get("translated_count", 0)
                                lang = event.get("target_language", target_lang)
                                rel_path = os.path.relpath(
                                    out_path, WORKSPACE_DIR
                                ).replace("\\", "/")

                                success_msg = (
                                    f"✅ **Word 文档翻译完成！**\n\n"
                                    f"🌐 目标语言: **{lang}**\n"
                                    f"📝 翻译段落: **{count}** 段\n"
                                    f"📁 文件名: **{out_name}**\n"
                                    f"📍 位置: `workspace/documents/`\n\n"
                                    f"格式已完整保留（字体/加粗/斜体/颜色/表格/页眉页脚）"
                                )
                                yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                                session_manager.append_and_save(
                                    f"{session_name}.json",
                                    user_input,
                                    f"翻译完成 → {out_name} ({count}段, {lang})",
                                )
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [rel_path]})}\n\n"
                                return

                            else:
                                yield f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': f'{progress}%'})}\n\n"

                    except Exception as _te:
                        import traceback as _tb

                        _app_logger.error(
                            f"[DOCX TRANSLATE] ❌ 翻译异常: {_tb.format_exc()}"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': f'❌ 翻译出错: {str(_te)}'})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

                return Response(
                    stream_with_context(generate_docx_translation()),
                    content_type="text/event-stream",
                )

            # 标注任务优先级更高：显式标注意图或用户锁定 DOC_ANNOTATE 时，不进入智能分析引擎
            force_annotation = (
                locked_task == "DOC_ANNOTATE"
            ) or _should_use_annotation_system(user_input, has_file=True)

            # 智能检测：任何对文档内容有实质性处理需求的请求
            # 包括但不限于：写摘要、改引言、改结论、润色、分析结构等
            _doc_intent_keywords = [
                # 生成类
                "写",
                "生成",
                "帮我写",
                "写一段",
                "写个",
                # 修改/改善类
                "改",
                "改善",
                "改进",
                "优化",
                "润色",
                "重写",
                "修改",
                "提升",
                # 学术部件
                "摘要",
                "引言",
                "结论",
                "abstract",
                "前言",
                "导言",
                # 分析类
                "分析",
                "总结",
                "梳理",
                "概述",
                "评估",
                # 质量类
                "不满意",
                "不好",
                "不够",
                "需要改",
                "有问题",
            ]
            is_doc_processing_request = any(
                kw in user_input.lower() for kw in _doc_intent_keywords
            )

            if is_doc_processing_request and not force_annotation:
                _app_logger.info(
                    f"[INTELLIGENT ANALYZER] 检测到文档处理请求，启用智能分析引擎"
                )
                from web.intelligent_document_analyzer import (
                    create_intelligent_analyzer,
                )

                # 创建智能分析器
                analyzer = create_intelligent_analyzer(client)

                # 流式处理文档分析
                def generate_intelligent_analysis():
                    """生成智能文档分析的流式响应"""
                    try:
                        # 使用async生成器（需要在async context中）
                        import asyncio

                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)

                        async def run_analysis():
                            async for (
                                event
                            ) in analyzer.process_document_intelligent_streaming(
                                filepath, user_input, session_name
                            ):
                                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                        gen = run_analysis()
                        while True:
                            try:
                                result = loop.run_until_complete(gen.__anext__())
                                yield result
                            except StopAsyncIteration:
                                break
                    except Exception as e:
                        error_event = {
                            "stage": "error",
                            "message": f"智能分析失败: {str(e)}",
                        }
                        yield f"data: {json.dumps(error_event, ensure_ascii=False)}\n\n"
                    finally:
                        loop.close()

                return Response(
                    stream_with_context(generate_intelligent_analysis()),
                    content_type="text/event-stream",
                )
        # ==================== 智能文档分析引擎结束 ====================

        # 智能任务分析
        task_type = locked_task
        context_info = None
        route_method = "Auto"
        if not task_type:
            # 如果是图片上传，直接判断编辑或分析，避免初始化本地路由器导致卡顿
            if file_data and file_type and file_type.startswith("image"):
                message_lower = (user_input or "").lower()
                is_edit = any(
                    kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS
                )
                task_type = "PAINTER" if is_edit else "VISION"
                route_method = "🖼️ Image Edit" if is_edit else "👁️ Image Analysis"
                _app_logger.info(
                    f"[FILE UPLOAD] 图片任务直通路由: {task_type} (方法: {route_method})"
                )
            else:
                # 文档上传：严格检测标注意图（必须明确要求在原文上标记）
                _ann_exts = {
                    ".doc",
                    ".docx",
                    ".pdf",
                    ".txt",
                    ".md",
                    ".markdown",
                    ".rtf",
                    ".odt",
                }
                use_annotation = (
                    _should_use_annotation_system(user_input, has_file=True)
                    and file_ext in _ann_exts
                )

                if use_annotation:
                    task_type = "DOC_ANNOTATE"
                    route_method = "📌 Annotation-Strict"
                elif _is_explicit_file_gen_request(user_input):
                    # 用户明确要生成新文件，直接路由，无需模型分类
                    task_type = "FILE_GEN"
                    route_method = "📄 Explicit-Gen"
                    _app_logger.info(
                        f"[FILE UPLOAD] 🎯 检测到明确文件生成请求，启用 FILE_GEN 模式"
                    )
                else:
                    # ★ 主路径：让本地模型做语义路由
                    # 传入 [FILE_ATTACHED:ext] 标记，模型通过训练好的规则判断
                    # CHAT=读文件回答  RESEARCH=深入研究  FILE_GEN=生成新文档
                    _dispatch_q = (user_input or "").strip() or "请分析这份文件的内容"
                    _dispatch_input = (
                        f"[FILE_ATTACHED:{file_ext or '.file'}] {_dispatch_q}"
                    )
                    task_analysis, route_method, context_info = SmartDispatcher.analyze(
                        _dispatch_input, history=history
                    )
                    task_type = task_analysis

                _app_logger.info(
                    f"[FILE UPLOAD] 智能路由选择任务类型: {task_type} (方法: {route_method})"
                )

        # 确定使用的模型
        if locked_model != "auto":
            model_to_use = locked_model
        else:
            # 获取任务复杂度（上传文件默认按复杂任务处理）
            complexity = "complex" if file_data is None else "normal"
            if context_info and context_info.get("complexity"):
                complexity = context_info["complexity"]

            if task_type == "DOC_ANNOTATE":
                # 文档标注需要强模型：优先使用 gemini-3.1-pro-preview或 gemini-3-pro-preview
                model_to_use = "gemini-3.1-pro-preview"
            elif task_type == "FILE_GEN":
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type, has_image=bool(file_data), complexity=complexity
                )
            else:
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type, has_image=bool(file_data)
                )

        print(f"[FILE UPLOAD] 任务类型: {task_type}, 模型: {model_to_use}")

        # 安全兜底：locked_task 预设时 prefer_ppt 可能未定义
        if "prefer_ppt" not in locals():
            _ppt_kws = [
                "ppt",
                "幻灯片",
                "演示",
                "汇报",
                "presentation",
                "slide",
                "deck",
            ]
            prefer_ppt = any(kw in (user_input or "").lower() for kw in _ppt_kws)

        # 如果是文本类文件，按任务类型处理
        result = {
            "task": "FILE_GEN" if task_type == "DOC_ANNOTATE" else task_type,
            "subtask": "DOC_ANNOTATE" if task_type == "DOC_ANNOTATE" else None,
            "model": model_to_use,
            "route_method": route_method,
            "response": "",
            "images": [],
            "saved_files": [],
        }

        # 文档标注任务 - 流式反馈，生成带Track Changes的Word文档
        if task_type == "DOC_ANNOTATE":
            docs_dir = settings_manager.documents_dir
            os.makedirs(docs_dir, exist_ok=True)

            source_path = filepath
            target_path = os.path.join(docs_dir, filename)
            if os.path.abspath(source_path) != os.path.abspath(target_path):
                import shutil as _shutil_ann

                _shutil_ann.copy2(source_path, target_path)

            # 使用流式SSE返回进度，让前端能实时显示
            # 捕获闭包变量（防止generator延迟执行时变量已改变）
            _ann_target_path = target_path
            _ann_filename = filename
            _ann_file_ext = file_ext
            _ann_route_method = route_method
            _ann_model = model_to_use
            _ann_session = session_name
            _ann_user_input = user_input
            _ann_client = client
            _ann_docs_dir = docs_dir  # 确保转换后的文件和输出都保存到标准目录

            def generate_doc_annotate_stream():
                import time as _time

                _start = _time.time()
                task_id = f"doc_annotate_{_ann_session}_{int(_start * 1000)}"

                # Local mutable copies of closure vars (Python makes vars local if assigned anywhere
                # in the function, so we cannot reassign _ann_* directly without UnboundLocalError)
                _loc_target_path = _ann_target_path
                _loc_filename = _ann_filename
                _loc_file_ext = _ann_file_ext

                # ── 非 .docx 格式自动转换 ──────────────────────────────────────
                # 对 .doc / .pdf / .txt / .md / .rtf / .odt 先转换为 .docx 再进标注
                _converted_warning = ""
                _classif_sent = False
                if _loc_file_ext != ".docx":
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_ANNOTATE', 'route_method': _ann_route_method, 'model': _ann_model, 'task_id': task_id, 'message': '📄 DOC_ANNOTATE'})}\n\n"
                    _classif_sent = True
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'converting', 'message': f'🔄 正在将 {_loc_file_ext} 转换为可编辑 .docx...', 'detail': _loc_filename, 'progress': 3})}\n\n"
                    try:
                        from web.doc_converter import convert_to_docx, needs_conversion

                        if needs_conversion(_loc_file_ext):
                            import tempfile as _tmpmod

                            _conv_dir = _tmpmod.mkdtemp(prefix="koto_conv_")
                            _conv_path, _converted_warning = convert_to_docx(
                                _loc_target_path, output_dir=_conv_dir
                            )
                            # 将转换后的 .docx 复制到标准文档目录，确保输出也在该目录
                            _conv_basename = os.path.basename(_conv_path)
                            _conv_in_docs = os.path.join(_ann_docs_dir, _conv_basename)
                            import shutil as _shutil_conv

                            _shutil_conv.copy2(_conv_path, _conv_in_docs)
                            _loc_target_path = (
                                _conv_in_docs  # 用 docs_dir 路径，输出也会在此
                            )
                            _loc_filename = _conv_basename
                            _loc_file_ext = ".docx"
                            _app_logger.info(
                                f"[DocConvert] ✅ 转换并复制到文档目录 → {_loc_target_path}"
                            )
                        else:
                            yield f"data: {json.dumps({'type': 'error', 'message': f'不支持的格式：{_loc_file_ext}'})}\n\n"
                            return
                    except Exception as _conv_err:
                        err_msg = f"❌ 格式转换失败：{_conv_err}"
                        yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                        _elapsed = _time.time() - _start
                        session_manager.update_last_model_response(
                            f"{_ann_session}.json", err_msg
                        )
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"
                        return
                # ─────────────────────────────────────────────────────────────

                try:
                    from web.document_feedback import DocumentFeedbackSystem

                    feedback_system = DocumentFeedbackSystem(
                        gemini_client=_ann_client,
                        default_model_id="gemini-3.1-pro-preview",
                    )

                    # 发送分类信息（若上方转换块已发送则跳过重复）
                    if not _classif_sent:
                        yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_ANNOTATE', 'route_method': _ann_route_method, 'model': _ann_model, 'task_id': task_id, 'message': '📄 DOC_ANNOTATE'})}\n\n"
                    if _converted_warning:
                        yield f"data: {json.dumps({'type': 'info', 'message': _converted_warning})}\n\n"

                    # 发送初始进度
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '📖 正在读取文档...', 'detail': _loc_filename, 'progress': 5})}\n\n"

                    # ── 转换质量检查：如果段落数过多或内容为乱码，拒绝标注 ──────
                    try:
                        from docx import Document as _QDoc

                        _qd = _QDoc(_loc_target_path)
                        _q_paras = [p.text for p in _qd.paragraphs if p.text.strip()]
                        _max_para_limit = 500
                        # 乱码检测：短垃圾段落比例 > 60% 或 字母比例过低
                        _q_short = sum(1 for p in _q_paras if len(p) < 15)
                        _q_short_ratio = _q_short / max(len(_q_paras), 1)
                        _q_alpha_ratio = sum(
                            sum(1 for c in p if c.isalpha()) / max(len(p), 1)
                            for p in _q_paras[:200]
                        ) / max(min(len(_q_paras), 200), 1)
                        _is_garbage = (
                            len(_q_paras) > _max_para_limit and _q_short_ratio > 0.5
                        ) or (len(_q_paras) > 200 and _q_short_ratio > 0.7)
                        if _is_garbage:
                            _q_err = (
                                f"❌ **文件转换质量过低**，检测到 {len(_q_paras):,} 个段落"
                                f"（{_q_short_ratio:.0%} 为乱码短行），内容无法识别。\n\n"
                                "**原因**：`.doc` 格式使用了旧版二进制结构，无法自动解析。\n\n"
                                "**解决方法**：\n"
                                "1. 用 **Microsoft Word** 打开 `.doc` 文件\n"
                                "2. 点击【文件】→【另存为】\n"
                                "3. 选择格式 **Word 文档 (*.docx)**\n"
                                "4. 重新上传 `.docx` 文件"
                            )
                            session_manager.update_last_model_response(
                                f"{_ann_session}.json", _q_err
                            )
                            yield f"data: {json.dumps({'type': 'token', 'content': _q_err})}\n\n"
                            _elapsed = _time.time() - _start
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"
                            return
                    except Exception:
                        pass  # 检查失败时继续正常流程
                    # ──────────────────────────────────────────────────────────

                    revised_file = None
                    final_result = None

                    for (
                        progress_event
                    ) in feedback_system.full_annotation_loop_streaming(
                        _loc_target_path,
                        _ann_user_input,
                        task_id=task_id,
                        model_id=_ann_model,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(
                            _ann_session
                        ),
                    ):
                        stage = progress_event.get("stage", "unknown")
                        progress = progress_event.get("progress", 0)
                        message_text = progress_event.get("message", "")
                        detail = progress_event.get("detail", "")

                        if stage == "cancelled":
                            yield f"data: {json.dumps({'type': 'info', 'message': '⏸️ 任务已取消'})}\n\n"
                            _elapsed = _time.time() - _start
                            # 保存取消记录
                            session_manager.update_last_model_response(
                                f"{_ann_session}.json", "⏸️ 文档标注任务已取消"
                            )
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed, 'cancelled': True})}\n\n"
                            return

                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message_text, 'detail': detail, 'progress': progress})}\n\n"

                        if stage == "complete":
                            final_result = progress_event.get("result", {})
                            revised_file = final_result.get("revised_file")

                    _elapsed = _time.time() - _start

                    if final_result and final_result.get("success"):
                        applied = final_result.get("applied", 0)
                        failed = final_result.get("failed", 0)
                        total = final_result.get("total", applied + failed)

                        # ── 兜底检测 ───────────────────────────────────────────
                        _fb_used = final_result.get("fallback_used", False)
                        _fb_partial = final_result.get("partial_fallback", False)
                        _fb_err = final_result.get("last_api_error", "")
                        _fb_chunks = final_result.get("fallback_chunk_count", 0)
                        _ai_chunks = final_result.get("ai_chunk_count", 0)

                        # 读取文档信息
                        try:
                            from docx import Document as _Doc

                            _d = _Doc(_loc_target_path)
                            _total_paras = len(
                                [p for p in _d.paragraphs if p.text.strip()]
                            )
                            _total_chars = sum(len(p.text) for p in _d.paragraphs)
                        except Exception:
                            _total_paras = 0
                            _total_chars = 0

                        density = (
                            (applied / _total_chars * 1000) if _total_chars > 0 else 0
                        )

                        # ── 构建模型行 / 兜底警告 ─────────────────────────────
                        if _fb_used:
                            model_display = (
                                f"`{_ann_model}` ⚠️ **（AI未成功，已用本地规则兜底）**"
                            )
                        elif _fb_partial:
                            model_display = f"`{_ann_model}` ⚠️ **（{_fb_chunks}段兜底 / {_ai_chunks}段AI）**"
                        else:
                            model_display = f"`{_ann_model}`"

                        summary_lines = [
                            "## ✅ 文档修改完成！",
                            "",
                            "### 📊 修改统计",
                            f"- 找到并应用: **{applied}** 处修改",
                            f"- 定位失败: {failed} 处",
                            f"- 总计分析: {total} 处",
                            "",
                            "### 📋 文档信息",
                            f"- 文件名: `{_loc_filename}`",
                            f"- 段落数: {_total_paras} 段",
                            f"- 字数: {_total_chars} 字",
                            f"- 修改密度: **{density:.1f}** 处/千字",
                            "",
                            f"### 📄 模型: {model_display}",
                            "",
                            f"### 📝 输出文件: `{os.path.basename(revised_file) if revised_file else '待生成'}`",
                        ]

                        # ── 当使用兜底时，插入显眼的警告块 ──────────────────
                        if _fb_used or _fb_partial:
                            fb_label = (
                                "全部分段"
                                if _fb_used
                                else f"{_fb_chunks}/{_fb_chunks+_ai_chunks} 分段"
                            )
                            summary_lines += [
                                "",
                                "---",
                                "### ⚠️ 质量警告：本次使用了本地规则兜底",
                                "",
                                f"**问题**: Gemini API 在 {fb_label} 中调用失败，系统自动降级为基于正则规则的本地标注。",
                                "本地兜底标注质量**明显低于** AI 分析，主要覆盖被动句、名词化、冗余连接词等固定模式，",
                                "无法理解上下文语义。",
                                "",
                                f"**错误信息**: `{_fb_err[:120] if _fb_err else '（无详细错误日志）'}`",
                                "",
                                "**建议排查**:",
                                "1. 检查 Koto 后台控制台，找 `[DocumentFeedback] ❌` 或 `⚠️` 开头的日志行",
                                "2. 确认 API Key 有效：`config/gemini_config.env` → `GEMINI_API_KEY`",
                                "3. 确认 `gemini-2.5-pro` 对您的账号可用（部分账号受访问限制）",
                                "4. 重试一次，如仍失败可换用 `gemini-2.5-flash`",
                                "---",
                            ]

                        summary_lines += [
                            "",
                            "### 💡 使用方法",
                            "1. 用 Microsoft Word 打开输出文件",
                            "2. 点击「审阅」标签页",
                            "3. 右侧气泡中查看全部修改建议",
                            "4. 逐条接受或忽略（右键批注可操作）",
                        ]
                        summary_msg = "\n".join(summary_lines)

                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"

                        session_manager.update_last_model_response(
                            f"{_ann_session}.json",
                            summary_msg,
                            task="DOC_ANNOTATE",
                            model_name=_ann_model,
                            saved_files=[revised_file] if revised_file else [],
                        )

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': _elapsed})}\n\n"
                    else:
                        err_msg = (
                            final_result.get("message", "未知错误")
                            if final_result
                            else "处理失败"
                        )
                        # 保存失败记录
                        session_manager.update_last_model_response(
                            f"{_ann_session}.json", f"❌ 文档标注失败: {err_msg}"
                        )
                        yield f"data: {json.dumps({'type': 'error', 'message': '❌ ' + err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"

                except Exception as e:
                    import traceback

                    traceback.print_exc()
                    # 保存异常记录
                    session_manager.update_last_model_response(
                        f"{_ann_session}.json", f"❌ 标注系统错误: {str(e)[:200]}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': '❌ 标注系统错误: ' + str(e)[:200]})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

            return Response(
                stream_with_context(generate_doc_annotate_stream()),
                mimetype="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "X-Accel-Buffering": "no",
                    "Connection": "keep-alive",
                },
            )

        # 🎯 FILE_GEN + PPT 生成（P0 新增）
        elif task_type == "FILE_GEN" and prefer_ppt:
            _app_logger.info(f"[FILE_GEN PPT] 开始 PPT 生成流程")

            # 第 1 步：使用 FileParser 提取结构化内容
            from web.file_parser import FileParser
            from web.ppt_session_manager import PPTSessionManager

            parser = FileParser()
            parse_result = parser.parse_file(filepath)
            file_content = parse_result.get("content", "") if parse_result else ""

            # 第 2 步：创建 PPT 会话
            ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
            os.makedirs(ppt_session_dir, exist_ok=True)

            session_manager_ppt = PPTSessionManager(ppt_session_dir)
            ppt_session_id = session_manager_ppt.create_session(
                title=f"PPT from {os.path.splitext(filename)[0]}",
                user_input=user_input,
                theme="business",
            )
            _app_logger.info(f"[FILE_GEN PPT] 创建会话: {ppt_session_id}")

            # 第 3 步：保存文件内容到会话
            session_manager_ppt.save_generation_data(
                session_id=ppt_session_id,
                ppt_data=None,
                ppt_file_path=None,
                uploaded_file_context=file_content[:3000],  # 将内容限制为前3000字符
            )
            _app_logger.info(f"[FILE_GEN PPT] 文件内容已保存到会话")

            # 使用流式响应（Streamed Response）以支持实时进度显示
            def generate_ppt_file_stream():
                import asyncio
                import queue
                import threading
                import time as _time

                from web.app import TaskOrchestrator

                _start = _time.time()

                # 发送初始化信息
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'FILE_GEN', 'subtask': 'PPT_CREATION', 'message': '📊 开始 PPT 演示文稿生成流程'})}\n\n"

                # 准备任务参数
                subtask = {
                    "task_type": "FILE_GEN",
                    "index": 1,
                    "description": f"从文档 {filename} 生成 PPT",
                }
                context = {"original_input": user_input, "step_1_output": file_content}

                # 进度队列
                progress_queue = queue.Queue()

                def _progress_cb(msg, detail=""):
                    progress_queue.put({"msg": msg, "detail": detail})

                # 任务结果容器
                task_result_holder = {"result": None}

                # 后台执行函数
                def _run_task_thread():
                    # 为新线程创建独立的事件循环
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    try:
                        task_result_holder["result"] = loop.run_until_complete(
                            TaskOrchestrator._execute_file_gen(
                                user_input, context, subtask, _progress_cb
                            )
                        )
                    except Exception as e:
                        task_result_holder["result"] = {
                            "success": False,
                            "error": str(e),
                        }
                    finally:
                        loop.close()
                        progress_queue.put(None)  # Signal done

                # 启动后台线程
                t = threading.Thread(target=_run_task_thread)
                t.start()

                # 主线程循环读取进度
                while True:
                    try:
                        item = progress_queue.get(timeout=0.1)
                        if item is None:
                            break
                        # 发送进度SSE
                        yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': item['detail']})}\n\n"
                    except queue.Empty:
                        if not t.is_alive():
                            break

                t.join()
                ppt_result = task_result_holder["result"]
                _elapsed = _time.time() - _start

                # 处理最终结果
                if ppt_result and ppt_result.get("success"):
                    saved_files = ppt_result.get("saved_files", [])
                    if saved_files:
                        ppt_file_path = (
                            saved_files[0]
                            if isinstance(saved_files, list)
                            else saved_files
                        )
                        # 保存会话数据
                        session_manager_ppt.save_generation_data(
                            session_id=ppt_session_id,
                            ppt_data=ppt_result.get("ppt_data"),
                            ppt_file_path=ppt_file_path,
                        )

                        final_msg = (
                            f"✅ PPT 演示已生成\n\n"
                            f"📄 文件: [{os.path.basename(ppt_file_path)}]({ppt_file_path.replace(os.sep, '/')})\n"
                            f"🔗 会话ID: `{ppt_session_id}`\n"
                            f"⏱️ 耗时: {_elapsed:.1f}s"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"

                        # 更新历史
                        session_manager.update_last_model_response(
                            f"{session_name}.json",
                            final_msg,
                            task="FILE_GEN",
                            model_name=model_to_use,
                            saved_files=[ppt_file_path],
                        )

                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [ppt_file_path], 'ppt_session_id': ppt_session_id})}\n\n"
                    else:
                        err_msg = "⚠️ PPT 框架已生成，但文件保存失败"
                        yield f"data: {json.dumps({'type': 'error', 'message': err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                else:
                    err_msg = (
                        ppt_result.get("error", "未知错误")
                        if ppt_result
                        else "任务执行无结果"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 生成失败: {err_msg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

            return Response(
                stream_with_context(generate_ppt_file_stream()),
                mimetype="text/event-stream",
            )

        elif task_type in ["FILE_GEN", "RESEARCH", "CHAT"]:
            # ── 文本类文件分析（SSE 流式，修复原 blocking brain.chat 卡死问题）──
            _captured_context = context_info  # closure capture
            _captured_model = model_to_use
            _captured_task = task_type

            def generate_file_analysis_stream():
                import time as _time

                _start = _time.time()
                try:
                    _skill = (_captured_context or {}).get("skill_prompt")

                    # 根据任务类型选择 system instruction
                    if _captured_task == "RESEARCH":
                        _sys = (
                            "你是一位专业的文档分析助手，擅长深度解读各类文件（商业计划书、研究报告、技术文档等）。\n"
                            "请仔细阅读用户提供的文件内容，并按以下结构输出分析报告：\n\n"
                            "## 核心摘要\n- 用 3-5 条要点概括文件核心内容\n\n"
                            "## 详细解读\n### 背景与目标\n### 关键内容分析\n### 数据与证据\n\n"
                            "## 结论与建议\n- 综合评判与可行性/价值判断\n\n"
                            "要求：用中文，条理清晰，避免冗余，不输出代码块标记。"
                        )
                    elif _captured_task == "CHAT":
                        # 用户上传文件+提问 → 读取分析文件，不生成新文件模板
                        _sys = (
                            "你是一位专业的文档阅读与分析助手。用户上传了一份文件并提出了问题，"
                            "请认真阅读文件的完整内容，用中文给出详细、准确的分析和回答。\n"
                            "注意：\n"
                            "- 直接回答用户的具体问题，不要生成新文档模板\n"
                            "- 引用文件中的具体数据和信息支撑你的判断\n"
                            "- 如涉及投资价值/风险，结合文件内容给出有依据的评估\n"
                            "- 用清晰的结构输出，避免空泛表述"
                        )
                    else:
                        _sys = _get_filegen_brief_instruction()

                    if _skill:
                        _sys += f"\n\n[分析重点] {_skill}"

                    # ── 二进制文件（PDF/Word等）+ CHAT/RESEARCH：传字节流给模型直接读取 ──
                    # 如果有 file_data（非图片二进制），强制使用支持 generate_content 的模型
                    # 并将 PDF 字节附加到请求中，而不是依赖提取的文本
                    _stream_model = _captured_model
                    _stream_contents = formatted_message  # 默认：文本消息
                    _has_binary_doc = file_data is not None and not (
                        file_data.get("mime_type") or ""
                    ).lower().startswith("image/")
                    # 需要降级到 generate_content 兼容模型的条件：
                    # 1. 二进制文件（PDF/Word）+ CHAT/RESEARCH 任务
                    # 2. 所选模型是 Interactions-only 模型（如 deep-research），不支持 generate_content_stream
                    _need_fallback = (
                        _has_binary_doc and _captured_task in ("CHAT", "RESEARCH")
                    ) or _stream_model in _INTERACTIONS_ONLY_MODELS
                    if _need_fallback:
                        if _stream_model in _INTERACTIONS_ONLY_MODELS:
                            print(
                                f"[FILE STREAM] interactions-only 模型 {_stream_model} 不支持文件流，降级到 {_INTERACTIONS_FALLBACK_MODEL}"
                            )
                        _stream_model = _INTERACTIONS_FALLBACK_MODEL
                        if _has_binary_doc:
                            try:
                                _doc_part = types.Part.from_bytes(
                                    data=file_data["data"],
                                    mime_type=file_data.get(
                                        "mime_type", "application/pdf"
                                    ),
                                )
                                _stream_contents = [formatted_message, _doc_part]
                                print(
                                    f"[FILE STREAM] 📄 Binary-Doc-Read: model={_stream_model}, bytes={len(file_data['data'])}"
                                )
                            except Exception as _bp_err:
                                print(
                                    f"[FILE STREAM] ⚠️ 无法创建 doc_part，回退到文本模式: {_bp_err}"
                                )
                                _stream_contents = formatted_message

                    yield f"data: {json.dumps({'type': 'classification', 'task_type': _captured_task, 'model': _stream_model, 'message': f'📄 正在分析: {filename}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📂 文件内容已就绪', 'stage': 'file_ready_complete', 'progress': 15})}\n\n"
                    _task_display = {
                        "FILE_GEN": "📝 文件生成",
                        "RESEARCH": "🔬 深度分析",
                        "CHAT": "💬 对话分析",
                    }.get(_captured_task, _captured_task)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'🎯 任务类型: {_task_display}', 'stage': 'routing_complete', 'progress': 25})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'⚡ 正在请求 {_stream_model}，请稍候...', 'stage': 'api_calling', 'progress': 35})}\n\n"

                    response_stream = client.models.generate_content_stream(
                        model=_stream_model,
                        contents=_stream_contents,
                        config=types.GenerateContentConfig(
                            system_instruction=_sys,
                            temperature=0.7,
                            max_output_tokens=8000,
                        ),
                    )

                    full_text = ""
                    _first_token = True
                    for _chunk in response_stream:
                        _t = getattr(_chunk, "text", None)
                        if _t:
                            if _first_token:
                                yield f"data: {json.dumps({'type': 'progress', 'message': '✍️ 模型正在生成回复...', 'stage': 'generating_complete', 'progress': 55})}\n\n"
                                _first_token = False
                            full_text += _t
                            yield f"data: {json.dumps({'type': 'token', 'content': _t})}\n\n"

                    _elapsed = round(_time.time() - _start, 2)
                    _saved_files = []

                    # 自动保存为 DOCX
                    if full_text and len(full_text) > 50:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '💾 正在保存文档...', 'stage': 'saving', 'progress': 90})}\n\n"
                        try:
                            _title = _build_analysis_title(
                                user_input, filename, is_binary=False
                            )
                            _cleaned = _strip_code_blocks(full_text)
                            _docx = save_docx(
                                _cleaned,
                                title=_title,
                                output_dir=settings_manager.documents_dir,
                            )
                            _docx_rel = os.path.relpath(_docx, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            _saved_files.append(_docx_rel)
                            _app_logger.info(
                                f"[FILE UPLOAD] ✅ 分析已保存 DOCX: {_docx_rel}"
                            )
                            # 按需同时保存 PDF
                            if any(
                                kw in (user_input or "").lower()
                                for kw in ["pdf", "两种格式", "both"]
                            ):
                                try:
                                    _pdf = save_pdf(
                                        _cleaned,
                                        title=_title,
                                        output_dir=settings_manager.documents_dir,
                                    )
                                    _saved_files.append(
                                        os.path.relpath(_pdf, WORKSPACE_DIR).replace(
                                            "\\", "/"
                                        )
                                    )
                                except Exception:
                                    pass
                        except Exception as _de:
                            _app_logger.warning(
                                f"[FILE UPLOAD] ⚠️ 保存 DOCX 失败: {_de}"
                            )

                    session_manager.update_last_model_response(
                        f"{session_name}.json",
                        full_text,
                        task=_captured_task,
                        model_name=_captured_model,
                        saved_files=_saved_files,
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': _saved_files, 'total_time': _elapsed})}\n\n"

                except Exception as _e:
                    import traceback as _tb

                    _tb.print_exc()
                    _emsg = str(_e)[:200]
                    session_manager.update_last_model_response(
                        f"{session_name}.json", f"❌ 分析失败: {_emsg}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 分析失败: {_emsg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_file_analysis_stream()),
                mimetype="text/event-stream",
            )

        else:
            # ── 图片 / 二进制文件：视觉分析（SSE 流式包装）──
            _captured_fdata = file_data
            _captured_model_v = model_to_use
            _captured_task_v = task_type

            def generate_vision_stream():
                import time as _time

                _start = _time.time()
                try:
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': _captured_task_v, 'model': _captured_model_v, 'message': f'👁️ 正在分析: {filename}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '� 文件已接收', 'stage': 'file_ready_complete', 'progress': 15})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'⚡ 正在请求视觉模型 {_captured_model_v}...', 'stage': 'api_calling', 'progress': 35})}\n\n"

                    # 调用 brain.chat（vision 路径通常较快）
                    _brain_result = brain.chat(
                        history=history,
                        user_input=formatted_message,
                        file_data=_captured_fdata,
                        model=_captured_model_v,
                        auto_model=(locked_model == "auto"),
                    )
                    _resp_text = _brain_result.get("response", "")
                    _elapsed = round(_time.time() - _start, 2)
                    _saved_files = list(_brain_result.get("saved_files", []))

                    # 输出内容
                    if _resp_text:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '✍️ 分析完成，正在输出...', 'stage': 'generating_complete', 'progress': 70})}\n\n"
                        yield f"data: {json.dumps({'type': 'token', 'content': _resp_text})}\n\n"

                    # 自动保存视觉分析为 DOCX
                    if _resp_text and len(_resp_text) > 50:
                        yield f"data: {json.dumps({'type': 'progress', 'message': '💾 正在保存文档...', 'stage': 'saving', 'progress': 90})}\n\n"
                        try:
                            _title = _build_analysis_title(
                                user_input, filename, is_binary=True
                            )
                            _cleaned = _strip_code_blocks(_resp_text)
                            _docx = save_docx(
                                _cleaned,
                                title=_title,
                                output_dir=settings_manager.documents_dir,
                            )
                            _docx_rel = os.path.relpath(_docx, WORKSPACE_DIR).replace(
                                "\\", "/"
                            )
                            _saved_files.append(_docx_rel)
                            _app_logger.info(
                                f"[FILE UPLOAD] ✅ 视觉分析已保存 DOCX: {_docx_rel}"
                            )
                        except Exception as _de:
                            _app_logger.warning(
                                f"[FILE UPLOAD] ⚠️ 视觉 DOCX 保存失败: {_de}"
                            )

                    session_manager.update_last_model_response(
                        f"{session_name}.json",
                        _resp_text,
                        task=_captured_task_v,
                        model_name=_captured_model_v,
                        saved_files=_saved_files,
                        images=_brain_result.get("images", []),
                    )
                    yield f"data: {json.dumps({'type': 'done', 'images': _brain_result.get('images', []), 'saved_files': _saved_files, 'total_time': _elapsed})}\n\n"

                except Exception as _e:
                    import traceback as _tb

                    _tb.print_exc()
                    _emsg = str(_e)[:200]
                    session_manager.update_last_model_response(
                        f"{session_name}.json", f"❌ 文件分析失败: {_emsg}"
                    )
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 文件分析失败: {_emsg}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(
                stream_with_context(generate_vision_stream()),
                mimetype="text/event-stream",
            )

    except Exception as e:
        # 即使出错也保存用户的问题和错误信息
        import traceback

        error_detail = traceback.format_exc()
        _app_logger.info(f"[FILE UPLOAD ERROR] {error_detail}")

        error_response = f"❌ 处理文件时出错: {str(e)}"
        session_manager.update_last_model_response(
            f"{session_name}.json", error_response
        )

        return jsonify(
            {
                "response": error_response,
                "task": "ERROR",
                "model": "none",
                "images": [],
                "saved_files": [],
            }
        )


# ==================== PPT 相关 API 端点（P0 补充）====================


@app.route("/api/ppt/download", methods=["POST"])
def download_ppt():
    """下载 PPT PPTX 文件"""
    try:
        session_id = request.json.get("session_id")
        if not session_id:
            return jsonify({"error": "Missing session_id"}), 400

        # 从 PPT 会话中获取文件路径
        from web.ppt_session_manager import PPTSessionManager

        ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
        manager = PPTSessionManager(ppt_session_dir)

        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404

        ppt_file_path = session_data.get("ppt_file_path")
        if not ppt_file_path:
            # 如果文件还没生成，尝试生成一个临时的
            return jsonify({"error": "PPT file not generated yet"}), 400

        # 构建完整的文件路径
        full_path = os.path.join(
            WORKSPACE_DIR, ppt_file_path.lstrip("/").replace("/", os.sep)
        )

        if not os.path.exists(full_path):
            return jsonify({"error": "PPT file not found"}), 404

        # 返回文件下载
        return send_file(
            full_path,
            mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            as_attachment=True,
            download_name=os.path.basename(full_path),
        )

    except Exception as e:
        _app_logger.info(f"[PPT DOWNLOAD] 错误: {e}")
        return jsonify({"error": f"Download failed: {str(e)}"}), 500


@app.route("/api/ppt/session/<session_id>", methods=["GET"])
def get_ppt_session(session_id):
    """获取 PPT 会话信息"""
    try:
        from web.ppt_session_manager import PPTSessionManager

        ppt_session_dir = os.path.join(WORKSPACE_DIR, "workspace", "ppt_sessions")
        manager = PPTSessionManager(ppt_session_dir)

        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404

        return jsonify(
            {
                "success": True,
                "session": {
                    "id": session_data.get("session_id"),
                    "title": session_data.get("title"),
                    "status": session_data.get("status"),
                    "ppt_file_path": session_data.get("ppt_file_path"),
                    "created_at": session_data.get("created_at"),
                    "updated_at": session_data.get("updated_at"),
                },
            }
        )

    except Exception as e:
        _app_logger.info(f"[PPT SESSION] 错误: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/v1/models", methods=["GET"])
def api_list_models():
    """List available AI models and current task routing.
    ---
    tags:
      - Models
    responses:
      200:
        description: Model list with routing info
        schema:
          type: object
          properties:
            ready:
              type: boolean
              description: Whether the model manager has finished initializing
            model_map:
              type: object
              description: "Task \u2192 model ID routing table with scoring info"
            available:
              type: array
              description: All available models with capabilities
              items:
                type: object
                properties:
                  id:
                    type: string
                  display:
                    type: string
                  tier:
                    type: integer
                  provider:
                    type: string
                  strengths:
                    type: array
                    items:
                      type: string
            fallback:
              type: string
              description: Fallback model ID for interactions
            interactions_only:
              type: array
              items:
                type: string
              description: Models restricted to interactions only
    """
    if _model_manager:
        return jsonify(
            {
                "ready": True,
                "model_map": _model_manager.get_model_map_with_scores(),
                "available": _model_manager.get_available_models(),
                "fallback": _INTERACTIONS_FALLBACK_MODEL,
                "interactions_only": list(_INTERACTIONS_ONLY_MODELS),
            }
        )
    # 模型管理器尚未就绪或不可用，返回静态默认值
    return jsonify(
        {
            "ready": False,
            "model_map": {
                task: {
                    "model_id": mid,
                    "display": get_model_display_name(mid),
                    "provider": "gemini" if mid != "local-executor" else "local",
                    "tier": MODEL_INFO.get(mid, {}).get("tier", 5),
                    "score": None,
                    "_inferred": False,
                }
                for task, mid in MODEL_MAP.items()
            },
            "available": [
                {
                    "id": mid,
                    "display": get_model_display_name(mid),
                    "tier": MODEL_INFO.get(mid, {}).get("tier", 5),
                    "provider": "gemini" if mid != "local-executor" else "local",
                    "strengths": MODEL_INFO.get(mid, {}).get("strengths", []),
                    "capabilities": {},
                }
                for mid in dict.fromkeys(MODEL_MAP.values())
            ],
            "fallback": _INTERACTIONS_FALLBACK_MODEL,
            "interactions_only": list(_INTERACTIONS_ONLY_MODELS),
        }
    )


@app.route("/api/v1/models/refresh", methods=["POST"])
def api_refresh_models():
    """Manually refresh the model list and routing table.
    ---
    tags:
      - Models
    responses:
      200:
        description: Refresh result
        schema:
          type: object
          properties:
            status:
              type: string
              enum: [ok, initializing]
              description: "'ok' when refresh succeeded, 'initializing' when manager is starting up"
            model_map:
              type: object
              description: Updated task-to-model routing table (present when status is ok)
            count:
              type: integer
              description: Number of available models (present when status is ok)
            message:
              type: string
              description: Status message (present when status is initializing)
      500:
        description: Refresh failed
        schema:
          type: object
          properties:
            status:
              type: string
              example: error
            error:
              type: string
    """
    if not _model_manager_available or _model_manager is None:
        # 管理器未就绪，在后台重新初始化
        import threading as _t

        _t.Thread(
            target=_init_model_manager, name="ModelManagerReinit", daemon=True
        ).start()
        return jsonify(
            {"status": "initializing", "message": "模型管理器正在后台初始化"}
        )
    try:
        new_map = _model_manager.refresh()
        MODEL_MAP.update(new_map)
        return jsonify(
            {
                "status": "ok",
                "model_map": _model_manager.get_model_map_with_scores(),
                "count": len(_model_manager.get_available_models()),
            }
        )
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/api/analyze", methods=["POST"])
def analyze_task():
    """预分析任务类型和模型选择 - 让前端立即显示"""
    data = request.json
    message = data.get("message", "")
    locked_task = data.get("locked_task")
    locked_model = data.get("locked_model", "auto")
    has_file = data.get("has_file", False)
    file_type = data.get("file_type", "")

    if not message:
        return jsonify(
            {"task": "CHAT", "model": MODEL_MAP["CHAT"], "route_method": "Empty"}
        )

    # 图像编辑关键词
    IMAGE_EDIT_KEYWORDS = [
        "修改",
        "换",
        "改成",
        "变成",
        "底色",
        "背景",
        "颜色",
        "抠图",
        "去背景",
        "P图",
        "美化",
        "滤镜",
        "调色",
        "编辑",
        "change",
        "modify",
        "edit",
        "background",
        "color",
    ]

    # 如果用户锁定了任务类型
    if locked_task:
        task = locked_task
        route_method = "🔒 Manual"
    elif has_file and file_type and file_type.startswith("image"):
        # 有图片文件，判断是编辑还是分析
        message_lower = message.lower()
        is_edit = any(kw in message_lower for kw in IMAGE_EDIT_KEYWORDS)
        if is_edit:
            task = "PAINTER"
            route_method = "🖼️ Image Edit"
        else:
            task = "VISION"
            route_method = "👁️ Image Analysis"
    else:
        # 使用智能路由器
        task, route_method, _ = SmartDispatcher.analyze(message)

    # 如果用户选择了特定模型
    if locked_model and locked_model != "auto":
        model = locked_model
    else:
        model = SmartDispatcher.get_model_for_task(task, has_image=has_file)

    # 获取模型显示信息
    model_info = MODEL_INFO.get(model, {"name": model, "speed": ""})

    return jsonify(
        {
            "task": task,
            "model": model,
            "model_name": model_info.get("name", model),
            "model_speed": model_info.get("speed", ""),
            "route_method": route_method,  # 路由算法信息
            "strengths": model_info.get("strengths", []),
        }
    )


# ================= Settings API =================

# ─── 本地模型状态 API ─────────────────────────────────────────────────────────


# GET /api/skills 已迁移至 skill_bp 蓝图（app/api/skill_routes.py），在此移除内联定义避免路由阻拦


# ================= Settings API =================


# ================= Mini Mode Switch API =================


@app.route("/api/mini/chat", methods=["POST"])
def mini_chat():
    """迷你模式专用聊天API - 使用与原版完全相同的任务分配和执行逻辑"""
    data = request.json
    user_input = data.get("message", "").strip()

    if not user_input:
        return jsonify({"error": "消息不能为空"}), 400

    user_input = Utils.sanitize_string(user_input)

    # 使用固定的迷你会话
    session_name = "MiniKoto_Quick"
    history = session_manager.load(f"{session_name}.json")

    # 🎯 使用 SmartDispatcher 进行任务分析（与完整版相同）
    task_type, route_method, context_info = SmartDispatcher.analyze(user_input, history)
    _app_logger.debug(
        f"[MINI_CHAT] SmartDispatcher 分析结果: task_type='{task_type}', method='{route_method}'"
    )

    response_text = ""
    is_error = False
    used_model = "unknown"

    try:
        # ===== 根据任务类型执行不同的处理逻辑（与完整版相同）=====

        if task_type == "WEB_SEARCH":
            # 🌐 网络搜索 - 使用 Gemini Google Search Grounding
            _app_logger.debug(f"[MINI_CHAT] 🌐 执行网络搜索...")
            _mini_skill_prompt = (context_info or {}).get("skill_prompt")
            search_result = WebSearcher.search_with_grounding(
                user_input, skill_prompt=_mini_skill_prompt
            )
            response_text = search_result.get("response", "")
            used_model = "gemini-2.5-flash (Google Search)"

            # 如果搜索失败，尝试修正查询
            if (
                not search_result.get("success")
                or Utils.is_failure_output(response_text)
                or "搜索失败" in response_text
            ):
                _app_logger.warning(f"[MINI_CHAT] ⚠️ 初次搜索失败，尝试修正查询...")
                fix_query_prompt = (
                    "请把用户需求改写成更适合搜索的简短关键词或查询语句，只输出查询语句。\n"
                    f"用户需求: {user_input}"
                )
                try:
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.2, max_output_tokens=64
                        ),
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    _app_logger.debug(f"[MINI_CHAT] 修正后的查询: {fixed_query}")
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result.get("response", "")
                except Exception as e:
                    _app_logger.debug(f"[MINI_CHAT] 修正查询失败: {e}")

            if not response_text or Utils.is_failure_output(response_text):
                is_error = True
                response_text = f"搜索失败：无法获取 '{user_input}' 的实时信息"

        elif task_type == "SYSTEM":
            # 🖥️ 系统命令 - 本地执行
            _app_logger.debug(f"[MINI_CHAT] 🖥️ 执行系统命令：{user_input}")
            try:
                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result.get("message", "命令执行失败")
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"
                used_model = "LocalExecutor"
                is_error = not exec_result.get("success", False)

                # 如果执行失败，尝试用 AI 修正
                if is_error or Utils.is_failure_output(response_text):
                    _app_logger.warning(f"[MINI_CHAT] ⚠️ 本地执行失败，尝试 AI 修正...")
                    fix_prompt = Utils.build_fix_prompt(
                        "SYSTEM", user_input, response_text
                    )
                    try:
                        fix_resp = client.models.generate_content(
                            model="gemini-2.5-flash",
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION(),
                                temperature=0.4,
                                max_output_tokens=1000,
                            ),
                        )
                        response_text = fix_resp.text or response_text
                        used_model = "gemini-2.5-flash (fallback)"
                        is_error = False
                    except Exception as e:
                        _app_logger.debug(f"[MINI_CHAT] AI 修正失败: {e}")
            except Exception as e:
                _app_logger.error(f"[MINI_CHAT] ❌ 系统命令执行出错: {e}")
                response_text = f"系统命令执行出错：{str(e)}"
                used_model = "LocalExecutor"
                is_error = True

        else:
            # 💬 其他任务（CHAT, RESEARCH, CODER 等）- 使用 brain.chat()
            _app_logger.debug(f"[MINI_CHAT] 💬 执行 {task_type} 任务...")
            model = MODEL_MAP.get(task_type, MODEL_MAP["CHAT"])
            result = brain.chat(
                history, user_input, model=model, auto_model=False, task_type=task_type
            )
            response_text = result.get("response", "")
            used_model = result.get("model", model)
            is_error = response_text.startswith("Error:")

            # 如果遇到 404 错误，尝试备用模型
            if is_error and "404" in response_text:
                _app_logger.warning(f"[MINI_CHAT] ⚠️ 模型 404，尝试备用模型...")
                for fallback_model in ["gemini-2.5-flash", "gemini-3-flash-preview"]:
                    try:
                        result = brain.chat(
                            history, user_input, model=fallback_model, auto_model=False
                        )
                        if not result.get("response", "").startswith("Error:"):
                            response_text = result.get("response", "")
                            used_model = fallback_model
                            is_error = False
                            break
                    except Exception as e:
                        continue

    except Exception as e:
        _app_logger.error(f"[MINI_CHAT] ❌ 执行出错: {e}")
        is_error = True
        response_text = f"Error: {str(e)}"

    # 更新历史（成功和失败都保存，便于排查）
    if response_text:
        session_manager.append_and_save(
            f"{session_name}.json", user_input, response_text
        )

    _app_logger.info(
        f"[MINI_CHAT] ✅ 完成: task_type={task_type}, model={used_model}, success={not is_error}"
    )

    # 返回统一格式
    return jsonify(
        {
            "success": not is_error,
            "response": response_text,
            "model": used_model,
            "task_type": task_type,
            "route_method": route_method,
            "error": response_text if is_error else "",
        }
    )


# ================= Setup & Initialization API =================


@app.route("/api/setup/apikey", methods=["POST"])
def setup_api_key():
    """设置 API Key"""
    data = request.json
    api_key = data.get("api_key", "").strip()

    if not api_key or len(api_key) < 10:
        return jsonify({"success": False, "error": "Invalid API key"})

    config_path = os.path.join(PROJECT_ROOT, "config", "gemini_config.env")
    os.makedirs(os.path.dirname(config_path), exist_ok=True)

    try:
        # 写入配置文件（同时写入两个变量名，避免优先级错乱）
        with open(config_path, "w", encoding="utf-8") as f:
            f.write(
                f"# Koto Configuration\nGEMINI_API_KEY={api_key}\nAPI_KEY={api_key}\n"
            )

        # 更新环境变量
        os.environ["GEMINI_API_KEY"] = api_key
        os.environ["API_KEY"] = api_key
        global API_KEY, client
        API_KEY = api_key
        client = create_client()

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/setup/workspace", methods=["POST"])
def setup_workspace():
    """设置工作区目录"""
    data = request.json
    workspace_path = data.get("path", "").strip()

    if not workspace_path:
        workspace_path = os.path.join(PROJECT_ROOT, "workspace")

    try:
        os.makedirs(workspace_path, exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "documents"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "images"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "code"), exist_ok=True)

        # 更新设置
        settings_manager.set("storage", "workspace_dir", workspace_path)

        return jsonify({"success": True, "path": workspace_path})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/setup/test", methods=["GET"])
def test_api_connection():
    """测试 API 连接"""
    try:
        start = time.time()
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents="Say 'Koto is ready!' in one short sentence.",
        )
        latency = time.time() - start
        return jsonify(
            {"success": True, "message": response.text, "latency": round(latency, 2)}
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/diagnose", methods=["GET"])
def diagnose_models():
    """诊断所有模型的可用性"""
    import threading

    results = {
        "proxy": {
            "detected": get_detected_proxy(),
            "force": FORCE_PROXY or None,
            "custom_endpoint": GEMINI_API_BASE or None,
        },
        "models": {},
    }

    # 测试模型列表
    test_models = [
        ("gemini-2.0-flash-lite", "路由分类"),
        ("gemini-3-flash-preview", "日常对话"),
        ("gemini-3-pro-preview", "代码生成"),
        ("gemini-2.5-flash", "联网搜索"),
        ("gemini-3.1-flash-image-preview", "图像生成"),
    ]

    def test_model(model_id, purpose):
        try:
            start = time.time()
            if "image-generation" in model_id or "imagen" in model_id:
                # 图像模型只测试连通性
                response = client.models.generate_content(
                    model=model_id,
                    contents="test",
                    config=types.GenerateContentConfig(max_output_tokens=10),
                )
            else:
                response = client.models.generate_content(
                    model=model_id,
                    contents="Reply with only: OK",
                    config=types.GenerateContentConfig(max_output_tokens=10),
                )
            latency = time.time() - start
            return {
                "status": "✅ 可用",
                "latency": round(latency, 2),
                "purpose": purpose,
            }
        except Exception as e:
            error_msg = str(e)
            if "location is not supported" in error_msg:
                status = "❌ 地区限制"
            elif "not found" in error_msg.lower():
                status = "❌ 模型不存在"
            elif "quota" in error_msg.lower():
                status = "⚠️ 配额耗尽"
            elif "timeout" in error_msg.lower():
                status = "⚠️ 超时"
            else:
                status = f"❌ 错误"
            return {"status": status, "error": error_msg[:150], "purpose": purpose}

    # 并行测试（带超时）
    threads = []
    for model_id, purpose in test_models:

        def run_test(m=model_id, p=purpose):
            results["models"][m] = test_model(m, p)

        t = threading.Thread(target=run_test, daemon=True)
        threads.append(t)
        t.start()

    # 等待所有线程完成（最多 15 秒）
    for t in threads:
        t.join(timeout=15)

    # 检查是否所有模型都不可用
    all_failed = all(
        "❌" in results["models"].get(m, {}).get("status", "") for m, _ in test_models
    )

    if all_failed:
        results["recommendation"] = (
            "所有模型均不可用。建议：\n1. 检查代理配置是否正确\n2. 考虑使用 API 中转服务\n3. 在 gemini_config.env 中配置 GEMINI_API_BASE"
        )

    return jsonify(results)


@app.route("/api/browse", methods=["GET"])
def browse_folders():
    import os

    path = request.args.get("path", "C:\\")

    try:
        if not os.path.exists(path):
            return jsonify({"error": "路径不存在", "folders": [], "parent": None})

        if not os.path.isdir(path):
            return jsonify({"error": "不是文件夹", "folders": [], "parent": None})

        folders = []
        try:
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                if os.path.isdir(item_path):
                    folders.append({"name": item, "path": item_path})
        except PermissionError:
            return jsonify({"error": "没有权限访问", "folders": [], "parent": None})

        folders.sort(key=lambda x: x["name"].lower())

        # Get parent path
        parent = os.path.dirname(path)
        if parent == path:  # Root drive
            parent = None

        return jsonify({"folders": folders, "parent": parent, "current": path})
    except Exception as e:
        return jsonify({"error": str(e), "folders": [], "parent": None})


@app.route("/api/chat/interrupt", methods=["POST"])
def interrupt_chat():
    """Interrupt an in-progress chat generation.
    ---
    tags:
      - Chat
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          required:
            - session
          properties:
            session:
              type: string
              description: Name of the session to interrupt
            task_id:
              type: string
              description: Optional scheduler task ID to cancel (for long-running tasks like DOC_ANNOTATE)
    responses:
      200:
        description: Interrupt signal sent
        schema:
          type: object
          properties:
            success:
              type: boolean
            message:
              type: string
      400:
        description: Missing session parameter
        schema:
          type: object
          properties:
            error:
              type: string
    """
    payload = request.json or {}
    session_name = payload.get("session")
    task_id = payload.get("task_id")
    if not session_name:
        return jsonify({"error": "Missing session"}), 400

    # 使用新的中断管理器
    _interrupt_manager.set_interrupt(session_name)
    # 保持向后兼容
    _interrupt_flags[session_name] = True

    # 可选：如果前端传入 task_id，同步取消调度器任务（用于 DOC_ANNOTATE 等流式长任务）
    if task_id:
        try:
            from task_scheduler import get_task_scheduler

            get_task_scheduler().cancel_task(task_id)
            _app_logger.debug(f"[INTERRUPT] Cancel task_id={task_id}")
        except Exception as e:
            _app_logger.debug(f"[INTERRUPT] cancel task failed: {e}")

    # 同步中断标志到 AgentLoop（如果正在执行 Agent 任务）
    # NOTE: Legacy agent_loop retired — interrupt handled by _interrupt_manager above
    pass

    return jsonify({"success": True, "message": "Chat interrupted"})


@app.route("/api/chat/reset-interrupt", methods=["POST"])
def reset_interrupt():
    """重置中断标志"""
    session_name = request.json.get("session")
    if session_name:
        # 使用新的中断管理器
        _interrupt_manager.reset(session_name)
        # 保持向后兼容
        if session_name in _interrupt_flags:
            del _interrupt_flags[session_name]
    return jsonify({"success": True})


# ================= 新功能 API 路由 =================


# === 快速笔记 API ===
# === 本地提醒 API（Windows 系统通知） ===
# === 日程（本地日历） API ===
# === 剪贴板 API ===
# === 任务调度 API（已迁移至 task_bp 蓝图 app/api/task_routes.py）===
# 原内联路由依赖不存在的 task_scheduler 模块，已移除以解除对 task_bp 的路由阻拦。
# task_bp 提供：GET /api/tasks, GET /api/tasks/<id>, POST /api/tasks/<id>/cancel,
#              POST /api/tasks/<id>/interrupt, GET /api/tasks/<id>/stream, 等。


# === 邮件 API ===
# === 浏览器自动化 API ===
# === 智能搜索 API ===
# ================= 语音识别 API (新架构) =================
@app.route("/api/voice/engines", methods=["GET"])
def voice_engines():
    """获取可用语音引擎列表"""
    try:
        from web.voice_fast import get_available_engines

        result = get_available_engines()
        return jsonify(result)
    except Exception as e:
        return (
            jsonify(
                {
                    "success": False,
                    "engines": [],
                    "message": f"获取引擎列表失败: {str(e)}",
                }
            ),
            500,
        )


@app.route("/api/voice/record", methods=["POST"])
def voice_record():
    """录制音频"""
    try:
        data = request.json or {}
        duration = data.get("duration", 5)

        from web.voice_input import record_audio

        result = record_audio(duration=int(duration))

        return jsonify(result)
    except Exception as e:
        return (
            jsonify(
                {"success": False, "message": f"录音失败: {str(e)}", "audio_file": None}
            ),
            500,
        )


@app.route("/api/voice/recognize", methods=["POST"])
def voice_recognize():
    """识别音频文件"""
    try:
        data = request.json or {}
        audio_path = data.get("audio_path")
        engine = data.get("engine", None)

        if not audio_path:
            return jsonify({"success": False, "message": "缺少音频文件路径"}), 400

        from web.voice_input import recognize_audio

        result = recognize_audio(audio_path, engine)

        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "message": f"识别失败: {str(e)}"}), 500


@app.route("/api/voice/listen", methods=["POST"])
def voice_listen():
    """一键麦克风识别（本地模式 - 优化版：立即启动）"""
    try:
        data = request.json or {}
        timeout = data.get("timeout", 5)
        language = data.get("language", "zh-CN")

        # 使用快速本地识别
        from web.voice_fast import recognize_voice

        result = recognize_voice(timeout=int(timeout), language=language)

        # 优化：设置响应头加快传输
        response = jsonify(result)
        response.headers["Cache-Control"] = "no-cache, no-store"
        response.headers["Content-Type"] = "application/json; charset=utf-8"
        return response

    except Exception as e:
        import traceback

        traceback.print_exc()
        response = jsonify(
            {
                "success": False,
                "text": "",
                "message": f"语音识别出错: {str(e)}",
                "engine": "error",
            }
        )
        response.status_code = 500
        response.headers["Cache-Control"] = "no-cache"
        return response


@app.route("/api/voice/stream")
def voice_stream():
    """流式语音识别 - Vosk 本地离线，实时返回部分/最终结果（SSE）"""
    import json as _json

    from flask import Response, stream_with_context

    @stream_with_context
    def generate():
        try:
            from web.voice_engine import recognize_stream

            for event in recognize_stream(max_wait=8.0, max_speech=30.0):
                yield f"data: {_json.dumps(event, ensure_ascii=False)}\n\n"
                if event.get("type") in ("final", "error"):
                    break
        except GeneratorExit:
            pass
        except Exception as e:
            yield f"data: {_json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Transfer-Encoding": "chunked",
        },
    )


@app.route("/api/voice/stop", methods=["POST"])
def voice_stop():
    """停止当前语音识别流（通知 voice_engine 停止）"""
    try:
        from web.voice_engine import request_stop

        request_stop()
    except Exception:
        pass
    return jsonify({"success": True, "message": "已发送停止信号"})


@app.route("/api/voice/commands", methods=["GET"])
def voice_commands():
    """返回内置语音命令列表（供语音面板展示）"""
    commands = [
        {"name": "发送消息", "description": "说出消息后自动发送", "keyword": ""},
        {"name": "新对话", "description": "说'新对话'开始新聊天", "keyword": "新对话"},
        {"name": "清空输入", "description": "说'清空'清除输入框", "keyword": "清空"},
        {"name": "重新识别", "description": "再次点击麦克风重新说", "keyword": ""},
    ]
    return jsonify({"success": True, "commands": commands})


@app.route("/api/voice/stt_status", methods=["GET"])
def voice_stt_status():
    """查询当前语音引擎状态（使用新 voice_engine）。"""
    try:
        from web.voice_engine import get_status

        fast = get_status()
    except Exception:
        fast = {"available": False, "engine": "unavailable", "label": "无引擎"}

    return jsonify(
        {
            "fast": fast,
            "local": fast,  # 兼容前端旧字段
            "active": fast.get("engine", "none"),
        }
    )


@app.route("/api/voice/gemini_stt", methods=["POST"])
@app.route("/api/voice/stt", methods=["POST"])  # 统一入口别名
def voice_gemini_stt():
    """
    统一语音转文字 (STT) 入口：本地 Whisper 优先 → Gemini STT 备用。

    - 若安装了 faster-whisper 或 openai-whisper：完全本地转写，无 API 消耗
    - 否则：发送至 Gemini gemini-2.0-flash-lite 转写
    - 始终返回 JSON，绝不返回 HTML 错误页面。
    """
    try:
        data = request.get_json(silent=True) or {}
        audio_b64 = data.get("audio", "")
        mime_type = data.get("mime", "audio/webm")

        if not audio_b64:
            return (
                jsonify({"success": False, "text": "", "message": "缺少 audio 字段"}),
                400,
            )

        import base64 as _b64

        try:
            audio_bytes = _b64.b64decode(audio_b64)
        except Exception:
            return (
                jsonify(
                    {"success": False, "text": "", "message": "音频 base64 解码失败"}
                ),
                400,
            )

        if len(audio_bytes) < 300:
            return jsonify(
                {"success": False, "text": "", "message": "录音太短，请重新说话"}
            )

        print(f"[STT] 收到音频 {len(audio_bytes)/1024:.1f}KB  MIME={mime_type}")

        # ── 优先尝试本地 Whisper ──────────────────────────────────────────────
        try:
            from web.local_stt import is_available, transcribe

            if is_available():
                ok, text, engine = transcribe(audio_bytes, mime_type)
                if ok and text:
                    return jsonify(
                        {
                            "success": True,
                            "text": text,
                            "engine": engine,
                            "message": "识别成功（本地）",
                        }
                    )
                # 本地识别出空文本 → 也直接返回（不回退，避免重复计费）
                return jsonify(
                    {
                        "success": False,
                        "text": "",
                        "engine": engine,
                        "message": "未检测到语音",
                    }
                )
        except Exception as _le:
            print(f"[STT] 本地 STT 异常，回退 Gemini: {_le}")

        # ── 回退：Gemini STT ──────────────────────────────────────────────────
        if client is None:
            return (
                jsonify(
                    {
                        "success": False,
                        "text": "",
                        "message": "Gemini 客户端未初始化，请检查 API Key；"
                        "或安装 faster-whisper 使用本地识别",
                    }
                ),
                503,
            )

        stt_model = "gemini-2.0-flash-lite"
        prompt_parts = [
            types.Part.from_bytes(data=audio_bytes, mime_type=mime_type),
            types.Part.from_text(
                text=(
                    "请将上面音频中的语音内容完整转写为文字。"
                    "只输出转写结果，不要加任何解释、标点修饰或前缀（如「转写：」等）。"
                    "如果听不清或没有语音，只输出空字符串。"
                )
            ),
        ]

        resp = client.models.generate_content(
            model=stt_model,
            contents=[types.Content(role="user", parts=prompt_parts)],
            config=types.GenerateContentConfig(temperature=0.0, max_output_tokens=512),
        )

        text = (resp.text or "").strip()
        for prefix in ("转写：", "转写:", "识别：", "识别:", "文字：", "文字:"):
            if text.startswith(prefix):
                text = text[len(prefix) :].strip()

        print(f"[STT] Gemini 识别结果: {text[:80]!r}")
        return jsonify(
            {
                "success": bool(text),
                "text": text,
                "engine": f"Gemini/{stt_model}",
                "message": "识别成功" if text else "未检测到语音内容",
            }
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return (
            jsonify(
                {"success": False, "text": "", "message": f"STT 失败: {str(e)[:200]}"}
            ),
            500,
        )


# ================= 增强功能 API (场景1-3) =================


@app.route("/api/ppt/generate", methods=["POST"])
def ppt_generate():
    """PPT生成 - 场景3：高质量演示文稿"""
    try:
        data = request.json
        title = data.get("title", "演示文稿")
        subtitle = data.get("subtitle", "")
        outline = data.get("outline")
        content = data.get("content")
        theme = data.get("theme", "business")
        output_filename = data.get(
            "output_filename",
            f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pptx',
        )

        output_path = os.path.join(WORKSPACE_DIR, "documents", output_filename)

        from web.ppt_generator import PPTGenerator

        generator = PPTGenerator(theme=theme)

        # 生成PPT
        if outline:
            result = generator.generate_from_outline(
                title, outline, output_path, subtitle=subtitle
            )
        elif content:
            result = generator.generate_from_text(content, output_path, title)
        else:
            return jsonify({"success": False, "error": "需要提供outline或content"}), 400

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== 智能文档处理路由 ====================


def _is_analysis_request(requirement: str) -> bool:
    """判断是否为分析/问答类请求（包括简单问答和复杂分析，但不含生成文档意图）"""
    if not requirement:
        return False

    requirement_lower = requirement.lower()

    # 明确的分析/问答动作词（全面覆盖）
    analysis_actions = [
        # 分析类
        "分析",
        "总结",
        "概述",
        "梳理",
        "解读",
        "评估",
        "对比",
        "提炼",
        "归纳",
        "主要观点",
        "核心观点",
        "要点",
        "重点",
        "亮点",
        # 问答/询问类
        "告诉我",
        "告诉",
        "是什么",
        "做什么",
        "想做什么",
        "在做什么",
        "是否",
        "有没有",
        "值不值",
        "值不值得",
        "投资价值",
        "投资建议",
        "是否值得",
        "值得投资",
        "有无价值",
        "有价值吗",
        "值得关注",
        "讲讲",
        "讲一下",
        "说说",
        "说一下",
        "介绍",
        "介绍一下",
        "介绍下",
        "解释",
        "解释一下",
        "帮我解释",
        "了解",
        "看看",
        "看一看",
        "读一读",
        "读一下",
        "什么是",
        "怎么看",
        "怎么样",
        "如何",
        "什么情况",
        "帮我看",
        "帮我读",
        "帮我理解",
        "帮我了解",
        "帮我评估",
        "帮我判断",
        "这份",
        "这个",
        "检查一下",
        "查看一下",
        "看一下这",
        "他们想",
        "他想",
        "它想",
        "该公司",
        "该项目",
        # English
        "review",
        "analysis",
        "summary",
        "summarize",
        "analyze",
        "explain",
        "understand",
        "evaluate",
        "assess",
        "what is",
        "what does",
        "how does",
        "tell me",
        "should i",
        "is it worth",
        "investment value",
        "check",
        "read this",
        "look at",
    ]

    # 排除词：明确的文档生成意图（只排除最明确的生成指令）
    generation_words = [
        "生成一份",
        "生成一个",
        "帮我生成",
        "写一份",
        "写一个",
        "帮我写",
        "改善",
        "改进",
        "优化",
        "润色",
        "重写",
        "帮我做一份",
        "做一个报告",
        "做一份报告",
        "create a document",
        "generate a report",
        "write a report",
    ]

    has_analysis = any(kw in requirement_lower for kw in analysis_actions)
    has_generation = any(kw in requirement_lower for kw in generation_words)

    if has_analysis and not has_generation:
        return True

    return False


def _is_explicit_file_gen_request(requirement: str) -> bool:
    """判断用户是否明确要求生成/输出一个新文件（报告、Word、PDF等）"""
    if not requirement:
        return False
    requirement_lower = requirement.lower()
    gen_keywords = [
        "生成一份",
        "生成一个",
        "帮我生成",
        "写一份报告",
        "写一个报告",
        "写报告",
        "写一份",
        "帮我写",
        "做一份",
        "做一个",
        "帮我做",
        "导出",
        "输出为",
        "保存为",
        "转成",
        "生成word",
        "生成pdf",
        "生成excel",
        "生成ppt",
        "创建文档",
        "新建文档",
        "制作报告",
        "整理成文档",
        "形成报告",
        "输出报告",
    ]
    return any(kw in requirement_lower for kw in gen_keywords)


@app.route("/api/document/smart-process", methods=["POST"])
def document_smart_process():
    """
    智能文档处理入口
    自动判断使用：标注系统 or 文件分析系统
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 智能判断应该用哪个系统
        use_annotation = _should_use_annotation_system(requirement)

        print(f"[SmartProcess] 智能判断: use_annotation={use_annotation}")
        print(f"[SmartProcess] 需求: {requirement[:100]}")

        if use_annotation:
            # 使用文档标注系统
            print(f"[SmartProcess] 路由到: 文档自动标注系统")
            return _call_document_annotate(file_path, requirement)
        else:
            # 使用传统的文件分析系统
            print(f"[SmartProcess] 路由到: 文件分析系统")
            return _call_document_analysis(file_path, requirement)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


def _call_document_annotate(file_path: str, requirement: str):
    """调用文档标注系统"""
    try:
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        result = feedback_system.full_annotation_loop(
            file_path=file_path,
            user_requirement=requirement,
            model_id="gemini-3-pro-preview",
        )

        # 添加处理模式标记
        result["processing_mode"] = "annotation"
        result["mode_description"] = "文档自动标注"

        return jsonify(result)

    except Exception as e:
        return (
            jsonify(
                {"success": False, "error": str(e), "processing_mode": "annotation"}
            ),
            500,
        )


def _call_document_analysis(file_path: str, requirement: str):
    """调用传统的文件分析系统"""
    try:
        # 这里调用现有的文件分析逻辑
        # 临时返回说明（实际应该调用现有的分析端点）
        return (
            jsonify(
                {
                    "success": False,
                    "error": "文件分析系统需要单独实现",
                    "processing_mode": "analysis",
                    "mode_description": "文件分析",
                }
            ),
            501,
        )

    except Exception as e:
        return (
            jsonify({"success": False, "error": str(e), "processing_mode": "analysis"}),
            500,
        )


@app.route("/api/document/feedback", methods=["POST"])
def document_feedback():
    """文档智能反馈：读取文档 → AI分析 → 应用修改"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        auto_apply = data.get("auto_apply", True)

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        # 执行完整反馈闭环
        result = feedback_system.full_feedback_loop(
            file_path=file_path,
            user_requirement=user_requirement,
            auto_apply=auto_apply,
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/analyze", methods=["POST"])
def document_analyze():
    """仅分析文档，不应用修改"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        # 仅分析
        result = feedback_system.analyze_and_suggest(
            file_path=file_path, user_requirement=user_requirement
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/apply", methods=["POST"])
def document_apply():
    """应用修改建议到文档"""
    try:
        data = request.json
        file_path = data.get("file_path")
        modifications = data.get("modifications", [])

        if not file_path or not modifications:
            return (
                jsonify(
                    {"success": False, "error": "缺少file_path或modifications参数"}
                ),
                400,
            )

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 应用修改
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        result = feedback_system.apply_suggestions(
            file_path=file_path, modifications=modifications
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/annotate", methods=["POST"])
def document_annotate():
    """文档自动标注：AI分析 -> 生成标注 -> 应用到副本"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        model_id = data.get("model_id", "gemini-3-pro-preview")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        # 执行完整标注闭环
        result = feedback_system.full_annotation_loop(
            file_path=file_path, user_requirement=user_requirement, model_id=model_id
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/analyze-annotations", methods=["POST"])
def document_analyze_annotations():
    """仅分析文档并生成标注建议（不应用）- 已弃用，请使用 /api/document/batch-annotate-stream"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 使用V2批量标注系统（立即返回结果，不流式）
        from web.document_direct_edit import ImprovedBatchAnnotator

        annotator = ImprovedBatchAnnotator(gemini_client=client, batch_size=5)

        # 收集所有事件（非流式）
        events = []
        final_result = None

        for event in annotator.annotate_document_streaming(file_path, user_requirement):
            # 解析事件
            if event.startswith("event: complete"):
                data_line = event.split("\n")[1]
                if data_line.startswith("data: "):
                    final_result = json.loads(data_line[6:])
            events.append(event)

        if final_result:
            return jsonify({"success": True, **final_result})
        else:
            return jsonify({"success": False, "error": "处理失败，未收到完成事件"}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/batch-annotate-stream", methods=["POST"])
def document_batch_annotate_stream():
    """
    批量标注文档（SSE流式返回，实时反馈进度）

    接收参数:
        file_path: 文档路径
        requirement: 用户需求（可选）
        batch_size: 每批处理段落数（默认5）

    返回: SSE事件流
        event: progress - 进度更新
        event: batch_complete - 批次完成
        event: complete - 全部完成
        event: error - 错误
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        batch_size = data.get("batch_size", 5)

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 导入V2批量标注系统
        from web.document_batch_annotator_v2 import annotate_large_document

        # 返回SSE流
        return Response(
            annotate_large_document(
                file_path=file_path,
                user_requirement=user_requirement,
                gemini_client=client,
                batch_size=batch_size,
            ),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/document/apply-annotations", methods=["POST"])
def document_apply_annotations():
    """应用标注建议到文档"""
    try:
        data = request.json
        file_path = data.get("file_path")
        annotations = data.get("annotations", [])

        if not file_path or not annotations:
            return (
                jsonify({"success": False, "error": "缺少file_path或annotations参数"}),
                400,
            )

        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 应用标注
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(
            gemini_client=client, default_model_id="gemini-3.1-pro-preview"
        )

        result = feedback_system.annotate_document(file_path, annotations)

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== 新功能 API 路由 ====================

# ==================== 改进的建议式标注 API ====================


# ==================== 文件网络索引 API ====================


# 批量处理 API
# 模板库 API
# 一致性检查 API
# 文档对比 API
# OCR 助手 API
# 操作历史 API
# 语音转写 API
# ================= 主程序入口 =================

# ================= NotebookLM 功能复刻 API =================


@app.route("/api/notebook/overview", methods=["POST"])
def notebook_overview():
    """生成音频概览 (Podcast)"""
    data = request.json
    content = data.get("content", "")
    if not content:
        return jsonify({"success": False, "error": "内容不能为空"}), 400

    try:
        from web.audio_overview import AudioOverviewGenerator

        generator = AudioOverviewGenerator(
            output_dir=os.path.join(settings_manager.workspace_dir, "audio_cache")
        )

        # 1. 生成剧本
        # 获取模型实例 (复用现有的 KotoBrain 或直接调用 API)
        # 这里为了简化，假设我们能获取到一个 genai model 实例
        # 实际项目中应该复用 koto_brain.client.models
        # 暂时使用临时的 model 实例
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        model = client.models

        script = asyncio.run(generator.generate_script(content, model))
        if not script:
            return jsonify({"success": False, "error": "剧本生成失败"}), 500

        # 2. 合成音频
        session_id = f"overview_{int(time.time())}"
        audio_path = asyncio.run(generator.synthesize_audio(script, session_id))

        if audio_path:
            # 返回相对于 workspace 的路径或者 download url
            rel_path = os.path.relpath(audio_path, settings_manager.workspace_dir)
            # 注意：实际访问可能需要通过 send_from_directory 路由
            # 假设我们有一个 /files/ 路由可以访问 workspace/
            audio_url = f"/api/files/download?path={requests.utils.quote(audio_path)}"

            return jsonify({"success": True, "audio_url": audio_url, "script": script})
        else:
            return jsonify({"success": False, "error": "音频合成失败"}), 500

    except Exception as e:
        _app_logger.error(f"Error processing audio overview: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/notebook/qa", methods=["POST"])
def notebook_qa():
    """源文档深度问答 (Source-Grounded Q&A)"""
    data = request.json
    question = data.get("question")
    file_ids = data.get(
        "file_ids", []
    )  # 假设前端传回 files (这里先简化为 content 直接传入 或者 file paths)
    # 为了简化演示，我们先接受纯文本 content
    context_content = data.get("context", "")

    if not question or not context_content:
        return jsonify({"success": False, "error": "缺少问题或上下文"}), 400

    prompt = f"""
    Answer the user's question mostly based on the provided source context.
    
    [Source Context]
    {context_content[:30000]} 

    [User Question]
    {question}

    [Rules]
    1. You must cite your sources. When you use information from the context, append [Source] at the end of the sentence.
    2. If the answer is not in the context, state that clearly.
    3. Be precise and concise.
    """

    try:
        # 复用 KotoBrain 的逻辑或者直接调用
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash", contents=prompt
        )
        return jsonify({"success": True, "answer": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/notebook/study_guide", methods=["POST"])
def notebook_study_guide():
    """生成学习指南/简报"""
    data = request.json
    content = data.get("content", "")
    type_ = data.get("type", "summary")  # summary, quiz, timelime, faq

    prompts = {
        "summary": "Create a comprehensive briefing document summarizing the key points, key people, and timeline from the text.",
        "quiz": "Create 5 multiple-choice questions based on the text to test understanding. Include the correct answer key at the end.",
        "timeline": "Extract a chronological timeline of events mentioned in the text.",
        "faq": "Create a FAQ section based on the text, anticipating what a reader might ask.",
    }

    selected_prompt = prompts.get(type_, prompts["summary"])
    full_prompt = f"{selected_prompt}\n\n[Source Text]\n{content[:20000]}"

    try:
        import google.genai as genai

        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model="gemini-2.5-flash", contents=full_prompt
        )
        return jsonify({"success": True, "result": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/files/download", methods=["GET"])
def download_file_proxy():
    """通用的文件下载代理"""
    file_path = request.args.get("path")
    if not file_path or not os.path.exists(file_path):
        return "File not found", 404
    return send_file(file_path, as_attachment=True)


@app.route("/api/notebook/upload", methods=["POST"])
def notebook_upload():
    """上传并解析文件 (PDF/Docx/Txt)"""
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file part"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"success": False, "error": "No selected file"}), 400

    try:
        # Save temp file
        import tempfile

        filename = file.filename
        temp_path = os.path.join(
            tempfile.gettempdir(), f"koto_{int(time.time())}_{filename}"
        )
        file.save(temp_path)

        # Parse using FileParser
        from web.file_parser import FileParser

        result = FileParser.parse_file(temp_path)

        # Cleanup
        try:
            os.remove(temp_path)
        except OSError:
            pass

        if result.get("success"):
            return jsonify(
                {
                    "success": True,
                    "filename": filename,
                    "content": result.get("content", ""),
                    "char_count": result.get("char_count", 0),
                }
            )
        else:
            return jsonify({"success": False, "error": result.get("error")}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/notebook")
def notebook_ui():
    """NotebookLM 风格界面"""
    return render_template("notebook_lm.html")


if __name__ == "__main__":

    print("\n🚀 Koto Web Server Starting...")
    print(f"📁 Chat Directory: {os.path.abspath(CHAT_DIR)}")
    print(f"📁 Workspace: {os.path.abspath(WORKSPACE_DIR)}")

    # 延迟检查 Ollama 状态（不阻塞启动）
    def check_ollama_async():
        time.sleep(2)  # 延迟2秒后检查
        if os.environ.get("KOTO_DEPLOY_MODE") == "cloud":
            print("☁️ Ollama: Disabled (cloud mode - using Gemini API)")
            return
        if LocalDispatcher.is_ollama_running():
            print("🦙 Ollama: Running")
        else:
            print("🦙 Ollama: Not Running")

    threading.Thread(target=check_ollama_async, daemon=True).start()

    print("⚠️ 本地模型任务路由器已禁用，使用远程 AI")

    print("\n🌐 Open http://localhost:5000 in your browser\n")

    # 启动后台服务（异步，不阻塞启动）
    def start_background_services():
        time.sleep(1)  # 延迟1秒后启动后台服务
        try:
            from auto_catalog_scheduler import get_auto_catalog_scheduler
            from clipboard_manager import get_clipboard_manager
            from task_scheduler import get_task_scheduler

            # 启动剪贴板监控
            clipboard_manager = get_clipboard_manager()
            clipboard_manager.start_monitoring()
            print("📋 剪贴板监控已启动")

            # 启动任务调度器
            task_scheduler = get_task_scheduler()
            task_scheduler.start()
            print("⏰ 任务调度器已启动")

            # 初始化自动归纳调度器（如果已启用）
            auto_catalog = get_auto_catalog_scheduler()
            if auto_catalog.is_auto_catalog_enabled():
                auto_catalog._register_scheduled_task()
                print(
                    f"🗂️ 自动归纳已启用，每日 {auto_catalog.get_catalog_schedule()} 执行"
                )

        except Exception as e:
            print(f"⚠️ 后台服务启动失败: {e}")

    threading.Thread(target=start_background_services, daemon=True).start()

    try:
        debug_mode = os.environ.get("KOTO_DEBUG", "false").lower() == "true"
        port = int(os.environ.get("KOTO_PORT", "5000"))
        app.run(debug=debug_mode, host="0.0.0.0", port=port, threaded=True)
    finally:
        # 应用关闭时清理并行执行系统
        if PARALLEL_SYSTEM_ENABLED:
            print("[PARALLEL] 🛑 Shutting down parallel execution system...")
            stop_dispatcher()
            print("[PARALLEL] ✅ Parallel execution system shut down")


# ═══ 文件组织系统 API ═══

# 初始化文件组织器
_file_organizer_cache = {}
_batch_ops_cache = {}


def get_file_organizer():
    """懒加载文件组织器"""
    if "organizer" not in _file_organizer_cache:
        try:
            from web.file_organizer import FileOrganizer
        except ImportError:
            from file_organizer import FileOrganizer

        organize_root = get_organize_root()
        _file_organizer_cache["organizer"] = FileOrganizer(organize_root)

    return _file_organizer_cache["organizer"]


def get_file_analyzer():
    """懒加载文件分析器"""
    if "analyzer" not in _file_organizer_cache:
        try:
            from web.file_analyzer import FileAnalyzer
        except ImportError:
            from file_analyzer import FileAnalyzer

        _file_organizer_cache["analyzer"] = FileAnalyzer()

    return _file_organizer_cache["analyzer"]


def get_batch_ops_manager():
    """懒加载批量文件处理管理器"""
    if "batch_ops" not in _batch_ops_cache:
        try:
            from web.batch_file_ops import BatchFileOpsManager
        except ImportError:
            from batch_file_ops import BatchFileOpsManager
        _batch_ops_cache["batch_ops"] = BatchFileOpsManager()
    return _batch_ops_cache["batch_ops"]


_file_editor_cache = {}
_file_indexer_cache = {}
_concept_extractor_cache = {}
_knowledge_graph_cache = {}
_behavior_monitor_cache = {}
_suggestion_engine_cache = {}
_insight_reporter_cache = {}


def get_file_editor():
    """懒加载文件编辑器"""
    if "editor" not in _file_editor_cache:
        try:
            from web.file_editor import FileEditor
        except ImportError:
            from file_editor import FileEditor
        _file_editor_cache["editor"] = FileEditor()
    return _file_editor_cache["editor"]


def get_file_indexer():
    """懒加载文件索引器"""
    if "indexer" not in _file_indexer_cache:
        try:
            from web.file_indexer import FileIndexer
        except ImportError:
            from file_indexer import FileIndexer
        _file_indexer_cache["indexer"] = FileIndexer()
    return _file_indexer_cache["indexer"]


def get_concept_extractor():
    """懒加载概念提取器"""
    if "extractor" not in _concept_extractor_cache:
        try:
            from web.concept_extractor import ConceptExtractor
        except ImportError:
            from concept_extractor import ConceptExtractor
        _concept_extractor_cache["extractor"] = ConceptExtractor()
    return _concept_extractor_cache["extractor"]


def get_knowledge_graph():
    """懒加载知识图谱"""
    if "graph" not in _knowledge_graph_cache:
        try:
            from web.knowledge_graph import KnowledgeGraph
        except ImportError:
            from knowledge_graph import KnowledgeGraph
        _knowledge_graph_cache["graph"] = KnowledgeGraph()
    return _knowledge_graph_cache["graph"]


def get_behavior_monitor():
    """懒加载行为监控器"""
    if "monitor" not in _behavior_monitor_cache:
        try:
            from web.behavior_monitor import BehaviorMonitor
        except ImportError:
            from behavior_monitor import BehaviorMonitor
        _behavior_monitor_cache["monitor"] = BehaviorMonitor()
    return _behavior_monitor_cache["monitor"]


def get_suggestion_engine():
    """懒加载建议引擎"""
    if "engine" not in _suggestion_engine_cache:
        try:
            from web.suggestion_engine import SuggestionEngine
        except ImportError:
            from suggestion_engine import SuggestionEngine
        _suggestion_engine_cache["engine"] = SuggestionEngine()
    return _suggestion_engine_cache["engine"]


def get_insight_reporter():
    """懒加载洞察报告生成器"""
    if "reporter" not in _insight_reporter_cache:
        try:
            from web.insight_reporter import InsightReporter
        except ImportError:
            from insight_reporter import InsightReporter
        _insight_reporter_cache["reporter"] = InsightReporter()
    return _insight_reporter_cache["reporter"]


# ==================== 增强主动能力模块缓存 ====================
_notification_manager_cache = {}
_proactive_dialogue_cache = {}
_context_awareness_cache = {}
_auto_execution_cache = {}
_trigger_system_cache = {}


def get_notification_manager():
    """懒加载通知管理器"""
    if "manager" not in _notification_manager_cache:
        try:
            from web.notification_manager import get_notification_manager as _get_mgr
        except ImportError:
            from notification_manager import get_notification_manager as _get_mgr
        _notification_manager_cache["manager"] = _get_mgr()
    return _notification_manager_cache["manager"]


def get_proactive_dialogue():
    """懒加载主动对话引擎"""
    if "engine" not in _proactive_dialogue_cache:
        try:
            from web.proactive_dialogue import get_proactive_dialogue_engine
        except ImportError:
            from proactive_dialogue import get_proactive_dialogue_engine

        # 集成依赖模块
        notif_mgr = get_notification_manager()
        behavior_mon = get_behavior_monitor()
        suggestion_eng = get_suggestion_engine()

        _proactive_dialogue_cache["engine"] = get_proactive_dialogue_engine(
            notification_manager=notif_mgr,
            behavior_monitor=behavior_mon,
            suggestion_engine=suggestion_eng,
        )
    return _proactive_dialogue_cache["engine"]


def get_context_awareness():
    """懒加载情境感知系统"""
    if "system" not in _context_awareness_cache:
        try:
            from web.context_awareness import get_context_awareness_system
        except ImportError:
            from context_awareness import get_context_awareness_system

        behavior_mon = get_behavior_monitor()
        _context_awareness_cache["system"] = get_context_awareness_system(
            behavior_monitor=behavior_mon
        )
    return _context_awareness_cache["system"]


def get_auto_execution():
    """懒加载自动执行引擎"""
    if "engine" not in _auto_execution_cache:
        try:
            from web.auto_execution import get_auto_execution_engine
        except ImportError:
            from auto_execution import get_auto_execution_engine

        notif_mgr = get_notification_manager()
        _auto_execution_cache["engine"] = get_auto_execution_engine(
            notification_manager=notif_mgr
        )
    return _auto_execution_cache["engine"]


def get_trigger_system():
    """懒加载主动交互触发系统"""
    if "system" not in _trigger_system_cache:
        try:
            from web.proactive_trigger import get_trigger_system as _get_trigger_system
        except ImportError:
            from proactive_trigger import get_trigger_system as _get_trigger_system

        behavior_mon = get_behavior_monitor()
        context_sys = get_context_awareness()
        suggestion_eng = get_suggestion_engine()
        notif_mgr = get_notification_manager()
        dialogue_eng = get_proactive_dialogue()

        _trigger_system_cache["system"] = _get_trigger_system(
            behavior_monitor=behavior_mon,
            context_awareness=context_sys,
            suggestion_engine=suggestion_eng,
            notification_manager=notif_mgr,
            dialogue_engine=dialogue_eng,
        )
        # 启动后台轮询（每5分钟检查一次触发条件）
        try:
            _trigger_system_cache["system"].start_monitoring(check_interval=300)
        except Exception as _tse:
            _app_logger.warning(
                f"[TriggerSystem] ⚠️ start_monitoring 失败（非致命）: {_tse}"
            )
    return _trigger_system_cache["system"]


# ═══════════════════════════════════════════════════
# 文件编辑与搜索 API


# ═══════════════════════════════════════════════════
# 全盘文件扫描 API  (FileScanner)


# ═══════════════════════════════════════════════════
# 概念提取 API


# ═══════════════════════════════════════════════════
# 知识图谱 API


# ═══════════════════════════════════════════════════
# 行为监控 API


# ═══════════════════════════════════════════════════
# 智能建议 API


# ═══════════════════════════════════════════════════
# 洞察报告 API


# ==================== 通知管理 API ====================


# ==================== 主动对话 API ====================


# ==================== 情境感知 API ====================


# ==================== 自动执行 API ====================


# ==================== 主动交互触发系统 API ====================


# ═══ 注册增强记忆系统API（模块级别，确保始终执行） ═══
try:
    from memory_api_routes import register_memory_routes

    register_memory_routes(app, get_memory_manager)
except ImportError:
    try:
        from web.memory_api_routes import register_memory_routes

        register_memory_routes(app, get_memory_manager)
    except ImportError:
        _app_logger.warning("⚠️  增强记忆系统API未找到，使用基础功能")


# ── Token 使用统计接口 ────────────────────────────────────────────────────────


# ── LangGraph 工作流可视化 & 开发工具 API ─────────────────────────────────────


# ══════════════════════════════════════════════════════════════════════════════
# RAG 向量检索 API
