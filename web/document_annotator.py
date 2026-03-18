#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档自动标注系统 - 核心模块
工作流：读取 -> 分析 -> 定位 -> 注入

默认采用方案A（Word 原生批注气泡）：
- 在原文位置添加右侧批注
- 格式无损，Word 原生体验
- 失败时自动回退为红色标注
"""

import json
import logging
import os
import shutil
from copy import deepcopy
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


class DocumentAnnotator:
    """文档自动标注系统"""

    def __init__(self, min_similarity: float = 0.8, annotation_mode: str = "comment"):
        """
        Args:
            min_similarity: 模糊匹配的相似度阈值 (0-1)
            annotation_mode: 标注方式 ("comment" | "highlight")
        """
        self.min_similarity = min_similarity
        self.annotation_mode = annotation_mode

    # ==================== Step 1: 文件预处理 ====================

    @staticmethod
    def prepare_document(file_path: str) -> Tuple[str, str]:
        """
        Step 1: 用户输入与预处理

        动作：
        1. 在后台复制一份副本（命名为 _revised.docx）
        2. 所有操作都在副本上进行，确保原件安全

        Args:
            file_path: 原Word文件路径

        Returns:
            (原文件路径, 副本路径)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 生成副本路径
        base_name = os.path.splitext(file_path)[0]
        ext = os.path.splitext(file_path)[1]
        revised_path = f"{base_name}_revised{ext}"

        # 复制文件
        shutil.copy2(file_path, revised_path)
        logger.info(f"[Annotator] 📋 已创建工作副本: {os.path.basename(revised_path)}")

        return file_path, revised_path

    # ==================== Step 1.5: 文本提取 ====================

    @staticmethod
    def extract_text_from_word(file_path: str) -> Dict[str, Any]:
        """
        提取Word文档的纯文本内容（包括表格）+ 格式信息

        Returns:
            {
                "full_text": "完整文本",
                "full_text_with_format": "带格式标记的文本",  # 新增
                "paragraphs": [
                    {
                        "index": 0,
                        "text": "段落文本",
                        "text_with_format": "<bold>标题</bold>正文",  # 新增
                        "runs": [{"text": "x", "bold": True, "size": 14}],  # 新增
                        "start_pos": 0,
                        "end_pos": 15,
                        "source": "body" | "table",
                        "table_idx": 0,
                        "row_idx": 0,
                        "cell_idx": 0
                    },
                    ...
                ]
            }
        """
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs_info = []
            full_text = ""
            full_text_with_format = ""  # 新增：带格式的文本
            para_global_idx = 0

            # 提取正文段落
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text

                # 跳过空段落
                if not text.strip():
                    continue

                start_pos = len(full_text)
                full_text += text + "\n"
                end_pos = len(full_text)

                # 新增：提取runs的格式信息
                runs_info = []
                text_with_format = ""
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue

                    run_data = {
                        "text": run_text,
                        "bold": run.bold if run.bold is not None else False,
                        "italic": run.italic if run.italic is not None else False,
                        "underline": (
                            run.underline if run.underline is not None else False
                        ),
                    }

                    # 提取字体大小
                    if run.font.size:
                        run_data["size_pt"] = run.font.size.pt

                    # 提取字体颜色
                    if run.font.color and run.font.color.rgb:
                        run_data["color"] = str(run.font.color.rgb)

                    # 提取字体名称
                    if run.font.name:
                        run_data["font_name"] = run.font.name

                    runs_info.append(run_data)

                    # 构建带格式标记的文本
                    formatted = run_text
                    if run.bold:
                        formatted = f"<b>{formatted}</b>"
                    if run.italic:
                        formatted = f"<i>{formatted}</i>"
                    if run.underline:
                        formatted = f"<u>{formatted}</u>"
                    text_with_format += formatted

                full_text_with_format += text_with_format + "\n"

                paragraphs_info.append(
                    {
                        "index": para_global_idx,
                        "text": text,
                        "text_with_format": text_with_format,  # 新增
                        "runs": runs_info,  # 新增
                        "start_pos": start_pos,
                        "end_pos": end_pos,
                        "para_obj": para,  # 保留原对象，用于后续修改
                        "source": "body",
                        "body_idx": para_idx,
                    }
                )
                para_global_idx += 1

            # 提取表格段落（同样提取格式信息）
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for cell_para_idx, cell_para in enumerate(cell.paragraphs):
                            text = cell_para.text
                            if not text.strip():
                                continue

                            start_pos = len(full_text)
                            full_text += text + "\n"
                            end_pos = len(full_text)

                            # 提取表格内的格式信息
                            runs_info = []
                            text_with_format = ""
                            for run in cell_para.runs:
                                run_text = run.text
                                if not run_text:
                                    continue

                                run_data = {
                                    "text": run_text,
                                    "bold": run.bold if run.bold is not None else False,
                                    "italic": (
                                        run.italic if run.italic is not None else False
                                    ),
                                }

                                if run.font.size:
                                    run_data["size_pt"] = run.font.size.pt
                                if run.font.color and run.font.color.rgb:
                                    run_data["color"] = str(run.font.color.rgb)

                                runs_info.append(run_data)

                                formatted = run_text
                                if run.bold:
                                    formatted = f"<b>{formatted}</b>"
                                if run.italic:
                                    formatted = f"<i>{formatted}</i>"
                                text_with_format += formatted

                            full_text_with_format += text_with_format + "\n"

                            paragraphs_info.append(
                                {
                                    "index": para_global_idx,
                                    "text": text,
                                    "text_with_format": text_with_format,  # 新增
                                    "runs": runs_info,  # 新增
                                    "start_pos": start_pos,
                                    "end_pos": end_pos,
                                    "para_obj": cell_para,  # 保留原对象
                                    "source": "table",
                                    "table_idx": table_idx,
                                    "row_idx": row_idx,
                                    "cell_idx": cell_idx,
                                    "cell_para_idx": cell_para_idx,
                                }
                            )
                            para_global_idx += 1

            return {
                "success": True,
                "file_path": file_path,
                "full_text": full_text,
                "full_text_with_format": full_text_with_format,  # 新增：带格式的文本
                "paragraphs": paragraphs_info,
                "total_chars": len(full_text),
            }

        except Exception as e:
            return {"success": False, "error": f"提取文本失败: {str(e)}"}

    # ==================== Step 3: 锚点定位 ====================

    @staticmethod
    def locate_text_in_runs(para_obj, target_text: str) -> Optional[Dict[str, Any]]:
        """
        在段落的Run级别精确定位文本

        Returns:
            {
                "found": True,
                "start_run_index": 0,  # 起始Run索引
                "end_run_index": 2,    # 结束Run索引
                "start_char_offset": 5,  # 在起始Run中的字符偏移
                "end_char_offset": 10,   # 在结束Run中的字符偏移
                "matched_text": "..."
            }
        """
        # 构建段落的完整文本和Run映射
        runs = para_obj.runs
        if not runs:
            return None

        # 构建文本到Run的映射
        full_text = ""
        run_map = []  # [(run_index, start_pos, end_pos, run_obj)]

        for i, run in enumerate(runs):
            start = len(full_text)
            full_text += run.text
            end = len(full_text)
            run_map.append((i, start, end, run))

        # 查找目标文本
        pos = full_text.find(target_text)
        if pos == -1:
            return None

        target_end = pos + len(target_text)

        # 找到起始和结束Run
        start_run = None
        end_run = None
        start_offset = 0
        end_offset = 0

        for run_idx, start_pos, end_pos, run_obj in run_map:
            # 目标文本开始位置
            if start_run is None and start_pos <= pos < end_pos:
                start_run = run_idx
                start_offset = pos - start_pos

            # 目标文本结束位置
            if start_pos < target_end <= end_pos:
                end_run = run_idx
                end_offset = target_end - start_pos
                break

        if start_run is not None and end_run is not None:
            return {
                "found": True,
                "start_run_index": start_run,
                "end_run_index": end_run,
                "start_char_offset": start_offset,
                "end_char_offset": end_offset,
                "matched_text": target_text,
                "match_type": "precise",
            }

        return None

    @staticmethod
    def _set_run_text(run_element, text: str) -> None:
        """设置run元素文本，保留原有格式"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for t in run_element.findall(qn("w:t")):
            run_element.remove(t)
        t = OxmlElement("w:t")
        if text and (text[0].isspace() or text[-1].isspace()):
            t.set(qn("xml:space"), "preserve")
        t.text = text
        run_element.append(t)

    @staticmethod
    def _split_run_at(para_obj, run_index: int, offset: int) -> int:
        """在指定run内按字符偏移拆分，返回右侧新run的索引"""
        runs = para_obj.runs
        if run_index >= len(runs):
            return run_index

        run = runs[run_index]
        text = run.text or ""
        if offset <= 0 or offset >= len(text):
            return run_index

        left_text = text[:offset]
        right_text = text[offset:]

        # 更新当前run为左侧文本
        run.text = left_text

        # 克隆run并设置右侧文本
        new_run = deepcopy(run._element)
        DocumentAnnotator._set_run_text(new_run, right_text)
        run._element.addnext(new_run)

        return run_index + 1

    @staticmethod
    def _isolate_target_runs(
        para_obj, precise_location: Dict[str, Any], target_text: str
    ) -> Tuple[int, int]:
        """拆分run，确保目标文本独占run范围"""
        start_idx = precise_location["start_run_index"]
        end_idx = precise_location["end_run_index"]
        start_offset = precise_location["start_char_offset"]
        end_offset = precise_location["end_char_offset"]

        if start_idx == end_idx:
            # 单run内：先拆起点，再拆终点
            right_idx = DocumentAnnotator._split_run_at(
                para_obj, start_idx, start_offset
            )
            DocumentAnnotator._split_run_at(para_obj, right_idx, len(target_text))
            return right_idx, right_idx

        # 多run：拆起点与终点
        right_idx = DocumentAnnotator._split_run_at(para_obj, start_idx, start_offset)
        if right_idx != start_idx:
            end_idx += 1  # 起点拆分后，终点索引右移
        DocumentAnnotator._split_run_at(para_obj, end_idx, end_offset)
        return right_idx, end_idx

    def locate_text_in_paragraphs(
        self, paragraphs_info: List[Dict], target_text: str
    ) -> Optional[Dict[str, Any]]:
        """
        Step 3: 锚点定位（The Locator）

        在Word段落中查找目标文本

        策略：
        1. 先尝试精确匹配（快速）
        2. 如果失败，启用模糊匹配（容错）

        Returns:
            {
                "found": True,
                "para_index": 0,
                "position": 5,  # 在段落中的位置
                "matched_text": "实际匹配的文本",
                "para_obj": 段落对象
            }
        """
        if not target_text or len(target_text.strip()) == 0:
            return None

        # 第1步：精确匹配
        for para_info in paragraphs_info:
            para_text = para_info["text"]

            if target_text in para_text:
                position = para_text.find(target_text)
                return {
                    "found": True,
                    "para_index": para_info["index"],
                    "position": position,
                    "matched_text": target_text,
                    "para_obj": para_info.get("para_obj"),
                    "full_para_text": para_text,
                    "match_type": "exact",
                }

        # 第2步：模糊匹配（容错处理）
        logger.warning(
            f"[Annotator] ⚠️ 精确匹配失败，启用模糊匹配: {target_text[:20]}..."
        )

        best_match = None
        best_ratio = 0

        for para_info in paragraphs_info:
            para_text = para_info["text"]

            # 计算相似度
            ratio = SequenceMatcher(None, target_text, para_text).ratio()

            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para_info

        # 如果相似度超过阈值，使用模糊匹配
        if best_ratio >= self.min_similarity and best_match:
            logger.info(f"[Annotator] ✓ 模糊匹配成功 (相似度: {best_ratio:.2%})")

            return {
                "found": True,
                "para_index": best_match["index"],
                "position": 0,
                "matched_text": target_text,
                "para_obj": best_match.get("para_obj"),
                "full_para_text": best_match["text"],
                "match_type": "fuzzy",
                "similarity": best_ratio,
            }

        return None

    # ==================== Step 4: 格式无损注入 ====================

    def inject_annotation_to_paragraph(
        self,
        doc,
        para_obj,
        position: int,
        original_text: str,
        suggestion: str,
        color: str = "FF0000",  # 红色
    ) -> bool:
        """
        Step 4: 格式无损注入（The Surgery）

        在找到的位置插入标注

        默认方案A：插入 Word 原生批注（精确标注original_text片段）
        失败时回退方案B：在原文后直接插入高亮文字
        格式：原文... [建议修改: suggestion]

        Args:
            para_obj: python-docx的段落对象
            position: 在段落中的位置（从0开始）
            original_text: 原文片段（用于精确定位）
            suggestion: 修改建议
            color: 高亮颜色 (RGB hex，不含#)

        Returns:
            是否成功
        """
        try:
            if self.annotation_mode == "comment":
                # 传递original_text以实现精确标注
                return self._inject_comment(doc, para_obj, suggestion, original_text)
        except Exception as e:
            logger.warning(f"[Annotator] ⚠️ 批注注入失败，回退高亮: {str(e)}")

        # 回退为高亮文字
        try:
            from docx.shared import RGBColor

            new_run = para_obj.add_run(f" [建议修改: {suggestion}]")
            new_run.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )
            new_run.font.bold = True
            return True
        except Exception as e:
            logger.info(f"[Annotator] ✗ 注入标注失败: {str(e)}")
            return False

    @staticmethod
    def _inject_comment(
        doc, para_obj, suggestion: str, target_text: str = None
    ) -> bool:
        """插入真正的Word原生批注（右侧Comment，不影响文档格式）

        Args:
            doc: Document对象
            para_obj: 段落对象
            suggestion: 批注建议内容
            target_text: 需要标注的具体文本片段（如果提供，会精确标注该片段）
        """
        try:
            import datetime

            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn

            # 1. 获取或创建comments部分
            comments_part = None
            for rel in doc.part.rels.values():
                if "comments" in rel.target_ref:
                    comments_part = rel.target_part
                    break

            if comments_part is None:
                # 创建新的comments.xml部分
                from docx.opc.packuri import PackURI
                from docx.opc.part import XmlPart

                # 创建comments part
                partname = PackURI("/word/comments.xml")
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

                # 创建空的comments XML结构
                comments_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            mc:Ignorable="w14 w15">
</w:comments>"""

                comments_part = XmlPart.load(
                    partname,
                    content_type,
                    comments_xml.encode("utf-8"),
                    doc.part.package,
                )

                # 添加关系
                doc.part.relate_to(comments_part, RT.COMMENTS)

            # 2. 生成唯一的批注ID
            comments_root = comments_part.element
            existing_ids = [
                int(c.get(qn("w:id")))
                for c in comments_root.findall(qn("w:comment"))
                if c.get(qn("w:id"))
            ]
            comment_id = str(max(existing_ids) + 1 if existing_ids else 0)

            # 3. 在段落中插入批注标记
            p = para_obj._element

            # 尝试精确定位目标文本
            precise_location = None
            if target_text:
                precise_location = DocumentAnnotator.locate_text_in_runs(
                    para_obj, target_text
                )

            if precise_location and precise_location["found"]:
                # 精确标注：只标注具体的文本片段
                start_idx, end_idx = DocumentAnnotator._isolate_target_runs(
                    para_obj, precise_location, target_text
                )
                runs = list(p.findall(qn("w:r")))

                # 在起始Run之前插入commentRangeStart
                if start_idx < len(runs):
                    commentRangeStart = parse_xml(
                        f'<w:commentRangeStart w:id="{comment_id}" '
                        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    )
                    start_run = runs[start_idx]
                    p.insert(list(p).index(start_run), commentRangeStart)

                    # 在结束Run之后插入commentRangeEnd
                    if end_idx < len(runs):
                        commentRangeEnd = parse_xml(
                            f'<w:commentRangeEnd w:id="{comment_id}" '
                            f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        )
                        end_run = runs[end_idx]
                        p.insert(list(p).index(end_run) + 1, commentRangeEnd)

                        # 在结束标记后插入commentReference
                        commentRef_run = parse_xml(
                            f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:rPr>
                                    <w:rStyle w:val="CommentReference"/>
                                </w:rPr>
                                <w:commentReference w:id="{comment_id}"/>
                            </w:r>"""
                        )
                        p.insert(list(p).index(end_run) + 2, commentRef_run)
                        logger.info(f"[Annotator] ✅ 精确标注: {target_text[:30]}...")
                    else:
                        raise Exception("End run not found")
                else:
                    raise Exception("Start run not found")
            else:
                # 回退：标注整段
                commentRangeStart = parse_xml(
                    f'<w:commentRangeStart w:id="{comment_id}" '
                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
                p.insert(0, commentRangeStart)

                commentRangeEnd = parse_xml(
                    f'<w:commentRangeEnd w:id="{comment_id}" '
                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
                p.append(commentRangeEnd)

                commentRef_run = parse_xml(
                    f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:rPr>
                            <w:rStyle w:val="CommentReference"/>
                        </w:rPr>
                        <w:commentReference w:id="{comment_id}"/>
                    </w:r>"""
                )
                p.append(commentRef_run)
                logger.warning(f"[Annotator] ⚠️ 回退到整段标注")

            # 4. 在comments.xml中添加批注内容
            date_str = datetime.datetime.now().isoformat()

            # 转义XML特殊字符
            suggestion_escaped = (
                suggestion.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
            )

            comment_xml = parse_xml(f"""<w:comment w:id="{comment_id}" 
                              w:author="Koto AI" 
                              w:initials="K" 
                              w:date="{date_str}"
                              xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:p>
                        <w:pPr>
                            <w:pStyle w:val="CommentText"/>
                        </w:pPr>
                        <w:r>
                            <w:rPr>
                                <w:rStyle w:val="CommentReference"/>
                            </w:rPr>
                            <w:annotationRef/>
                        </w:r>
                        <w:r>
                            <w:t xml:space="preserve">{suggestion_escaped}</w:t>
                        </w:r>
                    </w:p>
                </w:comment>""")

            comments_root.append(comment_xml)

            logger.info(f"[Annotator] ✅ Word批注已添加（ID: {comment_id}）")
            return True

        except Exception as e:
            logger.warning(f"[Annotator] ⚠️ Word批注插入失败: {str(e)}")
            import traceback

            traceback.print_exc()
            return False

    # ==================== 完整工作流 ====================

    def process_document(
        self, file_path: str, annotations: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """
        处理文档：定位 -> 注入

        Args:
            file_path: 工作副本文件路径（已复制）
            annotations: 标注列表
                [
                    {
                        "原文片段": "需要修改的文本",
                        "修改建议": "建议修改为..."
                    },
                    ...
                ]

        Returns:
            {
                "success": True,
                "applied": 成功应用的标注数,
                "failed": 失败的标注数,
                "details": [...]
            }
        """
        logger.info(f"[Annotator] 🔄 开始处理文档...")

        # 第1步：提取文本
        text_data = self.extract_text_from_word(file_path)
        if not text_data.get("success"):
            return {"success": False, "error": text_data.get("error")}

        paragraphs_info = text_data["paragraphs"]
        applied_count = 0
        failed_count = 0
        details = []

        # 第2步：逐个处理标注
        from docx import Document

        doc = Document(file_path)

        for anno in annotations:
            original_text = anno.get("原文片段", "").strip()
            suggestion = anno.get("修改建议", "").strip()

            if not original_text or not suggestion:
                failed_count += 1
                details.append({"status": "skip", "reason": "缺少原文片段或修改建议"})
                continue

            # 第3步：定位文本
            location = self.locate_text_in_paragraphs(paragraphs_info, original_text)

            if not location or not location.get("found"):
                failed_count += 1
                details.append(
                    {
                        "original": original_text[:30],
                        "suggestion": suggestion[:30],
                        "status": "failed",
                        "reason": "文本未找到",
                    }
                )
                continue

            # 第4步：获取段落对象（可能在表格中）
            para_info = paragraphs_info[location["para_index"]]

            if para_info["source"] == "body":
                # 正文段落
                para_obj = doc.paragraphs[para_info["body_idx"]]
            elif para_info["source"] == "table":
                # 表格段落
                table = doc.tables[para_info["table_idx"]]
                cell = table.rows[para_info["row_idx"]].cells[para_info["cell_idx"]]
                para_obj = cell.paragraphs[para_info["cell_para_idx"]]
            else:
                failed_count += 1
                details.append(
                    {
                        "original": original_text[:30],
                        "suggestion": suggestion[:30],
                        "status": "failed",
                        "reason": "未知段落来源",
                    }
                )
                continue

            # 第5步：注入标注
            success = self.inject_annotation_to_paragraph(
                doc=doc,
                para_obj=para_obj,
                position=location["position"],
                original_text=original_text,
                suggestion=suggestion,
            )

            if success:
                applied_count += 1
                details.append(
                    {
                        "original": original_text[:30],
                        "suggestion": suggestion[:30],
                        "status": "success",
                        "match_type": location.get("match_type"),
                    }
                )
            else:
                failed_count += 1
                details.append(
                    {
                        "original": original_text[:30],
                        "suggestion": suggestion[:30],
                        "status": "failed",
                        "reason": "注入失败",
                    }
                )

        # 第5步：保存修改
        try:
            doc.save(file_path)
            logger.info(f"[Annotator] 💾 文档已保存")
        except Exception as e:
            return {"success": False, "error": f"保存文档失败: {str(e)}"}

        return {
            "success": True,
            "applied": applied_count,
            "failed": failed_count,
            "total": len(annotations),
            "success_rate": (
                f"{applied_count / len(annotations) * 100:.1f}%"
                if annotations
                else "0%"
            ),
            "details": details,
        }

    # ==================== 完整闭环 ====================

    def annotate_document(
        self, file_path: str, annotations: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """
        完整闭环：预处理 -> 定位 -> 注入 -> 保存

        Args:
            file_path: 原始Word文件路径
            annotations: 标注列表

        Returns:
            {
                "success": True,
                "original_file": "...",
                "revised_file": "...",
                "applied": 5,
                "failed": 1
            }
        """
        logger.info("=" * 60)
        logger.info("📑 Koto 文档自动标注系统")
        logger.info("=" * 60)

        # Step 1: 预处理 - 创建副本
        try:
            original_file, revised_file = self.prepare_document(file_path)
        except Exception as e:
            return {"success": False, "error": str(e)}

        # Step 2: 处理文档 - 定位 + 注入
        result = self.process_document(revised_file, annotations)

        if result.get("success"):
            result["original_file"] = original_file
            result["revised_file"] = revised_file
            logger.info(f"[Annotator] ✅ 完成！已应用 {result['applied']} 个标注")
            return result
        else:
            return result


if __name__ == "__main__":
    logger.info("文档自动标注系统已准备就绪")
