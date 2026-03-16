#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word 文档翻译模块 - 格式保留翻译
读取 .docx 文件，通过 LLM 翻译所有文本，保留原始样式（字体/加粗/颜色/表格/页眉页脚）
"""

import os
import re
import json
from typing import Optional, Callable, Generator
import logging


# 语言代码映射（用于文件名后缀和提示语）

logger = logging.getLogger(__name__)

LANG_MAP = {
    "en": "English", "english": "English", "英文": "English", "英语": "English",
    "ja": "Japanese", "japanese": "Japanese", "日文": "Japanese", "日语": "Japanese",
    "ko": "Korean", "korean": "Korean", "韩文": "Korean", "韩语": "Korean",
    "fr": "French", "french": "French", "法文": "French", "法语": "French",
    "de": "German", "german": "German", "德文": "German", "德语": "German",
    "es": "Spanish", "spanish": "Spanish", "西班牙语": "Spanish",
    "ru": "Russian", "russian": "Russian", "俄文": "Russian", "俄语": "Russian",
    "ar": "Arabic", "arabic": "Arabic", "阿拉伯语": "Arabic",
    "zh-cn": "Chinese (Simplified)", "简体中文": "Chinese (Simplified)", "中文": "Chinese (Simplified)",
    "zh-tw": "Chinese (Traditional)", "繁体中文": "Chinese (Traditional)",
    "pt": "Portuguese", "portuguese": "Portuguese", "葡萄牙语": "Portuguese",
    "it": "Italian", "italian": "Italian", "意大利语": "Italian",
    "vi": "Vietnamese", "vietnamese": "Vietnamese", "越南语": "Vietnamese",
    "th": "Thai", "thai": "Thai", "泰语": "Thai",
}

# 文件名后缀
LANG_SUFFIX = {
    "English": "en", "Japanese": "ja", "Korean": "ko", "French": "fr",
    "German": "de", "Spanish": "es", "Russian": "ru", "Arabic": "ar",
    "Chinese (Simplified)": "zh-CN", "Chinese (Traditional)": "zh-TW",
    "Portuguese": "pt", "Italian": "it", "Vietnamese": "vi", "Thai": "th",
}


def detect_target_language(user_input: str) -> str:
    """从用户输入中识别目标语言，返回标准语言名称，默认 English"""
    text = user_input.lower()
    for key, lang in LANG_MAP.items():
        if key in text:
            return lang
    # 默认英文
    return "English"


def _collect_paragraphs(doc):
    """
    收集所有需要翻译的文本段落，返回 (source_text, setter) 列表。
    setter(translated) 将翻译结果写回文档，保留原始格式。
    """
    items = []
    # Store actual element OBJECTS (not id()) to prevent lxml proxy reuse causing
    # false dedup hits when old proxies are GC'd and addresses are recycled.
    seen_elements = set()

    def _make_para_setter(para):
        def setter(translated: str):
            if not translated or not translated.strip():
                return
            runs = para.runs
            if not runs:
                para.add_run(translated)
                return
            # 找第一个有实际文字的 run（text 非空）
            first_text_run = next((r for r in runs if r.text.strip()), None)
            if first_text_run is None:
                # 所有 run 都空（只有空格/制表符），写到第一个 run 保留其样式
                first_text_run = runs[0]
            # 把翻译写入目标 run，清空其余含文字的 run
            first_text_run.text = translated
            for r in runs:
                if r is not first_text_run and r.text.strip():
                    r.text = ''
        return setter

    def _add(para):
        """去重添加段落（避免合并单元格被处理两次）"""
        elem = para._element
        if elem in seen_elements:
            return
        seen_elements.add(elem)  # keep strong ref so proxy addr is never recycled
        txt = para.text.strip()
        if txt:
            items.append((txt, _make_para_setter(para)))

    # 正文段落
    for para in doc.paragraphs:
        _add(para)

    # 表格单元格段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(para)

    # 页眉页脚
    for section in doc.sections:
        for para in section.header.paragraphs:
            _add(para)
        for para in section.footer.paragraphs:
            _add(para)

    return items


_SEP = "|||---|||"  # 分隔符，不会出现在正常文本中


def _translate_batch_llm(texts: list, target_language: str, llm_client) -> list:
    """
    用 LLM 批量翻译一组文本。返回等长翻译列表。
    使用分隔符格式（比 JSON 更可靠）。
    """
    if not texts:
        return []

    # 用 SEP 拼接，要求 LLM 按同样格式返回
    joined = f"\n{_SEP}\n".join(texts)
    prompt = (
        f"Translate each text segment below to {target_language}.\n"
        f"Keep the EXACT same number of segments.\n"
        f"Use '{_SEP}' as the separator between segments — same as the input.\n"
        f"Do NOT add explanations, numbering, or extra lines.\n"
        f"Preserve punctuation, symbols, and whitespace structure.\n\n"
        f"{joined}"
    )

    try:
        from google.genai import types as gtypes
        resp = llm_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=gtypes.GenerateContentConfig(
                temperature=0.1,
                max_output_tokens=8192,
            )
        )
        raw = (resp.text or "").strip()

        # 去掉可能的 markdown 代码块包裹
        raw = re.sub(r'^```[\w]*\n?', '', raw)
        raw = re.sub(r'```$', '', raw.strip()).strip()

        parts = [p.strip() for p in raw.split(_SEP)]

        if len(parts) == len(texts):
            logger.info(f"[DocxTranslator] ✅ 批量翻译 {len(texts)} 段成功")
            return parts

        # 长度不匹配 → 逐条翻译作为兜底
        logger.warning(f"[DocxTranslator] ⚠️ 长度不匹配 (返回{len(parts)}，期望{len(texts)})，改为逐条翻译")
        return _translate_one_by_one(texts, target_language, llm_client)

    except Exception as e:
        logger.error(f"[DocxTranslator] ❌ 批量翻译异常: {e}，改为逐条翻译")
        return _translate_one_by_one(texts, target_language, llm_client)


def _translate_one_by_one(texts: list, target_language: str, llm_client) -> list:
    """逐条翻译兜底方案，保证翻译不丢失。"""
    results = []
    for text in texts:
        try:
            from google.genai import types as gtypes
            prompt = (
                f"Translate the following text to {target_language}.\n"
                "Return ONLY the translation, nothing else.\n\n"
                f"{text}"
            )
            resp = llm_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
                config=gtypes.GenerateContentConfig(temperature=0.1, max_output_tokens=512)
            )
            translated = (resp.text or "").strip()
            results.append(translated if translated else text)
        except Exception as e:
            logger.warning(f"[DocxTranslator] ⚠️ 单条翻译失败，保留原文: {e}")
            results.append(text)
    return results


def translate_docx_streaming(
    input_path: str,
    target_language: str,
    llm_client,
    output_dir: Optional[str] = None,
    batch_size: int = 8,
) -> Generator[dict, None, None]:
    """
    翻译 .docx 文件，流式返回进度事件。
    最终事件包含 output_path。

    Yields dict with keys:
        stage: 'reading' | 'translating' | 'saving' | 'complete' | 'error'
        message: str
        progress: int (0-100)
        output_path: str (仅 complete 时)
    """
    try:
        from docx import Document
    except ImportError:
        yield {"stage": "error", "message": "缺少 python-docx，请运行 pip install python-docx", "progress": 0}
        return

    # ── 读取文档 ──────────────────────────────────────────────────────────────
    yield {"stage": "reading", "message": "📖 正在读取文档结构...", "progress": 5}
    try:
        doc = Document(input_path)
    except Exception as e:
        yield {"stage": "error", "message": f"❌ 无法打开文档: {e}", "progress": 0}
        return

    # ── 收集需翻译文本 ──────────────────────────────────────────────────────
    items = _collect_paragraphs(doc)
    total = len(items)

    if total == 0:
        yield {"stage": "error", "message": "❌ 文档中没有可翻译的文本内容", "progress": 0}
        return

    yield {
        "stage": "reading",
        "message": f"📊 共找到 {total} 段文本，目标语言: {target_language}",
        "progress": 10,
    }

    # ── 分批翻译 ──────────────────────────────────────────────────────────────
    translated_count = 0
    for batch_start in range(0, total, batch_size):
        batch = items[batch_start: batch_start + batch_size]
        texts = [item[0] for item in batch]
        setters = [item[1] for item in batch]

        progress_pct = 10 + int((batch_start / total) * 80)
        yield {
            "stage": "translating",
            "message": f"🌐 翻译第 {batch_start+1}–{min(batch_start+len(batch), total)}/{total} 段...",
            "progress": progress_pct,
        }

        translated = _translate_batch_llm(texts, target_language, llm_client)

        for setter, t in zip(setters, translated):
            setter(t)
        translated_count += len(batch)

    yield {"stage": "saving", "message": "💾 正在保存翻译文档...", "progress": 92}

    # ── 确定输出路径 ──────────────────────────────────────────────────────────
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    lang_suffix = LANG_SUFFIX.get(target_language, target_language.lower()[:5])
    out_filename = f"{base_name}_{lang_suffix}.docx"

    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(input_path))
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, out_filename)
    try:
        doc.save(output_path)
    except Exception as e:
        yield {"stage": "error", "message": f"❌ 保存文件失败: {e}", "progress": 92}
        return

    yield {
        "stage": "complete",
        "message": f"✅ 翻译完成！共翻译 {translated_count} 段文本",
        "progress": 100,
        "output_path": output_path,
        "output_filename": out_filename,
        "translated_count": translated_count,
        "target_language": target_language,
    }
