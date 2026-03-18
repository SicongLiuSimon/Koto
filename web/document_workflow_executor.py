#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档工作流执行器 - Document Workflow Executor
自动识别文档中的工作流规划并执行

功能：
1. 读取Word/PDF/Markdown文档
2. 提取工作流步骤（实验步骤、计划、流程）
3. 自动分解为可执行任务
4. 按顺序执行并收集结果
5. 生成完整的执行报告
"""

import asyncio
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


class WorkflowStep:
    """工作流步骤"""

    def __init__(
        self,
        step_id: int,
        description: str,
        step_type: str,
        input_data: Any = None,
        expected_output: str = None,
    ):
        self.step_id = step_id
        self.description = description
        self.step_type = step_type  # VLM, SEARCH, CODE, FILE_GEN, etc.
        self.input_data = input_data
        self.expected_output = expected_output
        self.status = "pending"  # pending, running, completed, failed
        self.result = None
        self.error = None
        self.start_time = None
        self.end_time = None

    def to_dict(self):
        return {
            "step_id": self.step_id,
            "description": self.description,
            "step_type": self.step_type,
            "status": self.status,
            "result": self.result,
            "error": self.error,
            "duration": self._duration(),
        }

    def _duration(self):
        if self.start_time and self.end_time:
            return (self.end_time - self.start_time).total_seconds()
        return None


class DocumentWorkflowExecutor:
    """文档工作流执行器"""

    # 关键词映射到任务类型
    KEYWORD_TO_TASK = {
        # VLM相关
        "识别": "VLM",
        "分析图片": "VLM",
        "图像识别": "VLM",
        "看图": "VLM",
        "描述图片": "VLM",
        "视觉": "VLM",
        # 搜索相关
        "搜索": "WEB_SEARCH",
        "查找": "WEB_SEARCH",
        "查询": "WEB_SEARCH",
        "检索": "WEB_SEARCH",
        # 代码执行
        "运行代码": "CODE",
        "执行代码": "CODE",
        "计算": "CODE",
        "python": "CODE",
        # 文件生成
        "生成文档": "FILE_GEN",
        "创建文件": "FILE_GEN",
        "写入文件": "FILE_GEN",
        "生成报告": "FILE_GEN",
        "做ppt": "FILE_GEN",
        "做表格": "FILE_GEN",
        # 数据处理
        "处理数据": "DATA",
        "分析数据": "DATA",
        "统计": "DATA",
        # 通用
        "比较": "COMPARE",
        "对比": "COMPARE",
        "总结": "SUMMARY",
        "汇总": "SUMMARY",
    }

    def __init__(self, client, workspace_dir: str = "workspace"):
        self.client = client
        self.workspace_dir = workspace_dir
        self.steps: List[WorkflowStep] = []
        self.workflow_name = ""
        self.workflow_context = ""

    async def load_from_document(self, file_path: str) -> Dict[str, Any]:
        """
        从文档加载工作流
        支持 .docx, .md, .txt, .json
        """
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == ".docx":
                content = self._read_docx(file_path)
            elif file_ext == ".md":
                content = self._read_text(file_path)
            elif file_ext == ".txt":
                content = self._read_text(file_path)
            elif file_ext == ".json":
                return self._load_json_workflow(file_path)
            else:
                return {"success": False, "error": f"不支持的文件类型: {file_ext}"}

            # 使用LLM提取工作流步骤
            workflow_data = await self._extract_workflow_with_llm(content, file_path)

            if workflow_data.get("success"):
                self.workflow_name = workflow_data.get(
                    "name", os.path.basename(file_path)
                )
                self.workflow_context = workflow_data.get("context", content[:500])

                # 创建工作流步骤
                for i, step_info in enumerate(workflow_data.get("steps", []), 1):
                    step = WorkflowStep(
                        step_id=i,
                        description=step_info.get("description"),
                        step_type=step_info.get("type", "GENERAL"),
                        input_data=step_info.get("input"),
                        expected_output=step_info.get("expected_output"),
                    )
                    self.steps.append(step)

            return workflow_data

        except Exception as e:
            return {"success": False, "error": f"加载文档失败: {str(e)}"}

    def _read_docx(self, file_path: str) -> str:
        """读取Word文档"""
        try:
            from docx import Document

            doc = Document(file_path)

            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            # 读取表格
            for table in doc.tables:
                content.append("\n[表格]")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    content.append(" | ".join(cells))

            return "\n".join(content)
        except ImportError:
            raise Exception("需要安装 python-docx: pip install python-docx")

    def _read_text(self, file_path: str) -> str:
        """读取文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _load_json_workflow(self, file_path: str) -> Dict[str, Any]:
        """从JSON加载预定义的工作流"""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        self.workflow_name = data.get("name", "未命名工作流")
        self.workflow_context = data.get("context", "")

        for i, step_data in enumerate(data.get("steps", []), 1):
            step = WorkflowStep(
                step_id=i,
                description=step_data.get("description"),
                step_type=step_data.get("type"),
                input_data=step_data.get("input"),
                expected_output=step_data.get("expected_output"),
            )
            self.steps.append(step)

        return {"success": True, "steps_count": len(self.steps)}

    async def _extract_workflow_with_llm(
        self, content: str, file_name: str
    ) -> Dict[str, Any]:
        """使用LLM提取工作流步骤"""

        prompt = f"""
分析以下文档内容，提取其中描述的工作流、实验步骤或执行计划。

文档名称: {file_name}
文档内容:
{content[:3000]}  # 限制长度

请识别：
1. 工作流名称和目标
2. 按顺序列出的步骤
3. 每个步骤的类型（VLM视觉分析、WEB_SEARCH搜索、CODE代码执行、FILE_GEN文件生成、DATA数据处理、COMPARE对比、SUMMARY总结等）
4. 每个步骤的输入和预期输出

以JSON格式返回：
{{
    "name": "工作流名称",
    "context": "工作流背景和目标（50字内）",
    "steps": [
        {{
            "description": "步骤描述",
            "type": "任务类型（VLM/WEB_SEARCH/CODE/FILE_GEN等）",
            "input": "输入说明",
            "expected_output": "预期输出"
        }}
    ]
}}

如果文档不包含明确的工作流，返回：
{{
    "success": false,
    "reason": "未发现明确的工作流步骤"
}}
"""

        try:
            response = self.client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.3, response_mime_type="application/json"
                ),
            )

            result = json.loads(response.text)
            result["success"] = True
            return result

        except Exception as e:
            return {"success": False, "error": f"LLM解析失败: {str(e)}"}

    async def execute_workflow(self, task_orchestrator=None) -> Dict[str, Any]:
        """
        执行整个工作流

        Args:
            task_orchestrator: TaskOrchestrator实例（从app.py传入）
        """
        logger.info(f"\n{'='*70}")
        logger.info(f"🚀 开始执行工作流: {self.workflow_name}")
        logger.info(f"{'='*70}")
        logger.info(f"📋 总步骤数: {len(self.steps)}")
        logger.info(f"📝 背景: {self.workflow_context}\n")

        results = {
            "workflow_name": self.workflow_name,
            "start_time": datetime.now().isoformat(),
            "steps": [],
            "overall_status": "running",
        }

        for step in self.steps:
            logger.info(f"\n[步骤 {step.step_id}/{len(self.steps)}] {step.description}")
            logger.info(f"└─ 类型: {step.step_type}")
            logger.info(f"   ⏳ 执行中...")

            step.status = "running"
            step.start_time = datetime.now()

            try:
                # 执行步骤
                if task_orchestrator:
                    step_result = await self._execute_step_with_orchestrator(
                        step, task_orchestrator
                    )
                else:
                    step_result = await self._execute_step_standalone(step)

                step.result = step_result
                step.status = "completed"

                logger.info(f"   ✅ 完成")
                if isinstance(step_result, dict):
                    if step_result.get("output"):
                        logger.info(
                            f"   📄 输出: {str(step_result['output'])[:100]}..."
                        )

            except Exception as e:
                step.status = "failed"
                step.error = str(e)
                logger.error(f"   ❌ 失败: {e}")

                # 可选：失败后是否继续
                if not self._should_continue_on_error():
                    results["overall_status"] = "failed"
                    break

            finally:
                step.end_time = datetime.now()
                results["steps"].append(step.to_dict())

        # 完成
        results["end_time"] = datetime.now().isoformat()
        if results["overall_status"] == "running":
            results["overall_status"] = "completed"

        # 生成总结报告
        results["summary"] = self._generate_summary(results)

        logger.info(f"\n{'='*70}")
        logger.info(f"📊 工作流执行完成")
        logger.info(f"{'='*70}")
        logger.info(
            f"✅ 成功步骤: {sum(1 for s in results['steps'] if s['status']=='completed')}/{len(self.steps)}"
        )
        logger.error(
            f"❌ 失败步骤: {sum(1 for s in results['steps'] if s['status']=='failed')}/{len(self.steps)}"
        )

        return results

    async def _execute_step_with_orchestrator(
        self, step: WorkflowStep, orchestrator
    ) -> Dict[str, Any]:
        """使用TaskOrchestrator执行步骤"""

        # 构建任务请求
        task_input = step.input_data or step.description

        # 根据步骤类型调用不同的处理方法
        if step.step_type == "VLM":
            # VLM任务
            return await self._execute_vlm_step(step, orchestrator)

        elif step.step_type == "WEB_SEARCH":
            # 搜索任务
            result = await orchestrator.handle_search(task_input)
            return {"output": result}

        elif step.step_type == "CODE":
            # 代码执行
            result = await orchestrator.handle_code_execution(task_input)
            return {"output": result}

        elif step.step_type == "FILE_GEN":
            # 文件生成
            result = await orchestrator.handle_file_generation(task_input)
            return {"output": result}

        else:
            # 通用处理
            return await self._execute_step_standalone(step)

    async def _execute_vlm_step(
        self, step: WorkflowStep, orchestrator
    ) -> Dict[str, Any]:
        """执行VLM步骤"""

        # 查找图片文件
        image_path = None
        if step.input_data and isinstance(step.input_data, str):
            if os.path.exists(step.input_data):
                image_path = step.input_data

        # 如果没有指定图片，尝试查找最近的图片
        if not image_path:
            image_path = self._find_recent_image()

        if not image_path:
            return {"success": False, "error": "未找到图片文件"}

        # 读取图片
        with open(image_path, "rb") as f:
            image_data = f.read()

        # 调用VLM
        from google.genai import types

        response = self.client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=[
                {
                    "mime_type": f"image/{'jpeg' if image_path.endswith('.jpg') else 'png'}",
                    "data": image_data,
                },
                step.description,
            ],
            config=types.GenerateContentConfig(temperature=0.7, max_output_tokens=1000),
        )

        return {"success": True, "output": response.text, "image": image_path}

    async def _execute_step_standalone(self, step: WorkflowStep) -> Dict[str, Any]:
        """独立执行步骤（不依赖orchestrator）"""

        # 简化版执行
        if step.step_type == "VLM":
            return {"output": f"[VLM] {step.description} - 需要图片输入"}

        elif step.step_type == "WEB_SEARCH":
            return {"output": f"[搜索] {step.description} - 需要搜索引擎"}

        else:
            return {"output": f"[{step.step_type}] {step.description} - 待实现"}

    def _find_recent_image(self) -> Optional[str]:
        """查找最近的图片文件"""
        image_dirs = ["workspace/images", "workspace/uploads", "uploads", "."]

        for dir_path in image_dirs:
            if not os.path.exists(dir_path):
                continue

            images = []
            for ext in [".jpg", ".jpeg", ".png", ".gif"]:
                images.extend(Path(dir_path).glob(f"**/*{ext}"))

            if images:
                # 返回最新的图片
                latest = max(images, key=lambda p: p.stat().st_mtime)
                return str(latest)

        return None

    def _should_continue_on_error(self) -> bool:
        """步骤失败后是否继续"""
        # 可配置策略
        return True  # 默认继续

    def _generate_summary(self, results: Dict[str, Any]) -> str:
        """生成执行总结"""

        total = len(results["steps"])
        completed = sum(1 for s in results["steps"] if s["status"] == "completed")
        failed = sum(1 for s in results["steps"] if s["status"] == "failed")

        summary = f"""
工作流执行总结

名称: {results['workflow_name']}
状态: {results['overall_status']}
总步骤: {total}
成功: {completed}
失败: {failed}
成功率: {completed/total*100:.1f}%

详细结果:
"""

        for step in results["steps"]:
            status_icon = "✅" if step["status"] == "completed" else "❌"
            summary += f"\n{status_icon} 步骤{step['step_id']}: {step['description']}"
            if step.get("error"):
                summary += f"\n   错误: {step['error']}"

        return summary

    async def save_results(
        self, results: Dict[str, Any], output_dir: str = "workspace/workflows"
    ) -> str:
        """保存执行结果"""

        os.makedirs(output_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"workflow_{timestamp}.json"
        output_path = os.path.join(output_dir, filename)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)

        logger.info(f"\n📁 结果已保存: {output_path}")

        # 同时生成文本报告
        report_path = output_path.replace(".json", "_report.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(results.get("summary", ""))

        logger.info(f"📄 报告已保存: {report_path}")

        return output_path


# =============== 辅助函数 ===============


async def execute_document_workflow(
    file_path: str, client, task_orchestrator=None
) -> Dict[str, Any]:
    """
    快速执行文档工作流的便捷函数

    Args:
        file_path: 文档路径
        client: Gemini客户端
        task_orchestrator: TaskOrchestrator实例（可选）

    Returns:
        执行结果字典
    """

    executor = DocumentWorkflowExecutor(client)

    # 加载工作流
    load_result = await executor.load_from_document(file_path)

    if not load_result.get("success"):
        return load_result

    logger.info(f"\n✅ 工作流加载成功")
    logger.info(f"   名称: {executor.workflow_name}")
    logger.info(f"   步骤: {len(executor.steps)}个")

    # 执行工作流
    results = await executor.execute_workflow(task_orchestrator)

    # 保存结果
    output_path = await executor.save_results(results)

    return {
        "success": True,
        "workflow_name": executor.workflow_name,
        "steps_count": len(executor.steps),
        "results": results,
        "output_path": output_path,
    }


# =============== 命令行测试 ===============

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        logger.info("用法: python document_workflow_executor.py <document_path>")
        sys.exit(1)

    file_path = sys.argv[1]

    # 简单测试（不使用orchestrator）
    async def test():
        from google import genai

        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key)

        result = await execute_document_workflow(file_path, client)

        if result.get("success"):
            logger.info(f"\n🎉 工作流执行完成！")
            logger.info(f"   结果文件: {result['output_path']}")
        else:
            logger.error(f"\n❌ 执行失败: {result.get('error')}")

    asyncio.run(test())
