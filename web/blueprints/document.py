"""
Document processing blueprint.

Routes (10 total):
  POST /api/document/smart-process        — Smart routing to annotation or analysis
  POST /api/document/feedback             — Full AI feedback loop on a document
  POST /api/document/analyze              — Analyze a document without applying changes
  POST /api/document/apply                — Apply modification suggestions
  POST /api/document/annotate             — Full annotation loop (AI → annotate → copy)
  POST /api/document/analyze-annotations  — Analyze and return annotations (non-streaming)
  POST /api/document/batch-annotate-stream — Batch annotate via SSE stream
  POST /api/document/apply-annotations    — Apply annotation suggestions
  POST /api/document/suggest-stream       — Generate modification suggestions via SSE stream
  POST /api/document/apply-suggestions    — Apply user-accepted suggestions
"""

import json
import logging
import os
import re

from flask import Blueprint, Response, jsonify, request

_logger = logging.getLogger("koto.routes.document")

document_bp = Blueprint("document", __name__)


# ---------------------------------------------------------------------------
# Lazy imports – avoids circular dependency with web.app
# ---------------------------------------------------------------------------


def _get_client():
    from web.app import client

    return client


def _get_workspace_dir():
    from web.app import WORKSPACE_DIR

    return WORKSPACE_DIR


# ---------------------------------------------------------------------------
# Internal helpers (moved verbatim from web/app.py)
# ---------------------------------------------------------------------------


def _should_use_annotation_system(requirement: str, has_file: bool = False) -> bool:
    """
    严格判断是否使用文档标注系统（在原文上标红修改）

    标注系统仅适用于：用户明确要求在原文上做标记/批注/标红/Track Changes

    注意："修改"、"优化"、"改善"等词太宽泛，不能单独触发标注。
    只有与"在原文上"、"标出来"、"标红"等定位词组合才触发。
    """
    if not requirement:
        return False

    requirement_lower = requirement.lower()

    # 第一层：明确的标注/批注关键词 — 直接触发
    explicit_annotation = [
        "标注",
        "标记",
        "批注",
        "标出",
        "标红",
        "track changes",
        "批改",
    ]
    if any(kw in requirement_lower for kw in explicit_annotation):
        return True

    # 第二层：编辑意图 + 定位词组合才触发
    # "修改"单独出现 ≠ 标注，"修改+标出来" = 标注
    edit_words = ["修改", "改正", "纠正", "校对", "审校", "纠错"]
    location_words = [
        "在原文",
        "原文上",
        "标出",
        "标记出",
        "指出.*位置",
        "哪些地方",
        "哪些位置",
    ]
    has_edit = any(kw in requirement_lower for kw in edit_words)
    has_location = any(re.search(kw, requirement_lower) for kw in location_words)

    if has_edit and has_location:
        return True

    # 第三层：审查/修改+质量描述组合
    review_words = ["审查", "评审", "审核", "改善", "优化", "修改", "润色", "调整"]
    quality_words = ["不合适", "生硬", "翻译腔", "语序", "用词", "逻辑", "问题"]
    has_review = any(kw in requirement_lower for kw in review_words)
    has_quality = any(kw in requirement_lower for kw in quality_words)

    if has_review and has_quality:
        return True

    # 默认不触发 — 宁可漏判也不误判
    return False


def _call_document_annotate(file_path: str, requirement: str):
    """调用文档标注系统"""
    try:
        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        result = feedback_system.full_annotation_loop(
            file_path=file_path,
            user_requirement=requirement,
            model_id="gemini-3.1-pro-preview",
        )

        # 添加处理模式标记
        result["processing_mode"] = "annotation"
        result["mode_description"] = "文档自动标注"

        return jsonify(result)

    except Exception as e:
        return (
            jsonify(
                {"success": False, "error": str(e), "processing_mode": "annotation"}
            ),
            500,
        )


def _call_document_analysis(file_path: str, requirement: str):
    """调用传统的文件分析系统"""
    try:
        # 这里调用现有的文件分析逻辑
        # 临时返回说明（实际应该调用现有的分析端点）
        return (
            jsonify(
                {
                    "success": False,
                    "error": "文件分析系统需要单独实现",
                    "processing_mode": "analysis",
                    "mode_description": "文件分析",
                }
            ),
            501,
        )

    except Exception as e:
        return (
            jsonify({"success": False, "error": str(e), "processing_mode": "analysis"}),
            500,
        )


# ---------------------------------------------------------------------------
# Route handlers
# ---------------------------------------------------------------------------


@document_bp.route("/api/document/smart-process", methods=["POST"])
def document_smart_process():
    """Smart document processing entry point. Automatically routes to annotation or analysis system.
    ---
    tags:
      - Documents
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          required:
            - file_path
          properties:
            file_path:
              type: string
              description: Path to the document file
            requirement:
              type: string
              default: ""
              description: Processing requirement or instruction
    responses:
      200:
        description: Processing result (varies by routing target)
        schema:
          type: object
          properties:
            success:
              type: boolean
      400:
        description: Missing file_path parameter
        schema:
          type: object
          properties:
            success:
              type: boolean
              example: false
            error:
              type: string
      500:
        description: Processing error
        schema:
          type: object
          properties:
            success:
              type: boolean
              example: false
            error:
              type: string
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        # 智能判断应该用哪个系统
        use_annotation = _should_use_annotation_system(requirement)

        _logger.debug(f"[SmartProcess] 智能判断: use_annotation={use_annotation}")
        _logger.debug(f"[SmartProcess] 需求: {requirement[:100]}")

        if use_annotation:
            # 使用文档标注系统
            _logger.debug(f"[SmartProcess] 路由到: 文档自动标注系统")
            return _call_document_annotate(file_path, requirement)
        else:
            # 使用传统的文件分析系统
            _logger.debug(f"[SmartProcess] 路由到: 文件分析系统")
            return _call_document_analysis(file_path, requirement)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/feedback", methods=["POST"])
def document_feedback():
    """文档智能反馈：读取文档 → AI分析 → 应用修改"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        auto_apply = data.get("auto_apply", True)

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        # 执行完整反馈闭环
        result = feedback_system.full_feedback_loop(
            file_path=file_path,
            user_requirement=user_requirement,
            auto_apply=auto_apply,
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/analyze", methods=["POST"])
def document_analyze():
    """Analyze a document without applying modifications.
    ---
    tags:
      - Documents
    parameters:
      - in: body
        name: body
        required: true
        schema:
          type: object
          required:
            - file_path
          properties:
            file_path:
              type: string
              description: Path to the document (relative paths are resolved under workspace/documents)
            requirement:
              type: string
              default: ""
              description: User requirement or analysis focus
    responses:
      200:
        description: Analysis result with suggestions
        schema:
          type: object
          properties:
            success:
              type: boolean
            analysis:
              type: object
              description: Document analysis and improvement suggestions
      400:
        description: Missing file_path parameter
      404:
        description: File not found
      500:
        description: Analysis error
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        # 仅分析
        result = feedback_system.analyze_and_suggest(
            file_path=file_path, user_requirement=user_requirement
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/apply", methods=["POST"])
def document_apply():
    """应用修改建议到文档"""
    try:
        data = request.json
        file_path = data.get("file_path")
        modifications = data.get("modifications", [])

        if not file_path or not modifications:
            return (
                jsonify(
                    {"success": False, "error": "缺少file_path或modifications参数"}
                ),
                400,
            )

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 应用修改
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        result = feedback_system.apply_suggestions(
            file_path=file_path, modifications=modifications
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/annotate", methods=["POST"])
def document_annotate():
    """文档自动标注：AI分析 -> 生成标注 -> 应用到副本"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        model_id = data.get("model_id", "gemini-3.1-pro-preview")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        # 执行完整标注闭环
        result = feedback_system.full_annotation_loop(
            file_path=file_path, user_requirement=user_requirement, model_id=model_id
        )

        return jsonify(result)

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/analyze-annotations", methods=["POST"])
def document_analyze_annotations():
    """仅分析文档并生成标注建议（不应用）- 已弃用，请使用 /api/document/batch-annotate-stream"""
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 使用V2批量标注系统（立即返回结果，不流式）
        from web.document_direct_edit import ImprovedBatchAnnotator

        annotator = ImprovedBatchAnnotator(gemini_client=_get_client(), batch_size=5)

        # 收集所有事件（非流式）
        events = []
        final_result = None

        for event in annotator.annotate_document_streaming(file_path, user_requirement):
            # 解析事件
            if event.startswith("event: complete"):
                data_line = event.split("\n")[1]
                if data_line.startswith("data: "):
                    final_result = json.loads(data_line[6:])
            events.append(event)

        if final_result:
            return jsonify({"success": True, **final_result})
        else:
            return jsonify({"success": False, "error": "处理失败，未收到完成事件"}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/batch-annotate-stream", methods=["POST"])
def document_batch_annotate_stream():
    """
    批量标注文档（SSE流式返回，实时反馈进度）

    接收参数:
        file_path: 文档路径
        requirement: 用户需求（可选）
        batch_size: 每批处理段落数（默认5）

    返回: SSE事件流
        event: progress - 进度更新
        event: batch_complete - 批次完成
        event: complete - 全部完成
        event: error - 错误
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")
        batch_size = data.get("batch_size", 5)

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 导入V2批量标注系统
        from web.document_batch_annotator_v2 import annotate_large_document

        # 返回SSE流
        return Response(
            annotate_large_document(
                file_path=file_path,
                user_requirement=user_requirement,
                gemini_client=_get_client(),
                batch_size=batch_size,
            ),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/apply-annotations", methods=["POST"])
def document_apply_annotations():
    """应用标注建议到文档"""
    try:
        data = request.json
        file_path = data.get("file_path")
        annotations = data.get("annotations", [])

        if not file_path or not annotations:
            return (
                jsonify({"success": False, "error": "缺少file_path或annotations参数"}),
                400,
            )

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 应用标注
        from web.document_feedback import DocumentFeedbackSystem

        feedback_system = DocumentFeedbackSystem(gemini_client=_get_client())

        result = feedback_system.annotate_document(file_path, annotations)

        return jsonify(result)

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== 改进的建议式标注 API ====================


@document_bp.route("/api/document/suggest-stream", methods=["POST"])
def document_suggest_stream():
    """
    生成修改建议流（SSE）

    请求参数:
        file_path: 文档路径
        requirement: 用户需求（可选）

    返回: SSE事件流
        event: progress - 进度
        event: suggestion - 单个建议
        event: suggestions_complete - 所有建议完成
        event: complete - 完成
    """
    try:
        data = request.json
        file_path = data.get("file_path")
        user_requirement = data.get("requirement", "")

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 使用建议式标注器
        from web.suggestion_annotator import SuggestionAnnotator

        annotator = SuggestionAnnotator(batch_size=3)

        # 返回SSE流
        return Response(
            annotator.analyze_document_streaming(file_path, user_requirement),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@document_bp.route("/api/document/apply-suggestions", methods=["POST"])
def document_apply_suggestions():
    """
    根据用户选择应用修改建议

    请求参数:
        file_path: 原始文档路径
        suggestions: 用户的选择列表
            [
                {
                    "id": "s_5_0",
                    "原文": "在被记录的",
                    "修改": "在记录的",
                    "接受": True/False
                },
                ...
            ]

    返回:
        {
            "success": True,
            "output_file": "修改后的文件路径",
            "applied_count": 实际应用的修改数,
            "accepted_count": 用户接受的数量
        }
    """
    try:
        from docx import Document

        data = request.json
        file_path = data.get("file_path")
        suggestions = data.get("suggestions", [])

        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400

        workspace_dir = _get_workspace_dir()
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(workspace_dir, "documents", file_path)

        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": f"文件不存在: {file_path}"}), 404

        # 读取文档
        doc = Document(file_path)

        # 筛选用户接受的建议
        accepted_suggestions = [s for s in suggestions if s.get("接受", False)]

        applied_count = 0

        # 应用修改（直接在段落中查找并替换）
        for suggestion in accepted_suggestions:
            original = suggestion.get("原文", "")
            modified = suggestion.get("修改", "")

            if not original or not modified:
                continue

            # 在所有段落中查找并替换
            for para in doc.paragraphs:
                if original in para.text:
                    # 替换文本
                    full_text = para.text
                    new_text = full_text.replace(original, modified, 1)

                    if new_text != full_text:
                        # 清空并重新添加（保留格式）
                        para.clear()
                        para.add_run(new_text)
                        applied_count += 1
                        break  # 每个建议只应用一次

            # 检查表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if original in para.text:
                                full_text = para.text
                                new_text = full_text.replace(original, modified, 1)
                                if new_text != full_text:
                                    para.clear()
                                    para.add_run(new_text)
                                    applied_count += 1

        # 保存为新文件
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_accepted.docx"
        doc.save(output_path)

        return jsonify(
            {
                "success": True,
                "output_file": output_path,
                "applied_count": applied_count,
                "accepted_count": len(accepted_suggestions),
                "message": f"已应用 {applied_count} 处修改（用户接受了 {len(accepted_suggestions)} 个建议）",
            }
        )

    except Exception as e:
        import traceback

        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500
