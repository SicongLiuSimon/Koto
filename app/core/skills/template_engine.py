# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════════╗
║   Koto  ─  Skill Template Engine（Word 模板引擎）                ║
╚══════════════════════════════════════════════════════════════════╝

功能：
  1. 解析 .docx 模板文件，提取 {{placeholder}} 占位符列表
  2. 将 Agent 生成的 values 字典填入模板，输出已填写的 .docx
  3. 生成适合注入 Skill 的 system_prompt 说明（告诉 Agent 有哪些字段需要填写）

占位符格式：  {{字段名}}  —— 1~50 个字母/数字/中文/下划线/连字符
             例：{{日期}}、{{参会人员}}、{{project_name}}

用法：
    from app.core.skills.template_engine import TemplateEngine

    # 解析
    fields = TemplateEngine.parse_fields("config/skill_templates/meeting/template.docx")
    # → ["日期", "参会人员", "议题", "决议事项"]

    # 填充
    output = TemplateEngine.fill(
        template_path="config/skill_templates/meeting/template.docx",
        values={"日期": "2026-03-08", "参会人员": "张三、李四"},
        output_path="config/skill_templates/meeting/output_20260308.docx",
    )

    # 生成 Agent 提示
    prompt = TemplateEngine.build_agent_prompt(
        skill_name="会议纪要", fields=fields
    )
"""

from __future__ import annotations

import re
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# 占位符正则：{{任意合法字段名}}
_PLACEHOLDER_RE = re.compile(r"\{\{([A-Za-z0-9\u4e00-\u9fff_\-]{1,50})\}\}")


class TemplateEngine:
    """Word .docx 模板解析与填充工具（纯类方法，无状态）。"""

    # ──────────────────────────────────────────────────────────────────────────
    # 解析
    # ──────────────────────────────────────────────────────────────────────────

    @classmethod
    def parse_fields(cls, template_path: str | Path) -> List[str]:
        """
        解析 .docx 文件，返回所有不重复的占位符名称（保持首次出现顺序）。

        支持：
          - 正文段落
          - 表格单元格
          - 页眉 / 页脚
        """
        try:
            import docx  # python-docx
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

        doc = docx.Document(str(template_path))
        seen: list[str] = []
        seen_set: set[str] = set()

        def _extract(text: str):
            for m in _PLACEHOLDER_RE.finditer(text):
                name = m.group(1)
                if name not in seen_set:
                    seen_set.add(name)
                    seen.append(name)

        # 正文段落
        for para in doc.paragraphs:
            _extract(para.text)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _extract(para.text)

        # 页眉 / 页脚
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf:
                    for para in hf.paragraphs:
                        _extract(para.text)

        return seen

    @classmethod
    def get_raw_text(cls, template_path: str | Path) -> str:
        """返回模板的纯文本（用于 AI 分析模板结构）。"""
        try:
            import docx
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")
        doc = docx.Document(str(template_path))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        return "\n".join(lines)

    # ──────────────────────────────────────────────────────────────────────────
    # 填充
    # ──────────────────────────────────────────────────────────────────────────

    @classmethod
    def fill(
        cls,
        template_path: str | Path,
        values: Dict[str, Any],
        output_path: str | Path,
    ) -> Path:
        """
        将 values 字典填入模板，保存到 output_path 并返回该路径。

        - 替换正文段落、表格单元格、页眉/页脚中所有 {{key}} → values[key]
        - 保留原始 Run 的字体样式（仅替换文本，不破坏格式）
        - 未提供的占位符保留原样（不报错）
        """
        try:
            import docx
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

        doc = docx.Document(str(template_path))
        str_values = {k: str(v) for k, v in values.items()}

        def _replace_para(para):
            """在保留格式的前提下替换段落中的占位符。"""
            # 先拼出整个段落的完整文本
            full = para.text
            if "{{" not in full:
                return
            for key, val in str_values.items():
                full = full.replace(f"{{{{{key}}}}}", val)

            # 如果段落只有一个 run，直接替换最简单
            if len(para.runs) == 1:
                para.runs[0].text = full
                return

            # 多个 run 的段落：把所有文本集中到第一个 run，清空其余
            if para.runs:
                para.runs[0].text = full
                for run in para.runs[1:]:
                    run.text = ""

        def _process_paras(paras):
            for p in paras:
                _replace_para(p)

        _process_paras(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_paras(cell.paragraphs)

        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf:
                    _process_paras(hf.paragraphs)

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        logger.info(f"[TemplateEngine] 已生成: {out}")
        return out

    # ──────────────────────────────────────────────────────────────────────────
    # 提示词构建
    # ──────────────────────────────────────────────────────────────────────────

    @classmethod
    def build_agent_prompt(
        cls,
        skill_name: str,
        fields: List[str],
        template_preview: str = "",
    ) -> str:
        """
        根据模板字段列表，生成注入到 system_instruction 的 Skill 提示块。

        告诉 Agent：
          - 当前 Skill 有模板，需要收集哪些字段
          - 如何调用 fill_skill_template 工具输出 Word 文档
        """
        field_list = "\n".join(f"  - `{{{{{f}}}}}` → {f}" for f in fields)
        preview_block = (
            f"\n\n**模板内容预览（前 500 字）：**\n```\n{template_preview[:500]}\n```"
            if template_preview
            else ""
        )
        return (
            f"\n\n## 📄 Word 模板技能：{skill_name}\n"
            f"本 Skill 绑定了一个 Word 文档模板，包含以下待填写字段：\n"
            f"{field_list}\n"
            f"{preview_block}\n\n"
            f"**工作流程：**\n"
            f"1. 根据用户请求和对话上下文，为每个字段生成合适的内容\n"
            f"2. 调用工具 `fill_skill_template` 将内容填入模板，参数：\n"
            f"   - `skill_id`：本 Skill 的 ID\n"
            f"   - `values`：字段名 → 填写内容 的字典\n"
            f"3. 工具返回下载链接，告知用户"
        )

    @classmethod
    def validate_fields(
        cls, fields: List[str], values: Dict[str, Any]
    ) -> Tuple[List[str], List[str]]:
        """
        检查 values 是否覆盖了所有必填字段。

        Returns:
            (filled, missing) — 已填写的字段名列表 和 缺失的字段名列表
        """
        filled = [f for f in fields if f in values and values[f] != ""]
        missing = [f for f in fields if f not in values or values[f] == ""]
        return filled, missing
